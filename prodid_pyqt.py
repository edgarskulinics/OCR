import os
import threading
import json
import tempfile
import smtplib
import threading
import time
import io
import winreg
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from PIL import Image, ImageEnhance, ImageOps, ImageTk, ImageFilter, ImageDraw, ImageFont, ImageChops
import pytesseract
import tkinter as tk
from tkinter import filedialog, messagebox, Toplevel, Scale, HORIZONTAL, VERTICAL, ttk, simpledialog
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from reportlab.pdfgen import canvas
from PIL import Image, ImageEnhance, ImageOps, ImageTk, ImageFilter, ImageDraw, ImageFont, ImageChops
import pytesseract
import tkinter as tk
from tkinter import filedialog, messagebox, Toplevel, Scale, HORIZONTAL, VERTICAL, ttk, simpledialog
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from reportlab.pdfgen import canvas
# JAUNS IMPORTS PRIEKŠ IKONĀM UN PAPILDU FUNKCIONALITĀTES
import os
import qrcode
import numpy as np
import cv2
import sys
import urllib.parse
from docx import Document
from docx.shared import Inches
import time
import pypdf
import fitz
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from scan_settings_window import ScanSettingsWindow
import random
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4, letter, landscape, portrait
from reportlab.lib.units import inch
import uuid
import qrcode
import datetime
import tkinter as tk
from tkinter import ttk, messagebox
import calendar
from datetime import datetime
from ttkbootstrap import Window
import subprocess
import uuid
import qrcode
from reportlab.graphics.barcode import code128, code39, eanbc
from reportlab.lib.units import mm
from tkcalendar import Calendar
import numpy as np
import cv2
import sys
import urllib.parse
from docx import Document
from docx.shared import Inches
import time  # Added for camera scanning stability
import pypdf  # For PDF encryption
import fitz  # PyMuPDF for PDF to image conversion
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from scan_settings_window import ScanSettingsWindow  # JAUNS IMPORTS
import random  # Pievienots random modulis


def save_user_file(file_path, content=None, mode='w'):
    """Palīgfuncija failu saglabāšanai"""
    try:
        if content is None:
            # Ja nav satura, vienkārši izveido tukšu failu
            with open(file_path, mode, encoding='utf-8') as f:
                pass
        else:
            with open(file_path, mode, encoding='utf-8') as f:
                if isinstance(content, (str, bytes)):
                    f.write(content)
                else:
                    f.write(str(content))
        return True
    except Exception as e:
        print(f"Kļūda saglabājot failu {file_path}: {e}")
        return False


def load_user_file(file_path, mode='r'):
    """Palīgfuncija failu ielādei"""
    try:
        with open(file_path, mode, encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Kļūda ielādējot failu {file_path}: {e}")
        return None


def show_loading_screen(root):
    """Parāda premium ielādes logu ar modernu dizainu"""
    loading_window = tk.Toplevel(root)
    loading_window.title("Ielādē...")
    loading_window.geometry("600x400")
    loading_window.resizable(False, False)
    loading_window.configure(bg="#222831")
    loading_window.overrideredirect(True)  # Noņemam borderus un title bar

    # Centrēšana
    loading_window.update_idletasks()
    width = loading_window.winfo_width()
    height = loading_window.winfo_height()
    x = (loading_window.winfo_screenwidth() // 2) - (width // 2)
    y = (loading_window.winfo_screenheight() // 2) - (height // 2)
    loading_window.geometry(f"+{x}+{y}")

    # Padara neaizveramu
    loading_window.grab_set()
    loading_window.protocol("WM_DELETE_WINDOW", lambda: None)  # Atspējo aizvēršanu

    # Dizaina elementi
    bg_frame = tk.Frame(loading_window, bg="#393E46")
    bg_frame.pack(fill='both', expand=True, padx=50, pady=50)

    # Logo/tituls
    logo_frame = tk.Frame(bg_frame, bg="#393E46")
    logo_frame.pack(pady=(20, 10))
    tk.Label(logo_frame,
             text="PDF Vision Pro",
             font=("Helvetica", 24, "bold"),
             fg="#00ADB5", bg="#393E46").pack()
    tk.Label(logo_frame,
             text="Profesionāls OCR un PDF apstrādes rīks",
             font=("Helvetica", 10),
             fg="#EEEEEE", bg="#393E46").pack(pady=(0, 20))

    # Progress bars
    progress_container = tk.Frame(bg_frame, bg="#393E46")
    progress_container.pack(fill='x', padx=50, pady=20)

    # Pamata progress bars
    main_progress = ttk.Progressbar(progress_container,
                                    orient="horizontal",
                                    length=400,
                                    mode="determinate")
    main_progress.pack()

    # Animēta veidņa progress bars (dekors)
    style = ttk.Style()
    style.configure("Striped.Horizontal.TProgressbar",
                    background='#00ADB5',
                    troughcolor='#393E46',
                    lightcolor='#00ADB5',
                    darkcolor='#00848B',
                    relief='flat')

    deco_progress = ttk.Progressbar(progress_container,
                                    style="Striped.Horizontal.TProgressbar",
                                    orient="horizontal",
                                    length=400,
                                    mode="determinate")
    deco_progress.pack(pady=(10, 0))

    # Statusa teksts
    status_text = tk.Label(bg_frame,
                           text="Inicializē sistēmas komponentes...",
                           font=("Helvetica", 9),
                           fg="#EEEEEE", bg="#393E46")
    status_text.pack(pady=(10, 0))

    # Progress animācija
    def update_progress():
        current = main_progress['value']
        if current < 100:
            increment = random.uniform(0.5, 3)
            main_progress['value'] = current + increment
            deco_progress['value'] = current + increment

            phases = [
                "Ielādē OCR dzinēju...",
                "Inicializē dokumenta pārvaldnieku...",
                "Konfigurē lietotāja saskarni...",
                "Gatavojas darbam...",
                "Gandrīz gatavs..."
            ]
            if current % 25 == 0 and current <= 80:
                status_text.config(text=random.choice(phases))

            loading_window.after(100, update_progress)
        else:
            # Saglabājam loading loga pozīciju
            loading_x = loading_window.winfo_x()
            loading_y = loading_window.winfo_y()

            loading_window.destroy()

            # Atver galveno logu
            root.deiconify()

            # Maksimizējam logu bez pilnekrāna (lai ir redzamas ikonas)
            root.state('zoomed')  # Maina no fullscreen uz maximizētu

            # Pārvieto uz loading loga pozīciju
            root.geometry(f"+{loading_x}+{loading_y}")

    # Sākam animāciju
    loading_window.after(2000, update_progress)  # 1 sekunde delay pirms sākas

    return loading_window

    # Centrēšana
    loading_window.update_idletasks()
    width = loading_window.winfo_width()
    height = loading_window.winfo_height()
    x = (loading_window.winfo_screenwidth() // 2) - (width // 2)
    y = (loading_window.winfo_screenheight() // 2) - (height // 2)
    loading_window.geometry(f"+{x}+{y}")

    # Padara neaizveramu
    loading_window.grab_set()
    loading_window.protocol("WM_DELETE_WINDOW", lambda: None)  # Atspējo aizvēršanu

    # Dizaina elementi
    bg_frame = tk.Frame(loading_window, bg="#393E46")
    bg_frame.pack(fill='both', expand=True, padx=50, pady=50)

    # Logo/tituls
    logo_frame = tk.Frame(bg_frame, bg="#393E46")
    logo_frame.pack(pady=(20, 10))
    tk.Label(logo_frame,
             text="PDF Vision Pro",
             font=("Helvetica", 24, "bold"),
             fg="#00ADB5", bg="#393E46").pack()
    tk.Label(logo_frame,
             text="Profesionāls OCR un PDF apstrādes rīks",
             font=("Helvetica", 10),
             fg="#EEEEEE", bg="#393E46").pack(pady=(0, 20))

    # Progress bars
    progress_container = tk.Frame(bg_frame, bg="#393E46")
    progress_container.pack(fill='x', padx=50, pady=20)

    # Pamata progress bars
    main_progress = ttk.Progressbar(progress_container,
                                    orient="horizontal",
                                    length=400,
                                    mode="determinate")
    main_progress.pack()

    # Animēta veidņa progress bars (dekors)
    style = ttk.Style()
    style.configure("Striped.Horizontal.TProgressbar",
                    background='#00ADB5',
                    troughcolor='#393E46',
                    lightcolor='#00ADB5',
                    darkcolor='#00848B',
                    relief='flat')

    deco_progress = ttk.Progressbar(progress_container,
                                    style="Striped.Horizontal.TProgressbar",
                                    orient="horizontal",
                                    length=400,
                                    mode="determinate")
    deco_progress.pack(pady=(10, 0))

    # Statusa teksts
    status_text = tk.Label(bg_frame,
                           text="Inicializē sistēmas komponentes...",
                           font=("Helvetica", 9),
                           fg="#EEEEEE", bg="#393E46")
    status_text.pack(pady=(10, 0))

    # Progress animācija
    def update_progress():
        current = main_progress['value']
        if current < 100:
            increment = random.uniform(0.5, 3)  # Nejaušs solis
            main_progress['value'] = current + increment
            deco_progress['value'] = current + increment

            # Mainam statusa tekstu
            phases = [
                "Ielādē OCR dzinēju...",
                "Inicializē dokumenta pārvaldnieku...",
                "Konfigurē lietotāja saskarni...",
                "Gatavojas darbam...",
                "Gandrīz gatavs..."
            ]
            if current % 25 == 0 and current <= 80:
                status_text.config(text=random.choice(phases))

            loading_window.after(100, update_progress)  # Atjaunina ik pēc 100ms
        else:
            loading_window.destroy()

    # Sākam animāciju
    loading_window.after(2000, update_progress)  # 1 sekunde delay pirms sākas

    return loading_window


def register_file_association():
    """Reģistrē .pdf failu asociāciju ar šo programmu"""
    try:
        app_name = "MyPDFEditor"  # Nomainiet uz jūsu programmas nosaukumu
        exe_path = sys.executable

        # Reģistrējam paplašinājumu
        with winreg.CreateKey(winreg.HKEY_CLASSES_ROOT, ".pdf") as key:
            winreg.SetValue(key, "", winreg.REG_SZ, f"{app_name}.pdf")

        # Reģistrējam komandu
        with winreg.CreateKey(winreg.HKEY_CLASSES_ROOT, f"{app_name}.pdf\\shell\\open\\command") as key:
            winreg.SetValue(key, "", winreg.REG_SZ, f'"{exe_path}" "%1"')

        # Atjauninām lietotāja iestatījumus
        subprocess.run(['assoc', '.pdf={app_name}.pdf'], shell=True)
        subprocess.run(['ftype', f'{app_name}.pdf="{exe_path}" "%1"'], shell=True)

        messagebox.showinfo("Sekmēs", "PDF asociācijas veiksmīgi iestatītas")
    except Exception as e:
        messagebox.showerror("Kļūda", f"Neizdevās reģistrēt asociācijas:\n{e}")


# Pārbauda un importē OpenCV un NumPy slīpuma korekcijai
OPENCV_AVAILABLE = False
try:
    import numpy as np
    import cv2

    OPENCV_AVAILABLE = True
except ImportError:
    print(
        "OpenCV and NumPy not found. Deskew, HSV conversion, stitching, inpainting, face detection, and camera scanning functionality will be disabled.")

# Tesseract ceļš - pielāgo savam datoram
import sys
import os


# Funkcija, lai iegūtu resursu ceļu, neatkarīgi no tā, vai programma darbojas kā .exe vai Python skripts
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# Noklusējuma Tesseract ceļš. Pieņemam, ka tesseract.exe atradīsies mapē "Tesseract-OCR" blakus EXE.
# Ja tesseract.exe ir tieši blakus EXE, tad vienkārši "tesseract.exe"
# Tesseract konfigurācija portatīvajai versijai
DEFAULT_TESSERACT_CMD = resource_path(os.path.join("Tesseract-OCR", "tesseract.exe"))


# Konfigurē pytesseract, lai izmantotu portatīvo Tesseract
def configure_tesseract():
    """Konfigurē Tesseract portatīvajai izmantošanai"""
    tesseract_path = resource_path(os.path.join("Tesseract-OCR", "tesseract.exe"))
    tessdata_path = resource_path(os.path.join("Tesseract-OCR", "tessdata"))

    # Iestatām Tesseract ceļu
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    # Iestatām TESSDATA_PREFIX vides mainīgo
    os.environ['TESSDATA_PREFIX'] = tessdata_path

    return tesseract_path, tessdata_path


pytesseract.pytesseract.tesseract_cmd = DEFAULT_TESSERACT_CMD

# Definējiet ceļu uz iestatījumu failu
APP_SETTINGS_FILE = "app_settings.json"  # Varat mainīt uz citu ceļu, ja nepieciešams

# Noklusējuma kameras indekss
DEFAULT_CAMERA_INDEX = 1


class SettingsWindow(Toplevel):
    """Paplašināta iestatījumu klase ar e-pasta iestatījumiem un skenēšanas iestatījumiem"""

    def __init__(self, master, app_instance):
        super().__init__(master)
        self.app = app_instance
        self.title("Vispārīgie Iestatījumi")  # MAINĪTS TEKSTS
        # Pielāgo iestatījumu loga izmēru, lai tas labāk ietilptu mazākos ekrānos
        # Var izmantot arī relatīvus izmērus, piemēram, 80% no galvenā loga izmēra
        self.geometry("1000x900")  # Palielināts iestatījumu loga izmērs
        self.minsize(800, 800)  # Pievienots minimālais izmērs iestatījumu logam
        self.transient(master)
        self.grab_set()

        self.create_widgets()
        self.load_current_settings()

    def toggle_id_code_options(self):
        """Ieslēdz/izslēdz ID koda opciju rāmi atkarībā no izvēles rūtiņas stāvokļa."""
        if self.add_id_code_var.get():
            self.id_code_options_frame.grid()
        else:
            self.id_code_options_frame.grid_remove()

    def init_camera(self, force_camera_index=None):
        """Kameras inicializācija ar konkrētu kameras indeksu."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas", "Nepieciešams opencv-python.")
            return False

        # Ja kamera jau ir atvērta un nav pieprasīta konkrēta kamera
        if self.camera is not None and force_camera_index is None:
            return True

        # Atbrīvo esošo kameru, ja vajag mainīt
        if self.camera is not None:
            self.camera.release()
            self.camera = None
            self.camera_active = False

        try:
            # Nosaka kameras indeksu
            if force_camera_index is not None:
                camera_index = force_camera_index
                print(f"🎯 Piespiedu kārtā izmanto kameru: {camera_index}")
            elif hasattr(self, 'scan_camera_index'):
                camera_index = self.scan_camera_index.get()
                print(f"📋 Iestatījumos norādītā kamera: {camera_index}")
            else:
                camera_index = 0
                print("⚠️ Nav atrasts scan_camera_index, izmanto 0")

            print(f"🔍 Mēģina atvērt kameru {camera_index}")

            # Atver norādīto kameru
            self.camera = cv2.VideoCapture(camera_index)
            if not self.camera.isOpened():
                print(f"❌ Kamera {camera_index} nav pieejama")
                raise IOError(f"Kamera {camera_index} nav pieejama")

            print(f"✅ Veiksmīgi atvērta kamera {camera_index}")

            # Iestata kvalitāti
            self.camera.set(cv2.CAP_PROP_FRAME_WIDTH, 1920)
            self.camera.set(cv2.CAP_PROP_FRAME_HEIGHT, 1080)
            self.camera.set(cv2.CAP_PROP_BUFFERSIZE, 1)

            # Saglabā pašreizējo kameras indeksu
            self.current_camera_index = camera_index

            actual_width = int(self.camera.get(cv2.CAP_PROP_FRAME_WIDTH))
            actual_height = int(self.camera.get(cv2.CAP_PROP_FRAME_HEIGHT))
            print(f"📐 Kamera {camera_index}: {actual_width}x{actual_height}")

            self.camera_active = True
            return True

        except Exception as e:
            print(f"❌ Kameras {camera_index} kļūda: {e}")
            messagebox.showerror("Kameras kļūda", f"Nevar atvērt kameru {camera_index}: {e}")
            if self.camera:
                self.camera.release()
            self.camera = None
            self.camera_active = False
            return False

    def release_camera(self):
        """Atbrīvo kameras resursus."""
        if self.camera is not None:
            self.camera.release()
            self.camera = None
            self.camera_active = False

    def get_camera_frame(self):
        """Ātri iegūst kameras kadru priekšskatījumam."""
        if self.camera is None or not self.camera_active:
            return None

        try:
            ret, frame = self.camera.read()
            if not ret:
                return None

            # Samazina tikai priekšskatījumam (ātrāk)
            height, width = frame.shape[:2]
            if width > 800:  # Samazina tikai ja pārāk liels
                scale = 800 / width
                new_width = int(width * scale)
                new_height = int(height * scale)
                frame = cv2.resize(frame, (new_width, new_height), interpolation=cv2.INTER_LINEAR)

            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            return Image.fromarray(frame_rgb)
        except Exception as e:
            print(f"Kadra kļūda: {e}")
            return None

    def _select_folder_dialog(self, root_folder):
        """Atver dialogu mapju izvēlei."""
        dialog = Toplevel(self)
        dialog.title("Izvēlieties mērķa mapi")
        dialog.geometry("400x300")
        dialog.transient(self)
        dialog.grab_set()

        selected_folder = [None]  # Izmanto sarakstu, lai varētu mainīt no nested funkcijas

        # Treeview mapju attēlošanai
        tree = ttk.Treeview(dialog)
        tree.pack(fill="both", expand=True, padx=10, pady=10)

        def populate_tree(parent_item, folder_node):
            """Rekursīvi aizpilda koku ar mapēm."""
            for item in folder_node.get("contents", []):
                if item["type"] == "folder":
                    item_id = tree.insert(parent_item, "end", text=item["name"], values=[id(item)])
                    populate_tree(item_id, item)

        # Aizpilda koku
        root_id = tree.insert("", "end", text="Sakne", values=[id(root_folder)])
        populate_tree(root_id, root_folder)
        tree.item(root_id, open=True)

        def on_select():
            selection = tree.selection()
            if selection:
                item_id = selection[0]
                folder_id = tree.item(item_id, "values")[0]
                # Atrod mapi pēc ID
                selected_folder[0] = find_folder_by_id(root_folder, int(folder_id))
            dialog.destroy()

        def find_folder_by_id(folder, target_id):
            """Atrod mapi pēc ID."""
            if id(folder) == target_id:
                return folder
            for item in folder.get("contents", []):
                if item["type"] == "folder":
                    result = find_folder_by_id(item, target_id)
                    if result:
                        return result
            return None

        # Pogas
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill="x", padx=10, pady=5)

        ttk.Button(button_frame, text="Izvēlēties", command=on_select).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Atcelt", command=dialog.destroy).pack(side="right")

        dialog.wait_window()
        return selected_folder[0]

    def _is_descendant(self, potential_ancestor, potential_descendant):
        """Pārbauda, vai potential_descendant ir potential_ancestor apakšmape."""
        current = potential_ancestor
        while current:
            if current == potential_descendant:
                return True
            current = current.get("parent")
        return False

    def save_as_word(self):
        """Saglabā atlasīto PDF kā Word dokumentu."""
        selection = self.pdf_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet PDF failu.")
            return

        try:
            # Vienkārša implementācija - var uzlabot
            messagebox.showinfo("Funkcija", "Word eksportēšana vēl nav pilnībā implementēta.")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās saglabāt kā Word: {e}")

    def drag_start(self, event):
        """Sāk drag operāciju."""
        self.drag_data["x"] = event.x
        self.drag_data["y"] = event.y
        self.drag_data["item"] = self.pdf_listbox.nearest(event.y)

    def drag_motion(self, event):
        """Apstrādā vilkšanas kustību `pdf_listbox`."""
        if self.drag_data["item_index"] is not None:
            # Iegūst jauno pozīciju
            new_index = self.pdf_listbox.nearest(event.y)
            current_index = self.drag_data["item_index"]

            if new_index != current_index:
                # Pārvieto elementu pamatā esošajā datu struktūrā
                item_to_move = self.current_folder["contents"].pop(current_index)
                self.current_folder["contents"].insert(new_index, item_to_move)

                # Atjaunina vilkšanas datus ar jauno indeksu
                self.drag_data["item_index"] = new_index

                # Atjauno listbox vizuālo attēlojumu
                self.refresh_pdf_list()
                # Pārliecinās, ka pārvietotais elements joprojām ir atlasīts
                self.pdf_listbox.selection_set(new_index)
                self.pdf_listbox.activate(new_index)

    def drag_drop(self, event):
        """Beidz drag operāciju."""
        target_index = self.pdf_listbox.nearest(event.y)
        source_index = self.drag_data["item"]

        if source_index != target_index and 0 <= target_index < len(self.current_folder["contents"]):
            # Pārvieto elementu sarakstā
            item = self.current_folder["contents"].pop(source_index)
            self.current_folder["contents"].insert(target_index, item)
            self.refresh_pdf_list()


    def create_widgets(self):
        """Izveido iestatījumu loga elementus"""
        main_frame = ttk.Frame(self, padding=(10, 10))
        main_frame.pack(fill=BOTH, expand=True)

        # Izveido notebook ar vairākām cilnēm
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=BOTH, expand=True, padx=5, pady=5)

        # Galvenie iestatījumi
        general_frame = ttk.Frame(notebook, padding=10)
        notebook.add(general_frame, text="Vispārīgi")

        # OCR iestatījumi
        ocr_frame = ttk.Frame(notebook, padding=10)
        notebook.add(ocr_frame, text="OCR")

        # PDF iestatījumi
        pdf_frame = ttk.Frame(notebook, padding=10)
        notebook.add(pdf_frame, text="PDF")

        # E-pasta iestatījumi
        email_frame = ttk.Frame(notebook, padding=10)
        notebook.add(email_frame, text="E-pasts")

        # Aizpilda iestatījumu rāmjus
        self.create_general_settings(general_frame)
        self.create_ocr_settings(ocr_frame)
        self.create_pdf_settings(pdf_frame)
        self.create_email_settings(email_frame)
        # self.create_scan_settings(scan_frame) # IZDZĒSTS: Skenēšanas iestatījumu cilne

        # Pogu rāmis
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=X, pady=(10, 0))

        ttk.Button(button_frame, text="Saglabāt", command=self.save_settings,
                   bootstyle=PRIMARY).pack(side=RIGHT, padx=5)
        ttk.Button(button_frame, text="Atcelt", command=self.destroy,
                   bootstyle=SECONDARY).pack(side=RIGHT, padx=5)

    def create_general_settings(self, frame):
        """Izveido vispārīgos iestatījumus"""
        ttk.Label(frame, text="Programmas tēma:").grid(row=0, column=0, sticky=W, pady=2)
        self.theme_var = tk.StringVar()
        theme_combo = ttk.Combobox(frame, textvariable=self.theme_var,
                                   values=self.app.style.theme_names())
        theme_combo.grid(row=0, column=1, sticky=EW, padx=5, pady=2)
        theme_combo.bind("<<ComboboxSelected>>", self.change_app_theme)

        ttk.Label(frame, text="Noklusējuma saglabāšanas mape:").grid(row=1, column=0, sticky=W, pady=2)
        self.save_path_var = tk.StringVar()
        save_path_entry = ttk.Entry(frame, textvariable=self.save_path_var, width=40)
        save_path_entry.grid(row=1, column=1, sticky=EW, padx=5, pady=2)
        ttk.Button(frame, text="Pārlūkot...", command=self.browse_save_path).grid(row=1, column=2, padx=5)

        ttk.Label(frame, text="Programmas tēma:").grid(row=0, column=0, sticky=W, pady=2)
        self.theme_var = tk.StringVar()
        theme_combo = ttk.Combobox(frame, textvariable=self.theme_var,
                                   values=self.app.style.theme_names())
        theme_combo.grid(row=0, column=1, sticky=EW, padx=5, pady=2)
        theme_combo.bind("<<ComboboxSelected>>", self.change_app_theme)
        # Pievienojiet jaunu iestatījumu par autentifikāciju
        ttk.Label(frame, text="Iespējot autentifikāciju:").grid(row=1, column=0, sticky=W, pady=2)
        self.enable_auth_var = tk.BooleanVar(value=True)  # Noklusējuma vērtība
        ttk.Checkbutton(frame, variable=self.enable_auth_var).grid(row=1, column=1, sticky=W, padx=5)

        # Piešķir visiem elementiem vienādu svaru
        frame.columnconfigure(1, weight=1)

    def create_ocr_settings(self, frame):
        """Izveido OCR iestatījumus"""
        ttk.Label(frame, text="Tesseract ceļš:").grid(row=0, column=0, sticky=W, pady=2)
        self.tesseract_var = tk.StringVar()
        tesseract_entry = ttk.Entry(frame, textvariable=self.tesseract_var, width=40)
        tesseract_entry.grid(row=0, column=1, sticky=EW, padx=5, pady=2)
        ttk.Button(frame, text="Pārlūkot...", command=self.browse_tesseract_path).grid(row=0, column=2, padx=5)

        ttk.Label(frame, text="OCR valodas:").grid(row=1, column=0, sticky=NW, pady=5)
        lang_frame = ttk.Frame(frame)
        lang_frame.grid(row=1, column=1, columnspan=2, sticky=EW, pady=2)

        self.lang_vars = {}
        col = 0
        for lang, code in self.app.lang_options.items():
            var = tk.BooleanVar()
            cb = ttk.Checkbutton(lang_frame, text=lang, variable=var)
            cb.grid(row=0, column=col, sticky=W, padx=5)
            self.lang_vars[lang] = var
            col += 1

        ttk.Label(frame, text="Pārējie OCR parametri:").grid(row=2, column=0, sticky=W, pady=5)

        param_frame = ttk.Frame(frame)
        param_frame.grid(row=3, column=0, columnspan=3, sticky=EW)

        ttk.Label(param_frame, text="DPI:").grid(row=0, column=0, sticky=W)
        self.dpi_var = tk.IntVar()
        ttk.Spinbox(param_frame, from_=70, to=600, increment=10, textvariable=self.dpi_var, width=5).grid(row=0,
                                                                                                          column=1,
                                                                                                          padx=5)

        ttk.Label(param_frame, text="Confidence:").grid(row=0, column=2, sticky=W)
        self.conf_var = tk.IntVar()
        ttk.Spinbox(param_frame, from_=0, to=100, increment=5, textvariable=self.conf_var, width=5).grid(row=0,
                                                                                                         column=3,
                                                                                                         padx=5)

        frame.columnconfigure(1, weight=1)

    def create_pdf_settings(self, frame):
        """Izveido PDF iestatījumus"""
        ttk.Label(frame, text="PDF izvades kvalitāte:").grid(row=0, column=0, sticky=W, pady=2)
        self.pdf_qual_var = tk.StringVar()
        pdf_qual_combo = ttk.Combobox(frame, textvariable=self.pdf_qual_var,
                                      values=["Zema (60)", "Vidēja (85)", "Augsta (95)"])
        pdf_qual_combo.grid(row=0, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="Lapas izmērs:").grid(row=1, column=0, sticky=W, pady=2)
        self.page_size_var = tk.StringVar()
        page_size_combo = ttk.Combobox(frame, textvariable=self.page_size_var,
                                       values=self.app.orientation_options)
        page_size_combo.grid(row=1, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="Fonta izmērs:").grid(row=2, column=0, sticky=W, pady=2)
        self.font_size_var = tk.IntVar()
        ttk.Spinbox(frame, from_=5, to=20, increment=1, textvariable=self.font_size_var, width=5).grid(row=2, column=1,
                                                                                                       sticky=W, padx=5)

        ttk.Checkbutton(frame, text="Iekļaut meklējamo tekstu PDF", variable=tk.BooleanVar(value=True)).grid(row=3,
                                                                                                             column=0,
                                                                                                             columnspan=2,
                                                                                                             sticky=W,
                                                                                                             pady=5)

        # JAUNS: QR koda/Svītrkoda iestatījumi
        ttk.Label(frame, text="Pievienot ID kodu PDF:").grid(row=4, column=0, sticky=W, pady=2)
        self.add_id_code_var = tk.BooleanVar()
        ttk.Checkbutton(frame, variable=self.add_id_code_var, command=self.toggle_id_code_options).grid(row=4, column=1,
                                                                                                        sticky=W,
                                                                                                        padx=5, pady=2)

        self.id_code_options_frame = ttk.Frame(frame)
        self.id_code_options_frame.grid(row=5, column=0, columnspan=2, sticky=EW, padx=5, pady=2)

        ttk.Label(self.id_code_options_frame, text="Koda tips:").grid(row=0, column=0, sticky=W, pady=2)
        self.id_code_type_var = tk.StringVar()
        ttk.Radiobutton(self.id_code_options_frame, text="QR kods", variable=self.id_code_type_var, value="QR").grid(
            row=0, column=1, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Svītrkods (Code 128)", variable=self.id_code_type_var,
                        value="Barcode").grid(row=0, column=2, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Svītrkods (Code 39)", variable=self.id_code_type_var,
                        value="Code39").grid(row=0, column=3, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Svītrkods (EAN-13)", variable=self.id_code_type_var,
                        value="EAN13").grid(row=0, column=4, sticky=W, padx=5)

        ttk.Label(self.id_code_options_frame, text="Koda pozīcija:").grid(row=1, column=0, sticky=W, pady=2)
        self.id_code_position_var = tk.StringVar()
        ttk.Radiobutton(self.id_code_options_frame, text="Augšā pa labi", variable=self.id_code_position_var,
                        value="top_right").grid(row=1, column=1, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Apakšā pa labi", variable=self.id_code_position_var,
                        value="bottom_right").grid(row=1, column=2, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Apakšā pa kreisi", variable=self.id_code_position_var,
                        value="bottom_left").grid(row=1, column=3, sticky=W, padx=5)
        ttk.Radiobutton(self.id_code_options_frame, text="Augšā pa kreisi", variable=self.id_code_position_var,
                        value="top_left").grid(row=1, column=4, sticky=W, padx=5)

        frame.columnconfigure(1, weight=1)
        self.toggle_id_code_options()  # Sākotnējā stāvokļa iestatīšana

    def create_email_settings(self, frame):
        """Izveido e-pasta iestatījumus"""
        ttk.Label(frame, text="SMTP serveris:").grid(row=0, column=0, sticky=W, pady=2)
        self.smtp_server_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.smtp_server_var).grid(row=0, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="SMTP ports:").grid(row=1, column=0, sticky=W, pady=2)
        self.smtp_port_var = tk.IntVar()
        ttk.Entry(frame, textvariable=self.smtp_port_var).grid(row=1, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="Lietotājvārds:").grid(row=2, column=0, sticky=W, pady=2)
        self.email_user_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.email_user_var).grid(row=2, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="Parole:").grid(row=3, column=0, sticky=W, pady=2)
        self.email_pass_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.email_pass_var, show="*").grid(row=3, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="No adreses:").grid(row=4, column=0, sticky=W, pady=2)
        self.from_email_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.from_email_var).grid(row=4, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="Uz adresi (noklusējums):").grid(row=5, column=0, sticky=W, pady=2)
        self.to_email_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.to_email_var).grid(row=5, column=1, sticky=EW, padx=5, pady=2)

        self.use_ssl_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(frame, text="Izmantot SSL", variable=self.use_ssl_var).grid(row=6, column=0,
                                                                                    columnspan=2, sticky=W,
                                                                                    pady=5)

        ttk.Label(frame, text="E-pasta tēma:").grid(row=7, column=0, sticky=W, pady=2)
        self.email_subject_var = tk.StringVar()
        ttk.Entry(frame, textvariable=self.email_subject_var).grid(row=7, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="E-pasta teksts (Plain):").grid(row=8, column=0, sticky=NW, pady=2)
        self.email_body_plain_text = tk.Text(frame, height=5, width=40, wrap="word")
        self.email_body_plain_text.grid(row=8, column=1, sticky=EW, padx=5, pady=2)

        ttk.Label(frame, text="E-pasta teksts (HTML):").grid(row=9, column=0, sticky=NW, pady=2)
        self.email_body_html_text = tk.Text(frame, height=5, width=40, wrap="word")
        self.email_body_html_text.grid(row=9, column=1, sticky=EW, padx=5, pady=2)

        # Pārbaudes poga
        ttk.Button(frame, text="Pārbaudīt savienojumu", command=self.test_email_settings,
                   bootstyle=INFO).grid(row=10, column=0, columnspan=2, pady=10)

        frame.columnconfigure(1, weight=1)

    def test_email_settings(self):
        """Pārbauda e-pasta savienojumu ar norādītajiem iestatījumiem"""
        try:
            server = self.smtp_server_var.get()
            port = self.smtp_port_var.get()
            username = self.email_user_var.get()
            password = self.email_pass_var.get()
            use_ssl = self.use_ssl_var.get()

            if not server or not port or not username or not password:
                messagebox.showwarning("Trūkst datu", "Lūdzu, aizpildiet visus laukus!")
                return

            if use_ssl:
                smtp = smtplib.SMTP_SSL(server, port)
            else:
                smtp = smtplib.SMTP(server, port)
                smtp.starttls()

            smtp.login(username, password)
            smtp.quit()
            messagebox.showinfo("Pārbaude", "Savienojums ar SMTP serveri veiksmīgs!")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās pieslēgties SMTP serverim:\n{str(e)}")

    def browse_save_path(self):
        """Atver dialogu, lai izvēlētos saglabāšanas ceļu"""
        path = filedialog.askdirectory(title="Izvēlieties noklusējuma saglabāšanas mapi")
        if path:
            self.save_path_var.set(path)

    def browse_tesseract_path(self):
        """Atver dialogu, lai izvēlētos Tesseract ceļu"""
        path = filedialog.askopenfilename(title="Izvēlieties Tesseract.exe",
                                          filetypes=[("Executable files", "*.exe")])
        if path:
            self.tesseract_var.set(path)

    def change_app_theme(self, event=None):
        """Maina lietotnes tēmu"""
        selected_theme = self.theme_var.get()
        self.app.style.theme_use(selected_theme)
        if hasattr(self.app, 'current_image_index') and self.app.current_image_index != -1:
            self.app.on_file_select()

    def load_current_settings(self):
        """Ielādē pašreizējos iestatījumus"""
        self.theme_var.set(self.app.style.theme_use())
        self.save_path_var.set(self.app.default_save_path)
        self.tesseract_var.set(pytesseract.pytesseract.tesseract_cmd)
        self.pdf_qual_var.set(self.app.pdf_quality)
        self.page_size_var.set(self.app.orientation_var.get())
        self.font_size_var.set(self.app.fontsize_var.get())
        self.dpi_var.set(self.app.dpi_var.get())
        self.conf_var.set(self.app.confidence_var.get())

        # E-pasta iestatījumi
        self.smtp_server_var.set(self.app.settings.get("smtp_server", ""))
        self.smtp_port_var.set(self.app.settings.get("smtp_port", 465))
        self.email_user_var.set(self.app.settings.get("email_user", ""))
        self.email_pass_var.set(self.app.settings.get("email_pass", ""))
        self.from_email_var.set(self.app.settings.get("from_email", ""))
        self.to_email_var.set(self.app.settings.get("to_email", ""))
        self.use_ssl_var.set(self.app.settings.get("use_ssl", True))
        self.email_subject_var.set(self.app.settings.get("email_subject", "OCR PDF dokumenti"))
        self.email_body_plain_text.delete("1.0", tk.END)
        self.email_body_plain_text.insert(tk.END, self.app.settings.get("email_body_plain",
                                                                        "Sveiki,\n\nPielikumā atradīsiet OCR apstrādātos PDF dokumentus.\n\nAr cieņu,\nJūsu OCR PDF App"))
        self.email_body_html_text.delete("1.0", tk.END)
        self.email_body_html_text.insert(tk.END, self.app.settings.get("email_body_html",
                                                                       "<html><body><p>Sveiki,</p><p>Pielikumā atradīsiet OCR apstrādātos PDF dokumentus.</p><p>Ar cieņu,<br/>Jūsu OCR PDF App</p></body></html>"))

        # Valodu atzīmēšana
        # JAUNS: ID koda iestatījumi
        self.add_id_code_var.set(self.app.settings.get("add_id_code_to_pdf", False))
        self.id_code_type_var.set(self.app.settings.get("id_code_type", "QR"))
        self.id_code_position_var.set(self.app.settings.get("id_code_position", "bottom_right"))
        self.toggle_id_code_options()  # Atjaunina redzamību

        for lang_name, var in self.lang_vars.items():
            if lang_name in self.app.lang_vars:
                var.set(self.app.lang_vars[lang_name].get())

    def save_settings(self):
        """Saglabā iestatījumus"""
        # Galvenie iestatījumi
        self.app.style.theme_use(self.theme_var.get())
        self.app.default_save_path = self.save_path_var.get()

        # OCR iestatījumi
        pytesseract.pytesseract.tesseract_cmd = self.tesseract_var.get()
        self.app.dpi_var.set(self.dpi_var.get())
        self.app.confidence_var.set(self.conf_var.get())

        # PDF iestatījumi
        self.app.pdf_quality = self.pdf_qual_var.get()
        self.app.orientation_var.set(self.page_size_var.get())
        self.app.fontsize_var.set(self.font_size_var.get())

        self.app.scan_camera_index.set(self.camera_index_var.get())
        self.app.scan_camera_width.set(self.camera_width_var.get())
        self.app.scan_camera_height.set(self.camera_height_var.get())
        self.app.scan_min_contour_area.set(self.min_contour_area_var.get())

        # E-pasta iestatījumi
        self.app.settings["smtp_server"] = self.smtp_server_var.get()
        self.app.settings["smtp_port"] = self.smtp_port_var.get()
        self.app.settings["email_user"] = self.email_user_var.get()
        self.app.settings["email_pass"] = self.email_pass_var.get()
        self.app.settings["from_email"] = self.from_email_var.get()
        self.app.settings["to_email"] = self.to_email_var.get()
        self.app.settings["use_ssl"] = self.use_ssl_var.get()
        self.app.settings["email_subject"] = self.email_subject_var.get()
        self.app.settings["email_body_plain"] = self.email_body_plain_text.get("1.0", tk.END).strip()
        self.app.settings["email_body_html"] = self.email_body_html_text.get("1.0", tk.END).strip()

        # JAUNS: ID koda iestatījumi
        self.app.settings["add_id_code_to_pdf"] = self.add_id_code_var.get()
        self.app.settings["id_code_type"] = self.id_code_type_var.get()
        self.app.settings["id_code_position"] = self.id_code_position_var.get()

        # Valodu atzīmēšana

        for lang_name, var in self.lang_vars.items():
            if lang_name in self.app.lang_vars:
                self.app.lang_vars[lang_name].set(var.get())

        self.app.save_app_settings()
        messagebox.showinfo("Iestatījumi", "Iestatījumi veiksmīgi saglabāti!")
        self.destroy()

    def create_image_processing_widgets(self, parent):
        """Izveido attēlu apstrādes komponentus."""
        params_frame = ttk.Frame(parent)
        params_frame.grid(row=0, column=0, padx=10, pady=10)
        # DPI iestatījums
        self.dpi_var = tk.IntVar(value=300)  # Noklusējuma DPI
        ttk.Label(params_frame, text="DPI:").grid(row=0, column=0, sticky=tk.W)
        ttk.Spinbox(params_frame, from_=70, to=600, increment=10, textvariable=self.dpi_var, width=4).grid(row=0,

                                                                                                           column=1)
        # Pievienojiet citus komponentus šeit


class PDFEditor:
    def __init__(self, filepath):
        self.filepath = filepath
        self.document = None
        self.open()

    def open(self):
        if self.document is None:
            self.document = fitz.open(self.filepath)

    def close(self):
        if self.document:
            self.document.close()
            self.document = None

    def __enter__(self):
        self.open()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    def get_page_count(self):
        return len(self.document)

    def get_page_image(self, page_num, dpi=72):
        page = self.document.load_page(page_num)
        pix = page.get_pixmap(dpi=dpi)
        return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # Pievienojiet citas metodes šeit...


class FullscreenImageViewer(Toplevel):
    """Pilnekrāna attēlu skatītājs ar tālummaiņas un pārvietošanas funkcijām."""

    def __init__(self, master, image_pil):
        super().__init__(master)
        self.title("Attēla priekšskatījums")
        # Sākotnējais izmērs kā galvenajam logam, maksimizē logu
        self.geometry(f"{master.winfo_width()}x{master.winfo_height()}+0+0")
        self.state('zoomed')
        self.transient(master)
        self.grab_set()

        self.image_pil = image_pil
        self.zoom_factor = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.start_pan_x = 0
        self.start_pan_y = 0

        self.canvas = tk.Canvas(self, bg="black")
        self.canvas.pack(fill="both", expand=True)

        self.canvas.bind("<Configure>", self.on_canvas_resize)
        self.canvas.bind("<MouseWheel>", self.on_mouse_wheel)  # Windows/Linux
        self.canvas.bind("<Button-4>", self.on_mouse_wheel)  # MacOS
        self.canvas.bind("<Button-5>", self.on_mouse_wheel)  # MacOS
        self.canvas.bind("<ButtonPress-2>", self.on_pan_start)  # Middle mouse button
        self.canvas.bind("<B2-Motion>", self.on_pan_drag)
        self.canvas.bind("<ButtonRelease-2>", self.on_pan_end)

        self.canvas.bind("<ButtonPress-1>", self.on_selection_start)
        self.canvas.bind("<B1-Motion>", self.on_selection_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_selection_end)
        self.selection_rect = None
        self.selection_start_x = None
        self.selection_start_y = None

        self.display_image()

    def display_image(self):
        """Attēlo attēlu uz kanvasa, pielāgojot tālummaiņu un pārvietošanu."""
        if not hasattr(self, 'canvas') or not self.canvas.winfo_exists():
            return

        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()

        if canvas_width <= 1 or canvas_height <= 1:
            return

        img_width, img_height = self.image_pil.size
        scaled_width = int(img_width * self.zoom_factor)
        scaled_height = int(img_height * self.zoom_factor)

        # Pārmēro attēlu
        display_img = self.image_pil.resize((scaled_width, scaled_height), Image.LANCZOS)
        self.photo_image = ImageTk.PhotoImage(display_img)

        self.canvas.delete("all")

        # Aprēķina attēla pozīciju ar pārvietošanu
        x = (canvas_width - scaled_width) / 2 + self.pan_x
        y = (canvas_height - scaled_height) / 2 + self.pan_y

        self.canvas.create_image(x, y, anchor="nw", image=self.photo_image)
        self.canvas.image = self.photo_image  # Saglabā atsauci

    def on_canvas_resize(self, event):
        """Pielāgo attēla attēlošanu, ja kanvasa izmērs mainās."""
        self.display_image()

    def on_mouse_wheel(self, event):
        """Apstrādā peles rullīša notikumus tālummaiņai."""
        if event.num == 5 or event.delta == -120:  # Tuvināt
            self.zoom_factor = max(0.1, self.zoom_factor - 0.1)
        if event.num == 4 or event.delta == 120:  # Attālināt
            self.zoom_factor = min(5.0, self.zoom_factor + 0.1)
        self.display_image()

    def on_pan_start(self, event):
        """Sāk attēla pārvietošanu (pan)."""
        self.start_pan_x = event.x - self.pan_x
        self.start_pan_y = event.y - self.pan_y
        self.canvas.config(cursor="fleur")

    def on_pan_drag(self, event):
        """Pārvieto attēlu, velkot peli."""
        self.pan_x = event.x - self.start_pan_x
        self.pan_y = event.y - self.start_pan_y
        self.display_image()

    def on_pan_end(self, event):
        """Beidz attēla pārvietošanu."""
        self.canvas.config(cursor="arrow")

    def on_selection_start(self, event):
        """Sāk atlases taisnstūra zīmēšanu."""
        self.selection_start_x = self.canvas.canvasx(event.x)
        self.selection_start_y = self.canvas.canvasy(event.y)
        if self.selection_rect:
            self.canvas.delete(self.selection_rect)
        self.selection_rect = self.canvas.create_rectangle(self.selection_start_x, self.selection_start_y,
                                                           self.selection_start_x, self.selection_start_y,
                                                           outline="blue", width=2, dash=(5, 2))

    def on_selection_drag(self, event):
        """Atjaunina atlases taisnstūra izmērus, velkot peli."""
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.selection_rect, self.selection_start_x, self.selection_start_y, cur_x, cur_y)

    def on_selection_end(self, event):
        """Beidz atlases taisnstūra zīmēšanu (šeit varētu apstrādāt iezīmēto apgabalu)."""
        pass


class DocumentScanner:
    def __init__(self, app_instance):
        self.app = app_instance
        self.image_to_process = None
        self.original_image_pil = None
        self.processed_image_pil = None
        self.corners = []
        self.preview_window = None
        self.canvas = None
        self.photo_image = None
        self.corner_handles = []
        self.active_handle = None
        self.zoom_factor = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.start_pan_x = 0
        self.start_pan_y = 0

        # Reāllaika skenēšanas mainīgie
        self.live_scan_active = False
        self.live_scan_button = None
        self.save_auto_button = None
        self.live_detected_corners = []
        self.scan_job = None
        self.color_picker_mode = False

        # Auto-adjust mainīgie
        self.auto_adjust_active = False
        self.auto_adjust_button = None
        self.auto_adjust_job = None
        self.auto_adjust_progress_label = None
        self.live_scan_was_active_before_auto = False
        # PIEVIENO ŠĪS RINDAS:
        self.settings_history = []  # Iestatījumu vēsture
        self.history_listbox = None
        self.load_history_button = None
        self.delete_history_button = None

        # Automātiskās pielāgošanas mainīgie
        self.auto_adjust_active = False
        self.auto_adjust_job = None
        self.best_score = -1.0
        self.best_settings = {
            "brightness": 0,
            "contrast": 0,
            "gamma": 1.0
        }
        self.current_search_step = 0
        self.search_space = {
            "brightness": list(range(-100, 101, 10)),
            "contrast": [100],  # FIKSĒTS UZ 100
            "gamma": [round(x * 0.01, 2) for x in range(96, 151, 5)]  # 0.96..1.50 pa 0.05
        }
        self.search_combinations = []
        self.current_combination_index = 0

        # Reāllaika skenēšanas mainīgie
        self.live_scan_active = False

    def set_image(self, pil_image):
        self.original_image_pil = pil_image.copy()
        self.image_to_process = pil_image.copy()
        self.processed_image_pil = pil_image.copy()
        self.corners = []  # Reset corners for new image

    def toggle_fullscreen(self):
        """Pārslēdz starp pilnekrāna un parasto režīmu."""
        if hasattr(self, 'preview_window') and self.preview_window:
            current_state = self.preview_window.attributes('-fullscreen')
            self.preview_window.attributes('-fullscreen', not current_state)

            if not current_state:  # Ja ieslēdzam fullscreen
                self.preview_window.attributes('-topmost', True)
            else:  # Ja izslēdzam fullscreen
                self.preview_window.attributes('-topmost', False)
                self.preview_window.state('zoomed')  # Windows maximized

    def on_camera_change(self, event=None):
        """Apstrādā kameras maiņu dropdown."""
        try:
            new_camera_index = self.camera_var.get()
            current_index = getattr(self.app, 'current_camera_index', 0)

            if new_camera_index != current_index:
                print(f"🔄 Maina kameru no {current_index} uz {new_camera_index}")

                # Aptur skenēšanu
                was_scanning = self.live_scan_active
                if was_scanning:
                    self.stop_live_scan()

                # Maina kameru
                if self.app.init_camera(force_camera_index=new_camera_index):
                    print(f"✅ Kamera nomainīta uz {new_camera_index}")

                    # Atsāk skenēšanu ar jauno kameru
                    if was_scanning:
                        self.document_frozen = False
                        self.live_detected_corners = []
                        if self.save_auto_button:
                            self.save_auto_button.config(state="disabled", text="🔍 Meklē dokumentu...")
                        self.start_live_scan()
                else:
                    # Ja neizdevās, atgriež veco vērtību
                    self.camera_var.set(current_index)
                    messagebox.showerror("Kļūda", f"Neizdevās atvērt kameru {new_camera_index}")

        except Exception as e:
            print(f"Kameras maiņas kļūda: {e}")
            messagebox.showerror("Kļūda", f"Kameras maiņas kļūda: {e}")

    def capture_and_process_frame(self):
        """Saglabā pašreizējo dokumentu un jautā par turpināšanu."""
        if self.original_image_pil is None or not self.live_detected_corners:
            messagebox.showwarning("Nav attēla", "Nav atrasts dokuments, ko saglabāt.")
            return

        try:
            # Aptur skenēšanu uz laiku
            self.stop_live_scan()

            # Iegūst AUGSTAS KVALITĀTES kadru saglabāšanai
            hq_frame = self.app.get_camera_frame_hq()
            if hq_frame is None:
                # Ja neizdevās iegūt HQ, izmanto esošo
                hq_frame = self.original_image_pil

            print(f"📸 Saglabā HQ attēlu: {hq_frame.size}")

            # Pielieto dokumenta korekciju uz HQ attēlu
            img_cv = np.array(hq_frame.convert("RGB"))
            img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

            # Pārrēķina stūrus HQ attēlam
            if hq_frame.size != self.original_image_pil.size:
                # Mērogošanas koeficients
                scale_x = hq_frame.size[0] / self.original_image_pil.size[0]
                scale_y = hq_frame.size[1] / self.original_image_pil.size[1]

                # Mērogoti stūri HQ attēlam
                hq_corners = self.live_detected_corners.copy()
                hq_corners[:, 0] *= scale_x
                hq_corners[:, 1] *= scale_y
            else:
                hq_corners = self.live_detected_corners

            warped_cv = self.four_point_transform(img_cv, hq_corners)
            processed_image_pil = Image.fromarray(cv2.cvtColor(warped_cv, cv2.COLOR_BGR2RGB))

            # Pievieno attēlu sarakstam
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            image_name = f"Skenēts_dokuments_{timestamp}"

            self.app.images.append({
                "filepath": f"camera_scan_{timestamp}",
                "original_img": processed_image_pil.copy(),
                "processed_img": processed_image_pil.copy()
            })
            self.app.ocr_results.append(None)
            self.app.file_listbox.insert(tk.END, image_name)
            self.app.refresh_file_listbox()

            # Saglabāt pašreizējā loga atsauci
            current_window = self.preview_window

            # Jautā par turpināšanu ar kvalitātes info
            quality_info = f"Kvalitāte: {processed_image_pil.size[0]}x{processed_image_pil.size[1]}"
            response = messagebox.askyesno("Dokuments saglabāts",
                                           f"Dokuments '{image_name}' pievienots sarakstam!\n{quality_info}\n\nVai vēlaties skenēt vēl vienu dokumentu?")

            if response:
                # Novērst loga dubultu aizvēršanu
                self.preview_window = None

                # Atver jauno logu PIRMS vecā aizvēršanas
                self.app.scan_document_with_camera_fast()

                # Aizvērt veco logu BEZ kameras atbrīvošanas
                if current_window:
                    try:
                        current_window.unbind('<Return>')
                        current_window.unbind('<r>')
                        current_window.unbind('<R>')
                        current_window.destroy()
                    except Exception as e:
                        print(f"Kļūda vecā loga aizvēršanā: {e}")
            else:
                # Ja nevēlas turpināt, aizvērt logu un atbrīvot kameru
                self.preview_window = current_window  # Atjauno atsauci
                self.close_preview_window(release_camera=True)

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās saglabāt dokumentu: {e}")
            # Atsāk skenēšanu, ja bija kļūda
            if self.preview_window:
                self.start_live_scan()

    def refresh_camera_view(self):
        """Ātri atsvaidzina kameras skatu."""
        try:
            # Atiestatīt statusu
            self.document_frozen = False
            self.live_detected_corners = []

            # Vizuāla atgriezeniskā saite
            if self.refresh_camera_button:
                self.refresh_camera_button.config(text="✅ Atsvaidzināts!", bootstyle="success")
                self.preview_window.after(500, lambda: self.refresh_camera_button.config(
                    text="📷 Atsvaidzināt", bootstyle="info") if self.refresh_camera_button else None)

            # Atjaunināt pogas
            if self.save_auto_button:
                self.save_auto_button.config(state="disabled", text="🔍 Meklē dokumentu...")

            # Iegūt jaunu kadru
            camera_frame_pil = self.app.get_camera_frame()
            if camera_frame_pil:
                self.original_image_pil = camera_frame_pil.copy()
                self.image_to_process = camera_frame_pil.copy()
                self.display_live_scan_preview()

            current_cam = getattr(self.app, 'current_camera_index', 0)
            print(f"✅ Kamera {current_cam} atsvaidzināta!")

        except Exception as e:
            print(f"Kļūda kameras atsvaidzināšanā: {e}")

    def apply_image_enhancements(self, image_pil):
        """Pielieto attēla uzlabojumus."""
        from PIL import ImageEnhance

        enhanced = image_pil.copy()

        # Spilgtums
        brightness = 1.0 + (self.app.scan_brightness.get() / 100.0)
        if brightness != 1.0:
            enhancer = ImageEnhance.Brightness(enhanced)
            enhanced = enhancer.enhance(brightness)

        # Kontrasts
        contrast = 1.0 + (self.app.scan_contrast.get() / 100.0)
        if contrast != 1.0:
            enhancer = ImageEnhance.Contrast(enhanced)
            enhanced = enhancer.enhance(contrast)

        # Krāsu piesātinājums
        saturation = 1.0 + (self.app.scan_saturation.get() / 100.0)
        if saturation != 1.0:
            enhancer = ImageEnhance.Color(enhanced)
            enhanced = enhancer.enhance(saturation)

        # Gamma korekcija
        gamma = self.app.scan_gamma.get()
        if gamma != 1.0:
            enhanced = self.apply_gamma_correction(enhanced, gamma)

        return enhanced

    def apply_gamma_correction(self, image_pil, gamma):
        """Pielieto gamma korekciju."""
        import numpy as np

        # Konvertē uz numpy array
        img_array = np.array(image_pil, dtype=np.float32) / 255.0

        # Pielieto gamma korekciju
        corrected = np.power(img_array, 1.0 / gamma)

        # Konvertē atpakaļ uz PIL
        corrected = (corrected * 255).astype(np.uint8)
        return Image.fromarray(corrected)

    def apply_color_based_detection(self, img_cv):
        """Pielieto krāsu balstītu detekciju."""
        if not self.app.scan_use_color_detection.get():
            return img_cv

        # Konvertē mērķa krāsu uz HSV
        target_color_hex = self.app.scan_target_color.get()
        target_rgb = tuple(int(target_color_hex[i:i + 2], 16) for i in (1, 3, 5))
        target_bgr = target_rgb[::-1]  # RGB uz BGR
        target_hsv = cv2.cvtColor(np.uint8([[target_bgr]]), cv2.COLOR_BGR2HSV)[0][0]

        # Konvertē attēlu uz HSV
        hsv = cv2.cvtColor(img_cv, cv2.COLOR_BGR2HSV)

        # Definē krāsu diapazonu ar overflow aizsardzību
        tolerance = min(self.app.scan_color_tolerance.get(), 50)  # Ierobežo tolerance
        h_value = int(target_hsv[0])  # Konvertē uz int
        lower_bound = np.array([max(0, h_value - tolerance), 50, 50], dtype=np.uint8)
        upper_bound = np.array([min(179, h_value + tolerance), 255, 255], dtype=np.uint8)

        # Izveido masku
        mask = cv2.inRange(hsv, lower_bound, upper_bound)

        # Pielieto masku
        result = cv2.bitwise_and(img_cv, img_cv, mask=mask)

        return result

    def find_document_corners_enhanced(self):
        """Atrod dokumenta stūrus ar uzlabotu algoritmu."""
        print("find_document_corners_enhanced izsaukts")

        if not self.original_image_pil:
            print("Nav original_image_pil")
            return None
        """Uzlabota dokumentu stūru atrašana ar papildu apstrādi."""
        if self.image_to_process is None:
            return None

        # Pielieto attēla uzlabojumus
        enhanced_image = self.apply_image_enhancements(self.image_to_process)

        # Konvertē uz OpenCV formātu
        img_cv = np.array(enhanced_image.convert("RGB"))
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

        # Pielieto krāsu detekciju, ja ieslēgta
        if self.app.scan_use_color_detection.get():
            img_cv = self.apply_color_based_detection(img_cv)

        # Konvertē uz pelēktoņiem
        img_gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)

        # Gausa izplūšana
        kernel_size = self.app.scan_gaussian_blur_kernel.get()
        if kernel_size % 2 == 0:
            kernel_size += 1
        img_blur = cv2.GaussianBlur(img_gray, (kernel_size, kernel_size), 0)

        # Adaptīvā sliekšņošana
        block_size = self.app.scan_adaptive_thresh_block_size.get()
        if block_size % 2 == 0:
            block_size += 1
        C = self.app.scan_adaptive_thresh_c.get()
        img_thresh = cv2.adaptiveThreshold(img_blur, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                           cv2.THRESH_BINARY_INV, block_size, C)

        # Morfoloģiskās operācijas, ja ieslēgtas
        if self.app.scan_morphology_enabled.get():
            morph_kernel_size = self.app.scan_morphology_kernel_size.get()
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (morph_kernel_size, morph_kernel_size))
            img_thresh = cv2.morphologyEx(img_thresh, cv2.MORPH_CLOSE, kernel)

        # Canny malu detekcija
        canny_thresh1 = self.app.scan_canny_thresh1.get()
        canny_thresh2 = self.app.scan_canny_thresh2.get()
        edges = cv2.Canny(img_thresh, canny_thresh1, canny_thresh2)

        # Malu paplašināšana
        dilation_size = self.app.scan_edge_dilation.get()
        if dilation_size > 0:
            dilation_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (dilation_size, dilation_size))
            edges = cv2.dilate(edges, dilation_kernel, iterations=1)

        # Atrod kontūras (pārējais kods paliek tāds pats)
        contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
        contours = sorted(contours, key=cv2.contourArea, reverse=True)

        min_contour_area = self.app.scan_min_contour_area.get()
        aspect_ratio_min = self.app.scan_aspect_ratio_min.get()
        aspect_ratio_max = self.app.scan_aspect_ratio_max.get()

        for contour in contours:
            if cv2.contourArea(contour) < min_contour_area:
                continue

            peri = cv2.arcLength(contour, True)
            approx = cv2.approxPolyDP(contour, 0.02 * peri, True)

            if len(approx) == 4:
                x, y, w, h = cv2.boundingRect(approx)
                aspect_ratio = w / float(h)
                if aspect_ratio_min < aspect_ratio < aspect_ratio_max:
                    pts = approx.reshape(4, 2)
                    rect = np.zeros((4, 2), dtype="float32")

                    s = pts.sum(axis=1)
                    rect[0] = pts[np.argmin(s)]
                    rect[2] = pts[np.argmax(s)]

                    diff = np.diff(pts, axis=1)
                    rect[1] = pts[np.argmin(diff)]
                    rect[3] = pts[np.argmax(diff)]

                    self.corners = rect.tolist()
                    return self.corners

        self.corners = []
        return None

    def pick_color_from_image(self, event):
        """Atlasa krāsu no attēla."""
        if not self.original_image_pil or not hasattr(self, 'img_on_canvas_x'):
            return

        # Konvertē kanvasa koordinātas uz attēla koordinātām
        canvas_x = self.canvas.canvasx(event.x)
        canvas_y = self.canvas.canvasy(event.y)

        img_x = int((canvas_x - self.img_on_canvas_x) / self.zoom_factor)
        img_y = int((canvas_y - self.img_on_canvas_y) / self.zoom_factor)

        # Pārbauda, vai koordinātas ir attēla robežās
        img_w, img_h = self.original_image_pil.size
        if 0 <= img_x < img_w and 0 <= img_y < img_h:
            # Iegūst pikseļa krāsu
            pixel_color = self.original_image_pil.getpixel((img_x, img_y))
            if isinstance(pixel_color, int):  # Pelēktoņu attēls
                pixel_color = (pixel_color, pixel_color, pixel_color)

            # Konvertē uz hex formātu
            hex_color = "#{:02x}{:02x}{:02x}".format(pixel_color[0], pixel_color[1], pixel_color[2])
            self.app.scan_target_color.set(hex_color)

            messagebox.showinfo("Krāsa atlasīta", f"Atlasītā krāsa: {hex_color}")

            # Ja reāllaika skenēšana ir aktīva, atjauno
            if self.live_scan_active:
                self.display_live_scan_preview()

    def create_detection_visualization_enhanced(self):
        """Uzlabota detekcijas vizualizācija."""
        if not self.original_image_pil:
            return self.original_image_pil

        # Pielieto attēla uzlabojumus
        enhanced_image = self.apply_image_enhancements(self.original_image_pil)

        # Konvertē uz OpenCV formātu
        img_cv = np.array(enhanced_image.convert("RGB"))
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

        # Pielieto krāsu detekciju, ja ieslēgta
        if self.app.scan_use_color_detection.get():
            color_detected = self.apply_color_based_detection(img_cv.copy())
            # Parāda krāsu detekcijas rezultātu kā overlay
            img_cv = cv2.addWeighted(img_cv, 0.7, color_detected, 0.3, 0)

        # Pārējā detekcijas loģika...
        img_gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)

        kernel_size = self.app.scan_gaussian_blur_kernel.get()
        if kernel_size % 2 == 0:
            kernel_size += 1
        img_blur = cv2.GaussianBlur(img_gray, (kernel_size, kernel_size), 0)

        block_size = self.app.scan_adaptive_thresh_block_size.get()
        if block_size % 2 == 0:
            block_size += 1
        C = self.app.scan_adaptive_thresh_c.get()
        img_thresh = cv2.adaptiveThreshold(img_blur, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                           cv2.THRESH_BINARY_INV, block_size, C)

        if self.app.scan_morphology_enabled.get():
            morph_kernel_size = self.app.scan_morphology_kernel_size.get()
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (morph_kernel_size, morph_kernel_size))
            img_thresh = cv2.morphologyEx(img_thresh, cv2.MORPH_CLOSE, kernel)

        canny_thresh1 = self.app.scan_canny_thresh1.get()
        canny_thresh2 = self.app.scan_canny_thresh2.get()
        edges = cv2.Canny(img_thresh, canny_thresh1, canny_thresh2)

        # Malu paplašināšana
        dilation_size = self.app.scan_edge_dilation.get()
        if dilation_size > 0:
            dilation_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (dilation_size, dilation_size))
            edges = cv2.dilate(edges, dilation_kernel, iterations=1)

        # Kombinē oriģinālo attēlu ar malu detekciju
        edges_colored = cv2.cvtColor(edges, cv2.COLOR_GRAY2BGR)
        edges_colored[:, :, 0] = 0  # Noņem sarkano kanālu
        edges_colored[:, :, 1] = edges  # Zaļais kanāls malām
        edges_colored[:, :, 2] = 0  # Noņem zilo kanālu

        # Kombinē ar uzlaboto attēlu
        combined = cv2.addWeighted(img_cv, 0.6, edges_colored, 0.4, 0)

        # Zīmē atrastos stūrus, ja tie ir
        if self.live_detected_corners:
            pts = np.array(self.live_detected_corners, dtype=np.int32)
            cv2.polylines(combined, [pts], True, (0, 255, 255), 4)  # Dzeltena kontūra

            # Zīmē stūru punktus ar etiķetēm
            corner_labels = ["TL", "TR", "BR", "BL"]
            for i, corner in enumerate(self.live_detected_corners):
                center = tuple(map(int, corner))
                cv2.circle(combined, center, 12, (0, 0, 255), -1)  # Sarkani punkti
                cv2.circle(combined, center, 15, (255, 255, 255), 2)  # Balts kontūrs
                cv2.putText(combined, corner_labels[i],
                            (center[0] - 10, center[1] + 5),
                            cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 2)

        # Parāda krāsu mērķi, ja krāsu detekcija ir ieslēgta
        if self.app.scan_use_color_detection.get():
            target_color_hex = self.app.scan_target_color.get()
            target_rgb = tuple(int(target_color_hex[i:i + 2], 16) for i in (1, 3, 5))
            target_bgr = target_rgb[::-1]

            # Zīmē krāsu paraugu augšējā labajā stūrī
            cv2.rectangle(combined, (combined.shape[1] - 80, 10),
                          (combined.shape[1] - 10, 40), target_bgr, -1)
            cv2.rectangle(combined, (combined.shape[1] - 80, 10),
                          (combined.shape[1] - 10, 40), (255, 255, 255), 2)
            cv2.putText(combined, "Target", (combined.shape[1] - 75, 55),
                        cv2.FONT_HERSHEY_SIMPLEX, 0.4, (255, 255, 255), 1)

        # Konvertē atpakaļ uz PIL
        combined_rgb = cv2.cvtColor(combined, cv2.COLOR_BGR2RGB)
        return Image.fromarray(combined_rgb)

    def find_document_corners(self):
        """Atrod dokumenta stūrus (izmanto uzlaboto versiju)."""
        return self.find_document_corners_enhanced()

    def four_point_transform(self, image, pts):
        # Obtain a consistent order of the points and unpack them
        # individually
        rect = np.array(pts, dtype="float32")
        (tl, tr, br, bl) = rect

        # Compute the width of the new image, which will be the
        # maximum distance between bottom-right and bottom-left
        # x-coordinates or the top-right and top-left x-coordinates
        widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
        widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
        maxWidth = max(int(widthA), int(widthB))

        # Compute the height of the new image, which will be the
        # maximum distance between the top-right and bottom-right
        # y-coordinates or the top-left and bottom-left y-coordinates
        heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
        heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
        maxHeight = max(int(heightA), int(heightB))

        # Now that we have the dimensions of the new image, construct
        # the set of destination points to obtain a "birds eye view",
        # (i.e. top-down view) of the image, again specifying points
        # in the top-left, top-right, bottom-right, and bottom-left
        # order
        dst = np.array([
            [0, 0],
            [maxWidth - 1, 0],
            [maxWidth - 1, maxHeight - 1],
            [0, maxHeight - 1]], dtype="float32")

        # Compute the perspective transform matrix and then apply it
        M = cv2.getPerspectiveTransform(rect, dst)
        warped = cv2.warpPerspective(image, M, (maxWidth, maxHeight))

        # Return the warped image
        return warped

    def __init__(self, app_instance):
        self.app = app_instance
        self.image_to_process = None
        self.original_image_pil = None
        self.processed_image_pil = None
        self.corners = []
        self.preview_window = None
        self.canvas = None
        self.photo_image = None
        self.corner_handles = []
        self.active_handle = None
        self.zoom_factor = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.start_pan_x = 0
        self.start_pan_y = 0

        # PIEVIENOJIET ŠĪSRINDAS:
        # PIEVIENOJIET ŠĪSRINDAS:
        self.live_scan_active = False
        self.live_scan_button = None
        self.save_auto_button = None
        self.live_detected_corners = []
        self.scan_job = None

        # Auto-adjust mainīgie
        self.auto_adjust_active = False
        self.auto_adjust_button = None
        self.auto_adjust_job = None
        self.auto_adjust_progress_label = None
        self.best_score = -1.0
        self.best_settings = {
            "brightness": 0,
            "contrast": 0,
            "gamma": 1.0
        }
        self.current_search_step = 0
        self.search_space = {
            "brightness": list(range(-100, 101, 10)),
            "contrast": list(range(-100, 101, 10)),
            "gamma": [round(x * 0.1, 1) for x in range(5, 21, 1)]
        }
        self.search_combinations = []
        self.current_combination_index = 0
        self.color_picker_mode = False

    def toggle_live_scan(self):
        """Ieslēdz/izslēdz reāllaika skenēšanu."""
        if not self.live_scan_active:
            self.start_live_scan()
        else:
            self.stop_live_scan()

    def toggle_auto_adjust(self):
        """Ieslēdz/izslēdz automātiskās attēla pielāgošanas režīmu."""
        # Drošības pārbaude - ja kāds lauks trūkst, inicializē
        if not hasattr(self, 'auto_adjust_active'):
            self.auto_adjust_active = False
        if not hasattr(self, 'auto_adjust_job'):
            self.auto_adjust_job = None
        if not hasattr(self, 'best_score'):
            self.best_score = -1.0
        if not hasattr(self, 'best_settings'):
            self.best_settings = {"brightness": 0, "contrast": 0, "gamma": 1.0}
        if not hasattr(self, 'search_space'):
            self.search_space = {
                "brightness": list(range(-100, 101, 10)),
                "contrast": list(range(-100, 101, 10)),
                "gamma": [round(x * 0.1, 1) for x in range(5, 21, 1)]
            }
        if not hasattr(self, 'search_combinations'):
            self.search_combinations = []
        if not hasattr(self, 'current_combination_index'):
            self.current_combination_index = 0

        if not self.auto_adjust_active:
            self.start_auto_adjust()
        else:
            self.stop_auto_adjust()

    def start_auto_adjust(self):
        """Sāk automātisko attēla pielāgošanu."""
        if not self.original_image_pil:
            messagebox.showwarning("Nav attēla",
                                   "Lūdzu, vispirms ielādējiet attēlu, lai veiktu automātisko pielāgošanu.")
            return

            # Saglabāt live scan stāvokli un izslēgt to
            self.live_scan_was_active_before_auto = self.live_scan_active
            if self.live_scan_active:
                self.stop_live_scan()

        self.auto_adjust_active = True
        self.auto_adjust_button.config(text="Automātiskā pielāgošana (Iesl.)", bootstyle="success")
        self.save_auto_adjust_button.config(state="disabled")
        self.auto_adjust_progress_label.config(text="Progress: 0%")

        # Atiestatīt meklēšanas stāvokli
        self.best_score = 0.0  # Mainīts no -1.0 uz 0.0
        self.best_settings = {
            "brightness": self.app.scan_brightness.get(),
            "contrast": self.app.scan_contrast.get(),
            "gamma": self.app.scan_gamma.get()
        }
        self.current_search_step = 0

        # Izveidot visas iespējamās kombinācijas ar FIKSĒTU kontrastu 100
        from itertools import product
        fast_search_space = {
            "brightness": list(range(-80, 81, 20)),  # -80..80 pa 20 (9 vērtības)
            "contrast": [100],  # VIENMĒR 100!
            "gamma": [0.96, 1.0, 1.2, 1.4, 1.6, 1.8]  # 7 vērtības
        }
        self.search_combinations = list(product(
            fast_search_space["brightness"],
            fast_search_space["contrast"],
            fast_search_space["gamma"]
        ))
        self.current_combination_index = 0

        # Sākt meklēšanas ciklu
        self.auto_adjust_loop()

    def stop_auto_adjust(self):
        """Aptur automātisko attēla pielāgošanu."""
        self.auto_adjust_active = False
        self.auto_adjust_button.config(text="Automātiskā pielāgošana (Izsl.)", bootstyle="secondary")
        if self.auto_adjust_job:
            self.preview_window.after_cancel(self.auto_adjust_job)
            self.auto_adjust_job = None

        # Atjaunot slīdņus uz labākajiem atrastajiem iestatījumiem
        self.app.scan_brightness.set(self.best_settings["brightness"])
        self.app.scan_contrast.set(self.best_settings["contrast"])
        self.app.scan_gamma.set(self.best_settings["gamma"])
        self.on_realtime_change()  # Atjaunināt attēlu ar labākajiem iestatījumiem

        # Automātiski ieslēgt live scan, ja tas nebija ieslēgts pirms auto-adjust
        if not hasattr(self, 'live_scan_was_active_before_auto'):
            self.live_scan_was_active_before_auto = False

        if not self.live_scan_was_active_before_auto and not self.live_scan_active:
            # Ja live scan nebija ieslēgts pirms auto-adjust, ieslēgt to tagad
            self.start_live_scan()

        print(f"Auto-adjust pabeigts. Labākais rezultāts: {self.best_score}")

        if self.best_score is not None and self.best_score > 0.1:  # Samazināts slieksnis!
            if hasattr(self, 'save_auto_adjust_button') and self.save_auto_adjust_button:
                self.save_auto_adjust_button.config(state="normal")
            messagebox.showinfo("Automātiskā pielāgošana",
                                f"Automātiskā pielāgošana pabeigta! ✅\n"
                                f"Labākais rezultāts: {self.best_score:.3f}\n"
                                f"Spilgtums: {self.best_settings['brightness']}\n"
                                f"Kontrasts: {self.best_settings['contrast']}\n"
                                f"Gamma: {self.best_settings['gamma']}")
        else:
            messagebox.showwarning("Automātiskā pielāgošana",
                                   f"Automātiskā pielāgošana pabeigta, bet dokuments netika pietiekami labi atrasts.\n"
                                   f"Labākais rezultāts: {self.best_score:.3f}\n"
                                   f"Mēģiniet manuāli pielāgot iestatījumus.")

    def auto_adjust_loop(self):
        """Automātiskās pielāgošanas cikls."""
        if not self.auto_adjust_active or not self.preview_window:
            return

        if self.current_combination_index >= len(self.search_combinations):
            self.stop_auto_adjust()
            return

        # Iegūt nākamo parametru kombināciju
        brightness, contrast, gamma = self.search_combinations[self.current_combination_index]

        # Vizuāli atjaunināt slīdņus
        self.app.scan_brightness.set(brightness)
        self.app.scan_contrast.set(contrast)
        self.app.scan_gamma.set(gamma)
        self.on_realtime_change()  # Atjaunināt attēlu ar jaunajiem iestatījumiem

        # Novērtēt pašreizējo kombināciju
        try:
            print(f"\n--- Kombinācija {self.current_combination_index + 1}/{len(self.search_combinations)} ---")
            print(f"Iestatījumi: S={brightness}, K={contrast}, G={gamma}")

            score = self.evaluate_document_detection()
            print(f"Iegūtais novērtējums: {score}")

        except Exception as e:
            print(f"Kļūda novērtēšanā: {e}")
            import traceback
            traceback.print_exc()
            score = 0.0

        # Ja score ir labāks VAI vienāds bet ar augstāku kontrastu
        should_update = False

        if score > self.best_score:
            should_update = True
        elif score == self.best_score and score > 0:
            # Ja rezultāts vienāds, izvēlēties augstāko kontrastu
            if contrast > self.best_settings["contrast"]:
                should_update = True

        if should_update:
            self.best_score = score
            self.best_settings["brightness"] = brightness
            self.best_settings["contrast"] = contrast
            self.best_settings["gamma"] = gamma
        # Atjaunināt progresu
        progress_percent = (self.current_combination_index + 1) / len(self.search_combinations) * 100
        self.auto_adjust_progress_label.config(
            text=f"Progress: {progress_percent:.1f}% (Labākais: {self.best_score:.2f})")

        self.current_combination_index += 1
        self.auto_adjust_job = self.preview_window.after(20, self.auto_adjust_loop)  # Ļoti ātrs cikls

    def evaluate_document_detection(self):
        """
        Novērtē dokumenta atpazīšanas kvalitāti.
        Atgriež punktu skaitu (score), kur augstāks punkts nozīmē labāku atpazīšanu.
        """
        try:
            print("Sākam dokumenta novērtēšanu...")

            # Izmanto esošo find_document_corners_enhanced metodi
            detected_corners = self.find_document_corners_enhanced()

            if detected_corners and len(detected_corners) == 4:
                print(f"Atrasti 4 stūri: {detected_corners}")

                # Aprēķināt kontūras laukumu
                import cv2
                import numpy as np

                # Konvertēt stūrus uz numpy array
                corners_array = np.array(detected_corners, dtype=np.float32)
                area = cv2.contourArea(corners_array)

                # Iegūt attēla izmērus
                if hasattr(self, 'original_image_pil') and self.original_image_pil:
                    img_width, img_height = self.original_image_pil.size
                    img_area = img_width * img_height

                    # Normalizēts laukums (0-1)
                    normalized_area = area / img_area

                    # Pamata punkti par atrašanu
                    base_score = 1.0

                    # Papildu punkti par laukumu
                    area_score = normalized_area * 2.0  # Maksimums 2.0

                    # Papildu punkti par stūru kvalitāti
                    quality_score = self.evaluate_corner_quality(corners_array)

                    total_score = base_score + area_score + quality_score

                    print(
                        f"Novērtējums: pamata={base_score}, laukums={area_score:.3f}, kvalitāte={quality_score:.3f}, kopā={total_score:.3f}")

                    return total_score
                else:
                    print("Nav original_image_pil")
                    return 1.0  # Pamata punkti par atrašanu
            else:
                print(f"Stūri nav atrasti vai nav 4: {detected_corners}")
                return 0.0

        except Exception as e:
            print(f"Kļūda evaluate_document_detection: {e}")
            import traceback
            traceback.print_exc()
            return 0.0

    def evaluate_corner_quality(self, corners):
        """Novērtē stūru kvalitāti."""
        try:
            import cv2
            import numpy as np

            # Pārbaudīt, vai stūri veido taisnstūri
            # Aprēķināt malas garumu
            distances = []
            for i in range(4):
                p1 = corners[i]
                p2 = corners[(i + 1) % 4]
                dist = np.sqrt((p1[0] - p2[0]) ** 2 + (p1[1] - p2[1]) ** 2)
                distances.append(dist)

            # Pārbaudīt, vai pretējās malas ir līdzīgas
            if len(distances) == 4:
                side1_diff = abs(distances[0] - distances[2]) / max(distances[0], distances[2])
                side2_diff = abs(distances[1] - distances[3]) / max(distances[1], distances[3])

                # Jo mazāka atšķirība, jo labāk (maksimums 0.5 punkti)
                quality = (1.0 - side1_diff) * 0.25 + (1.0 - side2_diff) * 0.25

                return max(0.0, quality)

            return 0.0

        except Exception as e:
            print(f"Kļūda corner quality: {e}")
            return 0.0

    def save_auto_adjusted_settings(self):
        """Saglabā automātiski pielāgotos attēla uzlabošanas iestatījumus."""
        if self.best_score is not None and self.best_score > 0:
            # Iestatīt slīdņus uz labākajām vērtībām
            self.app.scan_brightness.set(self.best_settings["brightness"])
            self.app.scan_contrast.set(self.best_settings["contrast"])
            self.app.scan_gamma.set(self.best_settings["gamma"])

            # Saglabāt iestatījumus lietotnes konfigurācijā
            self.app.settings["scan_brightness"] = self.best_settings["brightness"]
            self.app.settings["scan_contrast"] = self.best_settings["contrast"]
            self.app.settings["scan_gamma"] = self.best_settings["gamma"]

            # Pievienot vēsturei
            import datetime
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            history_entry = {
                "timestamp": timestamp,
                "name": f"Auto_{timestamp[:16].replace(' ', '_').replace(':', '-')}",
                "brightness": self.best_settings["brightness"],
                "contrast": self.best_settings["contrast"],
                "gamma": self.best_settings["gamma"],
                "score": self.best_score,
                "type": "auto"
            }

            # Pievienot vēsturei (maksimums 20 ieraksti)
            self.settings_history.append(history_entry)
            if len(self.settings_history) > 20:
                self.settings_history.pop(0)  # Noņemt vecāko

            # Saglabāt vēsturi failā
            self.save_settings_history()
            self.update_history_display()

            # Saglabāt iestatījumus failā
            try:
                self.app.save_app_settings()
                messagebox.showinfo("Saglabāts",
                                    f"Automātiski pielāgotie iestatījumi ir saglabāti un pievienoti vēsturei:\n"
                                    f"Spilgtums: {self.best_settings['brightness']}\n"
                                    f"Kontrasts: {self.best_settings['contrast']}\n"
                                    f"Gamma: {self.best_settings['gamma']}\n"
                                    f"Rezultāts: {self.best_score:.2f}")
            except Exception as e:
                messagebox.showwarning("Saglabāšanas kļūda",
                                       f"Iestatījumi iestatīti, bet neizdevās saglabāt failā: {e}")

        if hasattr(self, 'save_auto_adjust_button') and self.save_auto_adjust_button:
            self.save_auto_adjust_button.config(state="disabled")

    def save_settings_history(self):
        """Saglabā iestatījumu vēsturi failā."""
        try:
            import json
            import os

            # Izveidot settings mapi, ja neeksistē
            settings_dir = os.path.join(os.path.dirname(__file__), "settings")
            if not os.path.exists(settings_dir):
                os.makedirs(settings_dir)

            history_file = os.path.join(settings_dir, "auto_adjust_history.json")

            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(self.settings_history, f, indent=2, ensure_ascii=False)

        except Exception as e:
            print(f"Kļūda saglabājot vēsturi: {e}")

    def load_settings_history(self):
        """Ielādē iestatījumu vēsturi no faila."""
        try:
            import json
            import os

            history_file = os.path.join(os.path.dirname(__file__), "settings", "auto_adjust_history.json")

            if os.path.exists(history_file):
                with open(history_file, 'r', encoding='utf-8') as f:
                    self.settings_history = json.load(f)
            else:
                self.settings_history = []

        except Exception as e:
            print(f"Kļūda ielādējot vēsturi: {e}")
            self.settings_history = []

    def update_history_display(self):
        """Atjaunina vēstures saraksta attēlojumu."""
        if not hasattr(self, 'history_listbox') or not self.history_listbox:
            return

        # Notīra sarakstu
        self.history_listbox.delete(0, tk.END)

        # Pievieno ierakstus (jaunākie augšā)
        for i, entry in enumerate(reversed(self.settings_history)):
            # Pārbauda, vai ir nosaukums
            name = entry.get("name", "Bez nosaukuma")
            entry_type = entry.get("type", "auto")
            type_symbol = "🔧" if entry_type == "manual" else "🤖"

            display_text = f"{type_symbol} {name} | {entry['timestamp'][:16]} | S:{entry['brightness']} K:{entry['contrast']} G:{entry['gamma']} | Rez:{entry['score']:.2f}"
            self.history_listbox.insert(tk.END, display_text)

    def auto_load_best_settings(self):
        """Automātiski ielādē labākos iestatījumus no vēstures."""
        if not self.settings_history:
            return

        # Atrast iestatījumus ar augstāko rezultātu
        best_entry = None
        best_score = -1.0

        for entry in self.settings_history:
            entry_score = entry.get("score", 0.0)
            if entry_score is not None and entry_score > best_score:
                best_score = entry_score
                best_entry = entry

        # Ja atrasts labs iestatījums, ielādēt to
        if best_entry and best_score > 0.5:  # Tikai ja rezultāts ir pietiekami labs
            try:
                self.app.scan_brightness.set(best_entry["brightness"])
                self.app.scan_contrast.set(best_entry["contrast"])
                self.app.scan_gamma.set(best_entry["gamma"])
                self.on_realtime_change()

                print(
                    f"Auto-ielādēti labākie iestatījumi: S:{best_entry['brightness']} K:{best_entry['contrast']} G:{best_entry['gamma']} (Rezultāts: {best_score:.2f})")

            except Exception as e:
                print(f"Kļūda auto-ielādējot iestatījumus: {e}")

    def load_selected_history(self):
        """Ielādē izvēlētos iestatījumus no vēstures."""
        if not hasattr(self, 'history_listbox') or not self.history_listbox:
            return

        selection = self.history_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav izvēles", "Lūdzu, izvēlieties iestatījumus no saraksta.")
            return

        # Iegūt izvēlēto ierakstu (saraksts ir apgriezts, tāpēc jāpārrēķina indekss)
        selected_index = len(self.settings_history) - 1 - selection[0]
        selected_entry = self.settings_history[selected_index]

        # Iestatīt slīdņus
        self.app.scan_brightness.set(selected_entry["brightness"])
        self.app.scan_contrast.set(selected_entry["contrast"])
        self.app.scan_gamma.set(selected_entry["gamma"])

        # Atjaunināt attēlu
        self.on_realtime_change()

        messagebox.showinfo("Ielādēts",
                            f"Iestatījumi ielādēti no {selected_entry['timestamp']}:\n"
                            f"Spilgtums: {selected_entry['brightness']}\n"
                            f"Kontrasts: {selected_entry['contrast']}\n"
                            f"Gamma: {selected_entry['gamma']}")

    def delete_selected_history(self):
        """Dzēš izvēlētos iestatījumus no vēstures."""
        if not hasattr(self, 'history_listbox') or not self.history_listbox:
            return

        selection = self.history_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav izvēles", "Lūdzu, izvēlieties iestatījumus dzēšanai.")
            return

        # Apstiprinājums
        if not messagebox.askyesno("Dzēst iestatījumus", "Vai tiešām vēlaties dzēst izvēlētos iestatījumus?"):
            return

        # Dzēst ierakstu (saraksts ir apgriezts, tāpēc jāpārrēķina indekss)
        selected_index = len(self.settings_history) - 1 - selection[0]
        deleted_entry = self.settings_history.pop(selected_index)

        # Saglabāt izmaiņas un atjaunināt attēlojumu
        self.save_settings_history()
        self.update_history_display()

        messagebox.showinfo("Dzēsts", f"Iestatījumi no {deleted_entry['timestamp']} ir dzēsti.")

    def save_current_settings_to_history(self):
        """Saglabā pašreizējos iestatījumus vēsturē ar lietotāja nosaukumu."""
        # Iegūt pašreizējos iestatījumus
        current_settings = {
            "brightness": self.app.scan_brightness.get(),
            "contrast": self.app.scan_contrast.get(),
            "gamma": self.app.scan_gamma.get()
        }

        # Prasīt nosaukumu
        from tkinter import simpledialog
        name = simpledialog.askstring("Iestatījumu nosaukums",
                                      "Ievadiet nosaukumu šiem iestatījumiem:",
                                      initialvalue=f"Manuāli_{len(self.settings_history) + 1}")

        if name:  # Ja lietotājs ievadīja nosaukumu
            import datetime
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # Novērtēt pašreizējos iestatījumus
            score = self.evaluate_document_detection() if hasattr(self,
                                                                  'original_image_pil') and self.original_image_pil else 0.0

            history_entry = {
                "timestamp": timestamp,
                "name": name,
                "brightness": current_settings["brightness"],
                "contrast": current_settings["contrast"],
                "gamma": current_settings["gamma"],
                "score": score,
                "type": "manual"
            }

            # Pievienot vēsturei
            self.settings_history.append(history_entry)
            if len(self.settings_history) > 50:  # Palielināts limits
                self.settings_history.pop(0)

            # Saglabāt un atjaunināt
            self.save_settings_history()
            self.update_history_display()

            messagebox.showinfo("Saglabāts", f"Iestatījumi '{name}' saglabāti vēsturē!")

    def rename_selected_history(self):
        """Pārdēvē izvēlēto vēstures ierakstu."""
        if not hasattr(self, 'history_listbox') or not self.history_listbox:
            return

        selection = self.history_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav izvēles", "Lūdzu, izvēlieties ierakstu pārdēvēšanai.")
            return

        # Iegūt izvēlēto ierakstu
        selected_index = len(self.settings_history) - 1 - selection[0]
        selected_entry = self.settings_history[selected_index]

        # Prasīt jaunu nosaukumu
        from tkinter import simpledialog
        current_name = selected_entry.get("name", "Bez nosaukuma")
        new_name = simpledialog.askstring("Pārdēvēt ierakstu",
                                          "Ievadiet jaunu nosaukumu:",
                                          initialvalue=current_name)

        if new_name and new_name != current_name:
            # Atjaunināt nosaukumu
            self.settings_history[selected_index]["name"] = new_name

            # Saglabāt izmaiņas un atjaunināt attēlojumu
            self.save_settings_history()
            self.update_history_display()

            messagebox.showinfo("Pārdēvēts", f"Ieraksts pārdēvēts uz '{new_name}'")

    def start_resize(self, event):
        """Sāk kreisās puses izmēra maiņu."""
        self.resize_start_x = event.x_root
        self.resize_active = True

    def do_resize(self, event):
        """Veic kreisās puses izmēra maiņu."""
        if not self.resize_active:
            return

        # Aprēķināt jauno platumu
        delta_x = event.x_root - self.resize_start_x
        new_width = self.left_panel_width + delta_x

        # Ierobežot minimālo un maksimālo platumu
        new_width = max(250, min(600, new_width))

        # Atjaunināt kreisās puses platumu
        if hasattr(self, 'preview_window') and self.preview_window:
            # Atrast kreiso frame un atjaunināt tā platumu
            for child in self.preview_window.winfo_children():
                if isinstance(child, ttk.Frame):
                    for subchild in child.winfo_children():
                        if isinstance(subchild, ttk.Frame) and subchild.winfo_reqwidth() > 200:
                            subchild.config(width=new_width)
                            break
                    break

    def end_resize(self, event):
        """Beidz kreisās puses izmēra maiņu."""
        if self.resize_active:
            # Saglabāt jauno platumu
            delta_x = event.x_root - self.resize_start_x
            self.left_panel_width = max(250, min(600, self.left_panel_width + delta_x))
            self.resize_active = False

    def start_live_scan(self):
        """Sāk reāllaika skenēšanu."""
        self.live_scan_active = True
        if self.live_scan_button:
            self.live_scan_button.config(text="⏹ Apturēt skenēšanu", bootstyle="danger")
        if self.save_auto_button:
            # Poga vienmēr aktīva - nav jāmaina stāvoklis
            pass

        self.live_scan_loop()

    def stop_live_scan(self):
        """Aptur reāllaika skenēšanu."""
        self.live_scan_active = False
        if self.live_scan_button:
            self.live_scan_button.config(text="📹 Ieslēgt skenēšanu", bootstyle="success")
        if self.save_auto_button:
            # Poga vienmēr aktīva - nav jāmaina stāvoklis
            pass

        if self.scan_job:
            self.preview_window.after_cancel(self.scan_job)
            self.scan_job = None

    def live_scan_loop(self):
        """Reāllaika skenēšanas cikls."""
        if not self.live_scan_active or not self.preview_window:
            if self.preview_window and not self.live_scan_active:
                self.live_scan_active = True
            else:
                return

        try:
            # Iegūst jaunu kadru no kameras
            camera_frame_pil = self.app.get_camera_frame()
            if camera_frame_pil:
                self.original_image_pil = camera_frame_pil.copy()
                self.image_to_process = camera_frame_pil.copy()

                # Atjauno video attēlojumu (vienmēr)
                self.display_live_scan_preview()

                # Meklē dokumentu tikai katru 3. reizi (ātrāk)
                if not hasattr(self, 'scan_counter'):
                    self.scan_counter = 0
                self.scan_counter += 1

                # Meklē dokumentu tikai ja nav "iesaldēts" un katru 3. reizi
                if not getattr(self, 'document_frozen', False) and self.scan_counter % 3 == 0:
                    detected_corners = self.find_document_corners_enhanced()

                    if detected_corners:
                        self.live_detected_corners = detected_corners
                        self.document_frozen = True
                        if self.save_auto_button:
                            self.save_auto_button.config(state="normal", bootstyle="success",
                                                         text="✅ Nospiediet ENTER vai šo pogu")
                    elif self.scan_counter % 3 == 0:  # Atjaunina pogu tikai kad meklē
                        self.live_detected_corners = []
                        if self.save_auto_button:
                            self.save_auto_button.config(state="disabled", text="🔍 Meklē dokumentu...")

            # Ātrāks cikls
            if self.preview_window:
                self.scan_job = self.preview_window.after(80, self.live_scan_loop)

        except Exception as e:
            print(f"Kļūda reāllaika skenēšanā: {e}")

    def display_live_scan_preview(self):
        """Atjauno priekšskatījuma attēlu (samazinātu ātrumam)."""
        if not self.preview_window or not self.original_image_pil:
            return

        try:
            canvas_width = self.preview_canvas.winfo_width()
            canvas_height = self.preview_canvas.winfo_height()

            if canvas_width <= 1 or canvas_height <= 1:
                return

            # Samazina TIKAI priekšskatījumam
            display_img = self.original_image_pil.copy()
            display_img.thumbnail((canvas_width, canvas_height), Image.Resampling.LANCZOS)

            # Stūri uz samazinātā attēla
            if self.live_detected_corners and getattr(self, 'document_frozen', False):
                orig_width, orig_height = self.original_image_pil.size
                display_width, display_height = display_img.size
                scale_x = display_width / orig_width
                scale_y = display_height / orig_height

                scaled_corners = self.live_detected_corners.copy()
                scaled_corners[:, 0] *= scale_x
                scaled_corners[:, 1] *= scale_y

                img_cv = cv2.cvtColor(np.array(display_img), cv2.COLOR_RGB2BGR)
                cv2.drawContours(img_cv, [scaled_corners.astype(int)], -1, (0, 255, 0), 3)
                display_img = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))

            self.preview_photo = ImageTk.PhotoImage(display_img)
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(
                canvas_width // 2, canvas_height // 2,
                image=self.preview_photo, anchor="center"
            )

        except Exception as e:
            print(f"Priekšskatījuma kļūda: {e}")

    def create_detection_visualization(self):
        """Izveido detekcijas vizualizāciju (izmanto uzlaboto versiju)."""
        return self.create_detection_visualization_enhanced()

    def save_auto_detected(self):
        """Saglabā automātiski atklātos stūrus."""
        if self.live_detected_corners:
            self.corners = self.live_detected_corners.copy()
            self.stop_live_scan()
            messagebox.showinfo("Saglabāts",
                                "Automātiski atklātie stūri ir saglabāti!\nTagad varat tos precizēt vai uzreiz pielietot korekciju.")
            self.display_image_on_canvas()
        else:
            messagebox.showwarning("Nav datu", "Nav automātiski atklātu stūru, ko saglabāt.")

    def toggle_live_scan(self):
        """Ieslēdz/izslēdz reāllaika skenēšanu."""
        if not self.live_scan_active:
            self.start_live_scan()
        else:
            self.stop_live_scan()

    def start_live_scan(self):
        """Sāk reāllaika skenēšanu."""
        self.live_scan_active = True
        self.document_frozen = False  # Pievienot šo rindu
        self.live_scan_loop()

    def stop_live_scan(self):
        """Aptur reāllaika skenēšanu."""
        self.live_scan_active = False
        if self.live_scan_button:
            self.live_scan_button.config(text="📹 Ieslēgt skenēšanu", bootstyle="success")
        if self.save_auto_button:
            # Poga vienmēr aktīva - nav jāmaina stāvoklis
            pass

        if self.scan_job:
            self.preview_window.after_cancel(self.scan_job)
            self.scan_job = None

    def live_scan_loop(self):
        """Reāllaika skenēšanas cikls."""
        if not self.live_scan_active or not self.preview_window:
            return

        try:
            # Mēģina atrast dokumenta stūrus
            detected_corners = self.find_document_corners()

            if detected_corners:
                self.live_detected_corners = detected_corners
                self.document_frozen = True

                # Pievienot skaņas signālu (neobligāti)
                try:
                    import winsound
                    winsound.Beep(1000, 200)  # 1000Hz, 200ms
                except:
                    pass  # Ja nav Windows vai nav winsound

                if self.save_auto_button:
                    self.save_auto_button.config(state="normal", bootstyle="success",
                                                 text="✅ Nospiediet ENTER vai šo pogu")

            # Atjauno attēlojumu ar reāllaika detekciju
            self.display_live_scan_preview()

            # Turpina ciklu
            self.scan_job = self.preview_window.after(300, self.live_scan_loop)

        except Exception as e:
            print(f"Kļūda reāllaika skenēšanā: {e}")
            self.stop_live_scan()

    def display_live_scan_preview(self):
        """Attēlo reāllaika skenēšanas priekšskatījumu."""
        if self.original_image_pil is None or self.canvas is None:
            return

        try:
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()
        except tk.TclError:
            return

        if canvas_width <= 1 or canvas_height <= 1:
            self.preview_window.after(50, self.display_live_scan_preview)
            return

        # Izveido vizualizācijas attēlu
        display_img = self.create_detection_visualization()

        img_width, img_height = display_img.size
        scaled_width = int(img_width * self.zoom_factor)
        scaled_height = int(img_height * self.zoom_factor)

        display_img_resized = display_img.resize((scaled_width, scaled_height), Image.LANCZOS)
        self.photo_image = ImageTk.PhotoImage(display_img_resized)

        self.canvas.delete("all")

        # Aprēķina attēla pozīciju
        self.img_on_canvas_x = (canvas_width - scaled_width) / 2 + self.pan_x
        self.img_on_canvas_y = (canvas_height - scaled_height) / 2 + self.pan_y

        self.canvas.create_image(self.img_on_canvas_x, self.img_on_canvas_y, anchor="nw", image=self.photo_image)
        self.canvas.image = self.photo_image

        # Status teksts
        status_text = "🔍 REĀLLAIKA SKENĒŠANA AKTĪVA"
        if self.live_detected_corners:
            status_text += "\n✅ DOKUMENTS ATRASTS!"
        else:
            status_text += "\n❌ Dokuments nav atrasts"

        self.canvas.create_text(
            10, 10, text=status_text, anchor="nw",
            fill="lime" if self.live_detected_corners else "red",
            font=("Arial", 12, "bold"), tags="status"
        )

    def on_realtime_change(self, *args):
        """Reāllaika iestatījumu maiņa."""
        # Ja reāllaika skenēšana ir aktīva, atjauno vizualizāciju
        if hasattr(self, 'live_scan_active') and self.live_scan_active:
            self.display_live_scan_preview()

    def choose_color(self):
        """Atver krāsu izvēles dialogu."""
        from tkinter import colorchooser
        color = colorchooser.askcolor(title="Izvēlieties dokumenta krāsu")
        if color[1]:  # Ja krāsa tika izvēlēta
            self.app.scan_target_color.set(color[1])
            self.on_realtime_change()

    def enable_color_picker(self):
        """Ieslēdz krāsu atlasīšanas režīmu."""
        self.color_picker_mode = True
        self.canvas.config(cursor="crosshair")
        messagebox.showinfo("Krāsu atlase", "Noklikšķiniet uz attēla, lai atlasītu dokumenta krāsu.")

    def pick_color_from_image(self, event):
        """Atlasa krāsu no attēla."""
        if not self.original_image_pil or not hasattr(self, 'img_on_canvas_x'):
            return

        # Konvertē kanvasa koordinātas uz attēla koordinātām
        canvas_x = self.canvas.canvasx(event.x)
        canvas_y = self.canvas.canvasy(event.y)

        img_x = int((canvas_x - self.img_on_canvas_x) / self.zoom_factor)
        img_y = int((canvas_y - self.img_on_canvas_y) / self.zoom_factor)

        # Pārbauda, vai koordinātas ir attēla robežās
        img_w, img_h = self.original_image_pil.size
        if 0 <= img_x < img_w and 0 <= img_y < img_h:
            # Iegūst pikseļa krāsu
            pixel_color = self.original_image_pil.getpixel((img_x, img_y))
            if isinstance(pixel_color, int):  # Pelēktoņu attēls
                pixel_color = (pixel_color, pixel_color, pixel_color)

            # Konvertē uz hex formātu
            hex_color = "#{:02x}{:02x}{:02x}".format(pixel_color[0], pixel_color[1], pixel_color[2])
            self.app.scan_target_color.set(hex_color)

            messagebox.showinfo("Krāsa atlasīta", f"Atlasītā krāsa: {hex_color}")

            # Ja reāllaika skenēšana ir aktīva, atjauno
            if self.live_scan_active:
                self.display_live_scan_preview()

    def close_preview_window(self, release_camera=True):
        """Aizver priekšskatījuma logu un aptur skenēšanu."""
        self.stop_live_scan()
        if self.preview_window:
            # Noņem visus taustiņu bindings
            try:
                self.preview_window.unbind('<Return>')
                self.preview_window.unbind('<r>')
                self.preview_window.unbind('<R>')
            except:
                pass
            try:
                self.preview_window.destroy()
            except:
                pass
            self.preview_window = None

        # Atbrīvo kameru tikai ja nepieciešams
        if release_camera:
            self.app.release_camera()

    def show_document_detection_preview(self):
        if self.original_image_pil is None:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        self.preview_window = Toplevel(self.app)
        self.preview_window.title("Dokumenta robežu korekcija un reāllaika detekcija")

        # Mēģina iegūt precīzus darba laukuma izmērus, lai izvairītos no uzdevumjoslas pārklāšanās
        try:
            import ctypes
            # Pārbauda, vai ir Windows operētājsistēma
            if sys.platform.startswith('win'):
                user32 = ctypes.windll.user32
                # SM_CXFULLSCREEN un SM_CYFULLSCREEN atgriež darba laukuma izmērus
                work_width = user32.GetSystemMetrics(16)  # SM_CXFULLSCREEN
                work_height = user32.GetSystemMetrics(17)  # SM_CYFULLSCREEN
                # Darba laukuma pozīcija (parasti 0,0)
                work_x = 0
                work_y = 0
                self.preview_window.geometry(f"{work_width}x{work_height}+{work_x}+{work_y}")
            else:
                # Citas OS (Linux, macOS) - izmanto standarta maksimizāciju
                self.preview_window.state('zoomed')
        except (ImportError, AttributeError, OSError) as e:
            print(f"Nevarēja izmantot ctypes Windows API: {e}. Izmanto standarta maksimizāciju.")
            self.preview_window.state(
                'zoomed')  # Atgriežas pie standarta maksimizācijas, ja ctypes nav pieejams vai rodas kļūda

        # Papildus, lai nodrošinātu, ka logs ir redzams un aktīvs
        self.preview_window.deiconify()  # Pārliecinās, ka logs ir redzams
        self.preview_window.lift()  # Paceļ logu virs citiem logiem
        self.preview_window.focus_force()  # Piešķir loga fokusu

        # Papildus, lai nodrošinātu, ka logs ir redzams un aktīvs
        self.preview_window.deiconify()  # Pārliecinās, ka logs ir redzams
        self.preview_window.lift()  # Paceļ logu virs citiem logiem
        self.preview_window.focus_force()  # Piešķir loga fokusu
        # Alternatīvi var izmantot:
        # self.preview_window.attributes('-fullscreen', True)  # Īsts fullscreen

        self.preview_window.transient(self.app)
        self.preview_window.grab_set()

        # Pievienot ESC taustiņu, lai izietu no pilnekrāna
        self.preview_window.bind('<Escape>', lambda event: self.close_preview_window())
        # Pievienot taustiņu atbalstu
        self.preview_window.bind('<Return>', lambda event: self.capture_and_process_frame())
        self.preview_window.bind('<r>', lambda event: self.refresh_camera_view())
        self.preview_window.bind('<R>', lambda event: self.refresh_camera_view())
        self.preview_window.focus_set()  # Nodrošina, ka logs var saņemt taustiņu nospiešanas

        # Galvenais konteiners
        main_container = ttk.Frame(self.preview_window)
        main_container.pack(fill="both", expand=True, padx=5, pady=5)

        # Kreisā puse - kontroles ar scrollbar (resizable)
        if not hasattr(self, 'left_panel_width'):
            self.left_panel_width = 380

        left_panel_container = ttk.Frame(main_container, width=self.left_panel_width)
        left_panel_container.pack(side="left", fill="y", padx=(0, 0))
        left_panel_container.pack_propagate(False)

        # Scrollable canvas kreisajam panelim
        left_canvas = tk.Canvas(left_panel_container, width=380, highlightthickness=0)
        left_scrollbar = ttk.Scrollbar(left_panel_container, orient="vertical", command=left_canvas.yview)
        left_panel = ttk.Frame(left_canvas)

        # Konfigurē scroll funkcionalitāti
        def configure_left_scroll_region(event):
            try:
                if left_canvas.winfo_exists():
                    left_canvas.configure(scrollregion=left_canvas.bbox("all"))
            except tk.TclError:
                pass

        left_panel.bind("<Configure>", configure_left_scroll_region)

        # Pievieno left_panel uz canvas
        left_canvas_frame = left_canvas.create_window((0, 0), window=left_panel, anchor="nw")

        # Konfigurē canvas izmēru
        def configure_left_canvas(event):
            try:
                if left_canvas.winfo_exists():
                    left_canvas.itemconfig(left_canvas_frame, width=event.width)
            except tk.TclError:
                pass

        left_canvas.bind('<Configure>', configure_left_canvas)
        left_canvas.configure(yscrollcommand=left_scrollbar.set)

        # Peles rullīša atbalsts kreisajam panelim
        def on_left_mousewheel(event):
            try:
                if left_canvas.winfo_exists() and self.preview_window.winfo_exists():
                    left_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            except (tk.TclError, AttributeError):
                pass

        def bind_left_mousewheel(event):
            try:
                if left_canvas.winfo_exists():
                    left_canvas.bind_all("<MouseWheel>", on_left_mousewheel)
            except tk.TclError:
                pass

        def unbind_left_mousewheel(event):
            try:
                left_canvas.unbind_all("<MouseWheel>")
            except tk.TclError:
                pass

        # Piesaista peles rullīti tikai kad pele ir virs kreisā paneļa
        left_canvas.bind('<Enter>', bind_left_mousewheel)
        left_canvas.bind('<Leave>', unbind_left_mousewheel)

        # Ievieto canvas un scrollbar
        left_canvas.pack(side="left", fill="both", expand=True)
        left_scrollbar.pack(side="right", fill="y")

        # Resizer handle starp kreiso un labo paneli
        resizer_frame = ttk.Frame(main_container, width=8, cursor="sb_h_double_arrow")
        resizer_frame.pack(side="left", fill="y", padx=2)
        resizer_frame.pack_propagate(False)

        # Vizuāls indikators resizer handle
        resizer_line = tk.Frame(resizer_frame, width=2, bg="#cccccc")
        resizer_line.pack(fill="y", expand=True, padx=3)

        # Resizer funkcionalitāte
        self.resize_active = False
        self.resize_start_x = 0

        def start_resize(event):
            self.resize_active = True
            self.resize_start_x = event.x_root
            resizer_frame.config(cursor="sb_h_double_arrow")

        def do_resize(event):
            if not self.resize_active:
                return

            delta_x = event.x_root - self.resize_start_x
            new_width = self.left_panel_width + delta_x

            # Ierobežot platumu
            new_width = max(250, min(600, new_width))

            # Atjaunināt kreisā paneļa platumu
            left_panel_container.config(width=new_width)

        def end_resize(event):
            if self.resize_active:
                delta_x = event.x_root - self.resize_start_x
                self.left_panel_width = max(250, min(600, self.left_panel_width + delta_x))
                left_panel_container.config(width=self.left_panel_width)
                self.resize_active = False
                resizer_frame.config(cursor="sb_h_double_arrow")

        # Piesaistīt notikumus
        resizer_frame.bind("<Button-1>", start_resize)
        resizer_frame.bind("<B1-Motion>", do_resize)
        resizer_frame.bind("<ButtonRelease-1>", end_resize)
        resizer_line.bind("<Button-1>", start_resize)
        resizer_line.bind("<B1-Motion>", do_resize)
        resizer_line.bind("<ButtonRelease-1>", end_resize)

        # Hover efekts
        def on_resizer_enter(event):
            resizer_line.config(bg="#999999")

        def on_resizer_leave(event):
            resizer_line.config(bg="#cccccc")

        resizer_frame.bind("<Enter>", on_resizer_enter)
        resizer_frame.bind("<Leave>", on_resizer_leave)
        resizer_line.bind("<Enter>", on_resizer_enter)
        resizer_line.bind("<Leave>", on_resizer_leave)

        # Labā puse - attēls
        right_panel = ttk.Frame(main_container)
        right_panel.pack(side="left", fill="both", expand=True)  # Mainīts no "right" uz "left"

        # === KREISĀ PANEĻA SATURS ===

        # Pamata kontroles
        basic_frame = ttk.LabelFrame(left_panel, text="Pamata kontroles", padding="10")
        basic_frame.pack(fill="x", pady=8)

        ttk.Button(basic_frame, text="🔄 Atiestatīt skatu", command=self.reset_view, width=25).pack(pady=2)
        ttk.Button(basic_frame, text="🔍+ Tuvināt", command=lambda: self.change_zoom(1.1), width=25).pack(pady=2)
        ttk.Button(basic_frame, text="🔍- Attālināt", command=lambda: self.change_zoom(0.9), width=25).pack(pady=2)
        ttk.Button(basic_frame, text="🤖 Auto noteikt", command=self.auto_detect_corners, width=25).pack(pady=2)

        # Reāllaika skenēšana
        scan_frame = ttk.LabelFrame(left_panel, text="Reāllaika skenēšana", padding="10")
        scan_frame.pack(fill="x", pady=8)

        self.live_scan_button = ttk.Button(scan_frame, text="📹 Ieslēgt skenēšanu",
                                           command=self.toggle_live_scan, bootstyle="success", width=25)
        self.live_scan_button.pack(pady=2)

        self.save_auto_button = ttk.Button(scan_frame, text="💾 Saglabāt auto ieskenēto",
                                           command=self.save_auto_detected, bootstyle="warning",
                                           state="normal", width=25)  # VIENMĒR AKTĪVA!
        self.save_auto_button.pack(pady=2)

        # Attēla uzlabojumi
        enhance_frame = ttk.LabelFrame(left_panel, text="Attēla uzlabojumi", padding="10")
        enhance_frame.pack(fill="x", pady=8)

        # Automātiskā pielāgošana
        self.auto_adjust_button = ttk.Button(enhance_frame, text="Automātiskā pielāgošana (Izsl.)",
                                             command=self.toggle_auto_adjust, bootstyle="secondary")
        self.auto_adjust_button.pack(fill="x", pady=5)

        self.save_auto_adjust_button = ttk.Button(enhance_frame, text="Saglabāt automātiski pielāgotos iestatījumus",
                                                  command=self.save_auto_adjusted_settings, bootstyle="info",
                                                  state="disabled")
        self.save_auto_adjust_button.pack(fill="x", pady=5)

        # Manuālā saglabāšana
        self.save_manual_button = ttk.Button(enhance_frame, text="Saglabāt pašreizējos iestatījumus vēsturē",
                                             command=self.save_current_settings_to_history, bootstyle="success")
        self.save_manual_button.pack(fill="x", pady=5)

        self.auto_adjust_progress_label = ttk.Label(enhance_frame, text="Progress: 0%", bootstyle="info")
        self.auto_adjust_progress_label.pack(fill="x", pady=2)

        # Iestatījumu vēstures sadaļa
        history_frame = ttk.LabelFrame(enhance_frame, text="Iestatījumu vēsture", padding=10)
        history_frame.pack(fill="both", expand=True, pady=5)

        # Saraksts ar saglabātajiem iestatījumiem
        history_list_frame = ttk.Frame(history_frame)
        history_list_frame.pack(fill="both", expand=True, pady=2)

        self.history_listbox = tk.Listbox(history_list_frame, height=4, font=("Arial", 9))
        self.history_listbox.pack(side="left", fill="both", expand=True)

        history_scrollbar = ttk.Scrollbar(history_list_frame, orient="vertical", command=self.history_listbox.yview)
        history_scrollbar.pack(side="right", fill="y")
        self.history_listbox.config(yscrollcommand=history_scrollbar.set)

        # Pogas vēstures pārvaldībai
        history_buttons_frame = ttk.Frame(history_frame)
        history_buttons_frame.pack(fill="x", pady=2)

        self.load_history_button = ttk.Button(history_buttons_frame, text="Ielādēt",
                                              command=self.load_selected_history, bootstyle="info")
        self.load_history_button.pack(side="left", padx=2)

        self.rename_history_button = ttk.Button(history_buttons_frame, text="Pārdēvēt",
                                                command=self.rename_selected_history, bootstyle="warning")
        self.rename_history_button.pack(side="left", padx=2)

        self.delete_history_button = ttk.Button(history_buttons_frame, text="Dzēst",
                                                command=self.delete_selected_history, bootstyle="danger")
        self.delete_history_button.pack(side="left", padx=2)

        # Ielādē saglabāto vēsturi
        self.load_settings_history()
        self.update_history_display()
        # Auto-ielādēt labākos iestatījumus, ja tie eksistē
        self.auto_load_best_settings()

        # Spilgtums
        ttk.Label(enhance_frame, text="Spilgtums:").pack(anchor="w")
        brightness_scale = ttk.Scale(enhance_frame, from_=-100, to=100, variable=self.app.scan_brightness,
                                     orient="horizontal", command=self.on_realtime_change)
        brightness_scale.pack(fill="x", pady=2)
        brightness_label = ttk.Label(enhance_frame, textvariable=self.app.scan_brightness)
        brightness_label.pack(anchor="w")

        # Kontrasts
        ttk.Label(enhance_frame, text="Kontrasts:").pack(anchor="w", pady=(5, 0))
        contrast_scale = ttk.Scale(enhance_frame, from_=-100, to=100, variable=self.app.scan_contrast,
                                   orient="horizontal", command=self.on_realtime_change)
        contrast_scale.pack(fill="x", pady=2)
        contrast_label = ttk.Label(enhance_frame, textvariable=self.app.scan_contrast)
        contrast_label.pack(anchor="w")

        # Gamma
        ttk.Label(enhance_frame, text="Gamma:").pack(anchor="w", pady=(5, 0))
        gamma_scale = ttk.Scale(enhance_frame, from_=0.1, to=3.0, variable=self.app.scan_gamma,
                                orient="horizontal", command=self.on_realtime_change)
        gamma_scale.pack(fill="x", pady=2)
        gamma_label = ttk.Label(enhance_frame, textvariable=self.app.scan_gamma)
        gamma_label.pack(anchor="w")

        # Krāsu detekcija
        color_frame = ttk.LabelFrame(left_panel, text="Krāsu detekcija", padding="10")
        color_frame.pack(fill="x", pady=8)

        color_check = ttk.Checkbutton(color_frame, text="Ieslēgt krāsu detekciju",
                                      variable=self.app.scan_use_color_detection,
                                      command=self.on_realtime_change)
        color_check.pack(anchor="w", pady=2)

        # Krāsu atlasītājs
        color_select_frame = ttk.Frame(color_frame)
        color_select_frame.pack(fill="x", pady=2)

        ttk.Label(color_select_frame, text="Mērķa krāsa:").pack(side="left")
        self.color_display = tk.Label(color_select_frame, width=3, height=1,
                                      bg=self.app.scan_target_color.get())
        self.color_display.pack(side="right", padx=5)

        ttk.Button(color_frame, text="🎨 Izvēlēties krāsu",
                   command=self.choose_color, width=25).pack(pady=2)
        ttk.Button(color_frame, text="👆 Atlasīt no attēla",
                   command=self.enable_color_picker, width=25).pack(pady=2)

        # Krāsu tolerance
        ttk.Label(color_frame, text="Krāsu tolerance:").pack(anchor="w", pady=(5, 0))
        tolerance_scale = ttk.Scale(color_frame, from_=1, to=100, variable=self.app.scan_color_tolerance,
                                    orient="horizontal", command=self.on_realtime_change)
        tolerance_scale.pack(fill="x", pady=2)
        tolerance_label = ttk.Label(color_frame, textvariable=self.app.scan_color_tolerance)
        tolerance_label.pack(anchor="w")

        # Detekcijas iestatījumi (kompaktāk)
        detection_frame = ttk.LabelFrame(left_panel, text="Detekcijas iestatījumi", padding="10")
        detection_frame.pack(fill="x", pady=8)

        ttk.Button(detection_frame, text="⚙️ Detalizēti iestatījumi",
                   command=lambda: self.app.show_scan_settings(self.preview_window), width=25).pack(pady=2)

        # Morfoloģija
        morph_check = ttk.Checkbutton(detection_frame, text="Morfoloģiskā apstrāde",
                                      variable=self.app.scan_morphology_enabled,
                                      command=self.on_realtime_change)
        morph_check.pack(anchor="w", pady=2)

        # === LABĀ PANEĻA SATURS ===

        # Canvas attēlam
        self.canvas = tk.Canvas(right_panel, bg="gray", cursor="fleur")
        self.canvas.pack(fill="both", expand=True)

        # Piesaista notikumus
        self.canvas.bind("<Configure>", self.on_canvas_resize)
        self.canvas.bind("<MouseWheel>", self.on_mouse_wheel)
        self.canvas.bind("<Button-4>", self.on_mouse_wheel)
        self.canvas.bind("<Button-5>", self.on_mouse_wheel)
        self.canvas.bind("<ButtonPress-1>", self.on_mouse_down)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_mouse_up)
        self.canvas.bind("<ButtonPress-2>", self.on_pan_start)
        self.canvas.bind("<B2-Motion>", self.on_pan_drag)
        self.canvas.bind("<ButtonRelease-2>", self.on_pan_end)

        # Apakšējās pogas
        button_frame = ttk.Frame(self.preview_window)
        button_frame.pack(fill="x", padx=5, pady=5)

        self.save_auto_button = ttk.Button(button_frame, text="🔍 Meklē dokumentu... (ENTER)",
                                           command=self.capture_and_process_frame,
                                           bootstyle="success", state="normal")  # VIENMĒR AKTĪVA!
        self.save_auto_button.pack(side="right", padx=5)

        # Kameras izvēles dropdown
        camera_frame = ttk.Frame(button_frame)
        camera_frame.pack(side="right", padx=5)

        ttk.Label(camera_frame, text="Kamera:", font=("Arial", 8)).pack(side="top")
        self.camera_var = tk.IntVar(value=getattr(self.app, 'current_camera_index', 0))
        self.camera_combo = ttk.Combobox(camera_frame, textvariable=self.camera_var,
                                         width=8, values=[0, 1, 2, 3], state="readonly")
        self.camera_combo.pack(side="top")
        self.camera_combo.bind('<<ComboboxSelected>>', self.on_camera_change)

        # Pievienot "Atsvaidzināt kameru" pogu
        self.refresh_camera_button = ttk.Button(button_frame, text="📷 Atsvaidzināt",
                                                command=self.refresh_camera_view,
                                                bootstyle="info")
        self.refresh_camera_button.pack(side="right", padx=5)

        ttk.Button(button_frame, text="❌ Pabeigt skenēšanu",
                   command=self.close_preview_window, bootstyle="danger").pack(side="right", padx=5)

        # Pievienot instrukciju tekstu
        instruction_label = ttk.Label(button_frame,
                                      text="ENTER - saglabāt | R - atsvaidzināt | Dropdown - mainīt kameru",
                                      font=("Arial", 9), foreground="blue")
        instruction_label.pack(side="left", padx=5)
        self.save_auto_button.pack(side="right", padx=5)

        # Inicializācija
        self.color_picker_mode = False
        self.preview_window.after(100, self.delayed_auto_detect)

    def on_realtime_change(self, *args):
        """Reāllaika iestatījumu maiņa."""
        # Atjauno krāsu displeja
        if hasattr(self, 'color_display'):
            self.color_display.config(bg=self.app.scan_target_color.get())

        # Ja reāllaika skenēšana ir aktīva, atjauno vizualizāciju
        if hasattr(self, 'live_scan_active') and self.live_scan_active:
            self.display_live_scan_preview()

        # JAUNS: Saglabā iestatījumus katru reizi, kad tie tiek mainīti
        self.app.save_app_settings()

    def choose_color(self):
        """Atver krāsu izvēles dialogu."""
        from tkinter import colorchooser
        color = colorchooser.askcolor(title="Izvēlieties dokumenta krāsu")
        if color[1]:  # Ja krāsa tika izvēlēta
            self.app.scan_target_color.set(color[1])
            self.on_realtime_change()

    def enable_color_picker(self):
        """Ieslēdz krāsu atlasīšanas režīmu."""
        self.color_picker_mode = True
        self.canvas.config(cursor="crosshair")
        messagebox.showinfo("Krāsu atlase", "Noklikšķiniet uz attēla, lai atlasītu dokumenta krāsu.")

    ''''def close_preview_window(self):
        """Droši aizver priekšskatījuma logu."""
        self.stop_live_scan()  # Aptur reāllaika skenēšanu

        # Notīra peles rullīša notikumus
        try:
            if hasattr(self, 'preview_window') and self.preview_window:
                self.preview_window.unbind_all("<MouseWheel>")
        except:
            pass

        if self.preview_window:
            self.preview_window.destroy()
            self.preview_window = None '''

    def on_mouse_down(self, event):
        """Apstrādā peles klikšķi."""
        # Ja krāsu atlasīšanas režīms ir aktīvs
        if hasattr(self, 'color_picker_mode') and self.color_picker_mode:
            self.pick_color_from_image(event)
            self.color_picker_mode = False
            self.canvas.config(cursor="fleur")
            return

        # Pārējā loģika stūru vilkšanai
        for handle in self.corner_handles:
            x1, y1, x2, y2 = self.canvas.coords(handle["id"])
            if x1 <= event.x <= x2 and y1 <= event.y <= y2:
                self.active_handle = handle
                self.canvas.config(cursor="hand2")
                break
        if self.active_handle is None:
            # Ja nav stūra marķieris, sāk pārvietošanu
            self.on_pan_start(event)

    def delayed_auto_detect(self):
        """Aizkavēta auto detekcija, lai nodrošinātu, ka canvas ir gatavs."""
        try:
            self.auto_detect_corners()
        except Exception as e:
            print(f"Kļūda auto detekcijā: {e}")
            # Ja auto detekcija neizdodas, iestatām noklusējuma stūrus
            if self.original_image_pil:
                img_w, img_h = self.original_image_pil.size
                margin = min(img_w, img_h) * 0.05
                self.corners = [
                    [margin, margin],
                    [img_w - margin, margin],
                    [img_w - margin, img_h - margin],
                    [margin, img_h - margin]
                ]
                self.display_image_on_canvas()

    def reset_view(self):
        self.zoom_factor = 1.0
        self.pan_x = 0
        self.pan_y = 0
        self.display_image_on_canvas()

    def change_zoom(self, factor):
        self.zoom_factor *= factor
        self.display_image_on_canvas()

    def auto_detect_corners(self):
        """Automātiski atrod dokumenta stūrus vai iestatīt manuāli."""
        if not self.corners:  # Ja nav iepriekš iestatīti stūri
            found_corners = self.find_document_corners()
            if not found_corners:
                messagebox.showinfo("Manuālā atlase",
                                    "Automātiski netika atrasts dokuments.\n\n" +
                                    "Tagad varat manuāli vilkt krāsainos stūru marķierus, " +
                                    "lai precīzi iezīmētu dokumenta robežas.\n\n" +
                                    "Instrukcijas:\n" +
                                    "• Vilkiet krāsainos apļus uz dokumenta stūriem\n" +
                                    "• Izmantojiet peles rullīti tālummaiņai\n" +
                                    "• Vilkiet ar vidējo pogu, lai pārvietotos")
                # Iestatīt noklusējuma stūrus uz visa attēla robežām
                img_w, img_h = self.original_image_pil.size
                margin = min(img_w, img_h) * 0.05  # 5% atkāpe no malām
                self.corners = [
                    [margin, margin],  # Augšā pa kreisi
                    [img_w - margin, margin],  # Augšā pa labi
                    [img_w - margin, img_h - margin],  # Apakšā pa labi
                    [margin, img_h - margin]  # Apakšā pa kreisi
                ]
            else:
                messagebox.showinfo("Automātiskā detekcija",
                                    "Dokuments veiksmīgi atrasts automātiski!\n\n" +
                                    "Jūs joprojām varat precizēt stūru pozīcijas, " +
                                    "vilkot krāsainos marķierus.")

        # Drošs izsaukums display_image_on_canvas
        try:
            self.display_image_on_canvas()
        except Exception as e:
            print(f"Kļūda attēlojot attēlu: {e}")
            # Mēģinām vēlreiz pēc īsa laika
            if self.canvas and self.preview_window:
                self.preview_window.after(100, self.display_image_on_canvas)

    def display_image_on_canvas(self):
        if self.original_image_pil is None or self.canvas is None:
            return

        # Pārbaudām, vai canvas joprojām eksistē
        try:
            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()
        except tk.TclError:
            # Canvas ir iznīcināts, izejam
            return

        if canvas_width <= 1 or canvas_height <= 1:
            # Canvas vēl nav gatavs, mēģinām vēlreiz pēc īsa laika
            self.canvas.after(50, self.display_image_on_canvas)
            return

        img_width, img_height = self.original_image_pil.size
        scaled_width = int(img_width * self.zoom_factor)
        scaled_height = int(img_height * self.zoom_factor)

        display_img = self.original_image_pil.resize((scaled_width, scaled_height), Image.LANCZOS)
        self.photo_image = ImageTk.PhotoImage(display_img)

        self.canvas.delete("all")
        self.corner_handles = []

        # Aprēķina attēla pozīciju ar pārvietošanu
        self.img_on_canvas_x = (canvas_width - scaled_width) / 2 + self.pan_x
        self.img_on_canvas_y = (canvas_height - scaled_height) / 2 + self.pan_y

        self.canvas.create_image(self.img_on_canvas_x, self.img_on_canvas_y, anchor="nw", image=self.photo_image)
        self.canvas.image = self.photo_image

        # Zīmē stūrus, ja tie ir definēti
        if len(self.corners) == 4:
            handle_size = max(8, int(12 / self.zoom_factor))

            # Zīmē līnijas starp stūriem
            points = []
            for corner in self.corners:
                x_on_canvas = self.img_on_canvas_x + corner[0] * self.zoom_factor
                y_on_canvas = self.img_on_canvas_y + corner[1] * self.zoom_factor
                points.extend([x_on_canvas, y_on_canvas])

            # Zīmē dokumenta kontūru
            self.canvas.create_polygon(points, outline="red", width=3, fill="", tags="corner_lines")

            # Zīmē stūru marķierus
            corner_colors = ["#FF6B6B", "#4ECDC4", "#45B7D1", "#96CEB4"]
            corner_labels = ["TL", "TR", "BR", "BL"]

            for i, corner in enumerate(self.corners):
                x_on_canvas = self.img_on_canvas_x + corner[0] * self.zoom_factor
                y_on_canvas = self.img_on_canvas_y + corner[1] * self.zoom_factor

                # Zīmē stūra marķieri
                handle_id = self.canvas.create_oval(
                    x_on_canvas - handle_size, y_on_canvas - handle_size,
                    x_on_canvas + handle_size, y_on_canvas + handle_size,
                    fill=corner_colors[i], outline="white", width=2, tags="corner_handle"
                )

                # Pievieno teksta etiķeti
                text_id = self.canvas.create_text(
                    x_on_canvas, y_on_canvas - handle_size - 15,
                    text=corner_labels[i], fill="white", font=("Arial", 10, "bold"),
                    tags="corner_label"
                )

                self.corner_handles.append({
                    "id": handle_id,
                    "text_id": text_id,
                    "index": i,
                    "original_x": corner[0],
                    "original_y": corner[1]
                })

            # Nodrošina, ka marķieri ir virspusē
            self.canvas.tag_raise("corner_handle")
            self.canvas.tag_raise("corner_label")

        # Pievieno instrukciju tekstu
        instruction_text = (
            "Instrukcijas:\n"
            "• Vilkiet krāsainos stūru marķierus\n"
            "• Peles rullītis: tālummaiņa\n"
            "• Vidējā poga: pārvietošana"
        )
        self.canvas.create_text(
            10, 10, text=instruction_text, anchor="nw",
            fill="yellow", font=("Arial", 10), tags="instructions"
        )

    def on_canvas_resize(self, event):
        self.display_image_on_canvas()

    def on_mouse_wheel(self, event):
        if event.num == 5 or event.delta == -120:  # Zoom out
            self.zoom_factor = max(0.1, self.zoom_factor * 0.9)
        if event.num == 4 or event.delta == 120:  # Zoom in
            self.zoom_factor = min(5.0, self.zoom_factor * 1.1)
        self.display_image_on_canvas()

    def on_pan_start(self, event):
        self.start_pan_x = event.x - self.pan_x
        self.start_pan_y = event.y - self.pan_y
        self.canvas.config(cursor="fleur")

    def on_pan_drag(self, event):
        self.pan_x = event.x - self.start_pan_x
        self.pan_y = event.y - self.start_pan_y
        self.display_image_on_canvas()

    def on_pan_end(self, event):
        self.canvas.config(cursor="arrow")

    def on_mouse_down(self, event):
        for handle in self.corner_handles:
            x1, y1, x2, y2 = self.canvas.coords(handle["id"])
            if x1 <= event.x <= x2 and y1 <= event.y <= y2:
                self.active_handle = handle
                self.canvas.config(cursor="hand2")
                break
        if self.active_handle is None:
            # If no handle is clicked, start pan
            self.on_pan_start(event)

    def on_mouse_drag(self, event):
        if self.active_handle:
            # Konvertē kanvasa koordinātas atpakaļ uz oriģinālā attēla koordinātām
            new_x_original = (event.x - self.img_on_canvas_x) / self.zoom_factor
            new_y_original = (event.y - self.img_on_canvas_y) / self.zoom_factor

            # Ierobežo koordinātas attēla robežās
            img_w, img_h = self.original_image_pil.size
            new_x_original = max(0, min(new_x_original, img_w))
            new_y_original = max(0, min(new_y_original, img_h))

            self.corners[self.active_handle["index"]] = [new_x_original, new_y_original]

            # Atjauno tikai aktīvo marķieri, lai uzlabotu veiktspēju
            self.update_active_corner_display()

        elif self.active_handle is None:
            # Turpina pārvietošanu, ja nav aktīvs marķieris
            self.on_pan_drag(event)

    def update_active_corner_display(self):
        """Atjauno tikai aktīvā stūra marķiera attēlojumu."""
        if not self.active_handle or not self.corners or not self.canvas:
            return

        try:
            i = self.active_handle["index"]
            corner = self.corners[i]

            x_on_canvas = self.img_on_canvas_x + corner[0] * self.zoom_factor
            y_on_canvas = self.img_on_canvas_y + corner[1] * self.zoom_factor

            handle_size = max(8, int(12 / self.zoom_factor))

            # Atjauno marķiera pozīciju
            self.canvas.coords(
                self.active_handle["id"],
                x_on_canvas - handle_size, y_on_canvas - handle_size,
                x_on_canvas + handle_size, y_on_canvas + handle_size
            )

            # Atjauno teksta pozīciju
            if "text_id" in self.active_handle:
                self.canvas.coords(
                    self.active_handle["text_id"],
                    x_on_canvas, y_on_canvas - handle_size - 15
                )

            # Atjauno kontūru
            if len(self.corners) == 4:
                points = []
                for corner in self.corners:
                    points.extend([
                        self.img_on_canvas_x + corner[0] * self.zoom_factor,
                        self.img_on_canvas_y + corner[1] * self.zoom_factor
                    ])

                # Atrod un atjauno kontūras līnijas
                for item in self.canvas.find_withtag("corner_lines"):
                    self.canvas.coords(item, *points)

        except tk.TclError:
            # Canvas ir iznīcināts, ignorējam
            pass
        except Exception as e:
            print(f"Kļūda atjaunojot stūra attēlojumu: {e}")

    def on_mouse_up(self, event):
        self.active_handle = None
        self.canvas.config(cursor="arrow")
        self.on_pan_end(event)  # Ensure pan cursor is reset

    def apply_document_correction(self):
        if self.original_image_pil is None or not self.corners:
            messagebox.showwarning("Kļūda", "Nav attēla vai nav definēti stūri.")
            return

        try:
            # Convert PIL image to OpenCV format
            img_cv = np.array(self.original_image_pil.convert("RGB"))
            img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

            # Apply perspective transform
            warped_cv = self.four_point_transform(img_cv, self.corners)

            # Convert back to PIL image
            self.processed_image_pil = Image.fromarray(cv2.cvtColor(warped_cv, cv2.COLOR_BGR2RGB))

            # Update the current image in the main app
            if self.app.current_image_index != -1:
                self.app.images[self.app.current_image_index]["processed_img"] = self.processed_image_pil
                self.app.show_image_preview(self.processed_image_pil)
                messagebox.showinfo("Korekcija veiksmīga", "Dokumenta robežas veiksmīgi koriģētas.")
            else:
                messagebox.showwarning("Korekcija veiksmīga",
                                       "Dokumenta robežas veiksmīgi koriģētas, bet attēls nav aktīvs galvenajā sarakstā.")

            self.preview_window.destroy()

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās pielietot dokumenta korekciju: {e}")


class OCRPDFApp(Window):
    """Galvenā lietojumprogrammas klase OCR un PDF ģenerēšanai."""

    def __init__(self):
        super().__init__(themename="darkly")

        # Paslēpt galveno logu
        self.withdraw()

        # Ielādes logs
        loading_window = show_loading_screen(self)
        self.wait_window(loading_window)

        # Jauni mainīgie attēla apgriešanai tieši uz kanvasa
        self.cropping_mode = False  # Norāda, vai apgriešanas režīms ir aktīvs
        self.crop_start_x = None
        self.crop_start_y = None
        self.crop_rect_id = None
        self.current_crop_coords = None  # Glabās pēdējās apgriešanas koordinātas
        # Rādīt galveno logu
        self.deiconify()

        # MAINĪTS: Izdarīta maksimizācija (nevis pilnekrāns)
        self.state('zoomed')  # Logs aizņems visu darba laukumu, bet saglabā kontroljoslas

        self.title("Advanced OCR uz PDF")
        # Sākotnējais izmērs un minimālais izmērs, kas labāk piemērots mazākiem ekrāniem
        self.geometry("1024x768")  # Samazināts noklusējuma izmērs
        self.minsize(800, 500)  # Samazināts minimālais izmērs
        self.settings = {}  # Inicializējiet settings kā tukšu vārdnīcu
        # JAUNS: Skenēšanas iestatījumu mainīgie tagad tiek inicializēti no self.settings
        self.scan_camera_index = tk.IntVar(value=self.settings.get("scan_camera_index", 1))
        self.scan_camera_width = tk.IntVar(value=self.settings.get("scan_camera_width", 1280))
        self.scan_camera_height = tk.IntVar(value=self.settings.get("scan_camera_height", 120))
        self.scan_min_contour_area = tk.IntVar(value=self.settings.get("scan_min_contour_area", 2500))
        self.scan_stable_threshold = tk.DoubleVar(value=self.settings.get("scan_stable_threshold", 0.8))
        self.scan_stability_tolerance = tk.DoubleVar(value=self.settings.get("scan_stability_tolerance", 0.01))
        self.scan_aspect_ratio_min = tk.DoubleVar(value=self.settings.get("scan_aspect_ratio_min", 0.4))
        self.scan_aspect_ratio_max = tk.DoubleVar(value=self.settings.get("scan_aspect_ratio_max", 2.3))
        self.scan_gaussian_blur_kernel = tk.IntVar(value=self.settings.get("scan_gaussian_blur_kernel", 9))
        self.scan_adaptive_thresh_block_size = tk.IntVar(value=self.settings.get("scan_adaptive_thresh_block_size", 11))
        self.scan_adaptive_thresh_c = tk.IntVar(value=self.settings.get("scan_adaptive_thresh_c", 3))
        self.scan_canny_thresh1 = tk.IntVar(value=self.settings.get("scan_canny_thresh1", 610))
        self.scan_canny_thresh2 = tk.IntVar(value=self.settings.get("scan_canny_thresh2", 190))
        self.scan_brightness = tk.IntVar(value=self.settings.get("scan_brightness", 0))
        self.scan_contrast = tk.IntVar(value=self.settings.get("scan_contrast", 0))
        self.scan_saturation = tk.IntVar(value=self.settings.get("scan_saturation", 0))
        self.scan_gamma = tk.DoubleVar(value=self.settings.get("scan_gamma", 1.0))
        self.scan_use_color_detection = tk.BooleanVar(value=self.settings.get("scan_use_color_detection", False))
        self.scan_target_color = tk.StringVar(value=self.settings.get("scan_target_color", "#FFFFFF"))
        self.scan_color_tolerance = tk.IntVar(value=self.settings.get("scan_color_tolerance", 30))
        self.scan_morphology_enabled = tk.BooleanVar(value=self.settings.get("scan_morphology_enabled", False))
        self.scan_morphology_kernel_size = tk.IntVar(value=self.settings.get("scan_morphology_kernel_size", 3))
        self.scan_edge_dilation = tk.IntVar(value=self.settings.get("scan_edge_dilation", 2))
        self.document_scanner = DocumentScanner(self)
        self.camera = None
        self.camera_active = False
        self.qr_code_frame_coords = None  # Pievienot šo rindu
        self.camera = None  # Kameras objekts (piem., cv2.VideoCapture)
        self.camera_active = False  # Kameras statusa karogs
        self.scan_settings = {}  # Šī rinda paliek, lai saglabātu tukšu vārdnīcu, kas tiks aizpildīta ar sync_scan_settings_from_vars
        # Iestatījumu faili joprojām tiek glabāti lietotāja profilā, jo tie ir lietotāja dati.
        # Tie netiek iekļauti ZIP arhīvā, jo tie ir mainīgi.
        self.settings_file = os.path.join(os.path.expanduser("~"), "AdvancedOCR_settings.json")
        self.scan_settings_file = os.path.join(os.path.expanduser("~"), "AdvancedOCR_scan_settings.json")
        self.pdf_archive_file = os.path.join(os.path.expanduser("~"), "AdvancedOCR_archive.json")
        self.scan_folder_path = tk.StringVar(
            value=os.path.join(os.path.expanduser("~"), "ScannedDocuments"))  # JAUNS: Skenēšanas mapes ceļš
        self.auto_scan_enabled = tk.BooleanVar(value=False)  # JAUNS: Automātiskās skenēšanas ieslēgšana/izslēgšana
        self.observer = None  # JAUNS: Watchdog observers

        # Skenēšanas iestatījumi
        self.scan_camera_index = tk.IntVar(value=1)
        self.scan_camera_width = tk.IntVar(value=1920)
        self.scan_camera_height = tk.IntVar(value=1080)
        self.current_camera_index = 1


        # JAUNS: Google Sheets iestatījumi
        self.google_sheet_id = tk.StringVar(value=self.settings.get("google_sheet_id", ""))
        self.google_sheet_name = tk.StringVar(value=self.settings.get("google_sheet_name", "OCR_Failu_Saraksts"))
        self.google_sheet_credentials_path = tk.StringVar(
            value=self.settings.get("google_sheet_credentials_path", "google_sheet_credentials.json"))
        self.google_sheet_service = None  # Tiks inicializēts pēc autentifikācijas
        self.google_drive_service = None  # Tiks inicializēts pēc autentifikācijas

        # JAUNS: Mainīgie PDF priekšskatījumam "Papildu rīki" cilnē
        # Šie mainīgie tagad atspoguļos self.images saraksta saturu
        self.additional_tools_pdf_preview_canvas = None
        self.additional_tools_pdf_preview_photo = None
        self.additional_tools_current_pdf_document = None  # Tiks ielādēts, ja atlasītais fails ir PDF
        self.additional_tools_current_pdf_page_count = 0
        self.additional_tools_current_pdf_page_index = 0
        self.additional_tools_pdf_preview_zoom_factor = 1.0
        self.additional_tools_pdf_preview_pan_x = 0
        self.additional_tools_pdf_preview_pan_y = 0
        self.additional_tools_pdf_preview_start_pan_x = 0
        self.additional_tools_pdf_preview_start_pan_y = 0
        self.additional_tools_pdf_page_label = None
        self.additional_tools_prev_page_button = None
        self.additional_tools_next_page_button = None

        self.title("Advanced OCR uz PDF")
        # Sākotnējais izmērs un minimālais izmērs, kas labāk piemērots mazākiem ekrāniem
        self.geometry("1024x768")  # Samazināts noklusējuma izmērs
        self.minsize(800, 500)  # Samazināts minimālais izmērs
        self.settings = {}  # Inicializējiet settings kā tukšu vārdnīcu
        self.scan_settings = {}  # JAUNS: Inicializējiet skenēšanas iestatījumus
        self.settings_file = os.path.join(os.path.expanduser("~"), "ocr_pdf_settings.json")
        self.scan_settings_file = os.path.join(os.path.expanduser("~"),
                                               "ocr_scan_settings.json")  # JAUNS: Skenēšanas iestatījumu fails
        self.pdf_archive_file = os.path.join(os.path.expanduser("~"), "ocr_pdf_archive.json")
        self.scan_folder_path = tk.StringVar(
            value=os.path.join(os.path.expanduser("~"), "ScannedDocuments"))  # JAUNS: Skenēšanas mapes ceļš
        self.auto_scan_enabled = tk.BooleanVar(value=False)  # JAUNS: Automātiskās skenēšanas ieslēgšana/izslēgšana
        self.observer = None  # JAUNS: Watchdog observers

        self.camera = None
        self.camera_active = False
        self.current_camera_index = 0  # Pievienot šo rindu

        # JAUNS: Attālinātās glabāšanas iestatījumi
        self.remote_storage_type = tk.StringVar(value=self.settings.get("remote_storage_type", "Local"))
        self.ftp_host = tk.StringVar(value=self.settings.get("ftp_host", ""))
        self.ftp_port = tk.IntVar(value=self.settings.get("ftp_port", 21))
        self.ftp_user = tk.StringVar(value=self.settings.get("ftp_user", ""))
        self.ftp_pass = tk.StringVar(value=self.settings.get("ftp_pass", ""))
        self.ftp_remote_path = tk.StringVar(value=self.settings.get("ftp_remote_path", "/"))
        self.ftp_use_sftp = tk.BooleanVar(value=self.settings.get("ftp_use_sftp", False))

        self.file_paths_to_open = []
        if len(sys.argv) > 1:
            self.file_paths_to_open = [arg for arg in sys.argv[1:] if arg.lower().endswith('.pdf')]

        # Iestatam callback pēc pilnīgas inicializācijas
        self.after(100, self.check_files_to_open)

        self.google_drive_folder_id = tk.StringVar(value=self.settings.get("google_drive_folder_id", ""))
        self.google_drive_credentials_path = tk.StringVar(
            value=self.settings.get("google_drive_credentials_path", "credentials.json"))
        self.google_drive_token_path = tk.StringVar(value=self.settings.get("google_drive_token_path", "token.json"))

        self.auto_upload_enabled = tk.BooleanVar(value=self.settings.get("auto_upload_enabled", False))
        self.auto_upload_target = tk.StringVar(
            value=self.settings.get("auto_upload_target", "Local"))  # Local, FTP, GoogleDrive

        self.file_listbox = tk.Listbox(self)  # Inicializē file_listbox
        self.file_listbox.bind('<<ListboxSelect>>', self.on_file_select)

        # JAUNS: Mainīgie PDF priekšskatījumam
        self.pdf_preview_canvas = None
        self.pdf_preview_photo = None
        self.current_pdf_document = None  # Lai glabātu atvērtu fitz dokumentu
        self.current_pdf_page_count = 0
        self.current_pdf_page_index = 0
        self.pdf_preview_zoom_factor = 1.0
        self.pdf_preview_pan_x = 0
        self.pdf_preview_pan_y = 0
        self.pdf_preview_start_pan_x = 0
        self.pdf_preview_start_pan_y = 0

        # Konfigurē krāsas
        self.file_listbox.configure(
            selectbackground='#d4edda',  # Zaļa atlases krāsa
            selectforeground='white'
        )

        self._selected_line_index = -1 # Inicializē atlasītās rindas indeksu

        # self.load_scan_settings()  # Šī rinda vairs nav nepieciešama, jo scan_settings tiek ielādēti caur app_settings
        self.load_app_settings()  # Ielādējiet galvenos iestatījumus
        # self.load_scan_settings()  # JAUNS: Ielādē skenēšanas iestatījumus
        # self.init_scan_settings()  # Inicializē skenēšanas iestatījumus ar ielādētajām vērtībām

        # self.gaussian_blur_kernel_var = tk.IntVar(value=self.scan_settings.get("scan_gaussian_blur_kernel", 5)) # Šī rinda vairs nav nepieciešama, jo tiek inicializēta init_scan_settings
        # Pievienojiet šo rindu, lai apstrādātu loga aizvēršanu
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.images = []
        self.ocr_results = []
        self.stop_processing = False
        self.default_save_path = r"C:\Users\edgar\Downloads\Advanced OCR"
        self.current_image_index = -1
        self.pdf_quality = "Vidēja"
        self.document_keywords = {
            "id_card": ["id karte", "personas apliecība", "identity card", "passport", "pase", "vadītāja apliecība",
                        "driver's license", "bankas karte", "credit card", "debit card"],
            # Pievienojiet citus atslēgvārdus, ja nepieciešams
        }

        self.internal_file_system = {"type": "folder", "name": "Sakne", "contents": []}
        self.current_folder = self.internal_file_system
        self.load_internal_file_system()

        self.lang_options = {
            "Latviešu (lav)": "lav", "Angļu (eng)": "eng", "Krievu (rus)": "rus",
            "Vācu (deu)": "deu", "Franču (fra)": "fra", "Spāņu (spa)": "spa",
            "Itāļu (ita)": "ita", "Lietuviešu (lit)": "lit", "Igauņu (est)": "est"
        }

        self.orientation_options = [
            "Auto", "Portrets", "Ainava", "A4 Portrets", "A4 Ainava",
            "Letter Portrets", "Letter Ainava", "Tāds pats kā attēls"
        ]

        self.lang_vars = {}
        for lang_name in self.lang_options.keys():
            self.lang_vars[lang_name] = tk.BooleanVar(value=(lang_name == "Angļu (eng)"))

        self.language_var = tk.StringVar(value="eng")
        self.output_dir_var = tk.StringVar(value=os.path.expanduser("~"))
        self.output_format_var = tk.StringVar(value=self.settings.get("output_format", "pdf"))
        self.psm_var = tk.IntVar(value=self.settings.get("psm", 3))
        self.oem_var = tk.IntVar(value=self.settings.get("oem", 3))

        self.brightness_var = tk.DoubleVar(value=1.0)
        self.contrast_var = tk.DoubleVar(value=1.0)
        self.sharpness_var = tk.DoubleVar(value=1.0)
        self.rotate_var = tk.IntVar(value=0)
        self.grayscale_var = tk.BooleanVar(value=True)
        self.deskew_var = tk.BooleanVar(value=False)
        self.remove_noise_var = tk.BooleanVar(value=False)
        self.invert_colors_var = tk.BooleanVar(value=False)
        self.edge_detection_var = tk.BooleanVar(value=False)
        self.binarize_var = tk.BooleanVar(value=False)

        # JAUNS: Skenēšanas iestatījumu mainīgie tagad tiek inicializēti no self.scan_settings
        self.scan_camera_index = tk.IntVar(value=self.scan_settings.get("scan_camera_index", 1))
        self.scan_camera_width = tk.IntVar(value=self.scan_settings.get("scan_camera_width", 1280))
        self.scan_camera_height = tk.IntVar(value=self.scan_settings.get("scan_camera_height", 120))
        self.scan_min_contour_area = tk.IntVar(value=self.scan_settings.get("scan_min_contour_area", 2500))
        self.scan_stable_threshold = tk.DoubleVar(value=self.scan_settings.get("scan_stable_threshold", 0.8))
        self.scan_stability_tolerance = tk.DoubleVar(value=self.scan_settings.get("scan_stability_tolerance", 0.01))
        self.scan_aspect_ratio_min = tk.DoubleVar(value=self.scan_settings.get("scan_aspect_ratio_min", 0.4))
        self.scan_aspect_ratio_max = tk.DoubleVar(value=self.scan_settings.get("scan_aspect_ratio_max", 2.3))
        self.scan_gaussian_blur_kernel = tk.IntVar(value=self.scan_settings.get("scan_gaussian_blur_kernel", 9))
        self.scan_adaptive_thresh_block_size = tk.IntVar(
            value=self.scan_settings.get("scan_adaptive_thresh_block_size", 11))
        self.scan_adaptive_thresh_c = tk.IntVar(value=self.scan_settings.get("scan_adaptive_thresh_c", 3))
        self.scan_canny_thresh1 = tk.IntVar(value=self.scan_settings.get("scan_canny_thresh1", 610))
        self.scan_canny_thresh2 = tk.IntVar(value=self.scan_settings.get("scan_canny_thresh2", 190))

        self.scan_settings = {}  # JAUNS: Inicializējiet skenēšanas iestatījumus
        self.document_scanner = DocumentScanner(self)
        self.camera = None
        self.camera_active = False
        self.qr_code_frame_coords = None  # Pievienot šo rindu
        self.camera = None  # Kameras objekts (piem., cv2.VideoCapture)
        self.camera_active = False  # Kameras statusa karogs

        # Pievienojiet šo, lai nodrošinātu, ka kamera tiek atbrīvota, kad lietotne tiek aizvērta
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.settings_file = os.path.join(os.path.expanduser("~"), "ocr_pdf_settings.json")

        self.init_scan_settings()  # Inicializē skenēšanas iestatījumus
        # self.document_scanner = DocumentScanner(self)  # Inicializē DocumentScanner

        self.create_widgets()
        self.configure_grid()
        self.create_menu()

        # Pielāgo iestatījumus no ielādētajām vērtībām
        self.orientation_var.set(self.settings.get("default_pdf_page_size"))
        self.fontsize_var.set(self.settings.get("default_pdf_font_size"))
        self.pdf_quality = self.settings.get("pdf_quality")
        pytesseract.pytesseract.tesseract_cmd = self.settings.get("tesseract_path")
        self.default_save_path = self.settings.get("default_save_path")

        # JAUNS: Skenēšanas iestatījumu mainīgo atjaunināšana vairs nav nepieciešama šeit, jo tie tiek ielādēti no `load_scan_settings`
        # un tiek atjaunināti caur `ScanSettingsWindow`

        # Atjaunina OCR valodu mainīgos
        for lang_name in self.lang_options.keys():
            if lang_name in self.settings.get("selected_ocr_languages"):
                self.lang_vars[lang_name].set(True)
            else:
                self.lang_vars[lang_name].set(False)

        # Atjaunina arī citus mainīgos, ja tie tiek ielādēti no settings
        self.output_format_var.set(self.settings.get("output_format"))
        self.psm_var.set(self.settings.get("psm"))
        self.oem_var.set(self.settings.get("oem"))
        self.language_var.set(self.settings.get("language"))
        self.output_dir_var.set(self.settings.get("output_dir"))

        # Atjaunina loga izmērus un pozīciju, nodrošinot, ka tas ir ekrāna robežās
        try:
            # Ielādē saglabātos izmērus
            saved_width = self.settings.get('window_width', 1024)
            saved_height = self.settings.get('window_height', 768)
            saved_x = self.settings.get('window_x', 0)
            saved_y = self.settings.get('window_y', 0)

            # Pārbauda ekrāna izmērus
            screen_width = self.winfo_screenwidth()
            screen_height = self.winfo_screenheight()

            # Pielāgo loga pozīciju, ja tas ir ārpus ekrāna
            if saved_x + saved_width > screen_width or saved_x < 0:
                saved_x = (screen_width - saved_width) // 2
            if saved_y + saved_height > screen_height or saved_y < 0:
                saved_y = (screen_height - saved_height) // 2

            # Nodrošina, ka loga izmēri nav mazāki par minimālajiem
            saved_width = max(saved_width,
                              self.winfo_reqwidth())  # winfo_reqwidth() atgriež minimālo nepieciešamo platumu
            saved_height = max(saved_height,
                               self.winfo_reqheight())  # winfo_reqheight() atgriež minimālo nepieciešamo augstumu

            self.geometry(f"{saved_width}x{saved_height}+{saved_x}+{saved_y}")
        except Exception as e:
            print(f"Nevarēja atjaunot loga izmērus/pozīciju: {e}")
            # Ja rodas kļūda, atiestata uz noklusējuma izmēriem un centrē
            self.geometry("1024x768")
            self.update_idletasks()  # Atjaunina, lai iegūtu pareizus izmērus centrēšanai
            x = (self.winfo_screenwidth() - self.winfo_width()) // 2
            y = (self.winfo_screenheight() - self.winfo_height()) // 2
            self.geometry(f"+{x}+{y}")

            # Pievienot metodes beigās
            self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def _get_physical_path_from_node(self, node):
        """Atgriež pilnu fizisko ceļu uz mapi vai failu no mezgla struktūras."""
        path_parts = []
        temp = node
        # Traverse up the parent chain until the root (internal_file_system)
        # Pievienots nosacījums, lai apstātos, ja temp ir None (aizsardzība)
        while temp and temp != self.internal_file_system:
            path_parts.insert(0, temp["name"])
            temp = temp.get("parent")
        # Construct the full path starting from default_save_path
        return os.path.join(self.default_save_path, *path_parts)

    def check_files_to_open(self):
        """Atver failus no komandrindas argumentiem"""
        if self.file_paths_to_open:
            for filepath in self.file_paths_to_open:
                if os.path.exists(filepath):
                    self.open_files(filepath)

    def show_document_detection_menu(self):
        """Parāda dokumentu detekcijas logu ar pašreizējo attēlu."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, ko apstrādāt.")
            return

        current_image_pil = self.images[self.current_image_index]["processed_img"]
        self.document_scanner.set_image(current_image_pil)
        self.document_scanner.show_document_detection_preview()

    def _get_physical_path_from_node(self, node):
        """Atgriež pilnu fizisko ceļu uz mapi no mezgla struktūras."""
        path_parts = []
        temp = node
        # Traverse up the parent chain until the root (internal_file_system)
        while temp and temp != self.internal_file_system:
            path_parts.insert(0, temp["name"])
            temp = temp.get("parent")
        # Construct the full path starting from default_save_path
        return os.path.join(self.default_save_path, *path_parts)

    def init_camera(self, force_camera_index=None):
        """Kameras inicializācija ar iespēju norādīt konkrētu kameru."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas", "Nepieciešams opencv-python.")
            return False

        # Ja kamera jau ir atvērta un nav pieprasīta konkrēta kamera
        if self.camera is not None and force_camera_index is None:
            return True

        # Atbrīvo esošo kameru, ja vajag mainīt
        if self.camera is not None:
            self.camera.release()
            self.camera = None
            self.camera_active = False

        try:
            # Nosaka kameras indeksu
            if force_camera_index is not None:
                camera_index = force_camera_index
                print(f"🎯 Piespiedu kārtā izmanto kameru: {camera_index}")
            elif hasattr(self, 'scan_camera_index'):
                camera_index = self.scan_camera_index.get()
                print(f"📋 Iestatījumos norādītā kamera: {camera_index}")
            else:
                camera_index = 0
                print("⚠️ Nav atrasts scan_camera_index, izmanto 0")

            print(f"🔍 Mēģina atvērt kameru {camera_index}")

            # Atver norādīto kameru
            self.camera = cv2.VideoCapture(camera_index)
            if not self.camera.isOpened():
                print(f"❌ Kamera {camera_index} nav pieejama")
                raise IOError(f"Kamera {camera_index} nav pieejama")

            print(f"✅ Veiksmīgi atvērta kamera {camera_index}")

            # Iestata kvalitāti
            self.camera.set(cv2.CAP_PROP_FRAME_WIDTH, 1920)
            self.camera.set(cv2.CAP_PROP_FRAME_HEIGHT, 1080)
            self.camera.set(cv2.CAP_PROP_BUFFERSIZE, 1)

            # Saglabā pašreizējo kameras indeksu
            self.current_camera_index = camera_index

            actual_width = int(self.camera.get(cv2.CAP_PROP_FRAME_WIDTH))
            actual_height = int(self.camera.get(cv2.CAP_PROP_FRAME_HEIGHT))
            print(f"📐 Kamera {camera_index}: {actual_width}x{actual_height}")

            self.camera_active = True
            return True

        except Exception as e:
            print(f"❌ Kameras {camera_index} kļūda: {e}")
            messagebox.showerror("Kameras kļūda", f"Nevar atvērt kameru {camera_index}: {e}")
            if self.camera:
                self.camera.release()
            self.camera = None
            self.camera_active = False
            return False

    def release_camera(self):
        """Atbrīvo kameras resursus."""
        if self.camera is not None:
            self.camera.release()
            self.camera = None
            self.camera_active = False

    def get_camera_frame(self):
        """Iegūst pašreizējo kadru no kameras kā PIL attēlu."""
        if self.camera is None or not self.camera_active:
            return None

        ret, frame = self.camera.read()
        if not ret:
            print("Neizdevās iegūt kadru no kameras.")
            return None

        # Pārveido OpenCV kadru par PIL attēlu
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        return Image.fromarray(frame_rgb)

    def scan_document_with_camera(self):
        """Ātri sāk dokumenta skenēšanu ar kameru."""

        # DEBUG: Pārbauda kameras iestatījumus
        print("🔍 DEBUG: Kameras iestatījumi:")
        print(f"scan_camera_index eksistē: {hasattr(self, 'scan_camera_index')}")
        if hasattr(self, 'scan_camera_index'):
            print(f"scan_camera_index vērtība: {self.scan_camera_index.get()}")

        # Vienkāršs progress bez animācijas
        loading_window = tk.Toplevel(self)
        loading_window.title("Kameru...")
        loading_window.geometry("200x60")
        loading_window.transient(self)
        loading_window.resizable(False, False)

        # Centrē
        loading_window.update_idletasks()
        x = (loading_window.winfo_screenwidth() // 2) - 100
        y = (loading_window.winfo_screenheight() // 2) - 30
        loading_window.geometry(f"200x60+{x}+{y}")

        label = ttk.Label(loading_window, text="Atver kameru...")
        label.pack(expand=True)
        loading_window.update()

        try:
            # Ātri inicializē kameru
            if not self.camera_active:
                if not self.init_camera():
                    loading_window.destroy()
                    return

            # Ātri iegūst kadru
            first_frame = self.get_camera_frame()
            loading_window.destroy()

            if first_frame:
                self.document_scanner.set_image(first_frame)
                self.document_scanner.document_frozen = False
                self.document_scanner.live_detected_corners = []
                self.document_scanner.show_document_detection_preview()
                self.document_scanner.start_live_scan()
            else:
                messagebox.showwarning("Kļūda", "Nav kameras kadra.")
                self.release_camera()

        except Exception as e:
            loading_window.destroy()
            messagebox.showerror("Kļūda", f"Kameras kļūda: {e}")

    def on_closing(self):
        """Apstrādā lietotnes aizvēršanu, atbrīvojot kameras resursus."""
        self.release_camera()
        self.save_app_settings()  # Saglabā iestatījumus pirms aizvēršanas
        self.destroy()

    def _display_pdf_page_on_canvas(self):
        """Attēlo pašreizējo PDF lapu uz priekšskatījuma kanvasa."""
        if not self.current_pdf_document or not self.pdf_preview_canvas:
            return

        try:
            # Iegūst kanvasa izmērus
            canvas_width = self.pdf_preview_canvas.winfo_width()
            canvas_height = self.pdf_preview_canvas.winfo_height()

            if canvas_width <= 1 or canvas_height <= 1:
                # Kanvass vēl nav gatavs, mēģinām vēlreiz pēc īsa laika
                self.after(50, self._display_pdf_page_on_canvas)
                return

            # Ielādē lapu
            page = self.current_pdf_document.load_page(self.current_pdf_page_index)

            # Konvertē lapu uz attēlu (PIL Image)
            # Izmantojam DPI, lai kontrolētu attēla kvalitāti/izmēru
            # Pielāgojam DPI, lai attēls ietilptu kanvasā, bet nebūtu pārāk liels
            # Noklusējuma DPI 72 ir labs priekšskatījumam

            # Aprēķina sākotnējo tālummaiņas koeficientu, lai lapa ietilptu kanvasā
            # ņemot vērā gan platumu, gan augstumu.
            fit_width_zoom = canvas_width / page.rect.width
            fit_height_zoom = canvas_height / page.rect.height

            # Izvēlas mazāko tālummaiņas koeficientu, lai visa lapa būtu redzama
            initial_fit_zoom = min(fit_width_zoom, fit_height_zoom)

            # Pielieto lietotāja definēto tālummaiņas faktoru virs sākotnējās pielāgošanas
            zoom_factor_for_render = self.pdf_preview_zoom_factor * initial_fit_zoom

            # Nodrošina minimālo tālummaiņu, lai attēls nebūtu pārāk mazs
            # Var pielāgot 0.1, ja nepieciešams, bet parasti nav vajadzīgs, ja initial_fit_zoom ir pareizi aprēķināts
            zoom_factor_for_render = max(0.1, zoom_factor_for_render)

            pix = page.get_pixmap(matrix=fitz.Matrix(zoom_factor_for_render, zoom_factor_for_render))
            img_pil = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Pārvēršam PIL attēlu par PhotoImage
            self.pdf_preview_photo = ImageTk.PhotoImage(img_pil)

            # Notīra kanvasu un attēlo jauno attēlu
            self.pdf_preview_canvas.delete("all")

            # Aprēķina attēla pozīciju ar pārvietošanu
            img_width, img_height = img_pil.size
            x = (canvas_width - img_width) / 2 + self.pdf_preview_pan_x
            y = (canvas_height - img_height) / 2 + self.pdf_preview_pan_y

            self.pdf_preview_canvas.create_image(x, y, anchor="nw", image=self.pdf_preview_photo)
            self.pdf_preview_canvas.image = self.pdf_preview_photo # Saglabā atsauci

            # Atjaunina lapas numura etiķeti
            self.pdf_page_label.config(text=f"Lapa: {self.current_pdf_page_index + 1}/{self.current_pdf_page_count}")

            # Atjaunina navigācijas pogu stāvokli
            self.prev_page_button.config(state=tk.NORMAL if self.current_pdf_page_index > 0 else tk.DISABLED)
            self.next_page_button.config(state=tk.NORMAL if self.current_pdf_page_index < self.current_pdf_page_count - 1 else tk.DISABLED)

        except Exception as e:
            print(f"Kļūda attēlojot PDF lapu: {e}")
            self.pdf_preview_canvas.delete("all")
            self.pdf_preview_canvas.create_text(
                self.pdf_preview_canvas.winfo_width() / 2, self.pdf_preview_canvas.winfo_height() / 2,
                text=f"Nevarēja ielādēt lapu:\n{e}", fill="red", font=("Helvetica", 12),
                justify="center"
            )
            self.pdf_page_label.config(text="Lapa: Kļūda")
            self.prev_page_button.config(state=tk.DISABLED)
            self.next_page_button.config(state=tk.DISABLED)


    def _load_pdf_for_preview(self, filepath):
        """Ielādē PDF dokumentu priekšskatījumam."""
        # Aizver iepriekšējo dokumentu, ja tāds ir
        if self.current_pdf_document:
            self.current_pdf_document.close()
            self.current_pdf_document = None

        self.pdf_preview_canvas.delete("all")
        self.pdf_preview_canvas.create_text(
            self.pdf_preview_canvas.winfo_width() / 2, self.pdf_preview_canvas.winfo_height() / 2,
            text="Ielādē...", fill="white", font=("Helvetica", 14)
        )
        self.pdf_page_label.config(text="Ielādē...")
        self.prev_page_button.config(state=tk.DISABLED)
        self.next_page_button.config(state=tk.DISABLED)
        self.update_idletasks() # Atjaunina UI

        try:
            self.current_pdf_document = fitz.open(filepath)
            self.current_pdf_page_count = self.current_pdf_document.page_count
            self.current_pdf_page_index = 0 # Sākam ar pirmo lapu
            self.pdf_preview_zoom_factor = 1.0 # Atiestatām tālummaiņu
            self.pdf_preview_pan_x = 0 # Atiestatām pārvietošanu
            self.pdf_preview_pan_y = 0

            self._display_pdf_page_on_canvas()

        except Exception as e:
            messagebox.showerror("PDF ielādes kļūda", f"Nevarēja ielādēt PDF priekšskatījumam:\n{e}")
            self.current_pdf_document = None
            self.current_pdf_page_count = 0
            self.current_pdf_page_index = 0
            self.pdf_preview_canvas.delete("all")
            self.pdf_preview_canvas.create_text(
                self.pdf_preview_canvas.winfo_width() / 2, self.pdf_preview_canvas.winfo_height() / 2,
                text=f"Nevarēja ielādēt PDF:\n{e}", fill="red", font=("Helvetica", 12),
                justify="center"
            )
            self.pdf_page_label.config(text="Lapa: Kļūda")
            self.prev_page_button.config(state=tk.DISABLED)
            self.next_page_button.config(state=tk.DISABLED)


    def _show_prev_pdf_page(self):
        """Parāda iepriekšējo PDF lapu priekšskatījumā."""
        if self.current_pdf_document and self.current_pdf_page_index > 0:
            self.current_pdf_page_index -= 1
            self._display_pdf_page_on_canvas()

    def _load_pdf_for_additional_tools_preview(self, filepath):
        """Ielādē PDF dokumentu priekšskatījumam "Papildu rīki" cilnē."""
        if self.additional_tools_current_pdf_document:
            self.additional_tools_current_pdf_document.close()
            self.additional_tools_current_pdf_document = None

        self.additional_tools_pdf_preview_canvas.delete("all")
        self.additional_tools_pdf_preview_canvas.create_text(
            self.additional_tools_pdf_preview_canvas.winfo_width() / 2, self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
            text="Ielādē...", fill="white", font=("Helvetica", 14)
        )
        if self.additional_tools_pdf_page_label:
            self.additional_tools_pdf_page_label.config(text="Ielādē...")
        if self.additional_tools_prev_page_button:
            self.additional_tools_prev_page_button.config(state=tk.DISABLED)
        if self.additional_tools_next_page_button:
            self.additional_tools_next_page_button.config(state=tk.DISABLED)
        self.update_idletasks()

        try:
            self.additional_tools_current_pdf_document = fitz.open(filepath)
            self.additional_tools_current_pdf_page_count = self.additional_tools_current_pdf_document.page_count
            self.additional_tools_current_pdf_page_index = 0
            self.additional_tools_pdf_preview_zoom_factor = 1.0
            self.additional_tools_pdf_preview_pan_x = 0
            self.additional_tools_pdf_preview_pan_y = 0

            self._display_pdf_page_on_additional_tools_canvas()

        except Exception as e:
            messagebox.showerror("PDF ielādes kļūda (Papildu rīki)", f"Nevarēja ielādēt PDF priekšskatījumam:\n{e}")
            self.additional_tools_current_pdf_document = None
            self.additional_tools_current_pdf_page_count = 0
            self.additional_tools_current_pdf_page_index = 0
            self.additional_tools_pdf_preview_canvas.delete("all")
            self.additional_tools_pdf_preview_canvas.create_text(
                self.additional_tools_pdf_preview_canvas.winfo_width() / 2, self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                text=f"Nevarēja ielādēt PDF:\n{e}", fill="red", font=("Helvetica", 12),
                justify="center"
            )
            if self.additional_tools_pdf_page_label:
                self.additional_tools_pdf_page_label.config(text="Lapa: Kļūda")
            if self.additional_tools_prev_page_button:
                self.additional_tools_prev_page_button.config(state=tk.DISABLED)
            if self.additional_tools_next_page_button:
                self.additional_tools_next_page_button.config(state=tk.DISABLED)

    def _display_pdf_page_on_additional_tools_canvas(self):
        """Attēlo pašreizējo PDF lapu uz priekšskatījuma kanvasa "Papildu rīki" cilnē."""
        if not self.additional_tools_current_pdf_document or not self.additional_tools_pdf_preview_canvas:
            return

        try:
            canvas_width = self.additional_tools_pdf_preview_canvas.winfo_width()
            canvas_height = self.additional_tools_pdf_preview_canvas.winfo_height()

            if canvas_width <= 1 or canvas_height <= 1:
                self.after(50, self._display_pdf_page_on_additional_tools_canvas)
                return

            page = self.additional_tools_current_pdf_document.load_page(self.additional_tools_current_pdf_page_index)

            fit_width_zoom = canvas_width / page.rect.width
            fit_height_zoom = canvas_height / page.rect.height
            initial_fit_zoom = min(fit_width_zoom, fit_height_zoom)
            zoom_factor_for_render = self.additional_tools_pdf_preview_zoom_factor * initial_fit_zoom
            zoom_factor_for_render = max(0.1, zoom_factor_for_render)

            pix = page.get_pixmap(matrix=fitz.Matrix(zoom_factor_for_render, zoom_factor_for_render))
            img_pil = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            self.additional_tools_pdf_preview_photo = ImageTk.PhotoImage(img_pil)

            self.additional_tools_pdf_preview_canvas.delete("all")

            img_width, img_height = img_pil.size
            x = (canvas_width - img_width) / 2 + self.additional_tools_pdf_preview_pan_x
            y = (canvas_height - img_height) / 2 + self.additional_tools_pdf_preview_pan_y

            self.additional_tools_pdf_preview_canvas.create_image(x, y, anchor="nw", image=self.additional_tools_pdf_preview_photo)
            self.additional_tools_pdf_preview_canvas.image = self.additional_tools_pdf_preview_photo

            if self.additional_tools_pdf_page_label:
                self.additional_tools_pdf_page_label.config(text=f"Lapa: {self.additional_tools_current_pdf_page_index + 1}/{self.additional_tools_current_pdf_page_count}")

            if self.additional_tools_prev_page_button:
                self.additional_tools_prev_page_button.config(state=tk.NORMAL if self.additional_tools_current_pdf_page_index > 0 else tk.DISABLED)
            if self.additional_tools_next_page_button:
                self.additional_tools_next_page_button.config(state=tk.NORMAL if self.additional_tools_current_pdf_page_index < self.additional_tools_current_pdf_page_count - 1 else tk.DISABLED)

        except Exception as e:
            print(f"Kļūda attēlojot PDF lapu (Papildu rīki): {e}")
            self.additional_tools_pdf_preview_canvas.delete("all")
            self.additional_tools_pdf_preview_canvas.create_text(
                self.additional_tools_pdf_preview_canvas.winfo_width() / 2, self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                text=f"Nevarēja ielādēt lapu:\n{e}", fill="red", font=("Helvetica", 12),
                justify="center"
            )
            if self.additional_tools_pdf_page_label:
                self.additional_tools_pdf_page_label.config(text="Lapa: Kļūda")
            if self.additional_tools_prev_page_button:
                self.additional_tools_prev_page_button.config(state=tk.DISABLED)
            if self.additional_tools_next_page_button:
                self.additional_tools_next_page_button.config(state=tk.DISABLED)

    def _show_prev_additional_tools_pdf_page(self):
        """
        Parāda iepriekšējo lapu/attēlu priekšskatījumā "Papildu rīki" cilnē.
        Navigē pa PDF lapām, ja atlasīts PDF, vai pa self.images sarakstu, ja atlasīts attēls.
        """
        if self.additional_tools_current_pdf_document:  # Ja pašlaik tiek rādīts PDF
            if self.additional_tools_current_pdf_page_index > 0:
                self.additional_tools_current_pdf_page_index -= 1
                self._display_pdf_page_on_additional_tools_canvas()
        elif self.current_image_index > 0:  # Ja pašlaik tiek rādīts attēls no self.images
            self.current_image_index -= 1
            self.file_listbox.selection_clear(0, tk.END)
            self.file_listbox.selection_set(self.current_image_index)
            self.file_listbox.activate(self.current_image_index)
            self.file_listbox.see(self.current_image_index)
            self._update_additional_tools_pdf_preview()  # Atjaunina priekšskatījumu
            self.show_image_preview(
                self.images[self.current_image_index]["processed_img"])  # Atjaunina arī attēlu apstrādes cilni

    def _show_next_additional_tools_pdf_page(self):
        """
        Parāda nākamo lapu/attēlu priekšskatījumā "Papildu rīki" cilnē.
        Navigē pa PDF lapām, ja atlasīts PDF, vai pa self.images sarakstu, ja atlasīts attēls.
        """
        if self.additional_tools_current_pdf_document:  # Ja pašlaik tiek rādīts PDF
            if self.additional_tools_current_pdf_page_index < self.additional_tools_current_pdf_page_count - 1:
                self.additional_tools_current_pdf_page_index += 1
                self._display_pdf_page_on_additional_tools_canvas()
        elif self.current_image_index < len(self.images) - 1:  # Ja pašlaik tiek rādīts attēls no self.images
            self.current_image_index += 1
            self.file_listbox.selection_clear(0, tk.END)
            self.file_listbox.selection_set(self.current_image_index)
            self.file_listbox.activate(self.current_image_index)
            self.file_listbox.see(self.current_image_index)
            self._update_additional_tools_pdf_preview()  # Atjaunina priekšskatījumu
            self.show_image_preview(
                self.images[self.current_image_index]["processed_img"])  # Atjaunina arī attēlu apstrādes cilni

    def _on_additional_tools_pdf_preview_canvas_resize(self, event):
        """Apstrādā PDF priekšskatījuma kanvasa izmēru maiņu "Papildu rīki" cilnē."""
        self._display_pdf_page_on_additional_tools_canvas()

    def _on_additional_tools_pdf_preview_mouse_wheel(self, event):
        """Apstrādā peles rullīša notikumus PDF priekšskatījuma tālummaiņai "Papildu rīki" cilnē."""
        if event.num == 5 or event.delta == -120:
            self.additional_tools_pdf_preview_zoom_factor = max(0.1, self.additional_tools_pdf_preview_zoom_factor - 0.1)
        if event.num == 4 or event.delta == 120:
            self.additional_tools_pdf_preview_zoom_factor = min(5.0, self.additional_tools_pdf_preview_zoom_factor + 0.1)
        self._display_pdf_page_on_additional_tools_canvas()

    def _on_additional_tools_pdf_preview_pan_start(self, event):
        """Sāk PDF priekšskatījuma pārvietošanu (pan) "Papildu rīki" cilnē."""
        self.additional_tools_pdf_preview_start_pan_x = event.x - self.additional_tools_pdf_preview_pan_x
        self.additional_tools_pdf_preview_start_pan_y = event.y - self.additional_tools_pdf_preview_pan_y
        self.additional_tools_pdf_preview_canvas.config(cursor="fleur")

    def _on_additional_tools_pdf_preview_pan_drag(self, event):
        """Pārvieto PDF priekšskatījumu, velkot peli "Papildu rīki" cilnē."""
        self.additional_tools_pdf_preview_pan_x = event.x - self.additional_tools_pdf_preview_start_pan_x
        self.additional_tools_pdf_preview_pan_y = event.y - self.additional_tools_pdf_preview_start_pan_y
        self._display_pdf_page_on_additional_tools_canvas()

    def _on_additional_tools_pdf_preview_pan_end(self, event):
        """Beidz PDF priekšskatījuma pārvietošanu "Papildu rīki" cilnē."""
        self.additional_tools_pdf_preview_canvas.config(cursor="arrow")

    def _clear_additional_tools_pdf_preview(self):
        """Notīra PDF priekšskatījumu un atbrīvo resursus "Papildu rīki" cilnē."""
        if self.additional_tools_current_pdf_document:
            self.additional_tools_current_pdf_document.close()
            self.additional_tools_current_pdf_document = None
        if self.additional_tools_pdf_preview_canvas:
            self.additional_tools_pdf_preview_canvas.delete("all")
        self.additional_tools_pdf_preview_photo = None
        self.additional_tools_current_pdf_page_count = 0
        self.additional_tools_current_pdf_page_index = 0
        if self.additional_tools_pdf_page_label:
            self.additional_tools_pdf_page_label.config(text="Lapa: 0/0")
        if self.additional_tools_prev_page_button:
            self.additional_tools_prev_page_button.config(state=tk.DISABLED)
        if self.additional_tools_next_page_button:
            self.additional_tools_next_page_button.config(state=tk.DISABLED)

    #def _load_pdf_for_additional_tools_preview_from_dialog(self):
        #"""Atver failu dialogu un ielādē PDF priekšskatījumam "Papildu rīki" cilnē."""
        #filepath = filedialog.askopenfilename(
            #title="Izvēlēties PDF failu priekšskatījumam",
            #filetypes=[("PDF faili", "*.pdf"), ("Visi faili", "*.*")]
        #)
        #if filepath:
            #self._load_pdf_for_additional_tools_pdf_preview(filepath)


    def _show_next_pdf_page(self):
        """Parāda nākamo PDF lapu priekšskatījumā."""
        if self.current_pdf_document and self.current_pdf_page_index < self.current_pdf_page_count - 1:
            self.current_pdf_page_index += 1
            self._display_pdf_page_on_canvas()

    def _on_pdf_preview_canvas_resize(self, event):
        """Apstrādā PDF priekšskatījuma kanvasa izmēru maiņu."""
        self._display_pdf_page_on_canvas()

    def _on_pdf_preview_mouse_wheel(self, event):
        """Apstrādā peles rullīša notikumus PDF priekšskatījuma tālummaiņai."""
        if event.num == 5 or event.delta == -120:  # Tuvināt
            self.pdf_preview_zoom_factor = max(0.1, self.pdf_preview_zoom_factor - 0.1)
        if event.num == 4 or event.delta == 120:  # Attālināt
            self.pdf_preview_zoom_factor = min(5.0, self.pdf_preview_zoom_factor + 0.1)
        self._display_pdf_page_on_canvas()

    def _on_pdf_preview_pan_start(self, event):
        """Sāk PDF priekšskatījuma pārvietošanu (pan)."""
        self.pdf_preview_start_pan_x = event.x - self.pdf_preview_pan_x
        self.pdf_preview_start_pan_y = event.y - self.pdf_preview_pan_y
        self.pdf_preview_canvas.config(cursor="fleur")

    def _on_pdf_preview_pan_drag(self, event):
        """Pārvieto PDF priekšskatījumu, velkot peli."""
        self.pdf_preview_pan_x = event.x - self.pdf_preview_start_pan_x
        self.pdf_preview_pan_y = event.y - self.pdf_preview_start_pan_y
        self._display_pdf_page_on_canvas()

    def _on_pdf_preview_pan_end(self, event):
        """Beidz PDF priekšskatījuma pārvietošanu."""
        self.pdf_preview_canvas.config(cursor="arrow")

    def _clear_pdf_preview(self):
        """Notīra PDF priekšskatījumu un atbrīvo resursus."""
        if self.current_pdf_document:
            self.current_pdf_document.close()
            self.current_pdf_document = None
        self.pdf_preview_canvas.delete("all")
        self.pdf_preview_photo = None
        self.current_pdf_page_count = 0
        self.current_pdf_page_index = 0
        self.pdf_page_label.config(text="Lapa: 0/0")
        self.prev_page_button.config(state=tk.DISABLED)
        self.next_page_button.config(state=tk.DISABLED)

    def load_app_settings(self):

        """Ielādē lietotnes iestatījumus no JSON faila"""
        # Sākumā iestata noklusējuma vērtības visiem iestatījumiem
        self.settings.setdefault("output_format", "pdf")
        self.settings.setdefault("psm", 3)
        self.settings.setdefault("oem", 3)
        self.settings.setdefault("default_pdf_page_size", "Auto")
        self.settings.setdefault("default_pdf_font_size", 7)
        self.settings.setdefault("pdf_quality", "Vidēja")
        self.settings.setdefault("tesseract_path", DEFAULT_TESSERACT_CMD)
        self.settings.setdefault("default_save_path", os.path.expanduser("~"))
        self.settings.setdefault("selected_ocr_languages", ["Angļu (eng)"])  # Noklusējums

        # E-pasta iestatījumu noklusējuma vērtības
        self.settings.setdefault("smtp_server", "")
        self.settings.setdefault("smtp_port", 465)
        self.settings.setdefault("email_user", "")
        self.settings.setdefault("email_pass", "")
        self.settings.setdefault("from_email", "")
        self.settings.setdefault("to_email", "")
        self.settings.setdefault("use_ssl", True)
        self.settings.setdefault("email_subject", "OCR PDF dokumenti")
        self.settings.setdefault("email_body_plain",
                                 "Sveiki,\n\nPielikumā atradīsiet OCR apstrādātos PDF dokumentus.\n\nAr cieņu,\nJūsu OCR PDF App")
        self.settings.setdefault("email_body_html",
                                 "<html><body><p>Sveiki,</p><p>Pielikumā atradīsiet OCR apstrādātos PDF dokumentus.</p><p>Ar cieņu,<br/>Jūsu OCR PDF App</p></body></html>")

        # Pārējie iestatījumi (logu izmēri utt.)
        self.settings.setdefault("window_width", 1024)  # Jaunais noklusējuma platums
        self.settings.setdefault("window_height", 768)  # Jaunais noklusējuma augstums
        self.settings.setdefault("window_x", 0)
        self.settings.setdefault("window_y", 0)
        self.settings.setdefault("scan_folder_path", os.path.join(os.path.expanduser("~"), "ScannedDocuments"))  # JAUNS
        self.settings.setdefault("auto_scan_enabled", False)  # JAUNS
        # JAUNS: Attālinātās glabāšanas noklusējuma vērtības
        self.settings.setdefault("remote_storage_type", "Local")
        self.settings.setdefault("ftp_host", "")
        self.settings.setdefault("ftp_port", 21)
        self.settings.setdefault("ftp_user", "")
        self.settings.setdefault("ftp_pass", "")
        self.settings.setdefault("ftp_remote_path", "/")
        self.settings.setdefault("ftp_use_sftp", False)
        self.settings.setdefault("google_drive_folder_id", "")
        self.settings.setdefault("google_drive_credentials_path", "credentials.json")
        self.settings.setdefault("google_drive_token_path", "token.json")
        self.settings.setdefault("auto_upload_enabled", False)
        self.settings.setdefault("auto_upload_target", "Local")
        # JAUNS: ID koda iestatījumu noklusējuma vērtības
        self.settings.setdefault("add_id_code_to_pdf", False)
        self.settings.setdefault("id_code_type", "QR")
        self.settings.setdefault("id_code_position", "bottom_right")  # MAINĪTS: no "bottom-right" uz "bottom_right"

        # JAUNS: ID koda iestatījumi
        self.settings["add_id_code_to_pdf"] = self.settings.get("add_id_code_to_pdf",
                                                                False)  # Jāpārliecinās, ka vērtība ir iestatīta
        self.settings["id_code_type"] = self.settings.get("id_code_type", "QR")
        self.settings["id_code_position"] = self.settings.get("id_code_position",
                                                              "bottom_right")  # MAINĪTS: no "bottom-right" uz "bottom_right"

        # JAUNS: Google Sheets iestatījumu noklusējuma vērtības
        self.settings.setdefault("google_sheet_id", "")
        self.settings.setdefault("google_sheet_name", "OCR_Failu_Saraksts")
        self.settings.setdefault("google_sheet_credentials_path", "google_sheet_credentials.json")


        # JAUNS: Attēla uzlabojumu iestatījumu noklusējuma vērtības
        self.settings.setdefault("scan_brightness", 0)
        self.settings.setdefault("scan_contrast", 0)
        self.settings.setdefault("scan_saturation", 0)
        self.settings.setdefault("scan_gamma", 1.0)
        self.settings.setdefault("scan_use_color_detection", False)
        self.settings.setdefault("scan_target_color", "#FFFFFF")
        self.settings.setdefault("scan_color_tolerance", 30)
        self.settings.setdefault("scan_morphology_enabled", False)
        self.settings.setdefault("scan_morphology_kernel_size", 3)
        self.settings.setdefault("scan_edge_dilation", 2)
        self.settings.setdefault("scan_camera_index", 0)
        self.settings.setdefault("scan_camera_width", 1280)
        self.settings.setdefault("scan_camera_height", 120)
        self.settings.setdefault("scan_min_contour_area", 2500)
        self.settings.setdefault("scan_stable_threshold", 0.8)
        self.settings.setdefault("scan_stability_tolerance", 0.01)
        self.settings.setdefault("scan_aspect_ratio_min", 0.4)
        self.settings.setdefault("scan_aspect_ratio_max", 2.3)
        self.settings.setdefault("scan_gaussian_blur_kernel", 9)
        self.settings.setdefault("scan_adaptive_thresh_block_size", 11)
        self.settings.setdefault("scan_adaptive_thresh_c", 3)
        self.settings.setdefault("scan_canny_thresh1", 610)
        self.settings.setdefault("scan_canny_thresh2", 190)

        if os.path.exists(self.settings_file):
            try:
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    loaded_settings = json.load(f)
                    self.settings.update(loaded_settings)
                return True
            except Exception as e:
                print(f"Nevarēja ielādēt iestatījumus: {e}")
                return False
        return False

    def load_scan_settings(self):
        """Ielādē skenēšanas iestatījumus no JSON faila."""
        try:
            if os.path.exists(self.scan_settings_file):
                with open(self.scan_settings_file, 'r', encoding='utf-8') as f:
                    self.scan_settings = json.load(f)
            else:
                # JAUNS: Ja fails neeksistē, iestatām noklusējuma vērtības
                self.scan_settings = {
                    "scan_camera_index": 0,
                    "scan_camera_width": 1280,
                    "scan_camera_height": 120,
                    "scan_min_contour_area": 2500,
                    "scan_stable_threshold": 0.8,
                    "scan_stability_tolerance": 0.01,
                    "scan_aspect_ratio_min": 0.4,
                    "scan_aspect_ratio_max": 2.3,
                    "scan_gaussian_blur_kernel": 9,
                    "scan_adaptive_thresh_block_size": 11,
                    "scan_adaptive_thresh_c": 3,
                    "scan_canny_thresh1": 610,
                    "scan_canny_thresh2": 190,
                    "scan_brightness": 0,
                    "scan_contrast": 0,
                    "scan_saturation": 0,
                    "scan_gamma": 1.0,
                    "scan_use_color_detection": False,
                    "scan_target_color": "#FFFFFF",
                    "scan_color_tolerance": 30,
                    "scan_morphology_enabled": False,
                    "scan_morphology_kernel_size": 3,
                    "scan_edge_dilation": 2
                }
        except Exception as e:
            print(f"Kļūda ielādējot skenēšanas iestatījumus: {e}")
            self.scan_settings = {}

    def save_app_settings(self):
        """Saglabā lietotnes iestatījumus JSON failā"""
        # Pārējie iestatījumi
        self.settings["output_format"] = self.output_format_var.get()
        self.settings["psm"] = self.psm_var.get()
        self.settings["oem"] = self.oem_var.get()
        self.settings["language"] = self.language_var.get()
        self.settings["output_dir"] = self.output_dir_var.get()
        self.settings["window_width"] = self.winfo_width()
        self.settings["window_height"] = self.winfo_height()
        self.settings["window_x"] = self.winfo_x()
        self.settings["window_y"] = self.winfo_y()
        self.settings["scan_folder_path"] = self.scan_folder_path.get()  # JAUNS
        self.settings["auto_scan_enabled"] = self.auto_scan_enabled.get()  # JAUNS
        # JAUNS: Attālinātās glabāšanas iestatījumi
        self.settings["remote_storage_type"] = self.remote_storage_type.get()
        self.settings["ftp_host"] = self.ftp_host.get()
        self.settings["ftp_port"] = self.ftp_port.get()
        self.settings["ftp_user"] = self.ftp_user.get()
        self.settings["ftp_pass"] = self.ftp_pass.get()
        self.settings["ftp_remote_path"] = self.ftp_remote_path.get()
        self.settings["ftp_use_sftp"] = self.ftp_use_sftp.get()
        self.settings["google_drive_folder_id"] = self.google_drive_folder_id.get()
        self.settings["google_drive_credentials_path"] = self.google_drive_credentials_path.get()
        self.settings["google_drive_token_path"] = self.google_drive_token_path.get()
        self.settings["auto_upload_enabled"] = self.auto_upload_enabled.get()
        self.settings["auto_upload_target"] = self.auto_upload_target.get()
        # JAUNS: ID koda iestatījumi
        self.settings["add_id_code_to_pdf"] = self.settings.get("add_id_code_to_pdf",
                                                                False)  # Jāpārliecinās, ka vērtība ir iestatīta
        self.settings["id_code_type"] = self.settings.get("id_code_type", "QR")
        self.settings["id_code_position"] = self.settings.get("id_code_position",
                                                              "bottom_right")  # MAINĪTS: no "bottom-right" uz "bottom_right"

        # JAUNS: Attēla uzlabojumu iestatījumu saglabāšana
        self.settings["scan_brightness"] = self.scan_brightness.get()
        self.settings["scan_contrast"] = self.scan_contrast.get()
        self.settings["scan_saturation"] = self.scan_saturation.get()
        self.settings["scan_gamma"] = self.scan_gamma.get()
        self.settings["scan_use_color_detection"] = self.scan_use_color_detection.get()
        self.settings["scan_target_color"] = self.scan_target_color.get()
        self.settings["scan_color_tolerance"] = self.scan_color_tolerance.get()
        self.settings["scan_morphology_enabled"] = self.scan_morphology_enabled.get()
        self.settings["scan_morphology_kernel_size"] = self.scan_morphology_kernel_size.get()
        self.settings["scan_edge_dilation"] = self.scan_edge_dilation.get()
        self.settings["scan_camera_index"] = self.scan_camera_index.get()
        self.settings["scan_camera_width"] = self.scan_camera_width.get()
        self.settings["scan_camera_height"] = self.scan_camera_height.get()
        self.settings["scan_min_contour_area"] = self.scan_min_contour_area.get()
        self.settings["scan_stable_threshold"] = self.scan_stable_threshold.get()
        self.settings["scan_stability_tolerance"] = self.scan_stability_tolerance.get()
        self.settings["scan_aspect_ratio_min"] = self.scan_aspect_ratio_min.get()
        self.settings["scan_aspect_ratio_max"] = self.scan_aspect_ratio_max.get()
        self.settings["scan_gaussian_blur_kernel"] = self.scan_gaussian_blur_kernel.get()
        self.settings["scan_adaptive_thresh_block_size"] = self.scan_adaptive_thresh_block_size.get()
        self.settings["scan_adaptive_thresh_c"] = self.scan_adaptive_thresh_c.get()
        self.settings["scan_canny_thresh1"] = self.scan_canny_thresh1.get()
        self.settings["scan_canny_thresh2"] = self.scan_canny_thresh2.get()

        # JAUNS: Google Sheets iestatījumi
        # Definējam StringVar mainīgos
        self.google_sheet_id = tk.StringVar(value=self.settings.get("google_sheet_id", ""))
        self.google_sheet_name = tk.StringVar(value=self.settings.get("google_sheet_name", "OCR_Failu_Saraksts"))
        self.google_sheet_credentials_path = tk.StringVar(
            value=self.settings.get("google_sheet_credentials_path", "google_sheet_credentials.json"))
        self.auto_upload_enabled = tk.BooleanVar(value=self.settings.get("auto_upload_enabled", False))
        self.remote_storage_type = tk.StringVar(value=self.settings.get("remote_storage_type", ""))
        self.google_drive_folder_id = tk.StringVar(value=self.settings.get("google_drive_folder_id", ""))
        # Tagad varam pievienot trace_add
        self.google_sheet_id.trace_add("write", lambda *args: self.save_app_settings())
        self.google_sheet_name.trace_add("write", lambda *args: self.save_app_settings())
        self.google_sheet_credentials_path.trace_add("write", lambda *args: self.save_app_settings())
        self.auto_upload_enabled.trace_add("write", lambda *args: self.save_app_settings())
        self.remote_storage_type.trace_add("write", lambda *args: self.save_app_settings())
        self.google_drive_folder_id.trace_add("write", lambda *args: self.save_app_settings())

        # Saglabā arī citus iestatījumus, kas tiek mainīti SettingsWindow
        self.settings["default_save_path"] = self.default_save_path
        self.settings["pdf_quality"] = self.pdf_quality
        self.settings["default_pdf_page_size"] = self.orientation_var.get()
        self.settings["default_pdf_font_size"] = self.fontsize_var.get()
        self.settings["tesseract_path"] = pytesseract.pytesseract.tesseract_cmd

        # Saglabā atlasītās OCR valodas
        selected_langs_codes = []
        for lang_name, var in self.lang_vars.items():
            if var.get():
                selected_langs_codes.append(lang_name)  # Saglabājam nosaukumu, nevis kodu
        self.settings["selected_ocr_languages"] = selected_langs_codes

        try:
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.settings, f, indent=4)
            print("Iestatījumi veiksmīgi saglabāti")  # Konsolei
            return True
        except Exception as e:
            print(f"Nevarēja saglabāt iestatījumus: {e}")
            return False

    def save_scan_settings(self):
        """JAUNS: Saglabā skenēšanas iestatījumus JSON failā."""
        self.scan_settings["scan_camera_index"] = self.scan_camera_index.get()
        self.scan_settings["scan_camera_width"] = self.scan_camera_width.get()
        self.scan_settings["scan_camera_height"] = self.scan_camera_height.get()
        self.scan_settings["scan_min_contour_area"] = self.scan_min_contour_area.get()
        self.scan_settings["scan_stable_threshold"] = self.scan_stable_threshold.get()
        self.scan_settings["scan_stability_tolerance"] = self.scan_stability_tolerance.get()
        self.scan_settings["scan_aspect_ratio_min"] = self.scan_aspect_ratio_min.get()
        self.scan_settings["scan_aspect_ratio_max"] = self.scan_aspect_ratio_max.get()
        self.scan_settings["scan_gaussian_blur_kernel"] = self.scan_gaussian_blur_kernel.get()
        self.scan_settings["scan_adaptive_thresh_block_size"] = self.scan_adaptive_thresh_block_size.get()
        self.scan_settings["scan_adaptive_thresh_c"] = self.scan_adaptive_thresh_c.get()
        self.scan_settings["scan_canny_thresh1"] = self.scan_canny_thresh1.get()
        self.scan_settings["scan_canny_thresh2"] = self.scan_canny_thresh2.get()

        try:
            with open(self.scan_settings_file, 'w', encoding='utf-8') as f:
                json.dump(self.scan_settings, f, indent=4)
            print("Skenēšanas iestatījumi veiksmīgi saglabāti")
            return True
        except Exception as e:
            print(f"Nevarēja saglabāt skenēšanas iestatījumus: {e}")
            return False

    def _flatten_file_system(self, node):
        """Rekursīvi pārveido koka struktūru par serializējamu dict, noņemot ciklisko 'parent' atsauci."""
        serializable_node = node.copy()
        serializable_node.pop("parent", None)  # Noņem 'parent' atsauci

        if serializable_node["type"] == "folder":
            # Rekursīvi apstrādā saturu
            serializable_node["contents"] = [self._flatten_file_system(item) for item in node["contents"]]

        return serializable_node

    def _unflatten_file_system(self, serializable_node, parent=None):
        """Rekursīvi pārveido serializējamu dict atpakaļ par koka struktūru, atjaunojot 'parent' atsauces."""
        node = serializable_node.copy()
        node["parent"] = parent  # Atjauno 'parent' atsauci

        if node["type"] == "folder":
            # Rekursīvi apstrādā saturu un nodod pašreizējo mezglu kā vecāku
            node["contents"] = [self._unflatten_file_system(item, node) for item in serializable_node["contents"]]

        return node

    def detect_and_decode_barcodes(self, img):
        """Atpazīst un atšifrē QR kodus un svītrkodus attēlā."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas",
                                   "QR kodu un svītrkodu atpazīšanai nepieciešams 'opencv-python' un 'pyzbar'.")
            return []

        from pyzbar.pyzbar import decode  # Pārliecinās, ka pyzbar ir pieejams

        img_cv = np.array(img)
        decoded_objects = decode(img_cv)

        decoded_texts = []
        for obj in decoded_objects:
            decoded_texts.append(obj.data.decode('utf-8'))  # Atšifrē QR kodu vai svītrkodu

        return decoded_texts

    def load_internal_file_system(self):
        """Ielādē iekšējo failu sistēmu no arhīva JSON faila."""
        # Sākumā inicializējam tukšu saknes mapi
        self.internal_file_system = {"type": "folder", "name": "Sakne", "contents": [], "parent": None}
        self.current_folder = self.internal_file_system

        if os.path.exists(self.pdf_archive_file):
            try:
                with open(self.pdf_archive_file, 'r', encoding='utf-8') as f:  # Pievienots encoding
                    loaded_data = json.load(f)
                if loaded_data:
                    # Ielādējam visu koka struktūru, sākot no saknes
                    self.internal_file_system = self._unflatten_file_system(loaded_data)
                    self.current_folder = self.internal_file_system  # Pēc ielādes vienmēr sākam no saknes
            except json.JSONDecodeError:
                messagebox.showwarning("Arhīva kļūda", "Neizdevās ielādēt PDF arhīvu. Fails ir bojāts vai tukšs.")
                # Ja fails ir bojāts, atiestatām uz tukšu sistēmu
                self.internal_file_system = {"type": "folder", "name": "Sakne", "contents": [], "parent": None}
                self.current_folder = self.internal_file_system
            except Exception as e:
                messagebox.showerror("Arhīva ielādes kļūda", f"Neizdevās ielādēt PDF arhīvu: {e}")
                # Ja rodas cita kļūda, atiestatām uz tukšu sistēmu
                self.internal_file_system = {"type": "folder", "name": "Sakne", "contents": [], "parent": None}
                self.current_folder = self.internal_file_system

        # Pēc ielādes sinhronizējam ar fizisko failu sistēmu
        #self.sync_with_physical_folders()

    def save_pdf_and_update_archive(self, pdf_filepath, file_node):
        """
        Saglabā PDF failu un atjaunina iekšējo arhīvu.
        Pēc tam augšupielādē Google Drive un atjaunina Google Sheet, ja iespējots.
        """
        import os
        from datetime import datetime

        # Pārliecināmies, ka fails eksistē
        if not os.path.exists(pdf_filepath):
            messagebox.showerror("Kļūda", f"Fails nav atrasts: {pdf_filepath}")
            return

        # Pievienojam vai atjaunojam faila mezglu iekšējā arhīvā
        if file_node not in self.current_folder["contents"]:
            self.current_folder["contents"].append(file_node)

        # Atjaunojam faila mezglu ar pamata informāciju
        file_node["filepath"] = pdf_filepath
        if "date" not in file_node or not file_node["date"]:
            file_node["date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        if "doc_id" not in file_node or not file_node["doc_id"]:
            import uuid
            file_node["doc_id"] = str(uuid.uuid4())[:8]  # Īss unikāls ID

        # Saglabājam arhīvu uz diska
        self.save_pdf_archive()

        # Ja ir ieslēgta automātiskā augšupielāde un mērķis ir Google Drive
        if self.auto_upload_enabled.get() and self.remote_storage_type.get() == "Google Drive":
            google_drive_folder_id = self.google_drive_folder_id.get()  # Iegūstam mērķa mapes ID

            # Augšupielādējam failu Google Drive
            file_id, web_view_link = self.upload_file_to_google_drive(pdf_filepath, google_drive_folder_id)

            if file_id and web_view_link:
                # Atjaunojam faila mezglu ar Google Drive informāciju
                file_node["google_drive_id"] = file_id
                file_node["google_drive_link"] = web_view_link

                # Saglabājam arhīvu ar jaunajiem datiem
                self.save_pdf_archive()

                # Aprēķinām iekšējā faila ceļu (mapju ceļu)
                internal_folder_path = self._get_internal_folder_path_for_node(file_node)

                # Sagatavojam datus Google Sheet atjaunināšanai
                sheet_file_info = {
                    "name": file_node.get("name", os.path.basename(pdf_filepath)),
                    "doc_id": file_node.get("doc_id", ""),
                    "filepath": pdf_filepath,
                    "assigned_id": file_node.get("assigned_id", ""),  # Ja jums ir šāds lauks
                    "date": file_node.get("date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                    "internal_folder_path": internal_folder_path,
                    "google_drive_id": file_id,
                    "google_drive_link": web_view_link
                }

                # Atjaunojam Google Sheet ierakstu
                self.update_google_sheet_entry(sheet_file_info)

                print(f"Fails '{file_node.get('name')}' augšupielādēts un Google Sheet atjaunināts.")
            else:
                print("Google Drive augšupielāde neizdevās vai saite nav pieejama.")
        else:
            print("Automātiskā augšupielāde nav ieslēgta vai mērķis nav Google Drive.")

        # Paziņojums lietotājam
        messagebox.showinfo("Saglabāts", f"Fails '{file_node.get('name')}' veiksmīgi saglabāts un arhivēts.")

    def _get_internal_folder_path_for_node(self, node):
        """Palīgmetode, lai iegūtu iekšējās mapes ceļu dotajam mezglam."""
        path_parts = []
        current = node.get("parent", None)
        while current and current != self.internal_file_system:
            if current["type"] == "folder":
                path_parts.insert(0, current["name"])
            current = current.get("parent", None)
        return "/".join(path_parts) if path_parts else "Sakne"



    def save_pdf_archive(self):
        """Saglabā PDF arhīva datus JSON failā."""
        try:
            # Pārveido koka struktūru par serializējamu dict
            serializable_data = self._flatten_file_system(self.internal_file_system)
            with open(self.pdf_archive_file, 'w', encoding='utf-8') as f:  # Pievienots encoding
                json.dump(serializable_data, f, indent=4,
                          ensure_ascii=False)  # ensure_ascii=False, lai atbalstītu latviešu burtus
        except Exception as e:
            messagebox.showerror("Arhīva saglabāšanas kļūda", f"Neizdevās saglabāt PDF arhīvu: {e}")

    def add_password_to_pdf(self, pdf_path):
        """Pievieno paroli PDF dokumentam"""
        try:
            # Iegūst paroli no lietotāja
            password = simpledialog.askstring("Parole", "Ievadiet jauno paroli:", show='*')
            if not password:
                return  # Lietotājs atcēla

            reader = pypdf.PdfReader(pdf_path)

            # Ja fails jau ir šifrēts, mēģinam to atšifrēt
            if reader.is_encrypted:
                try:
                    if not reader.decrypt(""):  # Mēģinam atšifrēt bez paroles
                        # Ja neizdodas, prasam esošo paroli
                        old_pass = simpledialog.askstring("Esošā parole",
                                                          "Dokuments jau ir aizsargāts. Ievadiet esošo paroli:",
                                                          show='*')
                        if not old_pass or not reader.decrypt(old_pass):
                            messagebox.showerror("Kļūda", "Nepareiza parole vai neizdevās atšifrēt!")
                            return
                except Exception as e:
                    messagebox.showerror("Kļūda", f"Atšifrēšanas kļūda: {e}")
                    return

            writer = pypdf.PdfWriter()

            # Pārkopējam visas lapas
            for page in reader.pages:
                writer.add_page(page)

            # Šifrējam ar jauno paroli
            writer.encrypt(password)

            # Saglabājam izmaiņas
            with open(pdf_path, "wb") as f:
                writer.write(f)

            messagebox.showinfo("Veiksmīgi", f"Parole veiksmīgi pievienota dokumentam!")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās pievienot paroli: {e}")

    def on_item_double_click(self, event=None):  # Pievienots event parametrs, lai varētu izmantot kā bind funkciju
        """
        Apstrādā dubultklikšķi uz failu saraksta elementa.
        """
        selection = self.pdf_listbox.curselection()
        if selection:
            index = selection[0]
            # Pārliecināmies, ka indekss ir derīgs pašreizējās mapes saturam
            if index < len(self.current_folder["contents"]):  # <--- PĀRBAUDI ŠO!
                selected_item = self.current_folder["contents"][index]  # <--- PĀRBAUDI ŠO!

                # Izsaucam open_selected_item, kas jau apstrādā gan failus, gan mapes
                self.open_selected_item(selected_item)

    def navigate_to_folder(self, folder_name):
        """
        Navigē uz norādīto mapi pēc nosaukuma
        """

        def find_folder_by_name(folder_contents, target_name):
            """Rekursīvi meklē mapi pēc nosaukuma"""
            for item in folder_contents:
                if item["type"] == "folder" and item["name"] == target_name:
                    return item
                elif item["type"] == "folder" and "contents" in item:
                    # Rekursīvi meklē apakšmapēs
                    result = find_folder_by_name(item["contents"], target_name)
                    if result:
                        return result
            return None

        # Meklē mapi sākot no saknes
        target_folder = find_folder_by_name(self.internal_file_system["contents"], folder_name)

        if target_folder:
            self.current_folder = target_folder
            self.update_pdf_list()
            print(f"Navigēts uz mapi: {folder_name}")

            # Atjaunina navigācijas ceļu
            self.update_navigation_path()
        else:
            print(f"Mape '{folder_name}' nav atrasta")
            messagebox.showwarning("Mape nav atrasta", f"Nevar atrast mapi: {folder_name}")

    def navigate_to_folder_by_object(self, folder_object):
        """
        Navigē uz konkrēto mapes objektu (nevis meklē pēc nosaukuma)
        """
        if folder_object and folder_object.get("type") == "folder":
            self.current_folder = folder_object
            self.update_pdf_list()
            print(f"Navigēts uz mapi: {folder_object.get('name', 'Nezināma mape')}")
            self.update_navigation_path()
        else:
            print("Kļūda: Nederīgs mapes objekts")
            messagebox.showerror("Kļūda", "Nevar navigēt uz norādīto mapi")

    def update_navigation_path(self):
        """Atjaunina navigācijas ceļa rādījumu"""
        path_parts = []
        current = self.current_folder

        # Iet atpakaļ pa vecāku ķēdi, lai izveidotu ceļu
        while current and current != self.internal_file_system:
            path_parts.insert(0, current.get("name", "Nezināma mape"))
            current = current.get("parent")

        # Pievieno saknes mapi
        path_parts.insert(0, "Sakne")

        # Atjaunina ceļa rādījumu (ja tāds eksistē)
        path_text = " > ".join(path_parts)
        print(f"Pašreizējais ceļš: {path_text}")

    def remove_password_from_pdf(self, pdf_path):
        """Noņem paroli no PDF dokumenta"""
        try:
            reader = pypdf.PdfReader(pdf_path)

            if not reader.is_encrypted:
                messagebox.showinfo("Info", "Dokuments jau nav aizsargāts ar paroli!")
                return

            # Mēģinam atšifrēt bez paroles
            if not reader.decrypt(""):
                # Ja neizdodas, prasam paroli
                password = simpledialog.askstring("Parole",
                                                  "Ievadiet dokumenta paroli:",
                                                  show='*')
                if not password or not reader.decrypt(password):
                    messagebox.showerror("Kļūda", "Nepareiza parole!")
                    return

            writer = pypdf.PdfWriter()

            # Pārkopējam visas lapas
            for page in reader.pages:
                writer.add_page(page)

            # Saglabājam bez paroles
            with open(pdf_path, "wb") as f:
                writer.write(f)

            messagebox.showinfo("Veiksmīgi", "Parole veiksmīgi noņemta!")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās noņemt paroli: {e}")

    def change_password_of_pdf(self, pdf_path):
        """Maina PDF dokumenta paroli"""
        try:
            reader = pypdf.PdfReader(pdf_path)

            if not reader.is_encrypted:
                messagebox.showinfo("Info", "Dokuments nav aizsargāts ar paroli!")
                return self.add_password_to_pdf(pdf_path)  # Pārslēdzam uz paroles pievienošanu

            # Mēģinam atšifrēt bez paroles
            if not reader.decrypt(""):
                # Ja neizdodas, prasam esošo paroli
                old_pass = simpledialog.askstring("Esošā parole",
                                                  "Ievadiet pašreizējo paroli:",
                                                  show='*')
                if not old_pass or not reader.decrypt(old_pass):
                    messagebox.showerror("Kļūda", "Nepareiza parole!")
                    return

            # Iegūstam jauno paroli
            new_pass = simpledialog.askstring("Jaunā parole",
                                              "Ievadiet jauno paroli:",
                                              show='*')
            if not new_pass:
                return  # Lietotājs atcēla

            writer = pypdf.PdfWriter()

            # Pārkopējam visas lapas
            for page in reader.pages:
                writer.add_page(page)

            # Šifrējam ar jauno paroli
            writer.encrypt(new_pass)

            # Saglabājam izmaiņas
            with open(pdf_path, "wb") as f:
                writer.write(f)

            messagebox.showinfo("Veiksmīgi", "Parole veiksmīgi nomainīta!")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās mainīt paroli: {e}")

    def split_pdf_to_pages(self, pdf_filepath):
        """
        Sadaļa PDF dokumentu atsevišķās lapās, katru saglabājot kā jaunu PDF
        ar OCR tekstu, automātiski izveidotā apakšmapē blakus oriģinālajam failam.
        """
        if not os.path.exists(pdf_filepath):
            messagebox.showerror("Kļūda", f"Fails nav atrasts: {pdf_filepath}")
            return

        try:
            # Izveido jaunu mapi blakus oriģinālajam PDF failam
            base_dir = os.path.dirname(pdf_filepath)
            file_name_without_ext = os.path.splitext(os.path.basename(pdf_filepath))[0]
            output_folder = os.path.join(base_dir, f"{file_name_without_ext}_pages")
            os.makedirs(output_folder, exist_ok=True)

            # Atrod oriģinālā PDF faila mezglu iekšējā failu sistēmā
            original_pdf_node = None
            original_pdf_node_index = -1
            for i, item in enumerate(self.current_folder["contents"]):
                if item["type"] == "file" and item["filepath"] == pdf_filepath:
                    original_pdf_node = item
                    original_pdf_node_index = i
                    break

            if not original_pdf_node:
                messagebox.showwarning("Kļūda", "Oriģinālais PDF fails nav atrasts iekšējā failu sistēmā.")
                return

            # Izveido jaunu mapes mezglu iekšējā failu sistēmā
            new_folder_name = f"{file_name_without_ext}_pages"
            new_folder_node = {
                "type": "folder",
                "name": new_folder_name,
                "contents": [],
                "parent": self.current_folder  # Jaunā mape atrodas pašreizējā mapē
            }

            # Pievieno jauno mapi iekšējai failu sistēmai tajā pašā vietā, kur bija oriģinālais fails
            self.current_folder["contents"].insert(original_pdf_node_index, new_folder_node)
            # Noņem oriģinālo failu no pašreizējās mapes, jo tas tiks pārvietots uz jauno mapi
            self.current_folder["contents"].pop(
                original_pdf_node_index + 1)  # +1, jo jaunā mape tika ievietota pirms tam

            print(f"Iekšējā failu sistēmā pievienota jauna mape: {new_folder_name}")

            # Pārvieto oriģinālo PDF failu uz jaunizveidoto mapi
            original_pdf_name = os.path.basename(pdf_filepath)
            new_original_pdf_path = os.path.join(output_folder, original_pdf_name)

            # Pārvieto fizisko failu
            if os.path.abspath(pdf_filepath) != os.path.abspath(new_original_pdf_path):
                os.rename(pdf_filepath, new_original_pdf_path)
                print(f"Oriģinālais PDF '{original_pdf_name}' pārvietots uz: {output_folder}")

            # Atjaunina oriģinālā faila mezglu, lai tas atrastos jaunajā mapē
            original_pdf_node["filepath"] = new_original_pdf_path
            original_pdf_node["parent"] = new_folder_node
            new_folder_node["contents"].append(original_pdf_node)  # Pievieno oriģinālo failu jaunajai mapei

            # Atveram PDF dokumentu ar pypdf un fitz
            pdf_reader = pypdf.PdfReader(new_original_pdf_path)  # Lasām no pārvietotā faila
            pdf_document = fitz.open(new_original_pdf_path)
            total_pages = len(pdf_reader.pages)

            if total_pages == 0:
                messagebox.showwarning("Info", "PDF dokuments nesatur lapas.")
                pdf_document.close()
                return

            # Parāda progresu
            progress_window = Toplevel(self)
            progress_window.title("Sadalīšana...")
            progress_window.geometry("300x100")
            progress_window.transient(self)
            progress_window.grab_set()
            progress_window.resizable(False, False)

            progress_label = ttk.Label(progress_window, text="Sadalīšana: 0%")
            progress_label.pack(pady=10)
            progress_bar = ttk.Progressbar(progress_window, orient="horizontal", length=250, mode="determinate")
            progress_bar.pack(pady=5)

            # Centrē progresu logu
            progress_window.update_idletasks()
            x = self.winfo_x() + (self.winfo_width() // 2) - (progress_window.winfo_width() // 2)
            y = self.winfo_y() + (self.winfo_height() // 2) - (progress_window.winfo_height() // 2)
            progress_window.geometry(f"+{x}+{y}")

            self.update_idletasks()

            for i in range(total_pages):
                fixed_page_number = i + 1
                # Failu nosaukumi lapām ar _page_XXX sufiksu
                page_file_name = f"{file_name_without_ext}_page_{fixed_page_number:03d}.pdf"
                output_pdf_path = os.path.join(output_folder, page_file_name)

                new_pdf_writer = pypdf.PdfWriter()
                new_pdf_writer.add_page(pdf_reader.pages[i])

                fitz_page = pdf_document.load_page(i)
                pix = fitz_page.get_pixmap(dpi=self.dpi_var.get())
                img_pil = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                ocr_text = pytesseract.image_to_string(img_pil, lang=self.language_var.get(),
                                                       config=f'--psm {self.psm_var.get()} --oem {self.oem_var.get()}')

                if self.include_text_var.get() and ocr_text:
                    try:
                        temp_text_pdf_path = tempfile.mktemp(suffix=".pdf")
                        a4_width, a4_height = A4
                        c = canvas.Canvas(temp_text_pdf_path, pagesize=A4)
                        c.setFont("Helvetica", self.fontsize_var.get())

                        text_lines = ocr_text.split('\n')
                        y_pos = a4_height - 50
                        for line in text_lines:
                            c.drawString(50, y_pos, line)
                            y_pos -= self.fontsize_var.get() * 1.2
                            if y_pos < 50:
                                break
                        c.save()

                        overlay_pdf_reader = pypdf.PdfReader(temp_text_pdf_path)
                        page_to_overlay = overlay_pdf_reader.pages[0]

                        new_pdf_writer.pages[0].merge_page(page_to_overlay)
                        os.remove(temp_text_pdf_path)
                    except Exception as e:
                        print(f"Kļūda pievienojot OCR tekstu lapai {fixed_page_number}: {e}")
                        messagebox.showwarning("OCR teksta kļūda",
                                               f"Neizdevās pievienot OCR tekstu lapai {fixed_page_number}: {e}")

                with open(output_pdf_path, "wb") as output_pdf_file:
                    new_pdf_writer.write(output_pdf_file)

                # Pievieno sadalīto lapu jaunās mapes saturam
                if new_folder_node:
                    doc_id = str(uuid.uuid4())[:8]
                    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                    page_file_node = {
                        "type": "file",
                        "name": page_file_name,  # Izmantojam jauno nosaukumu
                        "filepath": output_pdf_path,
                        "doc_id": doc_id,
                        "date": current_time,
                        "parent": new_folder_node,
                        "original_page_number": fixed_page_number,
                        "display_name": f"Lapa {fixed_page_number}"
                    }
                    new_folder_node["contents"].append(page_file_node)
                    print(f"Pievienota lapa '{os.path.basename(output_pdf_path)}' iekšējai mapes struktūrai.")

                progress_percent = int(((i + 1) / total_pages) * 100)
                progress_label.config(text=f"Sadalīšana: {progress_percent}%")
                progress_bar['value'] = progress_percent
                progress_window.update_idletasks()

            pdf_document.close()
            progress_window.destroy()

            messagebox.showinfo("Sadalīšana pabeigta",
                                f"PDF dokuments veiksmīgi sadalīts {total_pages} lapās.\nSaglabāts mapē: {output_folder}\n\nOriģinālais fails pārvietots uz šo mapi.")

            # Saglabā izmaiņas iekšējā failu sistēmā un atsvaidzina sarakstu
            self.save_pdf_archive()
            self.refresh_pdf_list()  # Atsvaidzina sarakstu, lai parādītu jauno struktūru

        except Exception as e:
            messagebox.showerror("Kļūda sadalīšanas laikā", f"Neizdevās sadalīt PDF dokumentu: {e}")
    def init_scan_settings(self):
        """
        Inicializē skenēšanas iestatījumus ar noklusējuma vērtībām,
        ja tās nav ielādētas no settings faila.
        Šī metode tagad galvenokārt nodrošina, ka tk.Variable vērtības
        tiek atjauninātas no self.settings, ja tās ir mainījušās.
        """
        # Pārliecināmies, ka tk.Variable mainīgie ir inicializēti ar vērtībām no self.settings
        self.scan_camera_index.set(self.settings.get("scan_camera_index", 0))
        self.scan_camera_width.set(self.settings.get("scan_camera_width", 1280))
        self.scan_camera_height.set(self.settings.get("scan_camera_height", 120))
        self.scan_min_contour_area.set(self.settings.get("scan_min_contour_area", 2500))
        self.scan_stable_threshold.set(self.settings.get("scan_stable_threshold", 0.8))
        self.scan_stability_tolerance.set(self.settings.get("scan_stability_tolerance", 0.01))
        self.scan_aspect_ratio_min.set(self.settings.get("scan_aspect_ratio_min", 0.4))
        self.scan_aspect_ratio_max.set(self.settings.get("scan_aspect_ratio_max", 2.3))
        self.scan_gaussian_blur_kernel.set(self.settings.get("scan_gaussian_blur_kernel", 9))
        self.scan_adaptive_thresh_block_size.set(self.settings.get("scan_adaptive_thresh_block_size", 11))
        self.scan_adaptive_thresh_c.set(self.settings.get("scan_adaptive_thresh_c", 3))
        self.scan_canny_thresh1.set(self.settings.get("scan_canny_thresh1", 610))
        self.scan_canny_thresh2.set(self.settings.get("scan_canny_thresh2", 190))
        self.scan_brightness.set(self.settings.get("scan_brightness", 0))
        self.scan_contrast.set(self.settings.get("scan_contrast", 0))
        self.scan_saturation.set(self.settings.get("scan_saturation", 0))
        self.scan_gamma.set(self.settings.get("scan_gamma", 1.0))
        self.scan_use_color_detection.set(self.settings.get("scan_use_color_detection", False))
        self.scan_target_color.set(self.settings.get("scan_target_color", "#FFFFFF"))
        self.scan_color_tolerance.set(self.settings.get("scan_color_tolerance", 30))
        self.scan_morphology_enabled.set(self.settings.get("scan_morphology_enabled", False))
        self.scan_morphology_kernel_size.set(self.settings.get("scan_morphology_kernel_size", 3))
        self.scan_edge_dilation.set(self.settings.get("scan_edge_dilation", 2))

        self.sync_scan_settings_from_vars()  # Sinhronizē self.scan_settings vārdnīcu ar tk.Variable vērtībām

    def sync_scan_settings_from_vars(self):
        """Sinhronizē self.scan_settings ar tk.IntVar/DoubleVar vērtībām."""
        self.scan_settings["scan_camera_index"] = self.scan_camera_index.get()
        self.scan_settings["scan_camera_width"] = self.scan_camera_width.get()
        self.scan_settings["scan_camera_height"] = self.scan_camera_height.get()
        self.scan_settings["scan_min_contour_area"] = self.scan_min_contour_area.get()
        self.scan_settings["scan_stable_threshold"] = self.scan_stable_threshold.get()
        self.scan_settings["scan_stability_tolerance"] = self.scan_stability_tolerance.get()
        self.scan_settings["scan_aspect_ratio_min"] = self.scan_aspect_ratio_min.get()
        self.scan_settings["scan_aspect_ratio_max"] = self.scan_aspect_ratio_max.get()
        self.scan_settings["scan_gaussian_blur_kernel"] = self.scan_gaussian_blur_kernel.get()
        self.scan_settings["scan_adaptive_thresh_block_size"] = self.scan_adaptive_thresh_block_size.get()
        self.scan_settings["scan_adaptive_thresh_c"] = self.scan_adaptive_thresh_c.get()
        self.scan_settings["scan_canny_thresh1"] = self.scan_canny_thresh1.get()
        self.scan_settings["scan_canny_thresh2"] = self.scan_canny_thresh2.get()

        # Attēlu apstrādes iestatījumi
        self.scan_settings["scan_brightness"] = self.scan_brightness.get()
        self.scan_settings["scan_contrast"] = self.scan_contrast.get()
        self.scan_settings["scan_saturation"] = self.scan_saturation.get()
        self.scan_settings["scan_gamma"] = self.scan_gamma.get()
        self.scan_settings["scan_use_color_detection"] = self.scan_use_color_detection.get()
        self.scan_settings["scan_target_color"] = self.scan_target_color.get()
        self.scan_settings["scan_color_tolerance"] = self.scan_color_tolerance.get()
        self.scan_settings["scan_morphology_enabled"] = self.scan_morphology_enabled.get()
        self.scan_settings["scan_morphology_kernel_size"] = self.scan_morphology_kernel_size.get()
        self.scan_settings["scan_edge_dilation"] = self.scan_edge_dilation.get()

    def create_widgets(self):
        """Izveido galvenās lietotnes logrīkus un cilnes."""
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # --- Attēlu apstrādes cilne ---
        self.image_processing_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.image_processing_tab, text="Attēlu apstrāde")
        self.create_image_processing_widgets(self.image_processing_tab)

        # --- Failu pārvaldības cilne ---
        self.file_management_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.file_management_tab, text="Failu pārvaldība")
        self.create_file_management_widgets(self.file_management_tab)

        # --- Papildu rīku cilne ---
        self.additional_tools_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.additional_tools_tab, text="Papildu rīki")
        self.create_additional_tools_widgets(self.additional_tools_tab)
        # --- Automatizācijas cilne ---
        self.automation_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.automation_tab, text="Automatizācija")
        self.create_automation_widgets(self.automation_tab)

    def create_image_processing_widgets(self, parent_frame):
        """Izveido logrīkus attēlu apstrādes cilnei."""

        # Augšējā rīkjosla ar ritjoslu (labots augstums)
        top_toolbar_container = ttk.Frame(parent_frame)
        top_toolbar_container.pack(fill="x", padx=5, pady=5)

        # Izveido kanvasu ar fiksētu augstumu (~40px, var pielāgot)
        self.top_toolbar_canvas = tk.Canvas(
            top_toolbar_container,
            highlightthickness=0,
            height=40  # Šo vērtību varat pielāgot pēc saviem vajadzībām
        )
        self.top_toolbar_canvas.pack(side=tk.TOP, fill=tk.X)

        # Lietojiet pack bez expand=True, lai novērstu nevēlamu augstuma pieaugumu

        # Izveido horizontālo ritjoslu kanvasam
        self.top_toolbar_scrollbar = ttk.Scrollbar(top_toolbar_container, orient=tk.HORIZONTAL,
                                                   command=self.top_toolbar_canvas.xview)
        self.top_toolbar_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Konfigurē kanvasu, lai tas izmantotu ritjoslu
        self.top_toolbar_canvas.configure(xscrollcommand=self.top_toolbar_scrollbar.set)
        self.top_toolbar_canvas.bind('<Configure>', lambda e: self.top_toolbar_canvas.configure(
            scrollregion=self.top_toolbar_canvas.bbox("all")))

        # Izveido rāmi, kurā tiks ievietotas visas pogas un citi elementi
        # Šis rāmis tiks ievietots kanvasā
        top_frame = ttk.Frame(self.top_toolbar_canvas, padding=1)
        self.top_toolbar_canvas.create_window((0, 0), window=top_frame, anchor="nw")

        # Svarīgi: Pārliecinieties, ka top_frame izmērs tiek atjaunināts, kad tā saturs mainās
        top_frame.bind("<Configure>",
                       lambda e: self.top_toolbar_canvas.configure(scrollregion=self.top_toolbar_canvas.bbox("all")))

        self.btn_open = ttk.Button(top_frame, text="Atvērt attēlus/PDF", command=self.open_files, bootstyle="primary")
        self.btn_open.pack(side=tk.LEFT, padx=2)

        # Pievienota poga kameras skenēšanai
        self.btn_scan_camera = ttk.Button(top_frame, text="Skenēt ar kameru", command=self.scan_document_with_camera,
                                          bootstyle="info", state=tk.NORMAL if OPENCV_AVAILABLE else tk.DISABLED)
        self.btn_scan_camera.pack(side=tk.LEFT, padx=2)

        self.btn_document_detection = ttk.Button(top_frame, text="Atlasīt dokumentu no attēla",
                                                 command=self.show_document_detection_menu,
                                                 bootstyle="warning")

        self.btn_document_detection.pack(side=tk.LEFT, padx=2)

        self.btn_settings = ttk.Button(top_frame, text="Vispārīgie Iestatījumi",
                                       command=self.show_settings)  # MAINĪTS TEKSTS
        self.btn_settings.pack(side=tk.LEFT, padx=2)

        self.btn_scan_settings = ttk.Button(top_frame, text="Skenēšanas Iestatījumi",
                                            command=self.show_scan_settingss)  # JAUNA POGA
        self.btn_scan_settings.pack(side=tk.LEFT, padx=2)

        self.btn_check_langs = ttk.Button(top_frame, text="Pārbaudīt valodas", command=self.check_ocr_languages)
        self.btn_check_langs.pack(side=tk.LEFT, padx=2)

        # OCR parametri
        params_frame = ttk.Frame(top_frame)
        params_frame.pack(side=tk.LEFT, padx=10)  # Šis rāmis joprojām izmanto pack, jo tas ir tiešais bērns top_frame

        ttk.Label(params_frame, text="DPI:").grid(row=0, column=0)
        self.dpi_var = tk.IntVar(value=300)
        ttk.Spinbox(params_frame, from_=70, to=600, increment=10, textvariable=self.dpi_var, width=4).grid(row=0,
                                                                                                           column=1,
                                                                                                           padx=2)

        ttk.Label(params_frame, text="Fonts:").grid(row=0, column=2)
        self.fontsize_var = tk.IntVar(value=7)
        ttk.Spinbox(params_frame, from_=5, to=20, increment=1, textvariable=self.fontsize_var, width=3).grid(row=0,
                                                                                                             column=3,
                                                                                                             padx=2)

        ttk.Label(params_frame, text="Konfidence:").grid(row=0, column=4)
        self.confidence_var = tk.IntVar(value=60)
        ttk.Spinbox(params_frame, from_=0, to=100, increment=5, textvariable=self.confidence_var, width=3).grid(row=0,
                                                                                                                column=5,
                                                                                                                padx=2)

        ttk.Label(params_frame, text="PSM:").grid(row=0, column=6)
        # self.psm_var = tk.IntVar(value=3) # Šī rinda vairs nav nepieciešama, jo definēta __init__
        ttk.Spinbox(params_frame, from_=0, to=13, increment=1, textvariable=self.psm_var, width=3).grid(row=0, column=7,
                                                                                                        padx=2)

        ttk.Label(params_frame, text="OEM:").grid(row=0, column=8)
        # self.oem_var = tk.IntVar(value=3) # Šī rinda vairs nav nepieciešama, jo definēta __init__
        ttk.Spinbox(params_frame, from_=0, to=3, increment=1, textvariable=self.oem_var, width=3).grid(row=0, column=9,
                                                                                                       padx=2)

        self.orientation_var = tk.StringVar(value="Auto")
        self.orientation_combo = ttk.Combobox(top_frame, values=self.orientation_options, state="readonly", width=15,
                                              textvariable=self.orientation_var)
        self.orientation_combo.pack(side=tk.LEFT, padx=5)

        self.include_text_var = tk.BooleanVar(value=tk.FALSE)
        ttk.Checkbutton(top_frame, text="Iekļaut meklējamo tekstu", variable=self.include_text_var).pack(side=tk.LEFT,
                                                                                                         padx=5)

        # Galvenā darba zona - izmantojam PanedWindow, lai nodrošinātu izmēru pielāgošanu
        main_paned_window = ttk.PanedWindow(parent_frame, orient=tk.HORIZONTAL)
        main_paned_window.pack(fill="both", expand=True, padx=5, pady=(0, 5))

        # Kreiso pusi (failu saraksts) ievietojam atsevišķā rāmī
        file_list_pane = ttk.Frame(main_paned_window)
        main_paned_window.add(file_list_pane,
                              weight=2)  # Piešķir nelielu svaru, lai sākumā būtu platāks

        # Labo pusi (attēla priekšskatījums un OCR teksts) ievietojam atsevišķā rāmī
        image_ocr_pane = ttk.Frame(main_paned_window)
        main_paned_window.add(image_ocr_pane, weight=1)  # Ļaujam tam izstiepties

        # Konfigurējam labās puses rāmi, lai tas izstieptos
        image_ocr_pane.columnconfigure(0, weight=1)
        image_ocr_pane.rowconfigure(0, weight=1)  # Rinda attēla priekšskatījumam
        image_ocr_pane.rowconfigure(1, weight=1)  # Rinda OCR tekstam

        # Failu saraksta rāmis
        file_list_container = ttk.Frame(file_list_pane)  # MAINĪTS: Vecāks tagad ir file_list_pane
        file_list_container.pack(fill="both", expand=True, padx=(0, 5))  # MAINĪTS: Izmantojam pack
        file_list_container.pack_propagate(False)  # Novērš rāmja izmēra maiņu
        file_list_container.rowconfigure(1, weight=1)  # Nodrošina, ka listbox izstiepjas

        ttk.Label(file_list_container, text="Atlasītie faili:").pack(fill="x")

        # Ritjosla failu sarakstam
        file_list_frame_with_scrollbar = ttk.Frame(file_list_container)
        file_list_frame_with_scrollbar.pack(fill="both", expand=True)  # Šis ir pareizi, jo ir iekš file_list_container

        self.file_listbox = tk.Listbox(file_list_frame_with_scrollbar, selectmode=tk.EXTENDED, exportselection=False)
        self.file_listbox.pack(side=tk.LEFT, fill="both", expand=True)

        file_list_scrollbar = ttk.Scrollbar(file_list_frame_with_scrollbar, orient=tk.VERTICAL,
                                            command=self.file_listbox.yview)
        file_list_scrollbar.pack(side=tk.RIGHT, fill="y")
        self.file_listbox.config(yscrollcommand=file_list_scrollbar.set)

        #self.file_listbox.bind("<Button-1>", self.file_list_drag_start)  # Kreisais klikšķis sāk vilkšanu
        #self.file_listbox.bind("<B1-Motion>", self.file_list_drag_motion)  # Vilkšanas kustība
        #self.file_listbox.bind("<ButtonRelease-1>", self.file_list_drag_drop)  # Nomešana
        self.file_listbox.bind("<Button-3>", self.show_file_context_menu)  # Labais klikšķis
        self.file_listbox.bind("<Button-1>", self.on_file_click)  # Pievienojam jaunu bind, lai apstrādātu vienu klikšķi
        self.file_listbox.bind("<<ListboxSelect>>",
                               lambda e: self.after(1, self.on_file_select_deferred))  # Aizkavēta atlase
        # Vilkšanas datus vairs nevajag, jo drag-and-drop ir atspējots.

        # JAUNS: OCR pogas un progresa josla failu saraksta rāmī
        ocr_controls_frame = ttk.Frame(file_list_container)
        ocr_controls_frame.pack(fill="x", pady=5)

        self.progress = ttk.Progressbar(ocr_controls_frame, orient="horizontal", mode="determinate")
        self.progress.pack(fill="x", expand=True, padx=(0, 5))

        ocr_buttons_inner_frame = ttk.Frame(ocr_controls_frame)
        ocr_buttons_inner_frame.pack(fill="x", pady=(5, 0))  # Neliela atstarpe starp progress bar un pogām

        self.btn_start = ttk.Button(ocr_buttons_inner_frame, text="Sākt OCR", command=self.start_processing,
                                    bootstyle="success")
        self.btn_start.pack(side=tk.LEFT, expand=True, padx=2)

        self.btn_stop = ttk.Button(ocr_buttons_inner_frame, text="Apturēt", command=self.stop_processing_func,
                                   state=tk.DISABLED,
                                   bootstyle="danger")
        self.btn_stop.pack(side=tk.LEFT, expand=True, padx=2)

        # Pogas failu kārtošanai
        move_buttons_frame = ttk.Frame(file_list_container)  # Izmantojam file_list_container kā vecāku
        move_buttons_frame.pack(fill="x", pady=5)

        ttk.Button(move_buttons_frame, text="↑ Uz augšu", command=self.move_file_up).pack(side=tk.LEFT, padx=2)
        ttk.Button(move_buttons_frame, text="↓ Uz leju", command=self.move_file_down).pack(side=tk.LEFT, padx=2)

        # Attēla priekšskatījuma rāmis
        preview_frame = ttk.Frame(image_ocr_pane)  # MAINĪTS: Vecāks tagad ir image_ocr_pane
        preview_frame.grid(row=0, column=0, sticky="nsew")  # MAINĪTS: Izmantojam grid iekš image_ocr_pane
        preview_frame.rowconfigure(1, weight=1)  # Nodrošina, ka canvas izstiepjas

        ttk.Label(preview_frame, text="Attēla priekšskatījums:").pack(fill="x")
        self.canvas = tk.Canvas(preview_frame, background="#222222")
        self.canvas.pack(fill="both", expand=True)  # Šis ir pareizi, jo ir iekš preview_frame
        self.canvas.bind("<Configure>", self.resize_canvas)
        self.canvas.bind("<MouseWheel>", self.on_canvas_mouse_wheel)  # Zoom
        self.canvas.bind("<Button-4>", self.on_canvas_mouse_wheel)  # MacOS
        self.canvas.bind("<Button-5>", self.on_canvas_mouse_wheel)  # MacOS
        self.canvas.bind("<ButtonPress-2>", self.on_canvas_pan_start)  # Pan
        self.canvas.bind("<B2-Motion>", self.on_canvas_pan_drag)
        self.canvas.bind("<ButtonRelease-2>", self.on_canvas_pan_end)
        self.canvas.bind("<ButtonPress-1>", self.on_canvas_selection_start)  # Selection
        self.canvas.bind("<B1-Motion>", self.on_canvas_selection_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_canvas_selection_end)
        self.canvas_selection_rect = None
        self.canvas_selection_start_x = None
        self.canvas_selection_start_y = None
        self.canvas_zoom_factor = 1.0
        self.canvas_pan_x = 0
        self.canvas_pan_y = 0
        self.canvas_start_pan_x = 0
        self.canvas_start_pan_y = 0

        # Attēlu apstrādes rīki (ar ritjoslu)
        # Izveido rāmi ar ritjoslu ap attēlu apstrādes rīkiem
        processing_tools_outer_frame = ttk.Frame(preview_frame)
        processing_tools_outer_frame.pack(fill="x", pady=5)  # Šis ir pareizi, jo ir iekš preview_frame

        processing_tools_canvas = tk.Canvas(processing_tools_outer_frame, highlightthickness=0)
        processing_tools_canvas.pack(side=tk.LEFT, fill="both", expand=True)

        processing_tools_scrollbar = ttk.Scrollbar(processing_tools_outer_frame, orient=tk.VERTICAL,
                                                   command=processing_tools_canvas.yview)
        processing_tools_scrollbar.pack(side=tk.RIGHT, fill="y")

        processing_tools_canvas.configure(yscrollcommand=processing_tools_scrollbar.set)
        processing_tools_canvas.bind('<Configure>', lambda e: processing_tools_canvas.configure(
            scrollregion=processing_tools_canvas.bbox("all")))

        image_processing_tools_frame = ttk.LabelFrame(processing_tools_canvas, text="Attēlu apstrāde", padding=5)
        processing_tools_canvas.create_window((0, 0), window=image_processing_tools_frame, anchor="nw")

        # Izveido rāmi slīdņiem un pogām
        processing_inner_frame = ttk.Frame(image_processing_tools_frame)
        processing_inner_frame.pack(fill="both", expand=True)  # Šis ir pareizi, jo ir iekš image_processing_tools_frame

        # Slīdņi
        sliders_frame = ttk.Frame(processing_inner_frame)
        sliders_frame.grid(row=0, column=0, sticky="nsew", padx=5,
                           pady=5)  # Šis ir pareizi, jo ir iekš processing_inner_frame

        ttk.Label(sliders_frame, text="Spilgtums:").grid(row=0, column=0, sticky="w")
        self.brightness_slider = tk.Scale(sliders_frame, from_=0.1, to=3.0, resolution=0.1, orient=tk.HORIZONTAL,
                                          variable=self.brightness_var, command=self.apply_image_filters)
        self.brightness_slider.grid(row=0, column=1, sticky="ew", padx=5)

        ttk.Label(sliders_frame, text="Kontrasts:").grid(row=1, column=0, sticky="w")
        self.contrast_slider = tk.Scale(sliders_frame, from_=0.1, to=3.0, resolution=0.1, orient=tk.HORIZONTAL,
                                        variable=self.contrast_var, command=self.apply_image_filters)
        self.contrast_slider.grid(row=1, column=1, sticky="ew", padx=5)

        ttk.Label(sliders_frame, text="Asums:").grid(row=2, column=0, sticky="w")
        self.sharpness_slider = tk.Scale(sliders_frame, from_=2.0, to=3.5, resolution=0.1, orient=tk.HORIZONTAL,
                                         variable=self.sharpness_var, command=self.apply_image_filters)
        self.sharpness_slider.grid(row=2, column=1, sticky="ew", padx=5)

        ttk.Label(sliders_frame, text="Rotācija:").grid(row=3, column=0, sticky="w")
        self.rotate_spinbox = ttk.Spinbox(sliders_frame, from_=-360, to=360, increment=90, textvariable=self.rotate_var,
                                          width=5, command=lambda: self.apply_image_filters(None))
        self.rotate_spinbox.grid(row=3, column=1, sticky="ew", padx=5)
        sliders_frame.columnconfigure(1, weight=1)

        # Papildu attēlu apstrādes pogas un čekboksi
        processing_buttons_frame = ttk.Frame(processing_inner_frame)
        processing_buttons_frame.grid(row=0, column=1, sticky="nsew", padx=10,
                                      pady=5)  # Šis ir pareizi, jo ir iekš processing_inner_frame

        ttk.Checkbutton(processing_buttons_frame, text="Pelēktoņi", variable=self.grayscale_var,
                        command=lambda: self.apply_image_filters(None)).pack(anchor="w", pady=2)
        ttk.Checkbutton(processing_buttons_frame, text="Slīpuma korekcija", variable=self.deskew_var,
                        command=lambda: self.apply_image_filters(None),
                        state=tk.NORMAL if OPENCV_AVAILABLE else tk.DISABLED).pack(anchor="w", pady=2)
        ttk.Checkbutton(processing_buttons_frame, text="Trokšņu samazināšana", variable=self.remove_noise_var,
                        command=lambda: self.apply_image_filters(None)).pack(anchor="w", pady=2)
        ttk.Checkbutton(processing_buttons_frame, text="Attēla negatīvs", variable=self.invert_colors_var,
                        command=lambda: self.apply_image_filters(None)).pack(anchor="w", pady=2)
        ttk.Checkbutton(processing_buttons_frame, text="Malu noteikšana", variable=self.edge_detection_var,
                        command=lambda: self.apply_image_filters(None)).pack(anchor="w", pady=2)
        ttk.Checkbutton(processing_buttons_frame, text="Binārizācija", variable=self.binarize_var,
                        command=lambda: self.apply_image_filters(None)).pack(anchor="w", pady=2)

        self.btn_toggle_crop_mode = ttk.Button(processing_buttons_frame, text="Apgriezt attēlu (vilkt)",
                                               command=self.toggle_cropping_mode)
        # JAUNS: Poga QR koda rāmja attēlošanai un inicializēšanai
        self.btn_toggle_qr_frame = ttk.Button(processing_buttons_frame, text="Rediģēt koda rāmi",
                                              command=self.toggle_qr_frame_display)
        self.btn_toggle_qr_frame.pack(anchor="w", pady=2)

        self.btn_toggle_crop_mode.pack(anchor="w", pady=2)

        ttk.Button(processing_buttons_frame, text="Pagriezt par 90°", command=self.rotate_90_degrees).pack(anchor="w",
                                                                                                           pady=2)
        ttk.Button(processing_buttons_frame, text="Spoguļot (Horiz.)",
                   command=lambda: self.flip_image(Image.FLIP_LEFT_RIGHT)).pack(anchor="w", pady=2)
        ttk.Button(processing_buttons_frame, text="Spoguļot (Vert.)",
                   command=lambda: self.flip_image(Image.FLIP_TOP_BOTTOM)).pack(anchor="w", pady=2)
        ttk.Button(processing_buttons_frame, text="Mainīt izmērus", command=self.resize_image_dialog).pack(anchor="w",
                                                                                                           pady=2)
        ttk.Button(processing_buttons_frame, text="Auto uzlabošana", command=self.auto_enhance_image).pack(anchor="w",
                                                                                                           pady=2)
        ttk.Button(processing_buttons_frame, text="Rādīt histogrammu", command=self.show_image_histogram).pack(
            anchor="w", pady=2)
        ttk.Button(processing_buttons_frame, text="Rādīt metadatus", command=self.show_image_metadata).pack(anchor="w",
                                                                                                            pady=2)
        ttk.Button(processing_buttons_frame, text="Rādīt krāsu paleti", command=self.show_color_palette).pack(
            anchor="w", pady=2)
        ttk.Button(processing_buttons_frame, text="Atvērt pilnekrāna priekšskatījumu",
                   command=self.open_fullscreen_preview).pack(anchor="w", pady=2)

        # OCR rezultātu rāmis
        text_frame = ttk.Frame(image_ocr_pane)  # MAINĪTS: Vecāks tagad ir image_ocr_pane
        text_frame.grid(row=1, column=0, sticky="nsew", padx=5,
                        pady=(0, 5))  # MAINĪTS: Izmantojam grid iekš image_ocr_pane
        text_frame.rowconfigure(1, weight=1)  # Nodrošina, ka text_ocr izstiepjas

        ttk.Label(text_frame, text="OCR rezultāts (rediģējams):").pack(fill="x")

        # Ritjosla OCR tekstam
        text_ocr_frame_with_scrollbar = ttk.Frame(text_frame)
        text_ocr_frame_with_scrollbar.pack(fill="both", expand=True)  # Šis ir pareizi, jo ir iekš text_frame

        self.text_ocr = tk.Text(text_ocr_frame_with_scrollbar, wrap="word")
        self.text_ocr.pack(side=tk.LEFT, fill="both", expand=True)

        text_ocr_scrollbar = ttk.Scrollbar(text_ocr_frame_with_scrollbar, orient=tk.VERTICAL,
                                           command=self.text_ocr.yview)
        text_ocr_scrollbar.pack(side=tk.RIGHT, fill="y")
        self.text_ocr.config(yscrollcommand=text_ocr_scrollbar.set)

        file_list_buttons_frame = ttk.Frame(file_list_container)
        file_list_buttons_frame.pack(fill="x", pady=5)  # Šis ir pareizi, jo ir iekš file_list_container
        ttk.Button(file_list_buttons_frame, text="Dzēst atlasīto", command=self.delete_selected_image,
                   bootstyle="danger").pack(fill="x")

    def toggle_cropping_mode(self):
        """Ieslēdz/izslēdz attēla apgriešanas režīmu uz kanvasa."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, ko apgriezt.")
            return

        self.cropping_mode = not self.cropping_mode
        if self.cropping_mode:
            self.btn_toggle_crop_mode.config(bootstyle="warning")
            self.canvas.config(cursor="cross")
            messagebox.showinfo("Apgriešanas režīms",
                                "Apgriešanas režīms ieslēgts. Velciet ar peli, lai atlasītu apgriešanas apgabalu.")
        else:
            self.btn_toggle_crop_mode.config(bootstyle="default")
            self.canvas.config(cursor="arrow")
            if self.crop_rect_id:
                self.canvas.delete(self.crop_rect_id)
                self.crop_rect_id = None
            messagebox.showinfo("Apgriešanas režīms", "Apgriešanas režīms izslēgts.")

    def toggle_qr_frame_display(self):
        """Ieslēdz/izslēdz QR koda/svītrkoda rāmja rediģēšanas režīmu."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, lai rediģētu QR rāmi.")
            return

        if hasattr(self, '_qr_edit_mode') and self._qr_edit_mode:
            # Izslēdz rediģēšanas režīmu
            self._qr_edit_mode = False
            self.btn_toggle_qr_frame.config(bootstyle="default", text="Rediģēt QR pozīciju")
            messagebox.showinfo("QR rāmis", "QR koda rediģēšanas režīms izslēgts.")
        else:
            # Ieslēdz rediģēšanas režīmu
            self._qr_edit_mode = True
            if not hasattr(self, 'qr_code_frame_coords') or self.qr_code_frame_coords is None:
                self._set_default_qr_frame_coords()
            self.btn_toggle_qr_frame.config(bootstyle="warning", text="Beigt QR rediģēšanu")
            messagebox.showinfo("QR rāmis",
                                "QR koda rediģēšanas režīms ieslēgts. Velciet rāmi, lai pārvietotu, vai velciet stūrus, lai mainītu izmēru.")

        self.show_image_preview(self.images[self.current_image_index]["processed_img"])

    def create_file_management_widgets(self, parent_frame):
        """Izveido logrīkus failu pārvaldības cilnei."""
        # Izveido rāmi ar ritjoslu ap visu failu pārvaldības cilnes saturu
        file_management_outer_frame = ttk.Frame(parent_frame)
        file_management_outer_frame.pack(fill="both", expand=True)

        # Pārliecinieties, ka šis rāmis aizņem visu pieejamo vietu
        file_management_outer_frame.columnconfigure(0, weight=1)
        file_management_outer_frame.rowconfigure(0, weight=1)

        # Iekšējais rāmis, kurā atradīsies viss failu pārvaldības saturs
        inner_file_management_frame = ttk.Frame(file_management_outer_frame)
        inner_file_management_frame.grid(row=0, column=0, sticky="nsew")  # Izmanto grid, lai aizņemtu visu vietu

        # Meklēšanas un filtrēšanas rāmis
        filter_frame = ttk.LabelFrame(inner_file_management_frame, text="Meklēšana un filtrēšana", padding=10)
        filter_frame.pack(fill="x", padx=10, pady=5)

        # Meklēšana
        ttk.Label(filter_frame, text="Meklēt:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.search_var = tk.StringVar()
        self.search_entry = ttk.Entry(filter_frame, textvariable=self.search_var, width=200)
        self.search_entry.grid(row=0, column=1, sticky="ew", padx=5, pady=2)
        self.search_entry.bind("<KeyRelease>", self.filter_pdf_list)

        # Datuma filtrēšana
        ttk.Label(filter_frame, text="No datuma:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.start_date_var = tk.StringVar()
        self.start_date_entry = ttk.Entry(filter_frame, textvariable=self.start_date_var, width=400)
        self.start_date_entry.grid(row=1, column=1, sticky="w", padx=5, pady=2)
        ttk.Button(filter_frame, text="Kalendārs", command=self.open_start_date_calendar).grid(row=1, column=2,
                                                                                               sticky="w", padx=2)

        ttk.Label(filter_frame, text="Līdz datumam:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        self.end_date_var = tk.StringVar()
        self.end_date_entry = ttk.Entry(filter_frame, textvariable=self.end_date_var, width=400)
        self.end_date_entry.grid(row=2, column=1, sticky="w", padx=5, pady=2)
        ttk.Button(filter_frame, text="Kalendārs", command=self.open_end_date_calendar).grid(row=2, column=2,
                                                                                             sticky="w", padx=2)

        # Pogu rāmis, lai tās būtu mazākas un centrētas
        button_row_frame = ttk.Frame(filter_frame)
        button_row_frame.grid(row=3, column=0, columnspan=3, pady=5)  # Izvieto rāmi visā platumā
        button_row_frame.columnconfigure(0, weight=1)  # Centra pogas
        button_row_frame.columnconfigure(1, weight=1)  # Centra pogas
        button_row_frame.columnconfigure(2, weight=1)  # Centra pogas

        ttk.Button(button_row_frame, text="Filtrēt", command=self.filter_pdf_list, width=15).grid(row=0, column=0,
                                                                                                  padx=5, sticky="e")
        ttk.Button(button_row_frame, text="Notīrīt filtrus", command=self.clear_pdf_filters, width=15).grid(row=0,
                                                                                                            column=2,
                                                                                                            padx=5,
                                                                                                            sticky="w")

        filter_frame.columnconfigure(1, weight=1)

        # JAUNS: PanedWindow, lai sadalītu apgabalu trīs rūtīs
        # Kreisā: PDF priekšskatījums
        # Vidējā: Failu saraksts
        # Labā: Darbību pogas
        self.file_management_paned_window = ttk.PanedWindow(inner_file_management_frame, orient=tk.HORIZONTAL)
        self.file_management_paned_window.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # --- 1. Rūts: PDF priekšskatījums ---
        pdf_preview_container = ttk.Frame(self.file_management_paned_window)
        self.file_management_paned_window.add(pdf_preview_container, weight=1)  # Svars, lai izstieptos

        ttk.Label(pdf_preview_container, text="PDF priekšskatījums:").pack(fill="x", pady=(0, 5))
        self.pdf_preview_canvas = tk.Canvas(pdf_preview_container, bg="gray", bd=2, relief="sunken")
        self.pdf_preview_canvas.pack(fill="both", expand=True)

        # Pievienojam peles notikumus priekšskatījuma kanvasam
        self.pdf_preview_canvas.bind("<Configure>", self._on_pdf_preview_canvas_resize)
        self.pdf_preview_canvas.bind("<MouseWheel>", self._on_pdf_preview_mouse_wheel)
        self.pdf_preview_canvas.bind("<Button-4>", self._on_pdf_preview_mouse_wheel)  # MacOS
        self.pdf_preview_canvas.bind("<Button-5>", self._on_pdf_preview_mouse_wheel)  # MacOS
        self.pdf_preview_canvas.bind("<ButtonPress-1>", self._on_pdf_preview_pan_start)
        self.pdf_preview_canvas.bind("<B1-Motion>", self._on_pdf_preview_pan_drag)
        self.pdf_preview_canvas.bind("<ButtonRelease-1>", self._on_pdf_preview_pan_end)

        # Navigācijas pogas priekšskatījumam
        pdf_preview_nav_frame = ttk.Frame(pdf_preview_container)
        pdf_preview_nav_frame.pack(fill="x", pady=(5, 0))

        self.prev_page_button = ttk.Button(pdf_preview_nav_frame, text="← Iepriekšējā",
                                           command=self._show_prev_pdf_page, state=tk.DISABLED)
        self.prev_page_button.pack(side=tk.LEFT, expand=True, padx=2)

        self.pdf_page_label = ttk.Label(pdf_preview_nav_frame, text="Lapa: 0/0")
        self.pdf_page_label.pack(side=tk.LEFT, expand=True, padx=2)

        self.next_page_button = ttk.Button(pdf_preview_nav_frame, text="Nākamā →", command=self._show_next_pdf_page,
                                           state=tk.DISABLED)
        self.next_page_button.pack(side=tk.LEFT, expand=True, padx=2)

        # --- 2. Rūts: Failu saraksts un mapju navigācija ---
        file_list_and_nav_container = ttk.Frame(self.file_management_paned_window)
        self.file_management_paned_window.add(file_list_and_nav_container, weight=2)  # Lielāks svars, lai būtu platāks

        # Šeit NENOLODZAM meklēšanas un filtrēšanas rāmja, jo tas jau ir augšpusē
        # Tātad nekas nav jāievieto šeit

        # PDF arhīva saraksts un mapju navigācija (tagad iekš file_list_and_nav_container)
        archive_list_container = ttk.LabelFrame(file_list_and_nav_container, text="Saglabātie PDF faili", padding=10)
        archive_list_container.pack(fill=BOTH, expand=True, padx=0, pady=5)  # Noņemam ārējās padx/pady

        # Mapju navigācijas rīkjosla
        folder_nav_frame = ttk.Frame(archive_list_container)
        folder_nav_frame.pack(fill="x", pady=(0, 5))

        self.back_button = ttk.Button(folder_nav_frame, text="Atpakaļ", command=self.go_back_folder, state=DISABLED)
        self.back_button.pack(side=LEFT, padx=2)

        self.refresh_page_button = ttk.Button(folder_nav_frame, text="Atsvaidzināt lapu", command=self.refresh_pdf_list)
        self.refresh_page_button.pack(side=LEFT, padx=(5, 10))

        self.current_path_label = ttk.Label(folder_nav_frame, text="/")
        self.current_path_label.pack(side=LEFT, padx=5, expand=True, fill=tk.X)

        # Ritjosla PDF sarakstam
        pdf_list_frame_with_scrollbar = ttk.Frame(archive_list_container)
        pdf_list_frame_with_scrollbar.pack(side=LEFT, fill="both", expand=True)

        # Izveidojam tk.Text logrīku failu sarakstam
        self.pdf_listbox = tk.Text(pdf_list_frame_with_scrollbar, wrap="none", exportselection=False,
                                   font=("TkDefaultFont", 10))  # Pielāgojiet fontu, ja nepieciešams
        self.pdf_listbox.pack(side=LEFT, fill="both", expand=True)

        # Ritjoslas tk.Text logrīkam
        pdf_list_scrollbar_y = ttk.Scrollbar(pdf_list_frame_with_scrollbar, orient="vertical",
                                             command=self.pdf_listbox.yview)
        pdf_list_scrollbar_y.pack(side=RIGHT, fill="y")
        self.pdf_listbox.config(yscrollcommand=pdf_list_scrollbar_y.set)

        pdf_list_scrollbar_x = ttk.Scrollbar(archive_list_container, orient="horizontal",
                                             command=self.pdf_listbox.xview)
        pdf_list_scrollbar_x.pack(side=BOTTOM, fill="x")  # Novietojam zem saraksta
        self.pdf_listbox.config(xscrollcommand=pdf_list_scrollbar_x.set)

        # Konfigurējam tagus iekrāsošanai
        self.pdf_listbox.tag_configure("highlight", background="yellow", foreground="black")
        self.pdf_listbox.tag_configure("normal", background="", foreground="")  # Noklusējuma stils
        # JAUNS: Tags atlasītajai rindai
        self.pdf_listbox.tag_configure("selected_line", background="#007bff",
                                       foreground="white")  # Zils fons, balts teksts

        # Pielāgojam notikumu piesaistes tk.Text logrīkam
        self.pdf_listbox.bind("<Double-Button-1>", self.on_text_double_click)  # Jauna funkcija dubultklikšķim
        self.pdf_listbox.bind("<ButtonRelease-3>", self.on_text_right_click)
        self.pdf_listbox.bind("<ButtonRelease-2>", self.on_text_right_click)
        self.pdf_listbox.bind("<Button-1>", self.on_text_single_click)  # Pievienojam jaunu bind, lai apstrādātu vienu klikšķi

        # Drag and Drop bindings (ja vēlaties saglabāt, bet tk.Text to neatbalsta tieši kā Listbox)
        # Šīs funkcijas (drag_start, drag_motion, drag_drop) būs jāpārraksta vai jāatspējo,
        # jo tās ir paredzētas Listbox. Es tās atstāju komentētas, lai neradītu kļūdas.
        # self.pdf_listbox.bind("<Button-1>", self.drag_start)
        # self.pdf_listbox.bind("<B1-Motion>", self.drag_motion)
        # self.pdf_listbox.bind("<ButtonRelease-1>", self.drag_drop)
        # self.drag_data = {"x": 0, "y": 0, "item": None, "index": None}

        # --- 3. Rūts: Darbību pogas ---
        pdf_buttons_frame = ttk.Frame(self.file_management_paned_window)
        self.file_management_paned_window.add(pdf_buttons_frame, weight=0)  # Mazs svars, lai būtu šaurāks

        # Izveidojam rāmi pogu grupēšanai un izkārtojumam
        button_grid_frame = ttk.Frame(pdf_buttons_frame, padding=5)
        button_grid_frame.pack(fill="both", expand=True)

        # Konfigurējam kolonnas, lai pogas izstieptos vienmērīgi
        button_grid_frame.columnconfigure(0, weight=1)
        button_grid_frame.columnconfigure(1, weight=1)

        # Pogu izveide un izkārtojums, izmantojot grid
        # Rinda 0
        ttk.Button(button_grid_frame, text="📂 Atvērt", command=self.open_selected_item,
                   bootstyle="primary").grid(row=0, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="📁 Atvērt mapē", command=self.open_pdf_location,
                   bootstyle="secondary").grid(row=0, column=1, sticky="ew", padx=2, pady=2)

        # Rinda 1
        ttk.Button(button_grid_frame, text="🗑️ Dzēst", command=self.delete_selected_item,
                   bootstyle="danger").grid(row=1, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="📧 Nosūtīt e-pastā", command=self.send_selected_pdfs_by_email,
                   bootstyle="info").grid(row=1, column=1, sticky="ew", padx=2, pady=2)

        # Rinda 2
        ttk.Button(button_grid_frame, text="➕ Izveidot mapi", command=self.create_new_folder_internal,
                   bootstyle="success").grid(row=2, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="➡️ Pārvietot uz...", command=self.move_selected_items,
                   bootstyle="warning").grid(row=2, column=1, sticky="ew", padx=2, pady=2)

        # Rinda 3
        ttk.Button(button_grid_frame, text="✏️ Pārdēvēt", command=self.rename_selected_item,
                   bootstyle="light").grid(row=3, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="📄 Saglabāt kā Word", command=self.save_as_word,
                   bootstyle="dark").grid(row=3, column=1, sticky="ew", padx=2, pady=2)

        # Papildu pogas (ja nepieciešams, var pievienot šeit)
        # Piemēram, PDF šifrēšana/atšifrēšana
        ttk.Button(button_grid_frame, text="🔒 Pievienot paroli",
                   command=lambda: self.add_password_to_pdf(self._get_selected_pdf_filepath()),
                   bootstyle="secondary").grid(row=4, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="🔓 Noņemt paroli",
                   command=lambda: self.remove_password_from_pdf(self._get_selected_pdf_filepath()),
                   bootstyle="secondary").grid(row=4, column=1, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="🔑 Mainīt paroli",
                   command=lambda: self.change_password_of_pdf(self._get_selected_pdf_filepath()),
                   bootstyle="secondary").grid(row=5, column=0, sticky="ew", padx=2, pady=2)
        ttk.Button(button_grid_frame, text="✂️ Sadalīt PDF",
                   command=lambda: self.split_pdf_to_pages(self._get_selected_pdf_filepath()),
                   bootstyle="secondary").grid(row=5, column=1, sticky="ew", padx=2, pady=2)
        self.refresh_pdf_list()  # Tagad self.current_folder ir inicializēts

    def _get_selected_pdf_filepath(self):
        """Atgriež atlasītā PDF faila ceļu, ja tāds ir, pretējā gadījumā None."""
        if hasattr(self, '_selected_line_index') and self._selected_line_index != -1:
            line_number = self._selected_line_index
            if 0 <= line_number < len(self._displayed_items):
                selected_item = self._displayed_items[line_number]
                if selected_item["type"] == "file" and selected_item["name"].lower().endswith(".pdf"):
                    return selected_item['filepath']
        messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet PDF failu, lai veiktu šo darbību.")
        return None

    def on_text_single_click(self, event):
        """Apstrādā vienu klikšķi uz tk.Text logrīka, lai atlasītu visu rindu un parādītu priekšskatījumu."""
        # Iegūst klikšķa pozīciju
        index = self.pdf_listbox.index(f"@{event.x},{event.y}")
        line_number = int(index.split(".")[0]) - 1  # 0-bāzēts rindas numurs

        # Noņem iepriekšējo atlasi no visām rindām
        self.pdf_listbox.tag_remove("selected_line", "1.0", tk.END)
        # Noņem noklusējuma "sel" tagu, ja tas ir aktīvs
        self.pdf_listbox.tag_remove("sel", "1.0", tk.END)

        # Pievieno "selected_line" tagu atlasītajai rindai
        start_index = f"{line_number + 1}.0"
        end_index = f"{line_number + 1}.end"
        self.pdf_listbox.tag_add("selected_line", start_index, end_index)

        # Saglabā atlasītās rindas indeksu, lai to varētu izmantot citās funkcijās
        self._selected_line_index = line_number

        # Izsauc priekšskatījuma funkciju
        if 0 <= line_number < len(self._displayed_items):
            selected_item = self._displayed_items[line_number]
            if selected_item["type"] == "file" and selected_item["name"].lower().endswith(".pdf"):
                filepath = selected_item['filepath']
                if os.path.exists(filepath):
                    self._load_pdf_for_preview(filepath)
                else:
                    messagebox.showwarning("Fails nav atrasts", "Atlasītais PDF fails nav atrasts diskā.")
                    self._clear_pdf_preview()
            else:
                self._clear_pdf_preview()
        else:
            self._clear_pdf_preview()
            self._selected_line_index = -1  # Atiestata, ja nekas nav atlasīts

        # Lai novērstu noklusējuma teksta atlases uzvedību, atgriež "break"
        return "break"

    def on_text_right_click(self, event):
        print("Labais klikšķis uztverts!")  # Debug izdruka

        index = self.pdf_listbox.index(f"@{event.x},{event.y}")
        line_number = int(index.split(".")[0]) - 1

        self.pdf_listbox.tag_remove("selected_line", "1.0", tk.END)
        self.pdf_listbox.tag_remove("sel", "1.0", tk.END)

        start_index = f"{line_number + 1}.0"
        end_index = f"{line_number + 1}.end"
        self.pdf_listbox.tag_add("selected_line", start_index, end_index)

        self._selected_line_index = line_number

        # Parāda konteksta izvēlni ar ekrāna koordinātām
        try:
            self.show_pdf_context_menu(event)
        except Exception as e:
            print(f"Kļūda parādot konteksta izvēlni: {e}")

        return "break"

    def create_additional_tools_widgets(self, parent_frame):
        """Izveido logrīkus papildu rīku cilnei ar uzlabotu vizuālo noformējumu."""
        main_frame = ttk.Frame(parent_frame, padding=10)
        main_frame.pack(fill=BOTH, expand=True)

        # Galvenais PanedWindow, lai sadalītu cilni trīs vertikālās rūtīs
        main_paned_window = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        main_paned_window.pack(fill=BOTH, expand=True)

        # --- Kreisā rūts: Attēlu analīzes un papildu attēlu rīki ---
        left_pane = ttk.Frame(main_paned_window)
        main_paned_window.add(left_pane, weight=1)  # Piešķir vienādu svaru

        # Ritjosla kreisajai rūtij
        left_canvas = tk.Canvas(left_pane, highlightthickness=0)
        left_canvas.pack(side=LEFT, fill="both", expand=True)
        left_scrollbar = ttk.Scrollbar(left_pane, orient="vertical", command=left_canvas.yview)
        left_scrollbar.pack(side=RIGHT, fill="y")
        left_canvas.configure(yscrollcommand=left_scrollbar.set)
        left_canvas.bind('<Configure>', lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all")))
        inner_left_frame = ttk.Frame(left_canvas)
        left_canvas.create_window((0, 0), window=inner_left_frame, anchor="nw")

        # Attēlu analīzes rīki
        image_analysis_frame = ttk.LabelFrame(inner_left_frame, text="Attēlu analīze", padding=10)
        image_analysis_frame.pack(fill=X, padx=5, pady=5)
        self._create_button_grid(image_analysis_frame, [
            ("Histogramma", self.show_image_histogram, "chart-bar-fill"),
            ("Metadati", self.show_image_metadata, "info-circle-fill"),
            ("Krāsu palete", self.show_color_palette, "palette-fill"),
            ("Attēla salīdzināšana", self.compare_images, "arrows-left-right"),
            ("Kvalitātes novērtēšana", self.evaluate_image_quality, "star-fill"),
            ("Teksta izvilkšana no apgabala", self.extract_text_from_region, "crop")
        ])

        # Papildu attēlu apstrādes rīki
        advanced_image_tools_frame = ttk.LabelFrame(inner_left_frame, text="Papildu attēlu rīki", padding=10)
        advanced_image_tools_frame.pack(fill=X, padx=5, pady=5)
        self._create_button_grid(advanced_image_tools_frame, [
            ("Krāsu konvertēšana", self.convert_color_space, "paint-bucket-fill"),
            ("Ūdenszīmes pievienošana", self.add_watermark, "water"),
            ("Attēla mozaīka", self.create_image_mosaic, "grid-fill"),
            ("Attēla salikšana (stitch)", self.stitch_images, "puzzle-fill"),
            ("Attēla atjaunošana (inpainting)", self.image_inpainting, "magic"),
            ("Attēla stilizācija", self.stylize_image, "brush-fill"),
            ("Ģeometriskās transformācijas", self.geometric_transformations, "bounding-box"),
            ("Konvertēt uz pelēktoņiem", self.convert_to_grayscale, "image-fill"),
            ("Pielietot sliekšņošanu", self.apply_thresholding, "brightness-high-fill"),
            ("Pielietot Gausa izplūšanu", self.apply_gaussian_blur, "blur"),
            ("Pielietot mediānas filtru", self.apply_median_filter, "filter-circle-fill"),
            ("Uzlabot asumu", self.sharpen_image, "plus-circle-fill"),
            ("Pagriezt par leņķi", self.rotate_image_by_angle, "arrow-clockwise"),
            ("Pievienot teksta pārklājumu", self.add_text_overlay, "text-paragraph"),
            ("Zīmēt taisnstūri", self.draw_rectangle_on_image, "square-fill"),
            ("Zīmēt apli", self.draw_circle_on_image, "circle-fill"),
            ("Izvilkt krāsu kanālus", self.extract_color_channels, "layers-fill"),
            ("Apvienot krāsu kanālus", self.merge_color_channels, "stack-fill"),
            ("Pielietot sēpijas filtru", self.apply_sepia_filter, "camera-fill"),
            ("Pielietot vinjetes efektu", self.apply_vignette_effect, "circle-half"),
            ("Pikselizēt attēlu", self.pixelate_image, "grid-3x3-gap-fill"),
            ("Noteikt sejas", self.detect_faces, "person-bounding-box")
        ])

        # --- Vidējā rūts: QR koda ģenerators ---
        middle_pane = ttk.Frame(main_paned_window)
        main_paned_window.add(middle_pane, weight=1)

        qr_generator_frame = ttk.LabelFrame(middle_pane, text="QR koda ģenerators", padding=10)
        qr_generator_frame.pack(fill=BOTH, expand=True, padx=5, pady=5)

        ttk.Label(qr_generator_frame, text="Teksts QR kodam:").pack(pady=5)
        self.qr_text_var = tk.StringVar()
        ttk.Entry(qr_generator_frame, textvariable=self.qr_text_var, width=40).pack(fill=X, padx=5, pady=2)

        ttk.Button(qr_generator_frame, text="Ģenerēt QR kodu", command=self.generate_qr_code,
                   bootstyle="success", image=self._get_icon("qr-code-scan"), compound=tk.LEFT).pack(pady=10)

        # QR koda priekšskatījums
        self.qr_canvas = tk.Canvas(qr_generator_frame, bg="white", bd=2, relief="sunken")
        self.qr_canvas.pack(fill=BOTH, expand=True, padx=5, pady=5)

        # --- Labā rūts: PDF priekšskatījums ---
        right_pane = ttk.Frame(main_paned_window)
        main_paned_window.add(right_pane, weight=1)

        pdf_preview_container = ttk.LabelFrame(right_pane, text="PDF priekšskatījums", padding=10)
        pdf_preview_container.pack(fill=BOTH, expand=True, padx=5, pady=5)

        # Šī poga vairs netiek izmantota tiešai faila ielādei, jo priekšskatījums tiks sinhronizēts
        # ar "Attēlu apstrāde" cilnes atlasi.
        # Ja vēlaties, varat to noņemt vai mainīt tās funkcionalitāti.
        # ttk.Button(pdf_preview_container, text="Atjaunināt PDF priekšskatījumu",
        #            command=self._update_additional_tools_pdf_preview,
        #            bootstyle="primary", image=self._get_icon("file-earmark-pdf-fill"), compound=tk.LEFT).pack(pady=5)
        self.additional_tools_pdf_preview_canvas = tk.Canvas(pdf_preview_container, bg="gray", bd=2, relief="sunken")
        self.additional_tools_pdf_preview_canvas.pack(fill="both", expand=True, pady=5)

        # Pievienojam peles notikumus priekšskatījuma kanvasam
        self.additional_tools_pdf_preview_canvas.bind("<Configure>",
                                                      self._on_additional_tools_pdf_preview_canvas_resize)
        self.additional_tools_pdf_preview_canvas.bind("<MouseWheel>", self._on_additional_tools_pdf_preview_mouse_wheel)
        self.additional_tools_pdf_preview_canvas.bind("<Button-4>", self._on_additional_tools_pdf_preview_mouse_wheel)
        self.additional_tools_pdf_preview_canvas.bind("<Button-5>", self._on_additional_tools_pdf_preview_mouse_wheel)
        self.additional_tools_pdf_preview_canvas.bind("<ButtonPress-1>",
                                                      self._on_additional_tools_pdf_preview_pan_start)
        self.additional_tools_pdf_preview_canvas.bind("<B1-Motion>", self._on_additional_tools_pdf_preview_pan_drag)
        self.additional_tools_pdf_preview_canvas.bind("<ButtonRelease-1>",
                                                      self._on_additional_tools_pdf_preview_pan_end)

        # Navigācijas pogas priekšskatījumam
        pdf_preview_nav_frame = ttk.Frame(pdf_preview_container)
        pdf_preview_nav_frame.pack(fill="x", pady=(5, 0))

        self.additional_tools_prev_page_button = ttk.Button(pdf_preview_nav_frame, text="← Iepriekšējā",
                                                            command=self._show_prev_additional_tools_pdf_page,
                                                            state=tk.DISABLED,
                                                            image=self._get_icon("arrow-left-circle-fill"),
                                                            compound=tk.LEFT)
        self.additional_tools_prev_page_button.pack(side=tk.LEFT, expand=True, padx=2)

        self.additional_tools_pdf_page_label = ttk.Label(pdf_preview_nav_frame, text="Lapa: 0/0")
        self.additional_tools_pdf_page_label.pack(side=tk.LEFT, expand=True, padx=2)

        self.additional_tools_next_page_button = ttk.Button(pdf_preview_nav_frame, text="Nākamā →",
                                                            command=self._show_next_additional_tools_pdf_page,
                                                            state=tk.DISABLED,
                                                            image=self._get_icon("arrow-right-circle-fill"),
                                                            compound=tk.RIGHT)
        self.additional_tools_next_page_button.pack(side=tk.LEFT, expand=True, padx=2)



    def _update_additional_tools_pdf_preview(self):
        """
        Atjaunina PDF priekšskatījumu "Papildu rīki" cilnē, pamatojoties uz
        pašreizējo atlasīto failu no self.images saraksta.
        """
        # Aizver iepriekšējo dokumentu, ja tāds ir
        if self.additional_tools_current_pdf_document:
            self.additional_tools_current_pdf_document.close()
            self.additional_tools_current_pdf_document = None

        # Notīra kanvasu un atiestata navigācijas pogas
        self._clear_additional_tools_pdf_preview()

        if self.current_image_index == -1 or not self.images:
            self.additional_tools_pdf_preview_canvas.create_text(
                self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                text="Nav atlasīts attēls vai PDF priekšskatījumam.",
                fill="white", font=("Helvetica", 12), justify="center"
            )
            return

        selected_item = self.images[self.current_image_index]
        filepath = selected_item.get("filepath")

        if not filepath or not os.path.exists(filepath):
            self.additional_tools_pdf_preview_canvas.create_text(
                self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                text="Atlasītais fails nav atrasts vai nav derīgs.",
                fill="red", font=("Helvetica", 12), justify="center"
            )
            return

        # Pārbauda, vai fails ir PDF
        if filepath.lower().endswith(".pdf"):
            self.additional_tools_pdf_preview_canvas.create_text(
                self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                text="Ielādē PDF...", fill="white", font=("Helvetica", 14)
            )
            if self.additional_tools_pdf_page_label:
                self.additional_tools_pdf_page_label.config(text="Ielādē...")
            self.update_idletasks()

            try:
                self.additional_tools_current_pdf_document = fitz.open(filepath)
                self.additional_tools_current_pdf_page_count = self.additional_tools_current_pdf_document.page_count
                self.additional_tools_current_pdf_page_index = 0  # Sākam ar pirmo lapu
                self.additional_tools_pdf_preview_zoom_factor = 1.0  # Atiestatām tālummaiņu
                self.additional_tools_pdf_preview_pan_x = 0  # Atiestatām pārvietošanu
                self.additional_tools_pdf_preview_pan_y = 0

                self._display_pdf_page_on_additional_tools_canvas()

            except Exception as e:
                messagebox.showerror("PDF ielādes kļūda (Papildu rīki)", f"Nevarēja ielādēt PDF priekšskatījumam:\n{e}")
                self.additional_tools_current_pdf_document = None
                self.additional_tools_current_pdf_page_count = 0
                self.additional_tools_current_pdf_page_index = 0
                self._clear_additional_tools_pdf_preview()
                self.additional_tools_pdf_preview_canvas.create_text(
                    self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                    self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                    text=f"Nevarēja ielādēt PDF:\n{e}", fill="red", font=("Helvetica", 12),
                    justify="center"
                )
        else:
            # Ja atlasītais fails nav PDF, parādām attēlu
            try:
                img_pil = selected_item.get("processed_img")
                if img_pil:
                    # Pielāgo attēlu kanvasa izmēram
                    canvas_width = self.additional_tools_pdf_preview_canvas.winfo_width()
                    canvas_height = self.additional_tools_pdf_preview_canvas.winfo_height()

                    if canvas_width <= 1 or canvas_height <= 1:
                        # Kanvass vēl nav gatavs, mēģinām vēlreiz pēc īsa laika
                        self.after(50, self._update_additional_tools_pdf_preview)
                        return

                    display_img = img_pil.copy()
                    display_img.thumbnail((canvas_width, canvas_height), Image.Resampling.LANCZOS)

                    self.additional_tools_pdf_preview_photo = ImageTk.PhotoImage(display_img)
                    self.additional_tools_pdf_preview_canvas.delete("all")
                    self.additional_tools_pdf_preview_canvas.create_image(
                        canvas_width // 2, canvas_height // 2,
                        image=self.additional_tools_pdf_preview_photo, anchor="center"
                    )
                    if self.additional_tools_pdf_page_label:
                        self.additional_tools_pdf_page_label.config(text="Attēls")
                    if self.additional_tools_prev_page_button:
                        self.additional_tools_prev_page_button.config(state=tk.DISABLED)
                    if self.additional_tools_next_page_button:
                        self.additional_tools_next_page_button.config(state=tk.DISABLED)
                else:
                    self.additional_tools_pdf_preview_canvas.create_text(
                        self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                        self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                        text="Nav attēla datu priekšskatījumam.",
                        fill="red", font=("Helvetica", 12), justify="center"
                    )
            except Exception as e:
                self.additional_tools_pdf_preview_canvas.create_text(
                    self.additional_tools_pdf_preview_canvas.winfo_width() / 2,
                    self.additional_tools_pdf_preview_canvas.winfo_height() / 2,
                    text=f"Kļūda attēlojot attēlu:\n{e}", fill="red", font=("Helvetica", 12),
                    justify="center"
                )

    def _get_icon(self, icon_name, size=16):
        """
        Ielādē ikonu no tkfontawesome, ttkbootstrap iebūvētajām ikonām vai no faila.
        Atgriež PhotoImage objektu.
        """
        # 1. Mēģina ielādēt ikonu no tkfontawesome
        try:
            # Piezīme: Font Awesome ikonu nosaukumi var atšķirties no Bootstrap ikonu nosaukumiem.
            # Jums būs jāatrod atbilstošie Font Awesome nosaukumi.
            # Piemēram, 'chart-bar-fill' varētu būt 'chart-bar' Font Awesome.
            # Šeit ir daži piemēri, kā varētu mapēt:
            fa_icon_map = {
                "chart-bar-fill": "chart-bar",
                "info-circle-fill": "info-circle",
                "palette-fill": "palette",  # Varbūt "paint-brush" vai "fill-drip"
                "arrows-left-right": "arrows-alt-h",  # Vai "exchange-alt"
                "star-fill": "star",
                "crop": "crop",
                "paint-bucket-fill": "fill-drip",  # Vai "paint-brush"
                "water": "water",
                "grid-fill": "th-large",  # Vai "grip-horizontal"
                "puzzle-fill": "puzzle-piece",
                "magic": "magic",
                "brush-fill": "brush",
                "bounding-box": "box",  # Vai "vector-square"
                "image-fill": "image",
                "brightness-high-fill": "sun",  # Vai "lightbulb"
                "blur": "smog",  # Vai "cloud"
                "filter-circle-fill": "filter",
                "plus-circle-fill": "plus-circle",
                "arrow-clockwise": "sync-alt",  # Vai "redo-alt"
                "text-paragraph": "paragraph",
                "square-fill": "square",
                "circle-fill": "circle",
                "layers-fill": "layer-group",
                "stack-fill": "stack-overflow",  # Vai "layer-group"
                "camera-fill": "camera",
                "circle-half": "adjust",  # Vai "circle-notch"
                "grid-3x3-gap-fill": "grip-horizontal",  # Vai "th"
                "person-bounding-box": "user",  # Vai "user-alt"
                "qr-code-scan": "qrcode",
                "arrow-left-circle-fill": "arrow-circle-left",
                "arrow-right-circle-fill": "arrow-circle-right",
                # Pievienojiet citus mapējumus šeit
            }

            fa_icon_name = fa_icon_map.get(icon_name,
                                           icon_name)  # Mēģina atrast mapējumu, ja nav, izmanto oriģinālo nosaukumu

            # tkfontawesome izmanto 'fa' prefiksu, ja ikona nav tieši atrasta
            # Pārbaudiet Font Awesome dokumentāciju par precīziem nosaukumiem
            return faw.icons.get(fa_icon_name, size=size)
        except Exception as e:
            # print(f"Nevarēja ielādēt ikonu '{icon_name}' no tkfontawesome: {e}") # Debugging
            pass  # Turpina meklēt citos avotos

        # 2. Mēģina ielādēt ikonu no ttkbootstrap iebūvētajām ikonām
        try:
            return ttk.PhotoImage(name=icon_name, size=size)
        except Exception as e:
            # print(f"Nevarēja ielādēt ikonu '{icon_name}' no ttkbootstrap: {e}") # Debugging
            pass  # Turpina meklēt failā

        # 3. Ja neizdodas, mēģina ielādēt no faila (pieņemot, ka ikonas ir 'icons' mapē)
        icon_path = resource_path(os.path.join("icons", f"{icon_name}.png"))
        if os.path.exists(icon_path):
            try:
                img = Image.open(icon_path)
                img = img.resize((size, size), Image.LANCZOS)
                return ImageTk.PhotoImage(img)
            except Exception as e:
                print(f"Kļūda ielādējot ikonu no faila '{icon_path}': {e}")
        else:
            print(f"Ikonas '{icon_name}' nav atrasta ne tkfontawesome, ne ttkbootstrap, ne failā.")
        return None  # Atgriež None, ja ikonu nevar ielādēt

    def _create_button_grid(self, parent_frame, buttons_data, cols=2):
        """
        Izveido pogu režģi dotajā rāmī.
        buttons_data: saraksts ar (teksts, komanda, ikonas_nosaukums) tuple.
        cols: kolonnu skaits režģī.
        """
        button_frame = ttk.Frame(parent_frame)
        button_frame.pack(fill=X, pady=5)
        for i in range(cols):
            button_frame.columnconfigure(i, weight=1)

        for idx, (text, command, icon_name) in enumerate(buttons_data):
            row = idx // cols
            col = idx % cols
            icon = self._get_icon(icon_name)
            btn = ttk.Button(button_frame, text=text, command=command, bootstyle="secondary")
            if icon:
                btn.config(image=icon, compound=tk.LEFT)
            btn.grid(row=row, column=col, sticky="ew", padx=2, pady=2)

    def create_automation_widgets(self, parent_frame):
        """Izveido logrīkus automatizācijas cilnei ar uzlabotu izkārtojumu."""
        main_frame = ttk.Frame(parent_frame, padding=15)
        main_frame.pack(fill=BOTH, expand=True)

        # Ritjosla visam automatizācijas saturam
        automation_canvas = tk.Canvas(main_frame, highlightthickness=0)
        automation_canvas.pack(side=LEFT, fill="both", expand=True)

        automation_scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=automation_canvas.yview)
        automation_scrollbar.pack(side=RIGHT, fill="y")

        automation_canvas.configure(yscrollcommand=automation_scrollbar.set)
        automation_canvas.bind('<Configure>',
                               lambda e: automation_canvas.configure(scrollregion=automation_canvas.bbox("all")))

        # Iekšējais rāmis, kurā atradīsies visi automatizācijas rīki
        inner_automation_frame = ttk.Frame(automation_canvas)
        automation_canvas.create_window((0, 0), window=inner_automation_frame, anchor="nw",
                                        width=automation_canvas.winfo_width())

        # Piesaistām inner_automation_frame platumu canvas platumam
        def _on_frame_configure(event):
            automation_canvas.itemconfig(automation_canvas.find_withtag("inner_frame_window"), width=event.width)
            automation_canvas.configure(scrollregion=automation_canvas.bbox("all"))

        inner_automation_frame.bind('<Configure>', _on_frame_configure)
        automation_canvas.bind('<Configure>', lambda e: automation_canvas.itemconfigure(
            automation_canvas.find_withtag("inner_frame_window"), width=e.width))
        automation_canvas.create_window((0, 0), window=inner_automation_frame, anchor="nw", tags="inner_frame_window")

        # --- Automātiskās skenēšanas uzraudzība ---
        scan_monitor_frame = ttk.LabelFrame(inner_automation_frame, text="Automātiskā skenēšanas mapes uzraudzība",
                                            padding=20)
        scan_monitor_frame.pack(fill=X, padx=10, pady=10)
        scan_monitor_frame.columnconfigure(1, weight=1)  # Ļauj ievades laukam izstiepties

        # --- Papildu automatizācijas rīki ---
        additional_automation_frame = ttk.LabelFrame(inner_automation_frame, text="Papildu automatizācijas rīki",
                                                     padding=20)
        additional_automation_frame.pack(fill=X, padx=10, pady=10)
        additional_automation_frame.columnconfigure(0, weight=1) # Lai pogas izstieptos
        additional_automation_frame.columnconfigure(1, weight=1) # Lai pogas izstieptos

        # Pogu režģis jaunajām funkcijām
        self._create_button_grid(additional_automation_frame, [
            ("Automātiska pārdēvēšana", self.show_auto_rename_dialog, "file-earmark-text"),
            ("Automātiska PDF apvienošana", self.show_auto_merge_dialog, "file-earmark-ruled"),
            ("Automātiska PDF sadalīšana", self.show_auto_split_dialog, "files"),
            ("Automātiska metadatu pievienošana", self.show_auto_metadata_dialog, "tags"),
            ("Automātiska dokumentu klasifikācija", self.show_auto_classify_dialog, "folder-symlink")
        ], cols=2) # Izmantojam 2 kolonnas, lai izskatītos labāk

        ttk.Label(scan_monitor_frame, text="Skenēšanas mapes ceļš:", font=("Helvetica", 10, "bold")).grid(row=0,
                                                                                                          column=0,
                                                                                                          sticky=W,
                                                                                                          pady=5,
                                                                                                          padx=5)
        self.scan_folder_entry = ttk.Entry(scan_monitor_frame, textvariable=self.scan_folder_path,
                                           font=("Helvetica", 10))
        self.scan_folder_entry.grid(row=0, column=1, sticky=EW, pady=5, padx=5)
        ttk.Button(scan_monitor_frame, text="Pārlūkot...", command=self.browse_scan_folder,
                   bootstyle="secondary", image=self._get_icon("folder-open"), compound=tk.LEFT).grid(row=0, column=2,
                                                                                                      padx=5, pady=5)

        ttk.Checkbutton(scan_monitor_frame, text="Iespējot automātisko skenēšanu", variable=self.auto_scan_enabled,
                        command=self.toggle_auto_scan, bootstyle="round-toggle").grid(row=1, column=0, columnspan=2,
                                                                                      sticky=W, pady=10, padx=5)

        self.auto_scan_status_label = ttk.Label(scan_monitor_frame, text="Statuss: Izslēgts", bootstyle="info",
                                                font=("Helvetica", 10, "italic"))
        self.auto_scan_status_label.grid(row=2, column=0, columnspan=3, sticky=W, pady=5, padx=5)

        # --- Attālinātās glabāšanas iestatījumi ---
        remote_storage_frame = ttk.LabelFrame(inner_automation_frame, text="Attālinātās glabāšanas iestatījumi",
                                              padding=20)
        remote_storage_frame.pack(fill=X, padx=10, pady=10)
        remote_storage_frame.columnconfigure(1, weight=1)  # Ļauj combobox izstiepties

        # JAUNS: Google Sheets integrācijas iestatījumi
        google_sheets_frame = ttk.LabelFrame(inner_automation_frame, text="Google Sheets integrācija", padding=20)
        google_sheets_frame.pack(fill=X, padx=10, pady=10)
        google_sheets_frame.columnconfigure(1, weight=1)
        ttk.Label(google_sheets_frame, text="Google Sheet ID:", font=("Helvetica", 10, "bold")).grid(row=0, column=0,
                                                                                                     sticky=W, pady=5,
                                                                                                     padx=5)
        ttk.Entry(google_sheets_frame, textvariable=self.google_sheet_id, font=("Helvetica", 10)).grid(row=0, column=1,
                                                                                                       sticky=EW,
                                                                                                       pady=5, padx=5)
        ttk.Button(google_sheets_frame, text="Pārlūkot...",
                   command=lambda: self.browse_file_path(self.google_sheet_id, "Izvēlēties Google Sheet ID"),
                   bootstyle="secondary").grid(row=0, column=2, padx=5, pady=5)
        ttk.Label(google_sheets_frame, text="Lapas nosaukums:", font=("Helvetica", 10, "bold")).grid(row=1, column=0,
                                                                                                     sticky=W, pady=5,
                                                                                                     padx=5)
        ttk.Entry(google_sheets_frame, textvariable=self.google_sheet_name, font=("Helvetica", 10)).grid(row=1,
                                                                                                         column=1,
                                                                                                         sticky=EW,
                                                                                                         pady=5, padx=5)
        ttk.Label(google_sheets_frame, text="Akreditācijas fails:", font=("Helvetica", 10, "bold")).grid(row=2,
                                                                                                         column=0,
                                                                                                         sticky=W,
                                                                                                         pady=5, padx=5)
        ttk.Entry(google_sheets_frame, textvariable=self.google_sheet_credentials_path, font=("Helvetica", 10)).grid(
            row=2, column=1, sticky=EW, pady=5, padx=5)
        ttk.Button(google_sheets_frame, text="Pārlūkot...",
                   command=lambda: self.browse_file_path(self.google_sheet_credentials_path,
                                                         "Izvēlēties Google Sheets akreditācijas failu",
                                                         [("JSON files", "*.json")]), bootstyle="secondary").grid(row=2,
                                                                                                                  column=2,
                                                                                                                  padx=5,
                                                                                                                  pady=5)
        ttk.Button(google_sheets_frame, text="Autentificēties Google", command=self.authenticate_google_apis,
                   bootstyle="primary").grid(row=3, column=0, columnspan=3, pady=10)
        self.google_auth_status_label = ttk.Label(google_sheets_frame, text="Statuss: Nav autentificēts",
                                                  bootstyle="info", font=("Helvetica", 10, "italic"))
        self.google_auth_status_label.grid(row=4, column=0, columnspan=3, sticky=W, pady=5, padx=5)
        ttk.Checkbutton(google_sheets_frame, text="Automātiski atjaunināt Google Sheet",
                        variable=self.auto_upload_enabled, command=self.toggle_google_sheet_update,
                        bootstyle="round-toggle").grid(row=5, column=0, columnspan=3, sticky=W, pady=10, padx=5)
        # Pievienojiet šo pogu, lai manuāli atjauninātu Google Sheet
        ttk.Button(google_sheets_frame, text="Manuāli atjaunināt Google Sheet",
                   command=self.update_google_sheet_from_archive, bootstyle="info").grid(row=6, column=0, columnspan=3,
                                                                                         pady=10)

        ttk.Label(remote_storage_frame, text="Glabāšanas veids:", font=("Helvetica", 10, "bold")).grid(row=0, column=0,
                                                                                                       sticky=W, pady=5,
                                                                                                       padx=5)
        self.remote_storage_type_combo = ttk.Combobox(remote_storage_frame, textvariable=self.remote_storage_type,
                                                      values=["Local", "FTP", "SFTP", "Google Drive"], state="readonly",
                                                      font=("Helvetica", 10))
        self.remote_storage_type_combo.grid(row=0, column=1, sticky=EW, pady=5, padx=5)
        self.remote_storage_type_combo.bind("<<ComboboxSelected>>", self.update_remote_storage_fields)

        # FTP/SFTP iestatījumi
        self.ftp_settings_frame = ttk.LabelFrame(remote_storage_frame, text="FTP/SFTP iestatījumi", padding=15)
        self.ftp_settings_frame.grid(row=1, column=0, columnspan=2, sticky=EW, padx=5, pady=10)
        self.ftp_settings_frame.columnconfigure(1, weight=1)
        self.ftp_settings_frame.columnconfigure(3, weight=1)

        ttk.Label(self.ftp_settings_frame, text="Host:", font=("Helvetica", 9)).grid(row=0, column=0, sticky=W, pady=2,
                                                                                     padx=5)
        ttk.Entry(self.ftp_settings_frame, textvariable=self.ftp_host, font=("Helvetica", 9)).grid(row=0, column=1,
                                                                                                   sticky=EW, pady=2,
                                                                                                   padx=5)
        ttk.Label(self.ftp_settings_frame, text="Port:", font=("Helvetica", 9)).grid(row=0, column=2, sticky=W, pady=2,
                                                                                     padx=5)
        ttk.Entry(self.ftp_settings_frame, textvariable=self.ftp_port, font=("Helvetica", 9)).grid(row=0, column=3,
                                                                                                   sticky=EW, pady=2,
                                                                                                   padx=5)

        ttk.Label(self.ftp_settings_frame, text="Lietotājvārds:", font=("Helvetica", 9)).grid(row=1, column=0, sticky=W,
                                                                                              pady=2, padx=5)
        ttk.Entry(self.ftp_settings_frame, textvariable=self.ftp_user, font=("Helvetica", 9)).grid(row=1, column=1,
                                                                                                   sticky=EW, pady=2,
                                                                                                   padx=5)
        ttk.Label(self.ftp_settings_frame, text="Parole:", font=("Helvetica", 9)).grid(row=1, column=2, sticky=W,
                                                                                       pady=2, padx=5)
        ttk.Entry(self.ftp_settings_frame, textvariable=self.ftp_pass, show="*", font=("Helvetica", 9)).grid(row=1,
                                                                                                             column=3,
                                                                                                             sticky=EW,
                                                                                                             pady=2,
                                                                                                             padx=5)

        ttk.Label(self.ftp_settings_frame, text="Attālā mape:", font=("Helvetica", 9)).grid(row=2, column=0, sticky=W,
                                                                                            pady=2, padx=5)
        ttk.Entry(self.ftp_settings_frame, textvariable=self.ftp_remote_path, font=("Helvetica", 9)).grid(row=2,
                                                                                                          column=1,
                                                                                                          columnspan=3,
                                                                                                          sticky=EW,
                                                                                                          pady=2,
                                                                                                          padx=5)

        ttk.Checkbutton(self.ftp_settings_frame, text="Izmantot SFTP (drošs savienojums)", variable=self.ftp_use_sftp,
                        bootstyle="round-toggle").grid(row=3, column=0, columnspan=4, sticky=W, pady=10, padx=5)
        ttk.Button(self.ftp_settings_frame, text="Pārbaudīt savienojumu", command=self.test_ftp_connection,
                   bootstyle="info", image=self._get_icon("plug-fill"), compound=tk.LEFT).grid(row=4, column=0,
                                                                                               columnspan=4, pady=5,
                                                                                               padx=5)

        # Google Drive iestatījumi
        self.google_drive_settings_frame = ttk.LabelFrame(remote_storage_frame, text="Google Drive iestatījumi",
                                                          padding=15)
        self.google_drive_settings_frame.grid(row=2, column=0, columnspan=2, sticky=EW, padx=5, pady=10)
        self.google_drive_settings_frame.columnconfigure(1, weight=1)

        ttk.Label(self.google_drive_settings_frame, text="Mapes ID:", font=("Helvetica", 9)).grid(row=0, column=0,
                                                                                                  sticky=W, pady=2,
                                                                                                  padx=5)
        ttk.Entry(self.google_drive_settings_frame, textvariable=self.google_drive_folder_id,
                  font=("Helvetica", 9)).grid(row=0, column=1, sticky=EW, pady=2, padx=5)

        ttk.Label(self.google_drive_settings_frame, text="Akreditācijas ceļš:", font=("Helvetica", 9)).grid(row=1,
                                                                                                            column=0,
                                                                                                            sticky=W,
                                                                                                            pady=2,
                                                                                                            padx=5)
        ttk.Entry(self.google_drive_settings_frame, textvariable=self.google_drive_credentials_path,
                  font=("Helvetica", 9)).grid(row=1, column=1, sticky=EW, pady=2, padx=5)
        ttk.Button(self.google_drive_settings_frame, text="Pārlūkot...", command=self.browse_google_credentials,
                   bootstyle="secondary", image=self._get_icon("file-earmark-text"), compound=tk.LEFT).grid(row=1,
                                                                                                            column=2,
                                                                                                            padx=5,
                                                                                                            pady=2)

        ttk.Label(self.google_drive_settings_frame, text="Token ceļš:", font=("Helvetica", 9)).grid(row=2, column=0,
                                                                                                    sticky=W, pady=2,
                                                                                                    padx=5)
        ttk.Entry(self.google_drive_settings_frame, textvariable=self.google_drive_token_path,
                  font=("Helvetica", 9)).grid(row=2, column=1, sticky=EW, pady=2, padx=5)
        ttk.Button(self.google_drive_settings_frame, text="Pārlūkot...", command=self.browse_google_token,
                   bootstyle="secondary", image=self._get_icon("key-fill"), compound=tk.LEFT).grid(row=2, column=2,
                                                                                                   padx=5, pady=2)

        ttk.Button(self.google_drive_settings_frame, text="Autorizēties Google Drive",
                   command=self.authorize_google_drive,
                   bootstyle="info", image=self._get_icon("google"), compound=tk.LEFT).grid(row=3, column=0,
                                                                                            columnspan=3, pady=5,
                                                                                            padx=5)

        # --- Automātiskās augšupielādes iestatījumi ---
        auto_upload_frame = ttk.LabelFrame(inner_automation_frame, text="Automātiskās augšupielādes iestatījumi",
                                           padding=20)
        auto_upload_frame.pack(fill=X, padx=10, pady=10)
        auto_upload_frame.columnconfigure(1, weight=1)

        ttk.Checkbutton(auto_upload_frame, text="Iespējot automātisko augšupielādi pēc OCR",
                        variable=self.auto_upload_enabled,
                        bootstyle="round-toggle").grid(row=0, column=0, columnspan=2, sticky=W, pady=10, padx=5)

        ttk.Label(auto_upload_frame, text="Augšupielādes mērķis:", font=("Helvetica", 10, "bold")).grid(row=1, column=0,
                                                                                                        sticky=W,
                                                                                                        pady=5, padx=5)
        self.auto_upload_target_combo = ttk.Combobox(auto_upload_frame, textvariable=self.auto_upload_target,
                                                     values=["Local", "FTP", "SFTP", "Google Drive"], state="readonly",
                                                     font=("Helvetica", 10))
        self.auto_upload_target_combo.grid(row=1, column=1, sticky=EW, pady=5, padx=5)

        # Sākotnējā lauku atjaunināšana
        self.update_remote_storage_fields()

        # Pievienojiet jaunas funkcijas, kas tiek izsauktas no pogām
        # Šīs funkcijas ir jāpievieno jūsu OCRPDFApp klasē
        # (Ja tās jau eksistē, tad tās nav jāpievieno, bet jāpārliecinās, ka tās ir pareizi implementētas)

    # --- Jaunas automatizācijas funkcijas ---

    def show_auto_rename_dialog(self):
        """Parāda dialogu automātiskai failu pārdēvēšanai."""
        dialog = Toplevel(self)
        dialog.title("Automātiska pārdēvēšana")
        dialog.geometry("400x250")
        dialog.transient(self)
        dialog.grab_set()

        ttk.Label(dialog, text="Pārdēvēt failus, pamatojoties uz:").pack(pady=10)

        rename_option = tk.StringVar(value="ocr_text")
        ttk.Radiobutton(dialog, text="OCR tekstu (pirmās 20 zīmes)", variable=rename_option, value="ocr_text").pack(
            anchor="w", padx=20)
        ttk.Radiobutton(dialog, text="Datumu un laiku", variable=rename_option, value="datetime").pack(anchor="w",
                                                                                                       padx=20)
        ttk.Radiobutton(dialog, text="Atslēgvārdu (ja atrasts)", variable=rename_option, value="keyword").pack(
            anchor="w", padx=20)

        ttk.Button(dialog, text="Sākt pārdēvēšanu",
                   command=lambda: self.perform_auto_rename(rename_option.get(), dialog)).pack(pady=20)

    def authenticate_google_apis(self):
        """Autentificējas Google Drive un Google Sheets API."""
        creds_path = self.google_sheet_credentials_path.get()
        if not os.path.exists(creds_path):
            messagebox.showerror("Kļūda", f"Akreditācijas fails nav atrasts: {creds_path}")
            self.google_auth_status_label.config(text="Statuss: Kļūda (fails nav atrasts)", bootstyle="danger")
            return

        try:
            from google.oauth2.service_account import Credentials
            from googleapiclient.discovery import build

            SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/spreadsheets']
            creds = Credentials.from_service_account_file(creds_path, scopes=SCOPES)

            self.google_drive_service = build('drive', 'v3', credentials=creds)
            self.google_sheet_service = build('sheets', 'v4', credentials=creds)

            # Pārbaude, vai autentifikācija ir veiksmīga
            # Mēģinām iegūt informāciju par Google Drive saknes mapi
            self.google_drive_service.files().get(fileId='root').execute()
            # Mēģinām iegūt informāciju par Google Sheet
            if self.google_sheet_id.get():
                self.google_sheet_service.spreadsheets().get(spreadsheetId=self.google_sheet_id.get()).execute()

            self.google_auth_status_label.config(text="Statuss: Autentificēts ✅", bootstyle="success")
            messagebox.showinfo("Autentifikācija", "Veiksmīgi autentificēts Google Drive un Google Sheets API!")
        except Exception as e:
            messagebox.showerror("Autentifikācijas kļūda", f"Neizdevās autentificēties Google API: {e}")
            self.google_auth_status_label.config(text=f"Statuss: Kļūda ({e})", bootstyle="danger")
            self.google_drive_service = None
            self.google_sheet_service = None

    def toggle_google_sheet_update(self):
        """Ieslēdz/izslēdz automātisko Google Sheet atjaunināšanu."""
        if self.auto_upload_enabled.get():
            self.update_google_sheet_from_archive()
        else:
            messagebox.showinfo("Automātiskā atjaunināšana", "Automātiskā Google Sheet atjaunināšana izslēgta.")

    def upload_file_to_google_drive(self, file_path, folder_id=None):
        """Augšupielādē failu Google Drive un atgriež faila ID un kopīgošanas saiti."""
        if not self.google_drive_service:
            messagebox.showerror("Kļūda", "Google Drive pakalpojums nav autentificēts.")
            return None, None

        try:
            from googleapiclient.http import MediaFileUpload

            file_name = os.path.basename(file_path)
            file_metadata = {'name': file_name}
            if folder_id:
                file_metadata['parents'] = [folder_id]

            media = MediaFileUpload(file_path, mimetype='application/pdf') # Pieņemam, ka augšupielādējam PDF
            file = self.google_drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()

            file_id = file.get('id')
            web_view_link = file.get('webViewLink')

            # Iestatīt faila kopīgošanas atļaujas (publiski pieejams ar saiti)
            self.google_drive_service.permissions().create(
                fileId=file_id,
                body={'type': 'anyone', 'role': 'reader'},
                fields='id'
            ).execute()

            print(f"Fails augšupielādēts: {file_name}, ID: {file_id}, Saite: {web_view_link}")
            return file_id, web_view_link
        except Exception as e:
            messagebox.showerror("Google Drive augšupielādes kļūda", f"Neizdevās augšupielādēt failu Google Drive: {e}")
            return None, None

    def update_google_sheet_entry(self, file_info):
        """Atjaunina Google Sheet ar faila informāciju."""
        if not self.google_sheet_service or not self.google_sheet_id.get() or not self.google_sheet_name.get():
            print("Google Sheet pakalpojums nav autentificēts vai iestatījumi nav konfigurēti.")
            return

        try:
            spreadsheet_id = self.google_sheet_id.get()
            range_name = f"{self.google_sheet_name.get()}!A:Z" # Meklējam visā lapā

            # Pārbaudām, vai lapa eksistē, ja nē, izveidojam to
            try:
                self.google_sheet_service.spreadsheets().get(spreadsheetId=spreadsheet_id).execute()
            except Exception:
                # Lapa neeksistē, mēģinām izveidot
                body = {
                    'requests': [{
                        'addSheet': {
                            'properties': {
                                'title': self.google_sheet_name.get()
                            }
                        }
                    }]
                }
                self.google_sheet_service.spreadsheets().batchUpdate(spreadsheetId=spreadsheet_id, body=body).execute()
                print(f"Izveidota jauna lapa: {self.google_sheet_name.get()}")
                # Pievienojam galvenes
                header_values = [
                    "Faila nosaukums", "Faila ID (iekšējais)", "Faila ceļš (lokālais)",
                    "Dokumenta ID (piešķirtais)", "Izveides datums", "Mapes ceļš (iekšējais)",
                    "Google Drive ID", "Google Drive Saite"
                ]
                self.google_sheet_service.spreadsheets().values().update(
                    spreadsheetId=spreadsheet_id,
                    range=f"{self.google_sheet_name.get()}!A1",
                    valueInputOption='RAW',
                    body={'values': [header_values]}
                ).execute()


            # Iegūstam esošos datus, lai atrastu, vai ieraksts jau eksistē
            result = self.google_sheet_service.spreadsheets().values().get(spreadsheetId=spreadsheet_id, range=range_name).execute()
            values = result.get('values', [])

            # Meklējam ierakstu pēc faila ID (iekšējais)
            row_index_to_update = -1
            if values:
                for i, row in enumerate(values):
                    if len(row) > 1 and row[1] == file_info.get("doc_id"): # Pārbaudām iekšējo faila ID
                        row_index_to_update = i
                        break

            # Sagatavojam datus
            row_data = [
                file_info.get("name", ""),
                file_info.get("doc_id", ""),
                file_info.get("filepath", ""),
                file_info.get("assigned_id", ""), # Ja jums ir šāds lauks
                file_info.get("date", ""),
                file_info.get("internal_folder_path", ""), # Jums būs jāaprēķina šis ceļš
                file_info.get("google_drive_id", ""),
                file_info.get("google_drive_link", "")
            ]

            if row_index_to_update != -1:
                # Atjauninām esošo rindu
                update_range = f"{self.google_sheet_name.get()}!A{row_index_to_update + 1}"
                self.google_sheet_service.spreadsheets().values().update(
                    spreadsheetId=spreadsheet_id,
                    range=update_range,
                    valueInputOption='RAW',
                    body={'values': [row_data]}
                ).execute()
                print(f"Google Sheet ieraksts atjaunināts rindā {row_index_to_update + 1}.")
            else:
                # Pievienojam jaunu rindu
                append_range = f"{self.google_sheet_name.get()}!A:Z"
                self.google_sheet_service.spreadsheets().values().append(
                    spreadsheetId=spreadsheet_id,
                    range=append_range,
                    valueInputOption='RAW',
                    body={'values': [row_data]}
                ).execute()
                print("Jauns ieraksts pievienots Google Sheet.")

        except Exception as e:
            messagebox.showerror("Google Sheet atjaunināšanas kļūda", f"Neizdevās atjaunināt Google Sheet: {e}")

    def update_google_sheet_from_archive(self):
        """Atjaunina visu Google Sheet, pamatojoties uz iekšējo failu sistēmu."""
        if not self.google_sheet_service or not self.google_sheet_id.get() or not self.google_sheet_name.get():
            messagebox.showwarning("Brīdinājums", "Google Sheet pakalpojums nav autentificēts vai iestatījumi nav konfigurēti.")
            return

        try:
            spreadsheet_id = self.google_sheet_id.get()
            sheet_name = self.google_sheet_name.get()

            # Notīrām esošos datus (izņemot galvenes rindu)
            clear_range = f"{sheet_name}!A2:Z" # Sākot no otrās rindas
            self.google_sheet_service.spreadsheets().values().clear(spreadsheetId=spreadsheet_id, range=clear_range).execute()

            # Pievienojam galvenes, ja tās nav
            header_values = [
                "Faila nosaukums", "Faila ID (iekšējais)", "Faila ceļš (lokālais)",
                "Dokumenta ID (piešķirtais)", "Izveides datums", "Mapes ceļš (iekšējais)",
                "Google Drive ID", "Google Drive Saite"
            ]
            self.google_sheet_service.spreadsheets().values().update(
                spreadsheetId=spreadsheet_id,
                range=f"{sheet_name}!A1",
                valueInputOption='RAW',
                body={'values': [header_values]}
            ).execute()

            all_files_info = []
            # Rekursīvi iegūstam visu failu informāciju no iekšējās failu sistēmas
            def collect_file_info(node, current_path=""):
                if node["type"] == "file":
                    file_info = {
                        "name": node.get("name", ""),
                        "doc_id": node.get("doc_id", ""),
                        "filepath": node.get("filepath", ""),
                        "assigned_id": node.get("assigned_id", ""), # Pievienojiet, ja jums ir šāds lauks
                        "date": node.get("date", ""),
                        "internal_folder_path": current_path,
                        "google_drive_id": node.get("google_drive_id", ""),
                        "google_drive_link": node.get("google_drive_link", "")
                    }
                    all_files_info.append(file_info)
                elif node["type"] == "folder":
                    new_path = os.path.join(current_path, node["name"]) if current_path else node["name"]
                    for item in node.get("contents", []):
                        collect_file_info(item, new_path)

            collect_file_info(self.internal_file_system, "")

            if all_files_info:
                rows_to_append = []
                for file_info in all_files_info:
                    rows_to_append.append([
                        file_info.get("name", ""),
                        file_info.get("doc_id", ""),
                        file_info.get("filepath", ""),
                        file_info.get("assigned_id", ""),
                        file_info.get("date", ""),
                        file_info.get("internal_folder_path", ""),
                        file_info.get("google_drive_id", ""),
                        file_info.get("google_drive_link", "")
                    ])

                self.google_sheet_service.spreadsheets().values().append(
                    spreadsheetId=spreadsheet_id,
                    range=f"{sheet_name}!A:Z",
                    valueInputOption='RAW',
                    body={'values': rows_to_append}
                ).execute()
                messagebox.showinfo("Google Sheet atjaunināšana", f"Google Sheet veiksmīgi atjaunināts ar {len(rows_to_append)} ierakstiem.")
            else:
                messagebox.showinfo("Google Sheet atjaunināšana", "Nav failu, ko pievienot Google Sheet.")

        except Exception as e:
            messagebox.showerror("Google Sheet atjaunināšanas kļūda", f"Neizdevās atjaunināt Google Sheet: {e}")

    def browse_file_path(self, tk_string_var, title, filetypes=None):
        """Atver failu pārlūka dialogu un iestata izvēlēto ceļu tk.StringVar."""
        if filetypes is None:
            filetypes = [("All files", "*.*")]
        filepath = filedialog.askopenfilename(title=title, filetypes=filetypes)
        if filepath:
            tk_string_var.set(filepath)

    def browse_folder_path(self, tk_string_var, title):
        """Atver mapes pārlūka dialogu un iestata izvēlēto ceļu tk.StringVar."""
        folderpath = filedialog.askdirectory(title=title)
        if folderpath:
            tk_string_var.set(folderpath)

    def perform_auto_rename(self, option, dialog):
        """Veic automātisku failu pārdēvēšanu."""
        dialog.destroy()
        if not self.images:
            messagebox.showwarning("Nav failu", "Nav ielādētu attēlu vai PDF, ko pārdēvēt.")
            return

        renamed_count = 0
        for i, item in enumerate(self.images):
            original_filepath = item.get("filepath")
            if not original_filepath or not os.path.exists(original_filepath):
                continue

            base_dir = os.path.dirname(original_filepath)
            file_ext = os.path.splitext(original_filepath)[1]
            new_name_base = ""

            if option == "ocr_text" and self.ocr_results[i]:
                ocr_text = self.ocr_results[i].strip()
                if ocr_text:
                    new_name_base = ocr_text[:20].replace('\n', '_').replace('/', '_').replace('\\', '_').strip()
                    if not new_name_base:  # Ja pēc tīrīšanas nekas nepaliek
                        new_name_base = "OCR_dokuments"
                else:
                    new_name_base = "Bez_OCR_teksta"
            elif option == "datetime":
                new_name_base = datetime.now().strftime("%Y%m%d_%H%M%S")
            elif option == "keyword":
                # Šeit varētu būt sarežģītāka loģika atslēgvārdu meklēšanai OCR rezultātos
                # Vienkāršības labad, pieņemsim, ka meklējam "invoice" vai "receipt"
                ocr_text = self.ocr_results[i].lower() if self.ocr_results[i] else ""
                if "invoice" in ocr_text:
                    new_name_base = "Rēķins"
                elif "receipt" in ocr_text:
                    new_name_base = "Čeks"
                else:
                    new_name_base = "Dokuments"
            else:
                new_name_base = "Pārdēvēts_fails"

            new_filepath = os.path.join(base_dir, f"{new_name_base}{file_ext}")

            # Pievienojam unikālu sufiksu, ja fails ar šādu nosaukumu jau eksistē
            counter = 1
            temp_filepath = new_filepath
            while os.path.exists(temp_filepath) and temp_filepath != original_filepath:
                temp_filepath = os.path.join(base_dir, f"{new_name_base}_{counter}{file_ext}")
                counter += 1
            new_filepath = temp_filepath

            try:
                os.rename(original_filepath, new_filepath)
                item["filepath"] = new_filepath  # Atjaunina filepath iekšējā sarakstā
                item["name"] = os.path.basename(new_filepath)  # Atjaunina nosaukumu
                self.file_listbox.delete(i)
                self.file_listbox.insert(i, item["name"])
                renamed_count += 1
            except Exception as e:
                print(f"Kļūda pārdēvējot {original_filepath}: {e}")

        self.refresh_file_listbox()  # Atsvaidzina failu sarakstu
        messagebox.showinfo("Pārdēvēšana pabeigta", f"Veiksmīgi pārdēvēti {renamed_count} faili.")

    def show_auto_merge_dialog(self):
        """Parāda dialogu automātiskai PDF apvienošanai."""
        dialog = Toplevel(self)
        dialog.title("Automātiska PDF apvienošana")
        dialog.geometry("400x200")
        dialog.transient(self)
        dialog.grab_set()

        ttk.Label(dialog, text="Apvienot visus atlasītos PDF failus vienā.").pack(pady=10)
        ttk.Label(dialog, text="Rezultāta faila nosaukums:").pack(pady=5)
        self.merged_pdf_name_var = tk.StringVar(value="Apvienotais_dokuments.pdf")
        ttk.Entry(dialog, textvariable=self.merged_pdf_name_var, width=40).pack(padx=10)

        ttk.Button(dialog, text="Sākt apvienošanu", command=lambda: self.perform_auto_merge(dialog)).pack(pady=20)

    def perform_auto_merge(self, dialog):
        """Veic automātisku PDF apvienošanu."""
        dialog.destroy()
        selected_indices = self.file_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vismaz divus PDF failus, ko apvienot.")
            return
        if len(selected_indices) < 2:
            messagebox.showwarning("Nepietiek failu", "Lūdzu, atlasiet vismaz divus PDF failus, ko apvienot.")
            return

        pdf_paths_to_merge = []
        for index in selected_indices:
            item = self.images[index]
            if item.get("filepath") and item["filepath"].lower().endswith(".pdf"):
                pdf_paths_to_merge.append(item["filepath"])
            else:
                messagebox.showwarning("Nederīgs fails",
                                       f"Fails '{item.get('name', 'Nezināms')}' nav PDF un tiks izlaists.")

        if len(pdf_paths_to_merge) < 2:
            messagebox.showwarning("Nepietiek PDF", "Pēc atlases filtrēšanas palika mazāk par diviem PDF failiem.")
            return

        output_filename = self.merged_pdf_name_var.get()
        if not output_filename.lower().endswith(".pdf"):
            output_filename += ".pdf"

        # Izvēlas saglabāšanas mapi
        save_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile=output_filename,
            title="Saglabāt apvienoto PDF kā..."
        )
        if not save_path:
            return  # Lietotājs atcēla

        try:
            pdf_merger = pypdf.PdfMerger()
            for pdf_path in pdf_paths_to_merge:
                pdf_merger.append(pdf_path)

            with open(save_path, "wb") as output_file:
                pdf_merger.write(output_file)
            pdf_merger.close()

            messagebox.showinfo("Apvienošana pabeigta", f"PDF faili veiksmīgi apvienoti: {save_path}")
            # Pēc apvienošanas varat piedāvāt ielādēt jauno failu
            self.open_files(save_path)

        except Exception as e:
            messagebox.showerror("Kļūda apvienošanā", f"Neizdevās apvienot PDF failus: {e}")

    def show_auto_split_dialog(self):
        """Parāda dialogu automātiskai PDF sadalīšanai."""
        dialog = Toplevel(self)
        dialog.title("Automātiska PDF sadalīšana")
        dialog.geometry("450x300")
        dialog.transient(self)
        dialog.grab_set()

        ttk.Label(dialog, text="Sadalīt atlasīto PDF failu pēc atslēgvārdiem vai lapu skaita.").pack(pady=10)

        ttk.Label(dialog, text="Sadalīšanas veids:").pack(anchor="w", padx=10)
        self.split_type_var = tk.StringVar(value="keyword")
        ttk.Radiobutton(dialog, text="Pēc atslēgvārdiem (katra lapa ar atslēgvārdu sāk jaunu PDF)",
                        variable=self.split_type_var, value="keyword").pack(anchor="w", padx=20)
        ttk.Radiobutton(dialog, text="Pēc fiksēta lapu skaita", variable=self.split_type_var, value="pages").pack(
            anchor="w", padx=20)

        ttk.Label(dialog, text="Atslēgvārdi (atdalīti ar komatu) / Lapu skaits:").pack(anchor="w", padx=10, pady=5)
        self.split_param_var = tk.StringVar()
        ttk.Entry(dialog, textvariable=self.split_param_var, width=50).pack(padx=10, fill="x")

        ttk.Button(dialog, text="Sākt sadalīšanu", command=lambda: self.perform_auto_split(dialog)).pack(pady=20)

    def perform_auto_split(self, dialog):
        """Veic automātisku PDF sadalīšanu."""
        dialog.destroy()
        selected_indices = self.file_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vienu PDF failu, ko sadalīt.")
            return
        if len(selected_indices) > 1:
            messagebox.showwarning("Pārāk daudz failu", "Lūdzu, atlasiet tikai vienu PDF failu, ko sadalīt.")
            return

        selected_item = self.images[selected_indices[0]]
        pdf_path = selected_item.get("filepath")
        if not pdf_path or not pdf_path.lower().endswith(".pdf"):
            messagebox.showwarning("Nederīgs fails", "Atlasītais fails nav PDF.")
            return

        split_type = self.split_type_var.get()
        split_param = self.split_param_var.get()

        if not split_param:
            messagebox.showwarning("Trūkst parametra", "Lūdzu, ievadiet atslēgvārdus vai lapu skaitu.")
            return

        output_folder = filedialog.askdirectory(title="Izvēlieties mapi sadalīto PDF saglabāšanai")
        if not output_folder:
            return  # Lietotājs atcēla

        try:
            reader = pypdf.PdfReader(pdf_path)
            total_pages = len(reader.pages)
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]

            if split_type == "pages":
                pages_per_split = int(split_param)
                if pages_per_split <= 0:
                    messagebox.showerror("Kļūda", "Lapu skaitam jābūt lielākam par 0.")
                    return

                for i in range(0, total_pages, pages_per_split):
                    writer = pypdf.PdfWriter()
                    for j in range(i, min(i + pages_per_split, total_pages)):
                        writer.add_page(reader.pages[j])

                    output_filepath = os.path.join(output_folder, f"{base_name}_part_{i // pages_per_split + 1}.pdf")
                    with open(output_filepath, "wb") as output_file:
                        writer.write(output_file)
                messagebox.showinfo("Sadalīšana pabeigta",
                                    f"PDF sadalīts {total_pages // pages_per_split + (1 if total_pages % pages_per_split != 0 else 0)} daļās.")

            elif split_type == "keyword":
                keywords = [k.strip().lower() for k in split_param.split(',')]
                current_writer = pypdf.PdfWriter()
                part_num = 1

                for i in range(total_pages):
                    page = reader.pages[i]
                    page_text = page.extract_text().lower() if page.extract_text() else ""

                    is_keyword_page = any(keyword in page_text for keyword in keywords)

                    if is_keyword_page and len(current_writer.pages) > 0:
                        # Saglabā iepriekšējo daļu, ja ir lapas
                        output_filepath = os.path.join(output_folder, f"{base_name}_part_{part_num}.pdf")
                        with open(output_filepath, "wb") as output_file:
                            current_writer.write(output_file)
                        part_num += 1
                        current_writer = pypdf.PdfWriter()  # Sāk jaunu daļu

                    current_writer.add_page(page)

                # Saglabā pēdējo daļu
                if len(current_writer.pages) > 0:
                    output_filepath = os.path.join(output_folder, f"{base_name}_part_{part_num}.pdf")
                    with open(output_filepath, "wb") as output_file:
                        current_writer.write(output_file)

                messagebox.showinfo("Sadalīšana pabeigta", f"PDF sadalīts {part_num} daļās pēc atslēgvārdiem.")

        except Exception as e:
            messagebox.showerror("Kļūda sadalīšanā", f"Neizdevās sadalīt PDF failu: {e}")

    def show_auto_metadata_dialog(self):
        """Parāda dialogu automātiskai metadatu pievienošanai/atjaunināšanai."""
        dialog = Toplevel(self)
        dialog.title("Automātiska metadatu pievienošana")
        dialog.geometry("450x350")
        dialog.transient(self)
        dialog.grab_set()

        ttk.Label(dialog, text="Pievienot/atjaunināt metadatus atlasītajiem PDF failiem.").pack(pady=10)

        form_frame = ttk.Frame(dialog)
        form_frame.pack(padx=10, pady=5, fill="x")
        form_frame.columnconfigure(1, weight=1)

        ttk.Label(form_frame, text="Virsraksts:").grid(row=0, column=0, sticky="w", pady=2)
        self.meta_title_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.meta_title_var).grid(row=0, column=1, sticky="ew", pady=2)

        ttk.Label(form_frame, text="Autors:").grid(row=1, column=0, sticky="w", pady=2)
        self.meta_author_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.meta_author_var).grid(row=1, column=1, sticky="ew", pady=2)

        ttk.Label(form_frame, text="Tēma:").grid(row=2, column=0, sticky="w", pady=2)
        self.meta_subject_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.meta_subject_var).grid(row=2, column=1, sticky="ew", pady=2)

        ttk.Label(form_frame, text="Atslēgvārdi (komats):").grid(row=3, column=0, sticky="w", pady=2)
        self.meta_keywords_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.meta_keywords_var).grid(row=3, column=1, sticky="ew", pady=2)

        ttk.Label(form_frame, text="Izveides datums (YYYY-MM-DD):").grid(row=4, column=0, sticky="w", pady=2)
        self.meta_creation_date_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=self.meta_creation_date_var).grid(row=4, column=1, sticky="ew", pady=2)

        ttk.Button(dialog, text="Sākt metadatu pievienošanu", command=lambda: self.perform_auto_metadata(dialog)).pack(
            pady=20)

    def perform_auto_metadata(self, dialog):
        """Veic automātisku metadatu pievienošanu/atjaunināšanu."""
        dialog.destroy()
        selected_indices = self.file_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet PDF failus, kam pievienot metadatus.")
            return

        metadata = {
            "/Title": self.meta_title_var.get(),
            "/Author": self.meta_author_var.get(),
            "/Subject": self.meta_subject_var.get(),
            "/Keywords": self.meta_keywords_var.get(),
        }
        creation_date_str = self.meta_creation_date_var.get()
        if creation_date_str:
            try:
                # Pypdf prasa datumu formātā "D:YYYYMMDDHHMMSSZ00'00"
                dt_obj = datetime.strptime(creation_date_str, "%Y-%m-%d")
                metadata["/CreationDate"] = dt_obj.strftime("D:%Y%m%d%H%M%S+00'00'")
            except ValueError:
                messagebox.showwarning("Nederīgs datums",
                                       "Izveides datums nav pareizā formātā (YYYY-MM-DD). Tas tiks ignorēts.")

        updated_count = 0
        for index in selected_indices:
            item = self.images[index]
            pdf_path = item.get("filepath")
            if not pdf_path or not pdf_path.lower().endswith(".pdf"):
                continue

            try:
                reader = pypdf.PdfReader(pdf_path)
                writer = pypdf.PdfWriter()

                for page in reader.pages:
                    writer.add_page(page)

                # Atjaunina esošos metadatus un pievieno jaunos
                existing_metadata = reader.metadata
                if existing_metadata:
                    for key, value in existing_metadata.items():
                        if key not in metadata:  # Saglabā esošos, ja nav jaunu vērtību
                            metadata[key] = value

                writer.add_metadata(metadata)

                with open(pdf_path, "wb") as output_file:
                    writer.write(output_file)
                updated_count += 1
            except Exception as e:
                print(f"Kļūda atjauninot metadatus failam {pdf_path}: {e}")

        messagebox.showinfo("Metadati atjaunināti", f"Metadati veiksmīgi atjaunināti {updated_count} failiem.")

    def show_auto_classify_dialog(self):
        """Parāda dialogu automātiskai dokumentu klasifikācijai un pārvietošanai."""
        dialog = Toplevel(self)
        dialog.title("Automātiska dokumentu klasifikācija")
        dialog.geometry("500x400")
        dialog.transient(self)
        dialog.grab_set()

        ttk.Label(dialog, text="Klasificēt atlasītos dokumentus, pamatojoties uz OCR tekstu, un pārvietot tos.").pack(
            pady=10)

        ttk.Label(dialog, text="Klasifikācijas noteikumi (Atslēgvārds:Mērķa_mape, katrs jaunā rindā):").pack(anchor="w",
                                                                                                             padx=10)
        self.classification_rules_text = tk.Text(dialog, height=8, width=50)
        self.classification_rules_text.pack(padx=10, pady=5, fill="both", expand=True)
        self.classification_rules_text.insert(tk.END,
                                              "rēķins:Rēķini\nčeks:Čeki\nlīgums:Līgumi\nID karte:Personu_dokumenti")

        ttk.Label(dialog, text="Noklusējuma mape (ja neatrod atbilstību):").pack(anchor="w", padx=10, pady=5)
        self.default_classify_folder_var = tk.StringVar(value="Neklasificēti")
        ttk.Entry(dialog, textvariable=self.default_classify_folder_var, width=50).pack(padx=10, fill="x")

        ttk.Button(dialog, text="Sākt klasifikāciju", command=lambda: self.perform_auto_classify(dialog)).pack(pady=20)

    def perform_auto_classify(self, dialog):
        """Veic automātisku dokumentu klasifikāciju un pārvietošanu."""
        dialog.destroy()
        selected_indices = self.file_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet dokumentus, ko klasificēt.")
            return

        rules_text = self.classification_rules_text.get("1.0", tk.END).strip()
        rules = {}
        for line in rules_text.split('\n'):
            if ":" in line:
                keyword, folder = line.split(':', 1)
                rules[keyword.strip().lower()] = folder.strip()

        if not rules:
            messagebox.showwarning("Nav noteikumu", "Lūdzu, ievadiet klasifikācijas noteikumus.")
            return

        default_folder_name = self.default_classify_folder_var.get()
        if not default_folder_name:
            default_folder_name = "Neklasificēti"

        classified_count = 0
        for index in selected_indices:
            item = self.images[index]
            filepath = item.get("filepath")
            if not filepath or not os.path.exists(filepath):
                continue

            ocr_text = self.ocr_results[index].lower() if self.ocr_results[index] else ""

            target_folder_name = default_folder_name
            for keyword, folder_name in rules.items():
                if keyword in ocr_text:
                    target_folder_name = folder_name
                    break  # Atrasts pirmais atbilstošais noteikums

            try:
                # Pārvieto failu fiziski
                base_dir = os.path.dirname(filepath)
                target_folder_path = os.path.join(base_dir, target_folder_name)
                os.makedirs(target_folder_path, exist_ok=True)  # Izveido mapi, ja tā neeksistē

                new_filepath = os.path.join(target_folder_path, os.path.basename(filepath))

                # Pievienojam unikālu sufiksu, ja fails ar šādu nosaukumu jau eksistē
                counter = 1
                temp_filepath = new_filepath
                while os.path.exists(temp_filepath) and temp_filepath != filepath:
                    temp_filepath = os.path.join(target_folder_path,
                                                 f"{os.path.splitext(os.path.basename(filepath))[0]}_{counter}{os.path.splitext(filepath)[1]}")
                    counter += 1
                new_filepath = temp_filepath

                os.rename(filepath, new_filepath)

                # Atjaunina iekšējo failu sistēmu (ja izmantojat)
                # Šī daļa ir sarežģītāka, jo jāatrod fails koka struktūrā un jāpārvieto
                # Vienkāršības labad, šis piemērs tikai pārvieto fizisko failu un atjaunina self.images
                item["filepath"] = new_filepath
                item["name"] = os.path.basename(new_filepath)
                self.file_listbox.delete(index)
                self.file_listbox.insert(index, item["name"])  # Atjaunina listbox ierakstu

                classified_count += 1
            except Exception as e:
                print(f"Kļūda klasificējot/pārvietojot {filepath}: {e}")

        self.refresh_file_listbox()  # Atsvaidzina failu sarakstu
        messagebox.showinfo("Klasifikācija pabeigta",
                            f"Veiksmīgi klasificēti un pārvietoti {classified_count} dokumenti.")

    def browse_google_credentials(self):
        """Atver failu dialogu Google Drive akreditācijas faila izvēlei."""
        filepath = filedialog.askopenfilename(
            title="Izvēlēties Google Drive akreditācijas failu (credentials.json)",
            filetypes=[("JSON faili", "*.json"), ("Visi faili", "*.*")]
        )
        if filepath:
            self.google_drive_credentials_path.set(filepath)
            messagebox.showinfo("Akreditācijas fails", f"Akreditācijas fails iestatīts uz: {filepath}")

    def browse_google_token(self):
        """Atver failu dialogu Google Drive token faila izvēlei."""
        filepath = filedialog.askopenfilename(
            title="Izvēlēties Google Drive token failu (token.json)",
            filetypes=[("JSON faili", "*.json"), ("Visi faili", "*.*")]
        )
        if filepath:
            self.google_drive_token_path.set(filepath)
            messagebox.showinfo("Token fails", f"Token fails iestatīts uz: {filepath}")

    def authorize_google_drive(self):
        """Autorizējas Google Drive API."""
        messagebox.showinfo("Autorizācija",
                            "Šī funkcija vēl nav pilnībā implementēta. Jums būs jāpievieno Google Drive API integrācijas loģika.")
        # Šeit būtu jāpievieno loģika, lai autorizētos Google Drive API, izmantojot credentials.json un saglabājot token.json
        # Piemēram, izmantojot Google API klienta bibliotēku:
        # from google.oauth2.credentials import Credentials
        # from google_auth_oauthlib.flow import InstalledAppFlow
        # from google.auth.transport.requests import Request
        # import pickle
        #
        # SCOPES = ['https://www.googleapis.com/auth/drive.file']
        # creds = None
        # if os.path.exists(self.google_drive_token_path.get()):
        #     with open(self.google_drive_token_path.get(), 'rb') as token:
        #         creds = pickle.load(token)
        # if not creds or not creds.valid:
        #     if creds and creds.expired and creds.refresh_token:
        #         creds.refresh(Request())
        #     else:
        #         flow = InstalledAppFlow.from_client_secrets_file(
        #             self.google_drive_credentials_path.get(), SCOPES)
        #         creds = flow.run_local_server(port=0)
        #     with open(self.google_drive_token_path.get(), 'wb') as token:
        #         pickle.dump(creds, token)
        # messagebox.showinfo("Autorizācija", "Google Drive autorizācija veiksmīga!")

    def update_remote_storage_fields(self, event=None):
        """Atjaunina attālinātās glabāšanas lauku redzamību atkarībā no izvēlētā veida."""
        selected_type = self.remote_storage_type.get()

        if selected_type in ["FTP", "SFTP"]:
            self.ftp_settings_frame.grid()
            self.google_drive_settings_frame.grid_remove()
        elif selected_type == "Google Drive":
            self.ftp_settings_frame.grid_remove()
            self.google_drive_settings_frame.grid()
        else:  # Local
            self.ftp_settings_frame.grid_remove()
            self.google_drive_settings_frame.grid_remove()

    def test_ftp_connection(self):
        """Pārbauda FTP/SFTP savienojumu."""
        host = self.ftp_host.get()
        port = self.ftp_port.get()
        user = self.ftp_user.get()
        password = self.ftp_pass.get()
        use_sftp = self.ftp_use_sftp.get()

        if not host or not port or not user or not password:
            messagebox.showwarning("Trūkst datu", "Lūdzu, aizpildiet visus FTP/SFTP laukus!")
            return

        try:
            if use_sftp:
                import paramiko
                with paramiko.SSHClient() as client:
                    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
                    client.connect(hostname=host, port=port, username=user, password=password, timeout=5)
                    sftp_client = client.open_sftp()
                    sftp_client.close()
                messagebox.showinfo("Savienojums", "SFTP savienojums veiksmīgs!")
            else:
                from ftplib import FTP
                with FTP() as ftp:
                    ftp.connect(host, port, timeout=5)
                    ftp.login(user, password)
                    ftp.quit()
                messagebox.showinfo("Savienojums", "FTP savienojums veiksmīgs!")
        except Exception as e:
            messagebox.showerror("Savienojuma kļūda", f"Neizdevās izveidot savienojumu:\n{e}")

    def browse_scan_folder(self):
        """Atver dialogu, lai izvēlētos skenēšanas mapi."""
        folder_selected = filedialog.askdirectory(title="Izvēlēties mapi automātiskai skenēšanai")
        if folder_selected:
            self.scan_folder_path.set(folder_selected)
            messagebox.showinfo("Skenēšanas mape", f"Skenēšanas mape iestatīta uz: {folder_selected}")

    def toggle_auto_scan(self):
        """Ieslēdz/izslēdz automātisko skenēšanu un Watchdog uzraudzību."""
        if self.auto_scan_enabled.get():
            folder_to_watch = self.scan_folder_path.get()
            if not os.path.isdir(folder_to_watch):
                messagebox.showerror("Kļūda", "Norādītā skenēšanas mape neeksistē vai nav derīga.")
                self.auto_scan_enabled.set(False)
                return

            self.event_handler = FileSystemEventHandler()
            self.event_handler.on_created = self.on_new_file_in_scan_folder
            self.observer = Observer()
            self.observer.schedule(self.event_handler, folder_to_watch, recursive=False)
            self.observer.start()
            self.auto_scan_status_label.config(text=f"Statuss: Aktīvs, uzrauga '{folder_to_watch}'",
                                               bootstyle="success")
            messagebox.showinfo("Automātiskā skenēšana",
                                f"Automātiskā skenēšana ieslēgta. Uzrauga mapi: {folder_to_watch}")
        else:
            if self.observer:
                self.observer.stop()
                self.observer.join()
                self.observer = None
            self.auto_scan_status_label.config(text="Statuss: Izslēgts", bootstyle="info")
            messagebox.showinfo("Automātiskā skenēšana", "Automātiskā skenēšana izslēgta.")

    def on_new_file_in_scan_folder(self, event):
        """Apstrādā jaunu failu parādīšanos skenēšanas mapē."""
        if event.is_directory:
            return
        filepath = event.src_path
        print(f"Jauns fails atrasts skenēšanas mapē: {filepath}")
        # Šeit varat pievienot loģiku, lai automātiski apstrādātu jauno failu, piemēram, pievienotu to OCR sarakstam
        # self.after(100, lambda: self.open_files(filepath)) # Var izsaukt open_files, lai pievienotu sarakstam
        # Vai arī automātiski veikt OCR un augšupielādi
        # self.process_and_upload_file(filepath)
        messagebox.showinfo("Jauns fails", f"Jauns fails atrasts skenēšanas mapē: {os.path.basename(filepath)}")

    def generate_qr_code(self):
        """Ģenerē QR kodu no ievadītā teksta un parāda to."""
        qr_text = self.qr_text_var.get()
        if not qr_text:
            messagebox.showwarning("Kļūda", "Lūdzu, ievadiet tekstu QR kodam.")
            return

        try:
            qr_img = qrcode.make(qr_text)
            qr_pil_img = qr_img.get_image()

            # Pielāgo izmēru, lai ietilptu kanvasā
            canvas_size = min(self.qr_canvas.winfo_width(), self.qr_canvas.winfo_height())
            if canvas_size == 1:  # Ja kanvass vēl nav inicializēts
                canvas_size = 200

            qr_pil_img = qr_pil_img.resize((canvas_size, canvas_size), Image.LANCZOS)
            self.qr_photo = ImageTk.PhotoImage(qr_pil_img)

            self.qr_canvas.delete("all")
            self.qr_canvas.create_image(0, 0, anchor="nw", image=self.qr_photo)
            self.qr_canvas.image = self.qr_photo  # Saglabā atsauci

            # Piedāvā saglabāt QR kodu
            save_path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG attēli", "*.png"), ("Visi faili", "*.*")],
                title="Saglabāt QR kodu kā"
            )
            if save_path:
                qr_img.save(save_path)
                messagebox.showinfo("QR kods", f"QR kods veiksmīgi saglabāts: {save_path}")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās ģenerēt QR kodu: {e}")

    def open_start_date_calendar(self):
        """Atver kalendāru sākuma datuma izvēlei."""
        self._open_calendar(self.start_date_var)

    def open_end_date_calendar(self):
        """Atver kalendāru beigu datuma izvēlei."""
        self._open_calendar(self.end_date_var)

    def _open_calendar(self, date_var):
        """Vienkāršs vizuāls kalendārs datuma izvēlei bez ārējām bibliotēkām."""

        class SimpleCalendar(tk.Toplevel):
            def __init__(self, parent, date_var):
                super().__init__(parent)
                self.title("Izvēlēties datumu")
                self.resizable(False, False)
                self.date_var = date_var
                self.parent = parent

                # Pašreizējais gads un mēnesis (var ielādēt no date_var, ja vēlaties)
                try:
                    dt = datetime.strptime(self.date_var.get(), "%Y-%m-%d")
                    self.year = dt.year
                    self.month = dt.month
                except Exception:
                    now = datetime.now()
                    self.year = now.year
                    self.month = now.month

                self.selected_day = None

                self._setup_widgets()
                self._populate_days()

                # Centrējam logu virs vecāka
                self.update_idletasks()
                x = self.parent.winfo_x() + (self.parent.winfo_width() // 2) - (self.winfo_width() // 2)
                y = self.parent.winfo_y() + (self.parent.winfo_height() // 2) - (self.winfo_height() // 2)
                self.geometry(f"+{x}+{y}")

                self.grab_set()
                self.focus_set()

            def _setup_widgets(self):
                # Augšējā josla ar mēneša un gada izvēli un pogām
                nav_frame = ttk.Frame(self)
                nav_frame.pack(padx=10, pady=5)

                self.prev_btn = ttk.Button(nav_frame, text="<", width=3, command=self._prev_month)
                self.prev_btn.grid(row=0, column=0)

                self.month_year_lbl = ttk.Label(nav_frame, text="", width=15, anchor="center")
                self.month_year_lbl.grid(row=0, column=1, columnspan=5)

                self.next_btn = ttk.Button(nav_frame, text=">", width=3, command=self._next_month)
                self.next_btn.grid(row=0, column=6)

                # Dienu nosaukumi
                days_frame = ttk.Frame(self)
                days_frame.pack(padx=10)

                self.day_labels = []
                for i, day_name in enumerate(["P", "O", "T", "C", "P", "S", "S"]):  # Pirmdiena līdz Svētdiena latviski
                    lbl = ttk.Label(days_frame, text=day_name, width=3, anchor="center", font=("Arial", 10, "bold"))
                    lbl.grid(row=0, column=i)
                    self.day_labels.append(lbl)

                # Rāmītis ar dienu pogām
                self.days_frame = ttk.Frame(self)
                self.days_frame.pack(padx=10, pady=5)

                self.day_buttons = []

                # Apakšā apstiprināšanas poga
                btn_frame = ttk.Frame(self)
                btn_frame.pack(pady=5)
                ok_btn = ttk.Button(btn_frame, text="Apstiprināt", command=self._on_ok, bootstyle="success")
                ok_btn.pack()

            def _populate_days(self):
                # Notīra iepriekšējās pogas
                for btn in self.day_buttons:
                    btn.destroy()
                self.day_buttons.clear()

                # Atjaunina mēneša un gada nosaukumu
                month_name = calendar.month_name[self.month]
                self.month_year_lbl.config(text=f"{month_name} {self.year}")

                # Iegūst pirmās dienas nedēļas dienu un dienu skaitu mēnesī
                cal = calendar.Calendar(firstweekday=0)  # Pirmdiena = 0
                month_days = list(cal.itermonthdays2(self.year, self.month))  # (diena, nedēļas diena)

                # Rindas un kolonnas izveide
                row = 0
                col = 0

                for day, weekday in month_days:
                    if day == 0:
                        # Dienas no iepriekšējā vai nākamā mēneša - tukšas vietas
                        lbl = ttk.Label(self.days_frame, text="", width=3)
                        lbl.grid(row=row, column=col)
                    else:
                        btn = ttk.Button(self.days_frame, text=str(day), width=3)
                        btn.grid(row=row, column=col, padx=1, pady=1)
                        btn.config(command=lambda d=day: self._on_day_selected(d))
                        self.day_buttons.append(btn)

                        # Ja šī diena ir atlasīta, izceļam
                        if (self.selected_day == day):
                            btn.state(["pressed"])
                        else:
                            btn.state(["!pressed"])

                    col += 1
                    if col > 6:
                        col = 0
                        row += 1

            def _on_day_selected(self, day):
                self.selected_day = day
                # Atjaunojam pogu stāvokli, lai izceltu atlasīto dienu
                for btn in self.day_buttons:
                    btn.state(["!pressed"])
                    if btn["text"] == str(day):
                        btn.state(["pressed"])

            def _prev_month(self):
                if self.month == 1:
                    self.month = 12
                    self.year -= 1
                else:
                    self.month -= 1
                self.selected_day = None
                self._populate_days()

            def _next_month(self):
                if self.month == 12:
                    self.month = 1
                    self.year += 1
                else:
                    self.month += 1
                self.selected_day = None
                self._populate_days()

            def _on_ok(self):
                if self.selected_day is None:
                    tk.messagebox.showwarning("Brīdinājums", "Lūdzu, izvēlieties datumu!")
                    return
                # Uzstāda datumu mainīgajā
                date_str = f"{self.year}-{self.month:02d}-{self.selected_day:02d}"
                self.date_var.set(date_str)
                self.destroy()
                self.parent.filter_pdf_list()

        # Izsaucam kalendāra logu
        SimpleCalendar(self, date_var)

    def filter_pdf_list(self, event=None):
        """Filtrē PDF sarakstu, pamatojoties uz meklēšanas terminu un datumu diapazonu, un iekrāso atbilstošos vārdus."""
        search_term = self.search_var.get().lower()
        start_date_str = self.start_date_var.get()
        end_date_str = self.end_date_var.get()

        self.pdf_listbox.config(state=tk.NORMAL)  # Atļaujam rediģēt
        self.pdf_listbox.delete("1.0", tk.END)  # Notīrām visu tekstu
        self.pdf_listbox.tag_remove("highlight", "1.0", tk.END)  # Notīrām iepriekšējos iekrāsojumus
        self.pdf_listbox.tag_remove("selected_line", "1.0", tk.END)  # Notīrām atlasi

        filtered_contents = []
        for item in self.current_folder["contents"]:
            match_search = True
            match_date = True

            # Meklēšana
            item_text_content = ""
            if item["type"] == "file":
                item_text_content = f"📄 {item['name']} ({item['date']})".lower()
                if search_term:
                    if search_term not in item_text_content and \
                            search_term not in item['filepath'].lower() and \
                            search_term not in item['doc_id'].lower():
                        match_search = False
            elif item["type"] == "folder":
                item_text_content = f"📁 {item['name']}".lower()
                if search_term:
                    if search_term not in item_text_content:
                        match_search = False

            # Datuma filtrēšana (tikai failiem)
            if item["type"] == "file" and (start_date_str or end_date_str):
                try:
                    entry_date = datetime.strptime(item['date'].split(" ")[0], "%Y-%m-%d").date()
                    if start_date_str:
                        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
                        if entry_date < start_date:
                            match_date = False
                    if end_date_str:
                        end_date = datetime.strptime(end_date_str, "%Y-%m-%d").date()
                        if entry_date > end_date:
                            match_date = False
                except ValueError:
                    pass  # Ignorē nederīgus datuma formātus

            if match_search and match_date:
                filtered_contents.append(item)

        # Šķirojam filtrēto saturu: vispirms mapes, tad faili, pēc tam alfabētiski
        # Šeit arī saglabājam oriģinālo failu prioritāti, ja tas ir sadalīts PDF mapē
        sorted_filtered_contents = []
        original_file_in_folder = None

        # Pārbaudām, vai pašreizējā mape ir sadalīta PDF mape un vai tajā ir oriģinālais fails
        if self.current_folder.get("name", "").endswith("_pages"):
            for item in filtered_contents:
                # Oriģinālais fails ir tas, kura nosaukums nav ar "_page_XXX" sufiksu
                if item["type"] == "file" and not "_page_" in item["name"]:
                    original_file_in_folder = item
                    break

            if original_file_in_folder:
                sorted_filtered_contents.append(original_file_in_folder)
                # Pievienojam pārējās lapas, šķirojot tās pēc lapas numura
                pages = sorted([
                    item for item in filtered_contents
                    if item["type"] == "file" and "_page_" in item["name"]
                ], key=lambda x: x.get("original_page_number", float('inf')))
                sorted_filtered_contents.extend(pages)

                # Pievienojam mapes, ja tādas ir
                folders = sorted([
                    item for item in filtered_contents
                    if item["type"] == "folder"
                ], key=lambda x: x["name"].lower())
                sorted_filtered_contents.extend(folders)
            else:
                # Ja nav oriģinālā faila vai nav sadalīta PDF mape, šķirojam kā parasti
                sorted_filtered_contents = sorted(filtered_contents,
                                                  key=lambda x: (0 if x["type"] == "folder" else 1, x["name"].lower()))
        else:
            # Ja nav sadalīta PDF mape, šķirojam kā parasti
            sorted_filtered_contents = sorted(filtered_contents,
                                              key=lambda x: (0 if x["type"] == "folder" else 1, x["name"].lower()))

        self._displayed_items = []  # Atjaunojam parādīto vienumu sarakstu

        for i, item in enumerate(sorted_filtered_contents):
            display_text = ""
            if item["type"] == "file":
                name_to_display = item.get("display_name", item['name'])
                display_text = f"{i + 1}. 📄 {name_to_display} ({item['date']})"
            elif item["type"] == "folder":
                display_text = f"{i + 1}. 📁 {item['name']}"

            start_index = self.pdf_listbox.index(tk.END)  # Iegūstam sākuma indeksu pirms ievietošanas
            self.pdf_listbox.insert(tk.END, display_text + "\n", "normal")  # Ievietojam tekstu ar jaunu rindu
            end_index = self.pdf_listbox.index(tk.END + "-1c")  # Iegūstam beigu indeksu pēc ievietošanas

            # Iekrāsojam atbilstošos vārdus, ja ir meklēšanas termins
            if search_term:
                start_pos = "1.0"
                while True:
                    start_pos = self.pdf_listbox.search(search_term, start_pos, stopindex=end_index, nocase=1)
                    if not start_pos:
                        break
                    end_pos = f"{start_pos}+{len(search_term)}c"
                    self.pdf_listbox.tag_add("highlight", start_pos, end_pos)
                    start_pos = end_pos

            self._displayed_items.append(item)  # Pievienojam vienumu sarakstam

        self.pdf_listbox.config(state=tk.DISABLED)  # Atkal atspējojam rediģēšanu

    def clear_pdf_filters(self):
        """Notīra visus PDF saraksta filtrus."""
        self.search_var.set("")
        self.start_date_var.set("")
        self.end_date_var.set("")
        self.refresh_pdf_list()

    def sync_current_folder_with_disk(self):
        """
        Sinhronizē pašreizējo iekšējās failu sistēmas mapi ar tās fizisko atbilstību diskā.
        Pievieno jaunus failus/mapes no diska un noņem tos, kas vairs neeksistē diskā.
        """
        current_physical_path = self._get_physical_path_from_node(self.current_folder)

        if not os.path.exists(current_physical_path):
            # Ja fiziskā mape vairs neeksistē, atgriežamies uz vecāku mapi
            messagebox.showwarning("Mape nav atrasta",
                                   f"Fiziskā mape '{current_physical_path}' vairs neeksistē. Atgriežamies uz iepriekšējo mapi.")
            self.go_back_folder()
            return

        # 1. Izveido sarakstu ar esošajiem vienumiem iekšējā struktūrā
        internal_items_map = {item["name"]: item for item in self.current_folder["contents"]}

        # 2. Pārbauda fiziskos failus/mapes diskā
        disk_items = set(os.listdir(current_physical_path))

        # Noņem vienumus no iekšējās struktūras, kas vairs neeksistē diskā
        items_to_remove_from_internal = []
        for name, item in internal_items_map.items():
            if name not in disk_items:
                items_to_remove_from_internal.append(item)

        for item_to_remove in items_to_remove_from_internal:
            if item_to_remove in self.current_folder[
                "contents"]:  # Pārbaude, lai izvairītos no kļūdām, ja vienums jau noņemts
                self.current_folder["contents"].remove(item_to_remove)
                print(f"Noņemts no iekšējās struktūras (neeksistē diskā): {item_to_remove['name']}")

        # Pievieno jaunus failus/mapes no diska, kas nav iekšējā struktūrā
        for name in disk_items:
            if name not in internal_items_map:
                item_path = os.path.join(current_physical_path, name)

                if os.path.isfile(item_path):
                    # Pievieno failu
                    doc_id = str(uuid.uuid4())[:8]
                    try:
                        creation_time = datetime.fromtimestamp(os.path.getctime(item_path))
                        date_str = creation_time.strftime("%Y-%m-%d %H:%M:%S")
                    except:
                        date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                    new_file = {
                        "type": "file",
                        "name": name,
                        "filepath": item_path,
                        "doc_id": doc_id,
                        "date": date_str,
                        "parent": self.current_folder
                    }
                    self.current_folder["contents"].append(new_file)
                    print(f"Pievienots jauns fails no diska: {name}")

                elif os.path.isdir(item_path):
                    # Pievieno mapi
                    new_folder = {
                        "type": "folder",
                        "name": name,
                        "contents": [],  # Satura ielāde notiks, kad tiks navigēts uz šo mapi
                        "parent": self.current_folder
                    }
                    self.current_folder["contents"].append(new_folder)
                    print(f"Pievienota jauna mape no diska: {name}")

        # Pēc sinhronizācijas saglabājam arhīvu
        self.save_pdf_archive()

    def refresh_pdf_list(self):
        """Atjaunina PDF sarakstu, parādot pašreizējās mapes saturu un sinhronizējot ar fizisko disku."""
        self.sync_current_folder_with_disk()
        self.pdf_listbox.config(state=tk.NORMAL)  # Atļaujam rediģēt, lai varētu ievietot tekstu
        self.pdf_listbox.delete("1.0", tk.END)  # Dzēšam visu tekstu

        if not hasattr(self, 'current_folder') or "contents" not in self.current_folder:
            print("Kļūda: current_folder nav pareizi inicializēts vai tam trūkst 'contents'.")
            self.pdf_listbox.config(state=tk.DISABLED)  # Atkal atspējojam rediģēšanu
            return

        # Šķirojam saturu: vispirms mapes, tad faili, pēc tam alfabētiski
        # Jaunizveidotajā mapē (pēc split_pdf_to_pages) oriģinālais fails jau būs pirmais
        # un lapas sekos, tāpēc šeit papildu šķirošana nav nepieciešama, ja vienumi jau ir pareizā secībā.
        # Ja vēlaties stingri nodrošināt oriģinālā faila prioritāti, varat to darīt šeit.
        # Piemēram, atdalīt oriģinālo failu, šķirot pārējos un tad salikt kopā.
        # Šobrīd pieņemam, ka `split_pdf_to_pages` jau sakārtoja `new_folder_node["contents"]`.
        sorted_contents = self.current_folder["contents"]  # Vairs nav nepieciešama papildu šķirošana šeit

        # Saglabājam sarakstu ar rādāmajiem vienumiem, lai varētu tos identificēt vēlāk
        self._displayed_items = []

        for i, item in enumerate(sorted_contents):
            display_text = ""
            if item["type"] == "file":
                # Pārbaudām, vai ir "display_name" (lapām) vai izmantojam "name"
                name_to_display = item.get("display_name", item['name'])
                display_text = f"{i + 1}. 📄 {name_to_display} ({item['date']})\n"
            elif item["type"] == "folder":
                display_text = f"{i + 1}. 📁 {item['name']}\n"

            self.pdf_listbox.insert(tk.END, display_text, "normal")  # Ievietojam tekstu ar noklusējuma tagu
            self._displayed_items.append(item)  # Pievienojam vienumu sarakstam

        self.pdf_listbox.config(state=tk.DISABLED)  # Atkal atspējojam rediģēšanu
        self.update_path_label()
        self.update_back_button_state()
        self.save_pdf_archive()  # Saglabā izmaiņas failu sistēmā

        # Pēc atsvaidzināšanas pielietojam filtrus, ja tādi ir
        self.filter_pdf_list()

    def on_text_double_click(self, event):
        """Apstrādā dubultklikšķi uz tk.Text logrīka, lai atvērtu vienumu."""
        # Izmanto iepriekš saglabāto atlases indeksu
        if hasattr(self, '_selected_line_index') and self._selected_line_index != -1:
            line_number = self._selected_line_index
            if 0 <= line_number < len(self._displayed_items):
                selected_item = self._displayed_items[line_number]
                self.open_selected_item(selected_item)
        else:
            # Ja nav iepriekšējas atlases, mēģina iegūt no klikšķa pozīcijas
            index = self.pdf_listbox.index(f"@{event.x},{event.y}")
            line_number = int(index.split(".")[0]) - 1  # Rindas numurs (0-bāzēts)
            if 0 <= line_number < len(self._displayed_items):
                selected_item = self._displayed_items[line_number]
                self.open_selected_item(selected_item)

    def update_path_label(self):
        """Atjaunina ceļa etiķeti, lai parādītu pašreizējo mapes ceļu."""
        path_parts = []
        temp_folder = self.current_folder
        # Pārliecināmies, ka mēs neejam tālāk par saknes mapi
        while temp_folder != self.internal_file_system and temp_folder.get("parent") is not None:
            path_parts.insert(0, temp_folder["name"])
            temp_folder = temp_folder["parent"]
        self.current_path_label.config(text="/".join([""] + path_parts) if path_parts else "/")

    def update_back_button_state(self):
        """Atjaunina pogas 'Atpakaļ' stāvokli."""
        if self.current_folder == self.internal_file_system:
            self.back_button.config(state=DISABLED)
        else:
            self.back_button.config(state=NORMAL)

    def go_back_folder(self):
        """Atgriežas uz iepriekšējo mapi."""
        if self.current_folder.get("parent"):
            self.current_folder = self.current_folder["parent"]
            self.refresh_pdf_list()

    def open_selected_item(self, item_to_open=None):
        """
        Atver atlasīto vienumu (failu vai mapi).
        Ja item_to_open ir None, tad ņem no iepriekš saglabātās atlases.
        """
        selected_item = None
        if item_to_open is None:
            if hasattr(self, '_selected_line_index') and self._selected_line_index != -1:
                line_number = self._selected_line_index
                if 0 <= line_number < len(self._displayed_items):
                    selected_item = self._displayed_items[line_number]

            if selected_item is None:
                messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vienumu no saraksta.")
                return
        else:
            selected_item = item_to_open  # Izmantojam padoto vienumu

        # Šis ir galvenais loģikas bloks, kas nosaka, vai atvērt failu vai navigēt uz mapi
        if selected_item["type"] == "file":
            filepath = selected_item['filepath']
            if os.path.exists(filepath):
                try:
                    os.startfile(filepath)  # Atver failu ar noklusējuma programmu
                except Exception as e:
                    messagebox.showerror("Kļūda", f"Neizdevās atvērt failu:\n{e}")
            else:
                messagebox.showwarning("Fails nav atrasts",
                                       "Fails nav atrasts norādītajā vietā. Iespējams, tas ir pārvietots vai dzēsts.")
        elif selected_item["type"] == "folder":
            # JA IR MAPE, TAD NAVIGĒ UZ TO PROGRAMMĀ
            self.current_folder = selected_item
            self.refresh_pdf_list()
            print(f"Navigēts uz mapi: {selected_item.get('name', 'Nezināma mape')}")
        else:
            messagebox.showwarning("Kļūda", "Nezināms vienuma tips.")

    def on_pdf_select(self, event=None):
        """Apstrādā PDF faila atlasi tk.Text logrīkā, ielādējot priekšskatījumu."""
        # Notīrām iepriekšējo atlasi, ja tāda bija
        self.pdf_listbox.tag_remove("sel", "1.0", tk.END)

        try:
            # Iegūstam pašreizējo atlasi
            selection_start = self.pdf_listbox.index(tk.SEL_FIRST)
            selection_end = self.pdf_listbox.index(tk.SEL_LAST)

            # Iegūstam atlasītās rindas numuru
            line_number = int(selection_start.split(".")[0]) - 1

            if 0 <= line_number < len(self._displayed_items):
                selected_item = self._displayed_items[line_number]

                # Pielietojam "sel" tagu atlasītajai rindai
                self.pdf_listbox.tag_add("sel", f"{line_number + 1}.0", f"{line_number + 1}.end")

                if selected_item["type"] == "file" and selected_item["name"].lower().endswith(".pdf"):
                    filepath = selected_item['filepath']
                    if os.path.exists(filepath):
                        self._load_pdf_for_preview(filepath)
                    else:
                        messagebox.showwarning("Fails nav atrasts", "Atlasītais PDF fails nav atrasts diskā.")
                        self._clear_pdf_preview()
                else:
                    self._clear_pdf_preview()  # Notīra priekšskatījumu, ja atlasīts nav PDF fails
            else:
                self._clear_pdf_preview()  # Notīra priekšskatījumu, ja nekas nav atlasīts vai atlase ir ārpus robežām
        except tk.TclError:
            # Nav aktīvas atlases, vai atlase ir tukša
            self._clear_pdf_preview()

    def open_pdf_location(self):
        """Atver mapes atrašanās vietu, kurā atrodas atlasītais PDF fails (sistēmā)."""
        selection = self.pdf_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet failu no saraksta.")
            return

        index = selection[0]
        selected_item = self.current_folder["contents"][index]

        if selected_item['type'] == 'file':
            self.open_pdf_file_by_path(selected_item['filepath'])
            if os.path.exists(filepath):
                try:
                    # Atver mapi un iezīmē failu
                    if os.name == 'nt':  # Windows
                        os.startfile(os.path.dirname(filepath))
                    elif os.name == 'posix':  # macOS, Linux
                        import sys
                        if sys.platform == 'darwin':  # macOS
                            subprocess.Popen(['open', '-R', filepath])
                        else:  # Linux
                            subprocess.Popen(['xdg-open', os.path.dirname(filepath)])
                except Exception as e:
                    messagebox.showerror("Kļūda", f"Neizdevās atvērt faila atrašanās vietu:\n{e}")
            else:
                messagebox.showwarning("Fails nav atrasts", "Fails nav atrasts norādītajā vietā.")
        else:
            messagebox.showwarning("Nav fails", "Atlasītais vienums nav fails.")

    def delete_selected_item(self):
        """Dzēš atlasītos vienumus (failus vai mapes) no iekšējās failu sistēmas un fiziski no diska."""
        # Izmanto iepriekš saglabāto atlases indeksu
        if not hasattr(self, '_selected_line_index') or self._selected_line_index == -1:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vienumu(s), ko dzēst.")
            return

        line_number = self._selected_line_index

        if not (0 <= line_number < len(self._displayed_items)):
            messagebox.showwarning("Nav atlasīts", "Nederīga atlase.")
            return

        item_to_delete = self._displayed_items[line_number]
        physical_path = self._get_physical_path_from_node(item_to_delete)

        confirm_msg = f"Vai tiešām vēlaties dzēst '{item_to_delete['name']}'?\n"
        confirm_msg += "Šī darbība neatgriezeniski dzēsīs failu/mapi arī no diska!"

        if not messagebox.askyesno("Dzēst vienumu", confirm_msg):
            return

        deleted_successfully = False
        if item_to_delete["type"] == "file":
            try:
                if os.path.exists(physical_path):
                    os.remove(physical_path)
                    print(f"Fiziski dzēsts fails: {physical_path}")
                else:
                    print(f"Fails neeksistē fiziski, dzēš tikai no programmas: {physical_path}")
                # Noņemam vienumu no current_folder["contents"]
                if item_to_delete in self.current_folder["contents"]:
                    self.current_folder["contents"].remove(item_to_delete)
                deleted_successfully = True
            except OSError as e:
                messagebox.showerror("Dzēšanas kļūda", f"Neizdevās dzēst failu {item_to_delete['name']}:\n{e}")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neparedzēta kļūda dzēšot failu {item_to_delete['name']}:\n{e}")
        elif item_to_delete["type"] == "folder":
            try:
                if os.path.exists(physical_path):
                    import shutil
                    shutil.rmtree(physical_path)
                    print(f"Fiziski dzēsta mape: {physical_path}")
                else:
                    print(f"Mape neeksistē fiziski, dzēš tikai no programmas: {physical_path}")
                # Noņemam vienumu no current_folder["contents"]
                if item_to_delete in self.current_folder["contents"]:
                    self.current_folder["contents"].remove(item_to_delete)
                deleted_successfully = True
            except OSError as e:
                messagebox.showerror("Dzēšanas kļūda", f"Neizdevās dzēst mapi {item_to_delete['name']}:\n{e}")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neparedzēta kļūda dzēšot mapi {item_to_delete['name']}:\n{e}")

        if deleted_successfully:
            self.refresh_pdf_list()
            messagebox.showinfo("Dzēsts", f"Vienums '{item_to_delete['name']}' veiksmīgi dzēsts.")
        else:
            messagebox.showinfo("Dzēšana", "Vienums netika dzēsts.")

    def _delete_folder_contents_from_disk(self, folder_node):
        """Rekursīvi dzēš mapes saturu no diska."""
        success = True
        for item in folder_node["contents"]:
            if item["type"] == "file":
                try:
                    if os.path.exists(item["filepath"]):
                        os.remove(item["filepath"])
                except Exception as e:
                    messagebox.showwarning("Dzēšanas kļūda", f"Neizdevās dzēst failu {item['name']} no diska: {e}")
                    success = False
            elif item["type"] == "folder":
                if not self._delete_folder_contents_from_disk(item):
                    success = False
        # Pēc satura dzēšanas mēģina dzēst pašu mapi, ja tā ir tukša
        try:
            # Pārbauda, vai mape ir tukša pirms dzēšanas
            if os.path.exists(os.path.join(self.default_save_path, folder_node["name"])):
                if not os.listdir(os.path.join(self.default_save_path, folder_node["name"])):
                    os.rmdir(os.path.join(self.default_save_path, folder_node["name"]))
        except Exception as e:
            messagebox.showwarning("Dzēšanas kļūda", f"Neizdevās dzēst mapi {folder_node['name']} no diska: {e}")
            success = False
        return success

    def send_selected_pdfs_by_email(self):
        """Nosūta atlasītos PDF failus, izmantojot SMTP iestatījumus."""
        # Iegūstam atlasītās rindas numuru
        try:
            selection_start = self.pdf_listbox.index(tk.SEL_FIRST)
            line_number = int(selection_start.split(".")[0]) - 1
        except tk.TclError:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet PDF failu(s), ko nosūtīt e-pastā.")
            return

        if not (0 <= line_number < len(self._displayed_items)):
            messagebox.showwarning("Nav atlasīts", "Nederīga atlase.")
            return

        selected_item = self._displayed_items[line_number]

        selected_filepaths = []
        if selected_item["type"] == "file":
            filepath = selected_item['filepath']
            if os.path.exists(filepath):
                selected_filepaths.append(filepath)
            else:
                messagebox.showwarning("Fails nav atrasts",
                                       f"Fails '{os.path.basename(filepath)}' nav atrasts un netiks pievienots.")
        else:
            messagebox.showwarning("Nav fails", f"Vienums '{selected_item['name']}' nav fails un netiks pievienots.")

        if not selected_filepaths:
            messagebox.showwarning("Nav failu", "Neviens derīgs fails nav atlasīts sūtīšanai.")
            return

        # Dialogs e-pasta adresāta izvēlei
        to_email_dialog = Toplevel(self)
        to_email_dialog.title("Nosūtīt e-pastu")
        to_email_dialog.transient(self)
        to_email_dialog.grab_set()

        ttk.Label(to_email_dialog, text="Nosūtīt uz:").pack(padx=10, pady=5)

        default_to_email = self.settings.get("to_email", "")
        to_email_var_dialog = tk.StringVar(value=default_to_email)

        ttk.Entry(to_email_dialog, textvariable=to_email_var_dialog, width=50).pack(padx=10, pady=5)

        def confirm_send():
            to_address = to_email_var_dialog.get()
            if not to_address:
                messagebox.showwarning("Trūkst adresāta", "Lūdzu, ievadiet e-pasta adresātu.")
                return
            to_email_dialog.destroy()
            self._send_email_with_attachments(selected_filepaths, to_address)

        ttk.Button(to_email_dialog, text="Nosūtīt", command=confirm_send, bootstyle=PRIMARY).pack(pady=10)
        ttk.Button(to_email_dialog, text="Atcelt", command=to_email_dialog.destroy, bootstyle=SECONDARY).pack(pady=5)
    def on_select(event):
        global selected_item  # Padara mainīgo globālu
        selection = event.widget.curselection()
        if selection:
            index = selection[0]
            selected_item = event.widget.get(index)  # Iegūst izvēlēto elementu

    def _send_email_with_attachments(self, filepaths, to_address):
        """Faktiski nosūta e-pastu ar pielikumiem."""
        smtp_server = self.settings.get("smtp_server")
        smtp_port = self.settings.get("smtp_port")
        email_user = self.settings.get("email_user")
        email_pass = self.settings.get("email_pass")
        from_email = self.settings.get("from_email")
        use_ssl = self.settings.get("use_ssl", True)
        email_subject = self.settings.get("email_subject", "OCR PDF dokumenti")
        email_body_plain = self.settings.get("email_body_plain",
                                             "Sveiki,\n\nPielikumā atradīsiet OCR apstrādātos PDF dokumentus.\n\nAr cieņu,\nJūsu OCR PDF App")
        email_body_html = self.settings.get("email_body_html",
                                            "<html><body><p>Sveiki,</p><p>Pielikumā atradīsiet OCR apstrādātos PDF dokumentus.</p><p>Ar cieņu,<br/>Jūsu OCR PDF App</p></body></html>")

        if not all([smtp_server, smtp_port, email_user, email_pass, from_email, to_address]):
            messagebox.showwarning("E-pasta iestatījumi",
                                   "Lūdzu, konfigurējiet e-pasta iestatījumus (SMTP serveris, ports, lietotājvārds, parole, sūtītāja un saņēmēja adreses) Iestatījumu logā.")
            return

        msg = MIMEMultipart('alternative')
        msg['From'] = from_email
        msg['To'] = to_address
        msg['Subject'] = email_subject

        part1 = MIMEText(email_body_plain, 'plain')
        msg.attach(part1)

        part2 = MIMEText(email_body_html, 'html')
        msg.attach(part2)

        for filepath in filepaths:
            try:
                filename = os.path.basename(filepath)
                with open(filepath, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                # Labojums: Pareizi kodē faila nosaukumu Content-Disposition galvenē
                part.add_header("Content-Disposition", f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}")
                msg.attach(part)
            except Exception as e:
                messagebox.showwarning("Pielikuma kļūda", f"Neizdevās pievienot failu {filename}: {e}")
                continue

        try:
            if use_ssl:
                server = smtplib.SMTP_SSL(smtp_server, smtp_port)
            else:
                server = smtplib.SMTP(smtp_server, smtp_port)
                server.starttls()

            server.login(email_user, email_pass)
            text = msg.as_string()
            server.sendmail(from_email, to_address, text)
            server.quit()
            messagebox.showinfo("E-pasts", "E-pasts veiksmīgi nosūtīts!")
        except Exception as e:
            messagebox.showerror("E-pasta sūtīšanas kļūda", f"Neizdevās nosūtīt e-pastu:\n{e}")

    def show_pdf_context_menu(self, event):
        """Parāda konteksta izvēlni PDF saraksta elementiem."""
        try:
            # Pārliecināmies, ka tiek atlasīts elements, uz kura uzklikšķināts
            # Iegūst klikšķa pozīciju rindas formātā (1-bāzēts)
            index = self.pdf_listbox.index(f"@{event.x},{event.y}")
            line_number = int(index.split(".")[0]) - 1  # 0-bāzēts rindas numurs

            # Noņem iepriekšējo atlasi
            self.pdf_listbox.tag_remove("selected_line", "1.0", tk.END)
            self.pdf_listbox.tag_remove("sel", "1.0", tk.END)

            # Pievieno atlasi atlasītajai rindai
            start_index = f"{line_number + 1}.0"
            end_index = f"{line_number + 1}.end"
            self.pdf_listbox.tag_add("selected_line", start_index, end_index)

            # Saglabā atlasīto indeksu
            self._selected_line_index = line_number

            # Iegūst atlasīto vienumu
            if 0 <= line_number < len(self._displayed_items):
                selected_item = self._displayed_items[line_number]
            else:
                selected_item = None

            context_menu = tk.Menu(self.pdf_listbox, tearoff=0)
            context_menu.add_command(label="Atvērt", command=lambda: self.open_pdf_file_by_path(selected_item['filepath']))
            context_menu.add_command(label="Atvērt mapē (sistēmā)", command=lambda: self.open_pdf_location())
            context_menu.add_command(label="Nosūtīt e-pastā", command=lambda: self.send_selected_pdfs_by_email())
            context_menu.add_separator()
            context_menu.add_command(label="Dzēst", command=lambda: self.delete_selected_item())
            context_menu.add_command(label="Pārdēvēt", command=lambda: self.rename_selected_item())
            context_menu.add_command(label="Pārvietot uz...", command=lambda: self.move_selected_items())
            context_menu.add_command(label="Saglabāt kā Word", command=lambda: self.save_as_word())
            context_menu.add_command(label="Pievienot paroli",
                                     command=lambda: self.add_password_to_pdf(selected_item['filepath']))
            context_menu.add_command(label="Noņemt paroli",
                                     command=lambda: self.remove_password_from_pdf(selected_item['filepath']))
            context_menu.add_command(label="Mainīt paroli",
                                     command=lambda: self.change_password_of_pdf(selected_item['filepath']))
            context_menu.add_separator()
            context_menu.add_command(label="Sadalīt PDF pa lapām",
                                     command=lambda: self.split_pdf_to_pages(selected_item['filepath']) if selected_item.get('type') == 'file' else None)

            context_menu.post(event.x_root, event.y_root)
        except Exception:
            pass  # Ja nav atlasīts nekas, ignorē

    def create_new_folder_internal(self):
        """Izveido jaunu mapi iekšējā failu sistēmā un fiziski diskā."""
        new_folder_name = simpledialog.askstring("Jauna mape", "Ievadiet jaunās mapes nosaukumu:", parent=self)
        if not new_folder_name:
            return

        # Pārbauda, vai mape ar šādu nosaukumu jau eksistē iekšējā struktūrā
        for item in self.current_folder["contents"]:
            if item["type"] == "folder" and item["name"] == new_folder_name:
                messagebox.showwarning("Mape jau eksistē",
                                       f"Mape ar nosaukumu '{new_folder_name}' jau eksistē šajā mapē.")
                return

        # Izveido fizisko ceļu jaunajai mapei
        current_physical_path = self._get_physical_path_from_node(self.current_folder)
        new_physical_folder_path = os.path.join(current_physical_path, new_folder_name)

        try:
            # Izveido fizisko mapi
            os.makedirs(new_physical_folder_path, exist_ok=True)
            print(f"Fiziski izveidota mape: {new_physical_folder_path}")

            # Ja fiziskā mape izveidota veiksmīgi, pievieno to iekšējai struktūrai
            new_folder = {"type": "folder", "name": new_folder_name, "contents": [], "parent": self.current_folder}
            self.current_folder["contents"].append(new_folder)
            self.refresh_pdf_list()
            messagebox.showinfo("Mape izveidota", f"Mape '{new_folder_name}' veiksmīgi izveidota.")
        except OSError as e:
            messagebox.showerror("Mapes izveides kļūda", f"Neizdevās izveidot mapi '{new_folder_name}':\n{e}")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neparedzēta kļūda veidojot mapi '{new_folder_name}':\n{e}")

    def move_selected_items(self):
        index = getattr(self, '_selected_line_index', -1)
        if index == -1:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vienumus.")
            return

        # Izveido mapju izvēles dialogu
        target_folder = self._select_folder_dialog(self.internal_file_system)

        if not target_folder:
            messagebox.showinfo("Pārvietošana", "Mērķa mape netika izvēlēta.")
            return

        # Pārbauda, vai mērķa mape nav pati pašreizējā mape
        if target_folder == self.current_folder:
            messagebox.showinfo("Pārvietošana", "Vienumi jau atrodas izvēlētajā mapē.")
            return

        moved_count = 0
        # Jāveido kopija, jo saraksts mainīsies dzēšot elementus
        # Atlasītie indeksi dilstošā secībā, lai pop() neietekmētu nākamos indeksus
        items_to_move_with_indices = sorted([(i, self.current_folder["contents"][i]) for i in selection], reverse=True)

        for original_index, item in items_to_move_with_indices:
            # Pārbauda, vai mērķa mape nav pati vienums vai tās apakšmape
            if item["type"] == "folder" and self._is_descendant(target_folder, item):
                messagebox.showwarning("Kļūda", f"Mapi '{item['name']}' nevar pārvietot uz tās paša apakšmapi.")
                continue

            # Pārbauda, vai mērķa mapē jau nav vienums ar tādu pašu nosaukumu
            # Ja ir, piedāvā pārdēvēt
            new_name = item["name"]
            name_exists = False
            for existing_item in target_folder["contents"]:
                if existing_item["name"] == new_name and existing_item["type"] == item["type"]:
                    name_exists = True
                    break

            if name_exists:
                response = messagebox.askyesno("Nosaukums jau eksistē",
                                               f"Mērķa mapē jau eksistē vienums ar nosaukumu '{new_name}'. Vai vēlaties to pārdēvēt?")
                if response:
                    temp_new_name = simpledialog.askstring("Pārdēvēt",
                                                           f"Ievadiet jauno nosaukumu vienumam '{new_name}':",
                                                           parent=self, initialvalue=new_name)
                    if temp_new_name:
                        new_name = temp_new_name
                    else:
                        continue  # Atcelt pārvietošanu šim vienumam
                else:
                    continue  # Atcelt pārvietošanu šim vienumam

            # Iegūst fiziskos ceļus
            old_physical_path = self._get_physical_path_from_node(item)
            target_physical_path = self._get_physical_path_from_node(target_folder)
            new_physical_path = os.path.join(target_physical_path, new_name)

            try:
                # Pārvieto fizisko failu/mapi
                if os.path.exists(old_physical_path):
                    os.rename(old_physical_path, new_physical_path)
                    print(f"Fiziski pārvietots: {old_physical_path} -> {new_physical_path}")
                else:
                    print(f"Brīdinājums: Fiziskais fails/mape neeksistē: {old_physical_path}. Pārvieto tikai iekšēji.")

                # Ja fiziskā pārvietošana veiksmīga, atjaunina iekšējo struktūru
                self.current_folder["contents"].pop(original_index)  # Izmanto original_index
                item["name"] = new_name  # Atjaunina nosaukumu, ja tas tika mainīts
                item["filepath"] = new_physical_path  # Atjaunina filepath failiem
                target_folder["contents"].append(item)
                item["parent"] = target_folder  # Atjaunina vecāka atsauci
                moved_count += 1

            except OSError as e:
                messagebox.showerror("Pārvietošanas kļūda", f"Neizdevās pārvietot '{item['name']}':\n{e}")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neparedzēta kļūda pārvietojot '{item['name']}':\n{e}")

        if moved_count > 0:
            self.refresh_pdf_list()
            messagebox.showinfo("Pārvietots", f"Veiksmīgi pārvietoti {moved_count} vienumi.")
        else:
            messagebox.showinfo("Pārvietošana", "Neviens vienums netika pārvietots.")



    def _select_folder_dialog(self, root_folder):
        """Atver dialogu mapes izvēlei iekšējā failu sistēmā."""
        dialog = Toplevel(self)
        dialog.title("Izvēlēties mapi")
        dialog.transient(self)
        dialog.grab_set()

        tree = ttk.Treeview(dialog, selectmode="browse")
        tree.pack(fill="both", expand=True, padx=10, pady=10)

        selected_folder = None

        def on_select():
            nonlocal selected_folder
            item_id = tree.selection()
            if item_id:
                # Piekļūst mapes objektam, kas saglabāts kā vērtība
                # tree.item(item_id, "values") atgriež tuple, kurā pirmais elements ir mūsu saglabātais dict
                item_data = tree.item(item_id, "values")
                if item_data and isinstance(item_data[0], dict) and item_data[0].get("type") == "folder":
                    selected_folder = item_data[0]
                    dialog.destroy()
                else:
                    messagebox.showwarning("Nederīga atlase", "Lūdzu, atlasiet mapi, nevis failu.")
            else:
                messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet mapi.")

        ttk.Button(dialog, text="Izvēlēties", command=on_select).pack(pady=5)
        ttk.Button(dialog, text="Atcelt", command=dialog.destroy).pack(pady=5)

        def populate_treeview_with_data(treeview, parent_node_id, folder_data):
            # Pievieno mapes objektu kā vērtību, lai to varētu atgūt on_select
            # Pārbaudām, vai tas nav saknes mezgls, lai nerādītu "Sakne" kā izvēles opciju
            if folder_data["name"] == "Sakne" and parent_node_id == "":
                node_id = treeview.insert(parent_node_id, "end", text="Sakne (Pašreizējā mape)", open=True,
                                          tags=("folder",),
                                          values=(folder_data,))
            else:
                node_id = treeview.insert(parent_node_id, "end", text=folder_data["name"], open=False, tags=("folder",),
                                          values=(folder_data,))  # Šeit saglabājam visu mapes dict kā vērtību

            for item in folder_data["contents"]:
                if item["type"] == "folder":
                    # Rekursīvi izsaucam funkciju katrai apakšmapei
                    populate_treeview_with_data(treeview, node_id, item)
                # Failus nepievienojam, jo dialogs ir paredzēts tikai mapju izvēlei

        # Ievieto saknes mapi un aizpilda to
        populate_treeview_with_data(tree, "", self.internal_file_system)

        self.wait_window(dialog)
        return selected_folder

    def _is_descendant(self, potential_descendant, potential_ancestor):
        """Pārbauda, vai potential_descendant ir potential_ancestor apakšmape."""
        current = potential_descendant
        while current.get("parent"):
            if current["parent"] == potential_ancestor:
                return True
            current = current["parent"]
        return False

    def rename_selected_item(self):
        """Pārdēvē atlasīto vienumu (failu vai mapi) iekšējā failu sistēmā un fiziski diskā."""
        selection = self.pdf_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet vienumu, ko pārdēvēt.")
            return

        index = selection[0]
        item = self.current_folder["contents"][index]

        old_name = item["name"]
        new_name = simpledialog.askstring("Pārdēvēt", f"Ievadiet jauno nosaukumu vienumam '{old_name}':",
                                          initialvalue=old_name, parent=self)

        if not new_name or new_name == old_name:
            return  # Lietotājs atcēla vai nosaukums nav mainīts

        # Pārbauda, vai jaunais nosaukums jau eksistē iekšējā struktūrā
        for existing_item in self.current_folder["contents"]:
            if existing_item["name"] == new_name and existing_item["type"] == item["type"]:
                messagebox.showwarning("Nosaukums jau eksistē",
                                       f"Vienums ar nosaukumu '{new_name}' jau eksistē šajā mapē.")
                return

        # Iegūst fiziskos ceļus
        current_physical_path = self._get_physical_path_from_node(self.current_folder)
        old_physical_path = os.path.join(current_physical_path, old_name)
        new_physical_path = os.path.join(current_physical_path, new_name)

        try:
            # Pārdēvē fizisko failu/mapi
            if os.path.exists(old_physical_path):
                os.rename(old_physical_path, new_physical_path)
                print(f"Fiziski pārdēvēts: {old_physical_path} -> {new_physical_path}")
            else:
                print(f"Brīdinājums: Fiziskais fails/mape neeksistē: {old_physical_path}. Pārdēvē tikai iekšēji.")

            # Ja fiziskā pārdēvēšana veiksmīga, atjaunina iekšējo struktūru
            item["name"] = new_name
            if item["type"] == "file":
                item["filepath"] = new_physical_path  # Atjaunina filepath failiem

            self.refresh_pdf_list()
            messagebox.showinfo("Pārdēvēts", f"Vienums veiksmīgi pārdēvēts uz '{new_name}'.")

        except OSError as e:
            messagebox.showerror("Pārdēvēšanas kļūda", f"Neizdevās pārdēvēt '{old_name}' uz '{new_name}':\n{e}")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neparedzēta kļūda pārdēvējot '{old_name}':\n{e}")

    from docx import Document
    from docx.shared import Inches
    from pdf2image import convert_from_path
    import tempfile
    import os
    from tkinter import filedialog, messagebox

    from docx import Document
    from docx.shared import Inches
    from pdf2image import convert_from_path
    import tempfile
    import os
    from tkinter import filedialog, messagebox

    def save_as_word(self):
        index = getattr(self, '_selected_line_index', -1)
        if index == -1:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet PDF failu.")
            return

        selected_item = self._displayed_items[index]
        if selected_item["type"] != "file" or not selected_item["name"].lower().endswith(".pdf"):
            messagebox.showwarning("Nepareizs fails", "Lūdzu, atlasiet PDF failu.")
            return

        pdf_path = selected_item["filepath"]

        # Mēģinām iegūt OCR tekstu, ja nav, tad None
        ocr_text = getattr(self, 'ocr_text', None)

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word dokumenti", "*.docx")],
            title="Saglabāt kā Word dokumentu"
        )
        if not save_path:
            return  # Lietotājs atcēla

        try:
            doc = Document()

            if ocr_text and ocr_text.strip():
                # Ja OCR teksts ir pieejams un nav tukšs, pievienojam to
                doc.add_paragraph(ocr_text)
            else:
                # Ja nav OCR teksta, var pievienot info vai atstāt tukšu
                doc.add_paragraph("[Nav pieejams OCR teksts]")

            # Konvertējam PDF pirmo lapu uz attēlu
            images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=200)

            if images:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img_file:
                    images[0].save(tmp_img_file.name, "PNG")
                    tmp_img_path = tmp_img_file.name

                doc.add_picture(tmp_img_path, width=Inches(6))

                os.remove(tmp_img_path)
            else:
                messagebox.showwarning("Brīdinājums", "Neizdevās konvertēt PDF lapu uz attēlu.")

            doc.save(save_path)
            messagebox.showinfo("Veiksmīgi", f"Dokuments saglabāts: {save_path}")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās saglabāt Word dokumentu:\n{e}")
    def drag_start(self, event):
        """Sāk vilkšanas operāciju `pdf_listbox`."""
        # Iegūst elementa indeksu, uz kura tika uzklikšķināts
        index = self.pdf_listbox.nearest(event.y)
        if index != -1:
            self.drag_data["item_index"] = index
            self.drag_data["start_y"] = event.y
            # Pārliecinās, ka elements ir atlasīts, lai vizuāli atspoguļotu vilkšanu
            self.pdf_listbox.selection_clear(0, tk.END)
            self.pdf_listbox.selection_set(index)
            self.pdf_listbox.activate(index)

    def drag_motion(self, event):
        """Apstrādā vilkšanas kustību `pdf_listbox`."""
        if self.drag_data["item_index"] is not None:
            # Iegūst jauno pozīciju
            new_index = self.pdf_listbox.nearest(event.y)
            current_index = self.drag_data["item_index"]

            if new_index != current_index:
                # Pārvieto elementu pamatā esošajā datu struktūrā
                item_to_move = self.current_folder["contents"].pop(current_index)
                self.current_folder["contents"].insert(new_index, item_to_move)

                # Atjaunina vilkšanas datus ar jauno indeksu
                self.drag_data["item_index"] = new_index

                # Atjauno listbox vizuālo attēlojumu
                self.refresh_pdf_list()
                # Pārliecinās, ka pārvietotais elements joprojām ir atlasīts
                self.pdf_listbox.selection_set(new_index)
                self.pdf_listbox.activate(new_index)

    def drag_drop(self, event):
        """Apstrādā nomešanas operāciju."""
        self.pdf_listbox.config(cursor="arrow")  # Atjauno kursoru
        if self.drag_data["item"]:
            target_index = self.pdf_listbox.nearest(event.y)
            if target_index is not None:
                target_item = self.current_folder["contents"][target_index]

                if target_item["type"] == "folder":
                    # Mēģina pārvietot uz mapi
                    if self.drag_data["item"] == target_item:
                        messagebox.showwarning("Pārvietošana", "Nevar pārvietot vienumu uz pašu sevi.")
                        self.drag_data["item"] = None
                        return

                    if self.drag_data["item"]["type"] == "folder" and self._is_descendant(target_item,
                                                                                          self.drag_data["item"]):
                        messagebox.showwarning("Kļūda", "Mapi nevar pārvietot uz tās paša apakšmapi.")
                        self.drag_data["item"] = None
                        return

                    # Pārbauda, vai mērķa mapē jau nav vienums ar tādu pašu nosaukumu
                    name_exists = False
                    for existing_item in target_item["contents"]:
                        if existing_item["name"] == self.drag_data["item"]["name"] and existing_item["type"] == \
                                self.drag_data["item"]["type"]:
                            name_exists = True
                            break

                    if name_exists:
                        response = messagebox.askyesno("Nosaukums jau eksistē",
                                                       f"Mērķa mapē jau eksistē vienums ar nosaukumu '{self.drag_data['item']['name']}'. Vai vēlaties to pārdēvēt?")
                        if response:
                            new_name = simpledialog.askstring("Pārdēvēt",
                                                              f"Ievadiet jauno nosaukumu vienumam '{self.drag_data['item']['name']}':",
                                                              parent=self)
                            if new_name:
                                self.drag_data["item"]["name"] = new_name
                            else:
                                self.drag_data["item"] = None  # Atcelt pārvietošanu
                                return
                        else:
                            self.drag_data["item"] = None  # Atcelt pārvietošanu
                            return

                    # Noņem no pašreizējās mapes
                    self.current_folder["contents"].pop(self.drag_data["index"])
                    # Pievieno mērķa mapei
                    target_item["contents"].append(self.drag_data["item"])
                    self.drag_data["item"]["parent"] = target_item  # Atjaunina vecāka atsauci
                    messagebox.showinfo("Pārvietots",
                                        f"Vienums '{self.drag_data['item']['name']}' pārvietots uz '{target_item['name']}'.")
                    self.refresh_pdf_list()

                elif target_item["type"] == "file":
                    # Mēģina pārvietot starp failiem (pārvietošana tajā pašā mapē)
                    if self.drag_data["item"]["type"] == "file":
                        # Pārvieto failu sarakstā
                        current_item = self.current_folder["contents"].pop(self.drag_data["index"])
                        self.current_folder["contents"].insert(target_index, current_item)
                        messagebox.showinfo("Pārvietots", f"Vienums '{current_item['name']}' pārvietots sarakstā.")
                        self.refresh_pdf_list()
                    else:
                        messagebox.showwarning("Pārvietošana", "Mapi nevar pārvietot uz failu.")
            else:
                # Nomešana tukšā vietā vai ārpus saraksta (varētu nozīmēt pārvietošanu uz pašreizējo mapi)
                # Šis scenārijs jau ir apstrādāts, ja vienums tiek nomests uz mapi
                pass
            self.drag_data["item"] = None  # Notīra vilkšanas datus

    def configure_grid(self):
        """Konfigurē galvenā loga režģa izkārtojumu (pašlaik netiek izmantots, jo ir Notebook)."""
        # Galvenais logs tagad izmanto Notebook, tāpēc rindu/kolonnu konfigurācija ir mazāk svarīga šeit
        # Tā vietā, katra cilne konfigurē savu iekšējo izkārtojumu
        pass

    def _show_folder_selection_dialog(self, suggested_category):
        """
        Parāda dialoga logu, lai lietotājs varētu izvēlēties vai izveidot mapi
        pirms PDF saglabāšanas.
        """
        dialog = Toplevel(self)
        dialog.title("Izvēlēties saglabāšanas mapi")
        dialog.transient(self)
        dialog.grab_set()

        selected_node = None  # Šeit tiks saglabāts izvēlētais mapes mezgls

        # Pašreizējā ceļa attēlošana
        current_path_label = ttk.Label(dialog, text="Pašreizējā mape: /")
        current_path_label.pack(fill="x", padx=10, pady=5)

        # Mapju saraksts
        folder_list_frame = ttk.Frame(dialog)
        folder_list_frame.pack(fill="both", expand=True, padx=10, pady=5)

        folder_listbox = tk.Listbox(folder_list_frame, selectmode=tk.SINGLE, exportselection=False)
        folder_listbox.pack(side=tk.LEFT, fill="both", expand=True)

        folder_scrollbar = ttk.Scrollbar(folder_list_frame, orient="vertical", command=folder_listbox.yview)
        folder_scrollbar.pack(side=tk.RIGHT, fill="y")
        folder_listbox.config(yscrollcommand=folder_scrollbar.set)

        # Iekšējā funkcija, lai atjauninātu sarakstu
        current_dialog_folder = self.internal_file_system  # Sāk ar saknes mapi dialogā

        def update_folder_listbox():
            folder_listbox.delete(0, tk.END)
            # Pievieno ".." opciju, ja nav saknes mapē
            if current_dialog_folder != self.internal_file_system:
                folder_listbox.insert(tk.END, "📁 .. (Atpakaļ)")

            # Pievieno mapes
            for item in current_dialog_folder["contents"]:
                if item["type"] == "folder":
                    folder_listbox.insert(tk.END, f"📁 {item['name']}")

            # Atjaunina ceļa etiķeti
            path_parts = []
            temp_folder = current_dialog_folder
            while temp_folder != self.internal_file_system:
                path_parts.insert(0, temp_folder["name"])
                temp_folder = temp_folder["parent"]
            current_path_label.config(text="Pašreizējā mape: /" + "/".join(path_parts))

        def on_folder_select(event):
            nonlocal current_dialog_folder
            selection = folder_listbox.curselection()
            if selection:
                index = selection[0]
                item_text = folder_listbox.get(index)

                if item_text.startswith("📁 .."):  # Atpakaļ
                    if current_dialog_folder.get("parent"):
                        current_dialog_folder = current_dialog_folder["parent"]
                        update_folder_listbox()
                else:  # Mape
                    folder_name = item_text[2:].strip()  # Noņem "📁 "
                    for item in current_dialog_folder["contents"]:
                        if item["type"] == "folder" and item["name"] == folder_name:
                            current_dialog_folder = item
                            update_folder_listbox()
                            break

        folder_listbox.bind("<Double-Button-1>", on_folder_select)

        # Ieteiktās kategorijas rāmis
        suggested_frame = ttk.LabelFrame(dialog, text="Ieteiktā kategorija", padding=5)
        suggested_frame.pack(fill="x", padx=10, pady=5)

        suggested_label = ttk.Label(suggested_frame, text=f"Programma iesaka: {suggested_category}")
        suggested_label.pack(side=tk.LEFT, padx=5)

        def use_suggested():
            nonlocal selected_node
            selected_node = self.get_or_create_folder(suggested_category)
            dialog.destroy()

        ttk.Button(suggested_frame, text="Izmantot ieteikto", command=use_suggested, bootstyle=SUCCESS).pack(
            side=tk.RIGHT, padx=5)

        # Pogas
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill="x", padx=10, pady=5)

        def create_new_folder_in_dialog():
            new_name = simpledialog.askstring("Jauna mape", "Ievadiet jaunās mapes nosaukumu:", parent=dialog)
            if new_name:
                # Pārbauda, vai mape ar šādu nosaukumu jau eksistē pašreizējā dialoga mapē
                for item in current_dialog_folder["contents"]:
                    if item["type"] == "folder" and item["name"] == new_name:
                        messagebox.showwarning("Mape jau eksistē",
                                               f"Mape ar nosaukumu '{new_name}' jau eksistē šajā mapē.", parent=dialog)
                        return
                new_folder = {"type": "folder", "name": new_name, "contents": [], "parent": current_dialog_folder}
                current_dialog_folder["contents"].append(new_folder)
                update_folder_listbox()
                self.save_pdf_archive()  # Saglabā izmaiņas arhīvā

        ttk.Button(button_frame, text="Izveidot jaunu mapi", command=create_new_folder_in_dialog).pack(side=tk.LEFT,
                                                                                                       padx=2)

        def confirm_selection():
            nonlocal selected_node
            selected_node = current_dialog_folder  # Izvēlētā mape ir pašreizējā dialoga mape
            dialog.destroy()

        ttk.Button(button_frame, text="Apstiprināt izvēli", command=confirm_selection).pack(side=tk.LEFT, padx=2)

        update_folder_listbox()  # Atjaunina sarakstu, kad logs tiek atvērts
        self.wait_window(dialog)
        return selected_node

        def update_folder_listbox():
            folder_listbox.delete(0, tk.END)
            # Pievieno ".." opciju, ja nav saknes mapē
            if current_dialog_folder != self.internal_file_system:
                folder_listbox.insert(tk.END, "📁 .. (Atpakaļ)")

            # Pievieno mapes
            for item in current_dialog_folder["contents"]:
                if item["type"] == "folder":
                    folder_listbox.insert(tk.END, f"📁 {item['name']}")

            # Atjaunina ceļa etiķeti
            path_parts = []
            temp_folder = current_dialog_folder
            while temp_folder != self.internal_file_system:
                path_parts.insert(0, temp_folder["name"])
                temp_folder = temp_folder["parent"]
            current_path_label.config(text="Pašreizējā mape: /" + "/".join(path_parts))

        def on_folder_select(event):
            nonlocal current_dialog_folder
            selection = folder_listbox.curselection()
            if selection:
                index = selection[0]
                item_text = folder_listbox.get(index)

                if item_text.startswith("📁 .."):  # Atpakaļ
                    if current_dialog_folder.get("parent"):
                        current_dialog_folder = current_dialog_folder["parent"]
                        update_folder_listbox()
                else:  # Mape
                    folder_name = item_text[2:].strip()  # Noņem "📁 "
                    for item in current_dialog_folder["contents"]:
                        if item["type"] == "folder" and item["name"] == folder_name:
                            current_dialog_folder = item
                            update_folder_listbox()
                            break

        folder_listbox.bind("<Double-Button-1>", on_folder_select)

        # Ieteiktās kategorijas rāmis
        suggested_frame = ttk.LabelFrame(dialog, text="Ieteiktā kategorija", padding=5)
        suggested_frame.pack(fill="x", padx=10, pady=5)

        suggested_label = ttk.Label(suggested_frame, text=f"Programma iesaka: {suggested_category}")
        suggested_label.pack(side=tk.LEFT, padx=5)

        def use_suggested():
            nonlocal selected_node
            selected_node = self.get_or_create_folder(suggested_category)
            dialog.destroy()

        ttk.Button(suggested_frame, text="Izmantot ieteikto", command=use_suggested, bootstyle=SUCCESS).pack(
            side=tk.RIGHT, padx=5)

        # Pogas
        button_frame = ttk.Frame(dialog)
        button_frame.pack(fill="x", padx=10, pady=5)

        def create_new_folder_in_dialog():
            new_name = simpledialog.askstring("Jauna mape", "Ievadiet jaunās mapes nosaukumu:", parent=dialog)
            if new_name:
                # Pārbauda, vai mape ar šādu nosaukumu jau eksistē pašreizējā dialoga mapē
                for item in current_dialog_folder["contents"]:
                    if item["type"] == "folder" and item["name"] == new_name:
                        messagebox.showwarning("Mape jau eksistē",
                                               f"Mape ar nosaukumu '{new_name}' jau eksistē šajā mapē.", parent=dialog)
                        return
                new_folder = {"type": "folder", "name": new_name, "contents": [], "parent": current_dialog_folder}
                current_dialog_folder["contents"].append(new_folder)
                update_folder_listbox()
                self.save_pdf_archive()  # Saglabā izmaiņas arhīvā

        ttk.Button(button_frame, text="Izveidot jaunu mapi", command=create_new_folder_in_dialog).pack(side=tk.LEFT,
                                                                                                       padx=2)

        def confirm_selection():
            nonlocal selected_node
            selected_node = current_dialog_folder  # Izvēlētā mape ir pašreizējā dialoga mape
            dialog.destroy()

        ttk.Button(button_frame, text="Apstiprināt izvēli", command=confirm_selection).pack(side=tk.LEFT, padx=2)

    def create_menu(self):
        """Izveido lietojumprogrammas izvēlni."""
        menu_bar = tk.Menu(self)

        file_menu = tk.Menu(menu_bar, tearoff=0)
        file_menu.add_command(label="Atvērt attēlus...", command=self.open_files)
        file_menu.add_command(label="Vispārīgie Iestatījumi...", command=self.show_settings)  # MAINĪTS TEKSTS
        file_menu.add_command(label="Skenēšanas Iestatījumi...", command=self.show_scan_settingss)  # JAUNA IZVĒLNE
        file_menu.add_separator()
        file_menu.add_command(label="Iziet", command=self.quit)
        menu_bar.add_cascade(label="Fails", menu=file_menu)

        help_menu = tk.Menu(menu_bar, tearoff=0)
        help_menu.add_command(label="Par programmu", command=self.show_about)
        help_menu.add_command(label="Pārbaudīt OCR valodas", command=self.check_ocr_languages)
        menu_bar.add_cascade(label="Palīdzība", menu=help_menu)

        self.config(menu=menu_bar)

    def scan_document_with_camera_fast(self):
        """Ļoti ātri atver jaunu skenēšanas logu."""
        try:
            # Bez progress loga - tieši iegūst kadru
            first_frame = self.get_camera_frame()
            if first_frame:
                # Tieši izveido jaunu scanner
                new_scanner = DocumentScanner(self)
                new_scanner.set_image(first_frame)
                new_scanner.document_frozen = False
                new_scanner.live_detected_corners = []

                # Tieši atver logu
                new_scanner.show_document_detection_preview()
                new_scanner.start_live_scan()

                self.document_scanner = new_scanner
            else:
                messagebox.showwarning("Kļūda", "Nav kameras kadra.")
                self.release_camera()
        except Exception as e:
            messagebox.showerror("Kļūda", f"Kļūda: {e}")
            self.scan_document_with_camera()

    def show_about(self):
        """Parāda informāciju par lietojumprogrammu."""
        messagebox.showinfo("Par programmu",
                            "Advanced OCR uz PDF\n\n"
                            "Versija: 2.3\n"
                            "Funkcijas:\n"
                            "- Augstas kvalitātes OCR teksts\n"
                            "- Meklējamais teksts PDF failos\n"
                            "- Vairāku valodu atbalsts\n"
                            "- Responsīvs lietotāja interfeiss ar cilnēm un ritjoslām\n"
                            "- Attēlu apstrādes rīki (spilgtums, kontrasts, asums, rotācija)\n"
                            "- Slīpuma korekcija (Deskew)\n"
                            "- Trokšņu samazināšana\n"
                            "- Attēla negatīvs, malu noteikšana, binārizācija\n"
                            "- Automātiska attēla uzlabošana\n"
                            "- Pielāgojami OCR parametri\n"
                            "- PDF lapas orientācijas un izmēra izvēle\n"
                            "- PDF izvades kvalitātes kontrole\n"
                            "- Iestatījumu saglabāšana un automātiska ielāde\n"
                            "- Attēlu dzēšana no saraksta\n"
                            "- Attēla priekšskatījuma tālummaiņa un pārvietošana\n"
                            "- Pilnekrāna attēla priekšskatījums\n"
                            "- Attēla histogramma, metadati, krāsu palete\n"
                            "- Unikāls dokumenta ID un QR kods PDF failos\n"
                            "- Integrēta PDF failu pārvaldības sistēma (arhīvs)\n"
                            "- Meklēšana un filtrēšana PDF arhīvā pēc datuma un teksta\n"
                            "- Iespēja nosūtīt PDF failus e-pastā, izmantojot noklusēto e-pasta klientu\n"
                            "- Mapju izveide un failu pārvietošana failu pārvaldības cilnē\n"
                            "- Automātiska dokumentu klasifikācija un saglabāšana atbilstošās mapēs")

    def show_settings(self):
        """Parāda vispārīgo iestatījumu logu."""
        if not hasattr(self, '_settings_window') or not self._settings_window.winfo_exists():
            self._settings_window = SettingsWindow(self, self)
        self._settings_window.lift()

    def show_scan_settingss(self):
        """JAUNS: Parāda skenēšanas iestatījumu logu."""
        if not hasattr(self, '_scan_settings_window') or not self._scan_settings_window.winfo_exists():
            self._scan_settings_window = ScanSettingsWindow(self, self)
        self._scan_settings_window.lift()

    def open_files(self, filepath=None):
        """Galvenā metode failu atvēršanai"""
        if filepath is None:
            filepaths = filedialog.askopenfilenames(
                title="Izvēlieties failus",
                filetypes=[("Attēli", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("PDF faili", "*.pdf")]  # Changed order
            )

        else:
            filepaths = [filepath]

        if not filepaths:
            return

        self.clear_files()

        for filepath in filepaths:
            if not os.path.exists(filepath):
                continue

            try:
                if filepath.lower().endswith('.pdf'):
                    with PDFEditor(filepath) as editor:
                        for page_num in range(editor.get_page_count()):
                            img = editor.get_page_image(page_num, self.dpi_var.get())
                            self.images.append({
                                "filepath": f"{filepath}||page{page_num}",
                                "original_img": img.copy(),
                                "processed_img": img.copy()
                            })
                            self.file_listbox.insert(tk.END, f"{os.path.basename(filepath)} (lapa {page_num + 1})")
                else:
                    img = Image.open(filepath)
                    self.images.append({
                        "filepath": filepath,
                        "original_img": img.copy(),
                        "processed_img": img.copy()
                    })
                    self.file_listbox.insert(tk.END, os.path.basename(filepath))

                # Saglabā failu lietotāja datu struktūrā
                # save_user_file(self.username, filepath)  # Komentēts, jo username nav definēts
            except Exception as e:
                print(f"Kļūda apstrādājot {filepath}: {e}")

        if self.images:
            self.file_listbox.select_set(0)
            self.on_file_select()

    def clear_files(self):
        """Notīra visus ielādētos attēlus un OCR rezultātus."""
        self.images.clear()
        self.file_listbox.delete(0, tk.END)
        self.ocr_results.clear()  # Notīra OCR rezultātus
        # Notīra visas iepriekšējās fona krāsas, ja tādas bija
        self.ocr_results.clear()
        self.text_ocr.delete("1.0", tk.END)
        self.canvas.delete("all")
        self.progress["value"] = 0
        self.current_image_index = -1
        self.reset_image_processing_vars()
        self.canvas_zoom_factor = 1.0
        self.canvas_pan_x = 0
        self.canvas_pan_y = 0
        # JAUNS: Notīra QR koda rāmja mainīgos
        self.qr_code_frame_id = None
        self.qr_code_handle_ids = []
        self.qr_code_active_handle = None
        self.qr_code_frame_coords = None
        self.qr_code_start_drag_x = None
        self.qr_code_start_drag_y = None
        self.qr_code_drag_mode = None
        self.btn_toggle_qr_frame.config(bootstyle="default")  # Atjauno pogas stilu
        self._qr_edit_mode = False

    def reset_image_processing_vars(self):
        """Atjauno attēlu apstrādes mainīgos uz noklusējuma vērtībām."""
        self.brightness_var.set(1.0)
        self.contrast_var.set(1.0)
        self.sharpness_var.set(1.0)
        self.rotate_var.set(0)
        self.grayscale_var.set(True)
        self.deskew_var.set(False)
        self.remove_noise_var.set(False)
        self.invert_colors_var.set(False)
        self.edge_detection_var.set(False)
        self.binarize_var.set(False)

    def delete_selected_image(self):
        """Dzēš atlasītos attēlus no saraksta."""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet attēlu(s), ko dzēst.")
            return

        # Sakārto indeksus dilstošā secībā, lai dzēšot nemainītos atlikušo elementu indeksi
        selected_indices = sorted(list(selection), reverse=True)

        if messagebox.askyesno("Dzēst attēlus",
                               f"Vai tiešām vēlaties dzēst {len(selected_indices)} atlasītos attēlus no saraksta?"):
            for index in selected_indices:
                # Pārbauda, vai indekss joprojām ir derīgs (ja saraksts ir mainījies)
                if index < len(self.images):
                    # Nav nepieciešams dzēst fizisko failu, jo tie ir tikai ielādēti attēli
                    # no dažādām vietām, nevis arhīva faili.
                    del self.images[index]
                    if len(self.ocr_results) > index:
                        del self.ocr_results[index]
                    self.file_listbox.delete(index)  # Dzēš no Listbox
                    # Pēc dzēšanas no Listbox, atjaunojam atlikušo vienumu krāsas, ja nepieciešams
                    # Tas nodrošina, ka, ja dzēšam vienumu, kas nav pēdējais,
                    # tad nākamie vienumi saglabā savas krāsas vai atjaunojas uz noklusējumu.
                    # Šeit mēs vienkārši atjaunojam visu sarakstu, lai nodrošinātu konsekvenci.
                    for index in selected_indices:
                        # Pārbauda, vai indekss joprojām ir derīgs (ja saraksts ir mainījies)
                        if index < len(self.images):
                            # Nav nepieciešams dzēst fizisko failu, jo tie ir tikai ielādēti attēli
                            # no dažādām vietām, nevis arhīva faili.
                            del self.images[index]
                            # Pārbauda, vai ocr_results saraksts ir pietiekami garš
                            if index < len(self.ocr_results):
                                del self.ocr_results[index]
                            # Listbox vienumi tiks atjaunināti ar refresh_file_listbox() pēc cikla

                    # Pēc visu atlasīto vienumu dzēšanas no datu struktūrām, atjaunojam Listbox
                    self.refresh_file_listbox()

                    if self.images:
                        # Pēc dzēšanas mēģina atlasīt pirmo atlikušo attēlu, ja tāds ir
                        new_index = 0
                        if new_index < len(self.images):
                            self.file_listbox.select_set(new_index)
                            self.on_file_select()
                        else:
                            self.clear_files()  # Ja saraksts ir tukšs
                    else:
                        self.clear_files()
                    messagebox.showinfo("Dzēsts", "Atlasītie attēli veiksmīgi dzēsti no saraksta.")

            if self.images:
                # Pēc dzēšanas mēģina atlasīt pirmo atlikušo attēlu, ja tāds ir
                new_index = 0
                if new_index < len(self.images):
                    self.file_listbox.select_set(new_index)
                    self.on_file_select()
                else:
                    self.clear_files()  # Ja saraksts ir tukšs
            else:
                self.clear_files()
            messagebox.showinfo("Dzēsts", "Atlasītie attēli veiksmīgi dzēsti no saraksta.")

    def on_file_click(self, event):
        """Apstrādā vienu klikšķi uz faila, lai to apskatītu."""
        # Pārbauda, vai nav nospiests Ctrl vai Shift, kas norāda uz vairāku atlasi
        if event.state & 0x4 or event.state & 0x1:  # 0x4 ir Ctrl, 0x1 ir Shift
            return  # Ļauj Listbox apstrādāt vairāku atlasi

        # Notīra iepriekšējo atlasi un atlasa tikai noklikšķināto vienumu
        self.file_listbox.selection_clear(0, tk.END)
        index = self.file_listbox.nearest(event.y)
        if index != -1:
            self.file_listbox.selection_set(index)
            self.on_file_select()  # Izsauc faila atlases funkciju

    def on_file_select_deferred(self):
        """Aizkavēta faila atlases apstrāde, lai izvairītos no konfliktiem ar vilkšanu."""
        # Šī funkcija tiek izsaukta pēc <<ListboxSelect>> notikuma.
        # Ja ir aktīva vilkšanas operācija, mēs neko nedarām.
        if self.file_drag_data["item_index"] is not None:
            return

        # Ja nav aktīvas vilkšanas, tad izsaucam parasto atlases loģiku.
        # Tomēr, lai nodrošinātu, ka vienmēr tiek parādīts atlasītais fails,
        # mēs izmantosim on_file_click, kas jau apstrādā vienu atlasi.
        # Šis <<ListboxSelect>> notikums joprojām ir noderīgs, ja lietotājs izmanto
        # tastatūras bulttaustiņus, lai pārvietotos pa sarakstu.
        selection = self.file_listbox.curselection()
        if selection:
            self.current_image_index = selection[0]
            self.apply_image_filters(None)
            if len(self.ocr_results) > self.current_image_index and self.ocr_results[
                self.current_image_index] is not None:
                self.text_ocr.delete("1.0", tk.END)
                self.text_ocr.insert(tk.END, self.ocr_results[self.current_image_index]['full_text'])
            else:
                self.text_ocr.delete("1.0", tk.END)
                self.text_ocr.insert(tk.END, "OCR rezultāts vēl nav pieejams.")
        else:
            self.current_image_index = -1
            self.canvas.delete("all")
            self.text_ocr.delete("1.0", tk.END)

    def move_file_up(self):
        """Pārvieto atlasīto failu sarakstā uz augšu."""
        selection = self.file_listbox.curselection()
        if not selection:
            return
        index = selection[0]
        if index > 0:
            # Pārvieto elementu self.images sarakstā
            item_to_move = self.images.pop(index)
            self.images.insert(index - 1, item_to_move)

            # Atjaunina pašreizējo attēla indeksu
            self.current_image_index = index - 1

            # Pilnībā atjauno listbox, lai atspoguļotu jauno secību
            self.refresh_file_listbox()

    def move_file_down(self):
        """Pārvieto atlasīto failu sarakstā uz leju."""
        selection = self.file_listbox.curselection()
        if not selection:
            return
        index = selection[0]
        if index < len(self.images) - 1:
            # Pārvieto elementu self.images sarakstā
            item_to_move = self.images.pop(index)
            self.images.insert(index + 1, item_to_move)

            # Atjaunina pašreizējo attēla indeksu
            self.current_image_index = index + 1

            # Pilnībā atjauno listbox, lai atspoguļotu jauno secību
            self.refresh_file_listbox()

    def file_list_drag_start(self, event):
        """Sāk vilkšanas operāciju attēlu sarakstā."""
        index = self.file_listbox.nearest(event.y)
        if index != -1:
            # Pārbauda, vai noklikšķinātais vienums ir atlasīts.
            # Ja nav, tad notīra iepriekšējo atlasi un atlasa tikai šo vienumu.
            if index not in self.file_listbox.curselection():
                self.file_listbox.selection_clear(0, tk.END)
                self.file_listbox.selection_set(index)
                self.on_file_select()  # Atjauno priekšskatījumu uz atlasīto failu

            self.file_drag_data["item_index"] = index
            self.file_drag_data["start_y"] = event.y
            self.file_listbox.config(cursor="fleur")  # Maina kursoru uz pārvietošanas ikonu

    def file_list_drag_motion(self, event):
        """Apstrādā vilkšanas kustību attēlu sarakstā."""
        if self.file_drag_data["item_index"] is not None:
            # Šeit varētu pievienot vizuālu atgriezenisko saiti, piemēram, zīmēt līniju
            # vai mainīt fona krāsu, kur fails tiks nomests.
            pass

    def file_list_drag_drop(self, event):
        """Apstrādā nomešanas operāciju attēlu sarakstā."""
        self.file_listbox.config(cursor="arrow")  # Atjauno kursoru
        if self.file_drag_data["item_index"] is not None:
            current_index = self.file_drag_data["item_index"]
            target_index = self.file_listbox.nearest(event.y)

            if target_index != -1 and current_index != target_index:
                # Pārvieto atlasītos vienumus
                selected_indices = sorted([i for i in self.file_listbox.curselection()], reverse=True)

                # Izveido sarakstu ar pārvietojamiem attēliem un OCR rezultātiem
                moved_images = []
                moved_ocr_results = []
                for idx in selected_indices:
                    moved_images.insert(0, self.images.pop(idx))
                    if len(self.ocr_results) > idx:
                        moved_ocr_results.insert(0, self.ocr_results.pop(idx))
                    else:
                        moved_ocr_results.insert(0, None)  # Ja OCR rezultāts vēl nav

                # Ievieto pārvietotos vienumus jaunajā pozīcijā
                for i, img_data in enumerate(moved_images):
                    insert_idx = target_index if target_index < current_index else target_index - (
                            len(selected_indices) - 1) + i
                    self.images.insert(insert_idx, img_data)
                    self.ocr_results.insert(insert_idx, moved_ocr_results[i])

                self.refresh_file_listbox()  # Atjauno Listbox vizuāli

                # Atjauno atlasi uz pārvietotajiem vienumiem
                new_selection_start = min(current_index, target_index)
                new_selection_end = new_selection_start + len(selected_indices) - 1
                for i in range(new_selection_start, new_selection_end + 1):
                    self.file_listbox.selection_set(i)

                # Pārliecinās, ka pašreizējais attēls joprojām ir redzams, ja tas tika pārvietots
                if self.current_image_index != -1:
                    old_current_image_filepath = self.images[self.current_image_index]["filepath"]
                    # Atrod jauno indeksu pašreizējam attēlam
                    for i, img_data in enumerate(self.images):
                        if img_data["filepath"] == old_current_image_filepath:
                            self.current_image_index = i
                            break
                    self.on_file_select()  # Atjauno priekšskatījumu

                messagebox.showinfo("Pārvietots", "Atlasītie faili veiksmīgi pārvietoti sarakstā.")

            self.file_drag_data["item_index"] = None  # Notīra vilkšanas datus

        # Pilnībā aizvietojiet esošo refresh_file_listbox metodi ar šo:

    def refresh_file_listbox(self):
        """
        Atjaunina failu sarakstu `file_listbox` no `self.images` saraksta.
        Nodrošina, ka vizuālais attēlojums atbilst pamatā esošajai datu struktūrai.
        """
        self.file_listbox.delete(0, tk.END)
        for i, img_data in enumerate(self.images):
            # Izmanto 'display_name', ja pieejams, citādi faila nosaukumu
            display_name = img_data.get("display_name", os.path.basename(img_data["filepath"]))
            self.file_listbox.insert(tk.END, f"{i + 1}. {display_name}")

        # Atjauno atlasi un ritināšanu, ja ir atlasīts attēls
        if self.current_image_index != -1 and self.current_image_index < len(self.images):
            self.file_listbox.selection_set(self.current_image_index)
            self.file_listbox.activate(self.current_image_index)
            self.file_listbox.see(self.current_image_index)

    def on_file_click(self, event):
        """Apstrādā vienu klikšķi uz faila, lai to apskatītu."""
        # Pārbauda, vai nav nospiests Ctrl vai Shift, kas norāda uz vairāku atlasi
        if event.state & 0x4 or event.state & 0x1:  # 0x4 ir Ctrl, 0x1 ir Shift
            return  # Ļauj Listbox apstrādāt vairāku atlasi

        # Notīra iepriekšējo atlasi un atlasa tikai noklikšķināto vienumu
        self.file_listbox.selection_clear(0, tk.END)
        index = self.file_listbox.nearest(event.y)
        if index != -1:
            self.file_listbox.selection_set(index)
            self.on_file_select()  # Izsauc faila atlases funkciju

    def on_file_select_deferred(self):
        """Aizkavēta faila atlases apstrāde, lai izvairītos no konfliktiem ar vilkšanu."""
        # Šī funkcija tiek izsaukta pēc <<ListboxSelect>> notikuma.
        # Ja ir aktīva vilkšanas operācija, mēs neko nedarām.
        if self.file_drag_data["item_index"] is not None:
            return

        # Ja nav aktīvas vilkšanas, tad izsaucam parasto atlases loģiku.
        # Tomēr, lai nodrošinātu, ka vienmēr tiek parādīts atlasītais fails,
        # mēs izmantosim on_file_click, kas jau apstrādā vienu atlasi.
        # Šis <<ListboxSelect>> notikums joprojām ir noderīgs, ja lietotājs izmanto
        # tastatūras bulttaustiņus, lai pārvietotos pa sarakstu.
        selection = self.file_listbox.curselection()
        if selection:
            self.current_image_index = selection[0]
            self.apply_image_filters(None)
            if len(self.ocr_results) > self.current_image_index and self.ocr_results[
                self.current_image_index] is not None:
                self.text_ocr.delete("1.0", tk.END)
                self.text_ocr.insert(tk.END, self.ocr_results[self.current_image_index]['full_text'])
            else:
                self.text_ocr.delete("1.0", tk.END)
                self.text_ocr.insert(tk.END, "OCR rezultāts vēl nav pieejams.")
        else:
            self.current_image_index = -1
            self.canvas.delete("all")
            self.text_ocr.delete("1.0", tk.END)

    def move_file_up(self):
        """Pārvieto atlasīto failu uz augšu sarakstā."""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet failu.")
            return

        index = selection[0]
        if index > 0:
            # Maina vietām ar iepriekšējo
            self.images[index], self.images[index - 1] = self.images[index - 1], self.images[index]
            if len(self.ocr_results) > index:
                self.ocr_results[index], self.ocr_results[index - 1] = self.ocr_results[index - 1], self.ocr_results[
                    index]

            self.refresh_file_listbox()
            self.file_listbox.select_set(index - 1)
            self.current_image_index = index - 1
            self.on_file_select()

    def move_file_down(self):
        """Pārvieto atlasīto failu uz leju sarakstā."""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showwarning("Nav atlasīts", "Lūdzu, atlasiet failu.")
            return

        index = selection[0]
        if index < len(self.images) - 1:
            # Maina vietām ar nākošo
            self.images[index], self.images[index + 1] = self.images[index + 1], self.images[index]
            if len(self.ocr_results) > index + 1:
                self.ocr_results[index], self.ocr_results[index + 1] = self.ocr_results[index + 1], self.ocr_results[
                    index]

            self.refresh_file_listbox()
            self.file_listbox.select_set(index + 1)
            self.current_image_index = index + 1
            self.on_file_select()

    def file_list_drag_start(self, event):
        """Sāk vilkšanas operāciju attēlu sarakstā."""
        index = self.file_listbox.nearest(event.y)
        if index != -1:
            # Pārbauda, vai noklikšķinātais vienums ir atlasīts.
            # Ja nav, tad notīra iepriekšējo atlasi un atlasa tikai šo vienumu.
            if index not in self.file_listbox.curselection():
                self.file_listbox.selection_clear(0, tk.END)
                self.file_listbox.selection_set(index)
                self.on_file_select()  # Atjauno priekšskatījumu uz atlasīto failu

            self.file_drag_data["item_index"] = index
            self.file_drag_data["start_y"] = event.y
            self.file_listbox.config(cursor="fleur")  # Maina kursoru uz pārvietošanas ikonu

    def file_list_drag_motion(self, event):
        """Apstrādā vilkšanas kustību attēlu sarakstā."""
        if self.file_drag_data["item_index"] is not None:
            # Šeit varētu pievienot vizuālu atgriezenisko saiti, piemēram, zīmēt līniju
            # vai mainīt fona krāsu, kur fails tiks nomests.
            pass

    def file_list_drag_drop(self, event):
        """Apstrādā nomešanas operāciju attēlu sarakstā."""
        self.file_listbox.config(cursor="arrow")  # Atjauno kursoru
        if self.file_drag_data["item_index"] is not None:
            current_index = self.file_drag_data["item_index"]
            target_index = self.file_listbox.nearest(event.y)

            if target_index != -1 and current_index != target_index:
                # Pārvieto atlasītos vienumus
                selected_indices = sorted([i for i in self.file_listbox.curselection()], reverse=True)

                # Izveido sarakstu ar pārvietojamiem attēliem un OCR rezultātiem
                moved_images = []
                moved_ocr_results = []
                for idx in selected_indices:
                    moved_images.insert(0, self.images.pop(idx))
                    if len(self.ocr_results) > idx:
                        moved_ocr_results.insert(0, self.ocr_results.pop(idx))
                    else:
                        moved_ocr_results.insert(0, None)  # Ja OCR rezultāts vēl nav

                # Ievieto pārvietotos vienumus jaunajā pozīcijā
                for i, img_data in enumerate(moved_images):
                    insert_idx = target_index if target_index < current_index else target_index - (
                            len(selected_indices) - 1) + i
                    self.images.insert(insert_idx, img_data)
                    self.ocr_results.insert(insert_idx, moved_ocr_results[i])

                self.refresh_file_listbox()  # Atjauno Listbox vizuāli

                # Atjauno atlasi uz pārvietotajiem vienumiem
                new_selection_start = min(current_index, target_index)
                new_selection_end = new_selection_start + len(selected_indices) - 1
                for i in range(new_selection_start, new_selection_end + 1):
                    self.file_listbox.selection_set(i)

                # Pārliecinās, ka pašreizējais attēls joprojām ir redzams, ja tas tika pārvietots
                if self.current_image_index != -1:
                    old_current_image_filepath = self.images[self.current_image_index]["filepath"]
                    # Atrod jauno indeksu pašreizējam attēlam
                    for i, img_data in enumerate(self.images):
                        if img_data["filepath"] == old_current_image_filepath:
                            self.current_image_index = i
                            break
                    self.on_file_select()  # Atjauno priekšskatījumu

                messagebox.showinfo("Pārvietots", "Atlasītie faili veiksmīgi pārvietoti sarakstā.")

            self.file_drag_data["item_index"] = None  # Notīra vilkšanas datus

    def refresh_file_listbox(self):
        """Atjauno failu saraksta Listbox vizuāli un atjaunina krāsas."""
        self.file_listbox.delete(0, tk.END)
        default_bg_color = self.style.colors.get("bg")  # Iegūst noklusējuma fona krāsu no tēmas

        for i, img_data in enumerate(self.images):
            display_name = ""
            # Pārbauda, vai tas ir PDF lapa vai atsevišķs attēls
            if "Lapa" in img_data["filepath"]:  # Vienkāršots veids, kā atpazīt PDF lapas
                display_name = img_data["filepath"]
            else:
                display_name = os.path.basename(img_data["filepath"])

            self.file_listbox.insert(tk.END, display_name)

            # Pārbauda, vai OCR ir pabeigts šim failam un iezīmē to zaļā krāsā
            # Pārliecinās, ka indekss ir derīgs ocr_results sarakstam
            if i < len(self.ocr_results) and self.ocr_results[i] is not None:
                self.file_listbox.itemconfig(i, {'bg': '#d4edda'})  # Gaiši zaļa krāsa
            else:
                # Atjauno noklusējuma krāsu, ja fails vēl nav apstrādāts
                self.file_listbox.itemconfig(i, {'bg': default_bg_color})

        # Atjauno iepriekšējo atlasi, ja tāda bija
        if self.current_image_index != -1 and self.current_image_index < len(self.images):
            self.file_listbox.select_set(self.current_image_index)
            self.file_listbox.activate(self.current_image_index)
            self.file_listbox.see(self.current_image_index)  # Pārvietojas uz atlasīto vienumu

    def show_file_context_menu(self, event):
        """Parāda konteksta izvēlni failu saraksta elementiem."""
        try:
            self.file_listbox.selection_clear(0, tk.END)
            self.file_listbox.selection_set(self.file_listbox.nearest(event.y))
            self.file_listbox.activate(self.file_listbox.nearest(event.y))

            context_menu = tk.Menu(self.file_listbox, tearoff=0)
            context_menu.add_command(label="Dzēst", command=self.delete_selected_image)
            context_menu.add_command(label="Atvērt faila atrašanās vietu", command=self.open_file_location)
            context_menu.post(event.x_root, event.y_root)
        except Exception:
            pass  # Ja nav atlasīts nekas, ignorē

    def open_file_location(self):
        """Atver mapes atrašanās vietu, kurā atrodas atlasītais attēla fails."""
        selection = self.file_listbox.curselection()
        if not selection:
            return
        index = selection[0]
        filepath = self.images[index]['filepath']
        if os.path.exists(filepath):
            try:
                if os.name == 'nt':  # Windows
                    os.startfile(os.path.dirname(filepath))
                elif os.name == 'posix':  # macOS, Linux
                    import subprocess
                    import sys
                    if sys.platform == 'darwin':  # macOS
                        subprocess.Popen(['open', '-R', filepath])
                    else:  # Linux
                        subprocess.Popen(['xdg-open', os.path.dirname(filepath)])
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās atvērt faila atrašanās vietu:\n{e}")
        else:
            messagebox.showwarning("Fails nav atrasts", "Attēla fails nav atrasts norādītajā vietā.")

    def on_file_select(self, event=None):
        """Apstrādā faila atlasi sarakstā."""
        if self.file_listbox.curselection():
            selected_index = self.file_listbox.curselection()[0]
            self.current_image_index = selected_index

            # Pārbauda, vai atlasītais fails ir attēls vai PDF
            if self.current_image_index < len(self.images):
                selected_item = self.images[self.current_image_index]
                if "processed_img" in selected_item and selected_item["processed_img"] is not None:
                    self.show_image_preview(selected_item["processed_img"])
                else:
                    # Ja nav apstrādāta attēla, mēģina parādīt oriģinālo
                    if "original_img" in selected_item and selected_item["original_img"] is not None:
                        self.show_image_preview(selected_item["original_img"])
                    else:
                        # Ja nav ne apstrādāta, ne oriģinālā attēla, notīra priekšskatījumu
                        self.canvas.delete("all")
                        self.canvas.create_text(self.canvas.winfo_width() / 2, self.canvas.winfo_height() / 2,
                                                text="Nav attēla priekšskatījumam", fill="white")

                # Atjaunina OCR teksta lauku
                if self.current_image_index < len(self.ocr_results) and self.ocr_results[
                    self.current_image_index] is not None:
                    self.text_ocr.delete("1.0", tk.END)
                    self.text_ocr.insert("1.0", self.ocr_results[self.current_image_index])
                else:
                    self.text_ocr.delete("1.0", tk.END)
                    self.text_ocr.insert("1.0", "OCR rezultāts nav pieejams.")

                # JAUNS: Atjaunina "Papildu rīki" cilnes priekšskatījumu
                self._update_additional_tools_pdf_preview()
            else:
                # Ja atlase ir ārpus saraksta robežām (piemēram, pēc dzēšanas)
                self.current_image_index = -1
                self.canvas.delete("all")
                self.text_ocr.delete("1.0", tk.END)
                self.text_ocr.insert("1.0", "Nav atlasīts fails.")
                self._clear_additional_tools_pdf_preview()  # Notīra arī papildu rīku priekšskatījumu
        else:
            self.current_image_index = -1
            self.canvas.delete("all")
            self.text_ocr.delete("1.0", tk.END)
            self.text_ocr.insert("1.0", "Nav atlasīts fails.")
            self._clear_additional_tools_pdf_preview()  # Notīra arī papildu rīku priekšskatījumu

    def apply_image_filters(self, event):
        """Pielieto attēlu apstrādes filtrus pašreizējam attēlam."""
        if self.current_image_index == -1:
            return

        img_data = self.images[self.current_image_index]
        img = img_data["original_img"].copy()

        rotation_angle = self.rotate_var.get()
        if rotation_angle != 0:
            img = img.rotate(rotation_angle, expand=True, fillcolor=(255, 255, 255) if img.mode == 'RGB' else 255)

        if self.grayscale_var.get():
            img = img.convert("L")

        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(self.brightness_var.get())

        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(self.contrast_var.get())

        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(self.sharpness_var.get())

        if self.deskew_var.get() and OPENCV_AVAILABLE:
            try:
                img_cv = np.array(img)
                if len(img_cv.shape) == 3:
                    img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
                img_cv = cv2.bitwise_not(img_cv)
                coords = np.column_stack(np.where(img_cv > 0))

                # Pārbaude, vai coords nav tukšs
                if coords.size > 0:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle
                    (h, w) = img_cv.shape[:2]
                    center = (w // 2, h // 2)
                    M = cv2.getRotationMatrix2D(center, angle, 1.0)
                    img_cv = cv2.warpAffine(img_cv, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
                    img_cv = cv2.bitwise_not(img_cv)
                    img = Image.fromarray(img_cv)
                else:
                    messagebox.showwarning("Slīpuma korekcija",
                                           "Netika atrasts pietiekami daudz teksta slīpuma korekcijai. Slīpuma korekcija netika veikta.")
                    self.deskew_var.set(False)  # Atspējo, ja nevar veikt
            except Exception as e:
                messagebox.showwarning("Slīpuma korekcijas kļūda", f"Neizdevās veikt slīpuma korekciju: {e}")
                self.deskew_var.set(False)
        elif self.deskew_var.get() and not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas",
                                   "Slīpuma korekcijai nepieciešams 'opencv-python' un 'numpy'. Lūdzu, instalējiet tās (pip install opencv-python numpy).")
            self.deskew_var.set(False)

        if self.remove_noise_var.get():
            if img.mode != 'L':
                img = img.convert("L")
            img = img.filter(ImageFilter.MedianFilter(size=3))

        if self.invert_colors_var.get():
            img = ImageOps.invert(img)

        if self.edge_detection_var.get():
            if img.mode != 'L':
                img = img.convert("L")
            img = img.filter(ImageFilter.FIND_EDGES)

        if self.binarize_var.get():
            if img.mode != 'L':
                img = img.convert("L")
            img = img.point(lambda x: 0 if x < 128 else 255, '1')

        img_data["processed_img"] = img
        self.show_image_preview(img)

    def _set_default_qr_frame_coords(self):
        """Iestata default QR koda rāmja koordinātas balstoties uz iestatījumiem."""
        if self.current_image_index == -1:
            return

        img_data = self.images[self.current_image_index]
        img_pil = img_data["processed_img"]
        img_width, img_height = img_pil.size

        # Default izmērs (10% no mazākās puses)
        qr_size = min(img_width, img_height) * 0.10
        margin = min(img_width, img_height) * 0.02

        # Iegūst pozīciju no iestatījumiem
        # Pārliecināmies, ka izmantojam pareizo atslēgu un noklusējuma vērtību
        id_code_position = self.settings.get("id_code_position", "bottom_right")

        print(f"DEBUG: Iestatītā QR pozīcija: {id_code_position}")  # Debug rinda

        if id_code_position == "top_left":  # MAINĪTS: no "top-left" uz "top_left"
            x1 = margin
            y1 = margin
            print(f"DEBUG: Izmanto top-left pozīciju: x1={x1}, y1={y1}")
        elif id_code_position == "top_right":
            x1 = img_width - qr_size - margin
            y1 = margin
            print(f"DEBUG: Izmanto top-right pozīciju: x1={x1}, y1={y1}")
        elif id_code_position == "bottom_left":  # MAINĪTS: no "bottom-left" uz "bottom_left"
            x1 = margin
            y1 = img_height - qr_size - margin
            print(f"DEBUG: Izmanto bottom-left pozīciju: x1={x1}, y1={y1}")
        elif id_code_position == "bottom_right":  # MAINĪTS: no "bottom-right" uz "bottom_right"
            x1 = img_width - qr_size - margin
            y1 = img_height - qr_size - margin
            print(f"DEBUG: Izmanto bottom-right pozīciju: x1={x1}, y1={y1}")
        else:
            # Fallback uz bottom-right, ja nav atpazīts
            x1 = img_width - qr_size - margin
            y1 = img_height - qr_size - margin
            print(f"DEBUG: Nezināma pozīcija '{id_code_position}', izmanto bottom-right: x1={x1}, y1={y1}")

        x2 = x1 + qr_size
        y2 = y1 + qr_size

        self.qr_code_frame_coords = (x1, y1, x2, y2)
        print(f"DEBUG: Finālās koordinātas: {self.qr_code_frame_coords}")

    def show_image_preview(self, img):
        """Parāda attēla priekšskatījumu uz kanvasa."""
        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()

        if canvas_width <= 1 or canvas_height <= 1:
            return

        img_width, img_height = img.size
        scaled_width = int(img_width * self.canvas_zoom_factor)
        scaled_height = int(img_height * self.canvas_zoom_factor)

        display_img = img.resize((scaled_width, scaled_height), Image.LANCZOS)
        self.photo_image = ImageTk.PhotoImage(display_img)

        self.canvas.delete("all")

        x = (canvas_width - scaled_width) / 2 + self.canvas_pan_x
        y = (canvas_height - scaled_height) / 2 + self.canvas_pan_y

        self.canvas.create_image(x, y, anchor="nw", image=self.photo_image)
        self.canvas.image = self.photo_image

        # JAUNS: Zīmē QR koda/svītrkoda rāmi, ja tas ir aktīvs
        # JAUNS: Zīmē QR koda/svītrkoda rāmi, ja tas ir aktīvs
        # JAUNS: Zīmē QR koda/svītrkoda rāmi, ja funkcija ir ieslēgta vai manuāli aktivizēta
        should_show_qr_frame = (self.settings.get("add_id_code_to_pdf", False) or
                                self.qr_code_frame_coords is not None)

        if should_show_qr_frame:
            # Ja nav manuāli iestatītas koordinātas, izmanto default pozīciju
            if not hasattr(self, 'qr_code_frame_coords') or self.qr_code_frame_coords is None:
                self._set_default_qr_frame_coords()

            # Notīra vecos QR rāmja elementus
            self.canvas.delete("qr_frame")
            self.canvas.delete("qr_handle")

            # Pārrēķina rāmja koordinātas uz kanvasa koordinātām
            x1_img, y1_img, x2_img, y2_img = self.qr_code_frame_coords

            # Pārrēķina rāmja koordinātas no oriģinālā attēla uz kanvasa koordinātām
            x1_canvas = x + x1_img * self.canvas_zoom_factor
            y1_canvas = y + y1_img * self.canvas_zoom_factor
            x2_canvas = x + x2_img * self.canvas_zoom_factor
            y2_canvas = y + y2_img * self.canvas_zoom_factor

            # Aprēķina kvadrātisku izmēru (aspect ratio 1:1)
            frame_width = x2_canvas - x1_canvas
            frame_height = y2_canvas - y1_canvas
            square_size = min(frame_width, frame_height)

            # Centrē kvadrātu rāmja ietvaros
            center_x = (x1_canvas + x2_canvas) / 2
            center_y = (y1_canvas + y2_canvas) / 2
            x1_canvas = center_x - square_size / 2
            y1_canvas = center_y - square_size / 2
            x2_canvas = center_x + square_size / 2
            y2_canvas = center_y + square_size / 2

            # Zīmē rāmi
            self.qr_code_frame_id = self.canvas.create_rectangle(
                x1_canvas, y1_canvas, x2_canvas, y2_canvas,
                outline="yellow", width=2, dash=(5, 2), tags="qr_frame"
            )

            # Zīmē stūru rokturus
            handle_size = 8
            self.qr_code_handle_ids = []
            handles = [
                (x1_canvas, y1_canvas, "nw"),
                (x2_canvas, y1_canvas, "ne"),
                (x2_canvas, y2_canvas, "se"),
                (x1_canvas, y2_canvas, "sw")
            ]
            for x_handle, y_handle, cursor_type in handles:
                handle_id = self.canvas.create_oval(
                    x_handle - handle_size, y_handle - handle_size,
                    x_handle + handle_size, y_handle + handle_size,
                    fill="cyan", outline="white", width=2, tags="qr_handle"
                )
                self.qr_code_handle_ids.append({"id": handle_id, "type": cursor_type})

            # Pārvieto QR elementus uz priekšu (virs attēla)
            self.canvas.tag_raise("qr_frame")
            self.canvas.tag_raise("qr_handle")

        # Notīra veco attēlu pirms jauna zīmēšanas
        self.canvas.delete("image")
        self.canvas.create_image(x, y, anchor="nw", image=self.photo_image, tags="image")
        self.canvas.image = self.photo_image

        # JAUNS: Zīmē QR koda/svītrkoda rāmi, ja funkcija ir ieslēgta vai manuāli aktivizēta
        should_show_qr_frame = (self.settings.get("add_id_code_to_pdf", False) or
                                self.qr_code_frame_coords is not None)

        if should_show_qr_frame:
            # Ja nav manuāli iestatītas koordinātas, izmanto default pozīciju
            if not hasattr(self, 'qr_code_frame_coords') or self.qr_code_frame_coords is None:
                self._set_default_qr_frame_coords()

            # Notīra vecos QR rāmja elementus
            self.canvas.delete("qr_frame")
            self.canvas.delete("qr_handle")

            # Pārrēķina rāmja koordinātas uz kanvasa koordinātām
            x1_img, y1_img, x2_img, y2_img = self.qr_code_frame_coords

            # Pārrēķina rāmja koordinātas no oriģinālā attēla uz kanvasa koordinātām
            x1_canvas = x + x1_img * self.canvas_zoom_factor
            y1_canvas = y + y1_img * self.canvas_zoom_factor
            x2_canvas = x + x2_img * self.canvas_zoom_factor
            y2_canvas = y + y2_img * self.canvas_zoom_factor

            # Aprēķina kvadrātisku izmēru (aspect ratio 1:1)
            frame_width = x2_canvas - x1_canvas
            frame_height = y2_canvas - y1_canvas
            square_size = min(frame_width, frame_height)

            # Centrē kvadrātu rāmja ietvaros
            center_x = (x1_canvas + x2_canvas) / 2
            center_y = (y1_canvas + y2_canvas) / 2
            x1_canvas = center_x - square_size / 2
            y1_canvas = center_y - square_size / 2
            x2_canvas = center_x + square_size / 2
            y2_canvas = center_y + square_size / 2

            # Zīmē rāmi
            self.qr_code_frame_id = self.canvas.create_rectangle(
                x1_canvas, y1_canvas, x2_canvas, y2_canvas,
                outline="yellow", width=3, dash=(5, 2), tags="qr_frame"
            )

            # Zīmē stūru rokturus
            handle_size = 10
            self.qr_code_handle_ids = []
            handles = [
                (x1_canvas, y1_canvas, "nw"),
                (x2_canvas, y1_canvas, "ne"),
                (x2_canvas, y2_canvas, "se"),
                (x1_canvas, y2_canvas, "sw")
            ]
            for x_handle, y_handle, cursor_type in handles:
                handle_id = self.canvas.create_oval(
                    x_handle - handle_size, y_handle - handle_size,
                    x_handle + handle_size, y_handle + handle_size,
                    fill="red", outline="white", width=3, tags="qr_handle"
                )
                self.qr_code_handle_ids.append({"id": handle_id, "type": cursor_type})

            # SVARĪGI: Pārvieto QR elementus uz priekšu (virs attēla)
            self.canvas.tag_raise("qr_frame")
            self.canvas.tag_raise("qr_handle")

            # Papildu pārbaude - pārvieto vēlreiz, lai būtu droši, ka ir virs
            for handle_data in self.qr_code_handle_ids:
                self.canvas.tag_raise(handle_data["id"])
            if self.qr_code_frame_id:
                self.canvas.tag_raise(self.qr_code_frame_id)

    def resize_canvas(self, event):
        """Pielāgo kanvasa izmērus un atjaunina attēla priekšskatījumu."""
        if self.current_image_index != -1:
            img = self.images[self.current_image_index]["processed_img"]
            self.show_image_preview(img)

    def on_canvas_mouse_wheel(self, event):
        """Apstrādā peles rullīša notikumus kanvasa tālummaiņai."""
        if self.current_image_index == -1: return
        if event.num == 5 or event.delta == -120:  # Tuvināt
            self.canvas_zoom_factor = max(0.1, self.canvas_zoom_factor - 0.1)
        if event.num == 4 or event.delta == 120:  # Attālināt
            self.canvas_zoom_factor = min(5.0, self.canvas_zoom_factor + 0.1)
        self.show_image_preview(self.images[self.current_image_index]["processed_img"])

    def on_canvas_pan_start(self, event):
        """Sāk kanvasa attēla pārvietošanu (pan)."""
        if self.current_image_index == -1: return
        self.canvas_start_pan_x = event.x - self.canvas_pan_x
        self.canvas_start_pan_y = event.y - self.canvas_pan_y
        self.canvas.config(cursor="fleur")

    def on_canvas_pan_drag(self, event):
        """Pārvieto kanvasa attēlu, velkot peli."""
        if self.current_image_index == -1: return
        self.canvas_pan_x = event.x - self.canvas_start_pan_x
        self.canvas_pan_y = event.y - self.canvas_start_pan_y
        self.show_image_preview(self.images[self.current_image_index]["processed_img"])

    def on_canvas_pan_end(self, event):
        """Beidz kanvasa attēla pārvietošanu."""
        if self.current_image_index == -1: return
        self.canvas.config(cursor="arrow")

    def on_canvas_selection_start(self, event):
        """Sāk atlases taisnstūra zīmēšanu vai QR koda rāmja vilkšanu uz kanvasa."""
        if self.current_image_index == -1: return

        # Pārbauda QR rāmja mijiedarbību tikai rediģēšanas režīmā
        if hasattr(self, '_qr_edit_mode') and self._qr_edit_mode:
            item = self.canvas.find_closest(event.x, event.y)[0]
            tags = self.canvas.gettags(item)

            if "qr_handle" in tags:
                self.qr_code_active_handle = next((h for h in self.qr_code_handle_ids if h["id"] == item), None)
                if self.qr_code_active_handle:
                    self.qr_code_drag_mode = 'resize'
                    self.qr_code_start_drag_x = event.x
                    self.qr_code_start_drag_y = event.y
                    cursor_map = {
                        "nw": "top_left_corner",
                        "ne": "top_right_corner",
                        "se": "bottom_right_corner",
                        "sw": "bottom_left_corner"
                    }
                    cursor_name = cursor_map.get(self.qr_code_active_handle["type"], "sizing")
                    self.canvas.config(cursor=cursor_name)
                    return
            elif "qr_frame" in tags:
                self.qr_code_drag_mode = 'move'
                self.qr_code_start_drag_x = event.x
                self.qr_code_start_drag_y = event.y
                self.canvas.config(cursor="fleur")
                return

        # Ja nav QR koda rāmja mijiedarbība, tad turpina ar apgriešanas vai atlases režīmu
        if self.cropping_mode:
            self.crop_start_x = self.canvas.canvasx(event.x)
            self.crop_start_y = self.canvas.canvasy(event.y)
            if self.crop_rect_id:
                self.canvas.delete(self.crop_rect_id)
            self.crop_rect_id = self.canvas.create_rectangle(self.crop_start_x, self.crop_start_y,
                                                             self.crop_start_x, self.crop_start_y,
                                                             outline="red", width=2, dash=(5, 2))
        else:
            self.canvas_selection_start_x = self.canvas.canvasx(event.x)
            self.canvas_selection_start_y = self.canvas.canvasy(event.y)
            if self.canvas_selection_rect:
                self.canvas.delete(self.canvas_selection_rect)
            self.canvas_selection_rect = self.canvas.create_rectangle(self.canvas_selection_start_x,
                                                                      self.canvas_selection_start_y,
                                                                      self.canvas_selection_start_x,
                                                                      self.canvas_selection_start_y,
                                                                      outline="blue", width=2, dash=(5, 2))

    def on_canvas_selection_drag(self, event):
        """Atjaunina atlases taisnstūra vai QR koda rāmja izmērus/pozīciju uz kanvasa, velkot peli."""
        if self.current_image_index == -1: return

        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)

        if self.qr_code_drag_mode == 'move' and self.qr_code_frame_coords:
            dx = (cur_x - self.qr_code_start_drag_x) / self.canvas_zoom_factor
            dy = (cur_y - self.qr_code_start_drag_y) / self.canvas_zoom_factor

            x1, y1, x2, y2 = self.qr_code_frame_coords
            self.qr_code_frame_coords = (x1 + dx, y1 + dy, x2 + dx, y2 + dy)

            self.qr_code_start_drag_x = cur_x
            self.qr_code_start_drag_y = cur_y
            # Atjauno tikai attēla priekšskatījumu bez pilnas pārzīmēšanas
            img_data = self.images[self.current_image_index]
            self.show_image_preview(img_data["processed_img"])



        elif self.qr_code_drag_mode == 'resize' and self.qr_code_active_handle and self.qr_code_frame_coords:

            x1, y1, x2, y2 = self.qr_code_frame_coords

            dx = (cur_x - self.qr_code_start_drag_x) / self.canvas_zoom_factor

            dy = (cur_y - self.qr_code_start_drag_y) / self.canvas_zoom_factor

            handle_type = self.qr_code_active_handle["type"]

            # Aprēķina jauno izmēru, saglabājot kvadrātisku formu

            if handle_type in ["nw", "se"]:

                # Diagonālie stūri - izmanto vidējo no dx un dy

                delta = (dx + dy) / 2

                if handle_type == "nw":

                    x1 += delta

                    y1 += delta

                else:  # se

                    x2 += delta

                    y2 += delta

            elif handle_type in ["ne", "sw"]:

                # Pretējie diagonālie stūri

                delta = (dx - dy) / 2

                if handle_type == "ne":

                    x2 += delta

                    y1 -= delta

                else:  # sw

                    x1 -= delta

                    y2 += delta

            # Nodrošina minimālo izmēru

            min_size = min(self.images[self.current_image_index]["processed_img"].size) * 0.05

            if (x2 - x1) < min_size or (y2 - y1) < min_size:
                return

            self.qr_code_frame_coords = (x1, y1, x2, y2)

            self.qr_code_start_drag_x = cur_x

            self.qr_code_start_drag_y = cur_y

            img_data = self.images[self.current_image_index]

            self.show_image_preview(img_data["processed_img"])


        elif self.cropping_mode and self.crop_rect_id:
            self.canvas.coords(self.crop_rect_id, self.crop_start_x, self.crop_start_y, cur_x, cur_y)
        elif self.canvas_selection_rect:
            self.canvas.coords(self.canvas_selection_rect, self.canvas_selection_start_x, self.canvas_selection_start_y,
                               cur_x, cur_y)

    def on_canvas_selection_end(self, event):
        """Beidz atlases taisnstūra vai QR koda rāmja zīmēšanu/vilkšanu uz kanvasa."""
        if self.current_image_index == -1: return

        self.canvas.config(cursor="arrow")  # Atjauno noklusējuma kursoru
        self.qr_code_drag_mode = None
        self.qr_code_active_handle = None
        self.qr_code_start_drag_x = None
        self.qr_code_start_drag_y = None

        if self.cropping_mode:
            end_x = self.canvas.canvasx(event.x)
            end_y = self.canvas.canvasy(event.y)

            x1, y1 = min(self.crop_start_x, end_x), min(self.crop_start_y, end_y)
            x2, y2 = max(self.crop_start_x, end_x), max(self.crop_start_y, end_y)

            img_data = self.images[self.current_image_index]
            img_pil = img_data["processed_img"]
            img_width, img_height = img_pil.size

            canvas_width = self.canvas.winfo_width()
            canvas_height = self.canvas.winfo_height()

            scaled_width = int(img_width * self.canvas_zoom_factor)
            scaled_height = int(img_height * self.canvas_zoom_factor)
            img_on_canvas_x = (canvas_width - scaled_width) / 2 + self.canvas_pan_x
            img_on_canvas_y = (canvas_height - scaled_height) / 2 + self.canvas_pan_y

            original_x1 = int((x1 - img_on_canvas_x) / self.canvas_zoom_factor)
            original_y1 = int((y1 - img_on_canvas_y) / self.canvas_zoom_factor)
            original_x2 = int((x2 - img_on_canvas_x) / self.canvas_zoom_factor)
            original_y2 = int((y2 - img_on_canvas_y) / self.canvas_zoom_factor)

            original_x1 = max(0, min(original_x1, img_width))
            original_y1 = max(0, min(original_y1, img_height))
            original_x2 = max(0, min(original_x2, img_width))
            original_y2 = max(0, min(original_y2, img_height))

            self.current_crop_coords = (original_x1, original_y1, original_x2, original_y2)

            if messagebox.askyesno("Apgriezt attēlu", "Vai vēlaties apgriezt attēlu ar atlasīto apgabalu?"):
                self.perform_crop()

            if self.crop_rect_id:
                self.canvas.delete(self.crop_rect_id)
                self.crop_rect_id = None
            self.toggle_cropping_mode()

        else:
            # Esošais kods parastai atlasei
            pass

    def perform_crop(self):
        """Veic attēla apgriešanu, pamatojoties uz saglabātajām koordinātām."""
        if self.current_image_index == -1 or not self.current_crop_coords:
            messagebox.showwarning("Kļūda", "Nav attēla vai apgriešanas koordinātu.")
            return

        img_data = self.images[self.current_image_index]
        img_pil = img_data["processed_img"]

        try:
            cropped_img = img_pil.crop(self.current_crop_coords)
            img_data["processed_img"] = cropped_img
            self.show_image_preview(cropped_img)
            messagebox.showinfo("Apgriešana", "Attēls veiksmīgi apgriezts.")
        except Exception as e:
            messagebox.showerror("Apgriešanas kļūda", f"Neizdevās apgriezt attēlu: {e}")
        finally:
            self.current_crop_coords = None  # Notīra koordinātas pēc apgriešanas

        # JAUNS: Mainīgie QR koda/svītrkoda rāmja attēlošanai un mijiedarbībai
        self.qr_code_frame_id = None  # Kanvasa ID QR koda rāmim
        self.qr_code_handle_ids = []  # Kanvasa ID rāmja stūru rokturiem
        self.qr_code_active_handle = None  # Aktīvais rokturis vilkšanas laikā
        self.qr_code_frame_coords = None  # (x1, y1, x2, y2) koordinātas QR koda rāmim attēla oriģinālajās koordinātās
        self.qr_code_start_drag_x = None  # Sākuma X koordināta vilkšanas laikā
        self.qr_code_start_drag_y = None  # Sākuma Y koordināta vilkšanas laikā
        self.qr_code_drag_mode = None  # 'move' vai 'resize'

    def open_fullscreen_preview(self):
        """Atver pašreizējo attēlu pilnekrāna priekšskatījuma logā."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu priekšskatījumam.")
            return
        img = self.images[self.current_image_index]["processed_img"]
        FullscreenImageViewer(self, img)

    def start_processing(self):
        """Sāk OCR apstrādes procesu atsevišķā pavedienā."""
        if not self.images:
            messagebox.showwarning("Nav attēlu", "Lūdzu, vispirms atlasiet attēlus!")
            return

        self.btn_start.config(state=DISABLED)
        self.btn_stop.config(state=NORMAL)
        self.progress.config(maximum=len(self.images))
        self.progress["value"] = 0
        self.ocr_results = [None] * len(self.images)
        self.stop_processing = False

        threading.Thread(target=self.process_images_thread, daemon=True).start()

    def stop_processing_func(self):
        """Aptur OCR apstrādes procesu."""
        self.stop_processing = True
        self.btn_stop.config(state=DISABLED)

    def process_images_thread(self):
        """OCR apstrādes loģika, kas darbojas atsevišķā pavedienā."""
        langs = []
        for lang_name, var in self.lang_vars.items():
            if var.get():
                langs.append(self.lang_options[lang_name])

        if not langs:
            langs = ["eng"]

        lang_param = "+".join(langs)
        tess_config = f"--psm {self.psm_var.get()} --oem {self.oem_var.get()}"
        confidence_threshold = self.confidence_var.get() / 100.0

        for i, img_data in enumerate(self.images):
            if self.stop_processing:
                messagebox.showinfo("Apstrāde apturēta", "OCR apstrāde tika apturēta pēc lietotāja pieprasījuma.")
                break

            filepath = img_data["filepath"]
            img_to_ocr = img_data["processed_img"]

            try:
                img_proc = self.preprocess_image_for_ocr(img_to_ocr, self.dpi_var.get())

                data = pytesseract.image_to_data(
                    img_proc,
                    lang=lang_param,
                    config=tess_config,
                    output_type=pytesseract.Output.DICT
                )

                processed_words = []
                full_text_lines = []
                current_line = []
                current_line_num = None

                for j in range(len(data["level"])):
                    if data["level"][j] == 5:
                        word = data["text"][j]
                        conf = float(data["conf"][j])

                        if word.strip() and conf >= confidence_threshold * 100:
                            x, y, w, h = data["left"][j], data["top"][j], data["width"][j], data["height"][j]
                            line_num = data["line_num"][j]

                            processed_words.append({
                                "text": word, "x": x, "y": y, "w": w, "h": h, "line_num": line_num
                            })

                            if current_line_num is None:
                                current_line_num = line_num
                            elif line_num != current_line_num:
                                full_text_lines.append(" ".join(current_line))
                                current_line = []
                                current_line_num = line_num
                            current_line.append(word)

                if current_line:
                    full_text_lines.append(" ".join(current_line))

                full_text = "\n".join(full_text_lines)
                is_empty_ocr = (full_text.strip() == "")  # Pārbauda, vai OCR rezultāts ir tukšs

                is_empty_ocr = (full_text.strip() == "")  # Pārbauda, vai OCR rezultāts ir tukšs

                self.ocr_results[i] = {
                    "full_text": full_text,
                    "word_data": processed_words,
                    "doc_id": str(uuid.uuid4())[:8].upper(),
                    "is_empty_ocr": is_empty_ocr  # Pievieno jaunu lauku
                }
                self.after(0, self.update_ocr_text, i)
                if hasattr(self, '_mark_file_as_processed'):
                    self.after(0, lambda: self._mark_file_as_processed(i))
                else:
                    print(f"Apstrādāts fails {i}")



            except Exception as e:
                self.ocr_results[i] = {
                    "full_text": f"[OCR kļūda failā {os.path.basename(filepath)}] {str(e)}",
                    "word_data": [],
                    "doc_id": None
                }
                self.after(0, self.update_ocr_text, i)
                # Aizstāt ar šo kodu
                if hasattr(self, '_mark_file_as_processed'):
                    self.after(0, lambda: self._mark_file_as_processed(i))
                else:
                    print(f"Apstrādāts fails {i}")
                    # Šeit varat pievienot failu apstrādes loģiku

            self.after(0, lambda: self.progress.config(value=i + 1))

        if not self.stop_processing and any(res is not None for res in self.ocr_results):
            # Pārbauda, vai apstrāde tika sākta no kameras skenēšanas
            if hasattr(self, '_camera_scan_in_progress') and self._camera_scan_in_progress:
                self.after(0, self.save_pdf, True)  # Automātiski saglabā PDF
                self._camera_scan_in_progress = False
            else:
                self.after(0, self.save_pdf)  # Piedāvā saglabāt PDF
        else:
            if hasattr(self, '_camera_scan_in_progress') and self._camera_scan_in_progress:
                self._camera_scan_in_progress = False

        self.after(0, lambda: self.btn_start.config(state=NORMAL))
        self.after(0, lambda: self.btn_stop.config(state=DISABLED))

        # Pievieno šo kodu pēc OCR apstrādes
        decoded_barcodes = self.detect_and_decode_barcodes(img_to_ocr)
        if decoded_barcodes:
            self.ocr_results[i]["decoded_barcodes"] = decoded_barcodes  # Pievieno atšifrētos datus rezultātiem

    def preprocess_image_for_ocr(self, img, dpi=300):
        """Veic attēla priekšapstrādi OCR vajadzībām (pelēktoņi, automātiskais kontrasts, DPI pielāgošana)."""
        if img.mode != 'L':
            img = img.convert("L")
        img = ImageOps.autocontrast(img)
        original_dpi = img.info.get("dpi", (72, 72))[0]
        if original_dpi < dpi:
            scale_factor = dpi / original_dpi
            new_size = (int(img.width * scale_factor), int(img.height * scale_factor))
            img = img.resize(new_size, Image.LANCZOS)
        return img

    def _mark_file_as_processed(self, index):
        """
        Iezīmē failu kā apstrādātu, mainot fona krāsu
        :param index: faila indekss sarakstā
        """
        try:
            if 0 <= index < self.file_listbox.size():  # Pārbauda derīgu indeksu
                self.file_listbox.itemconfig(index, {'bg': '#49be25'})  # Gaiši zaļa krāsa
                print(f"Fails ar indeksu {index} ir apstrādāts.")
        except Exception as e:
            print(f"Kļūda atzīmējot failu: {e}")

    def update_ocr_text(self, index):
        """Atjaunina OCR teksta lauku ar jaunākajiem rezultātiem."""
        selection = self.file_listbox.curselection()
        if selection and selection[0] == index and self.ocr_results[index]:
            self.text_ocr.delete("1.0", tk.END)
            self.text_ocr.insert(tk.END, self.ocr_results[index]["full_text"])

        # Atrodiet klases OCRPDFApp beigu daļu (pirms pēdējās "if __name__" rindas)
        # un ievietojiet šo jauno metodi:

        def _mark_file_as_processed(self, index):
            """Iezīmē failu kā apstrādātu, mainot fona krāsu"""
            if 0 <= index < self.file_listbox.size():  # Pārbauda derīgu indeksu
                self.file_listbox.itemconfig(index, {'bg': '#d4edda'})
                print(f"Fails ar indeksu {index} ir apstrādāts.")

    def classify_document(self, ocr_text):
        """Klasificē dokumentu, pamatojoties uz OCR tekstu un nosaka, vai tas ir sensitīvs."""
        ocr_text_lower = ocr_text.lower()

        # Pārbauda, vai dokuments ir sensitīvs
        is_sensitive = False
        for keyword in self.document_keywords["id_card"]:
            if keyword in ocr_text_lower:
                is_sensitive = True
                break

        if "pavadzīme" in ocr_text_lower or "delivery note" in ocr_text_lower or "shipping document" in ocr_text_lower:
            category = "Pavadzīmes"
        elif "rēķins" in ocr_text_lower or "invoice" in ocr_text_lower or "bill" in ocr_text_lower:
            category = "Rēķini"
        elif "līgums" in ocr_text_lower or "contract" in ocr_text_lower or "agreement" in ocr_text_lower:
            category = "Līgumi"
        elif "kvīts" in ocr_text_lower or "receipt" in ocr_text_lower or "payment confirmation" in ocr_text_lower:
            category = "Kvītis"
        elif "protokols" in ocr_text_lower or "minutes" in ocr_text_lower or "report" in ocr_text_lower:
            category = "Protokoli"
        elif "pasūtījums" in ocr_text_lower or "order" in ocr_text_lower or "purchase order" in ocr_text_lower:
            category = "Pasūtījumi"
        elif "darba laika uzskaite" in ocr_text_lower or "timesheet" in ocr_text_lower:
            category = "Darba laika uzskaite"
        elif "personāla dokuments" in ocr_text_lower or "personnel document" in ocr_text_lower or "employee record" in ocr_text_lower:
            category = "Personāla dokumenti"
        elif is_sensitive:  # Ja ir sensitīvs, bet nav citā kategorijā
            category = "Sensitīvi dokumenti"
        else:
            category = "Nekategorizēti"  # Ja nav atrasta neviena atbilstoša kategorija

        return category, is_sensitive

    def get_or_create_folder(self, folder_name):
        """Atgriež mapes objektu vai izveido jaunu, ja tā neeksistē."""
        for item in self.internal_file_system["contents"]:
            if item["type"] == "folder" and item["name"] == folder_name:
                return item
        # Ja mape neeksistē, izveido to saknes mapē
        new_folder = {"type": "folder", "name": folder_name, "contents": [], "parent": self.internal_file_system}
        self.internal_file_system["contents"].append(new_folder)
        return new_folder

    def save_pdf(self, auto_save=False):
        """Saglabā apstrādātos attēlus kā PDF failu ar meklējamu tekstu."""
        # Pārbauda, vai ir kādi rezultāti (arī tukši OCR rezultāti ir derīgi saglabāšanai)
        if not any(res is not None for res in self.ocr_results):
            messagebox.showwarning("Nav datu", "Nav neviena apstrādāta attēla saglabāšanai!")
            return

        # Pārbauda, vai visi rezultāti ir tukši
        all_empty = all(
            res is None or (not res["word_data"] and res.get("is_empty_ocr", False))
            for res in self.ocr_results
        )

        if all_empty:
            # Ja visi ir tukši, joprojām ļauj saglabāt, bet brīdina
            result = messagebox.askyesno("Tukši OCR rezultāti",
                                         "Nevienam attēlam nav atrasts teksts. Vai vēlaties saglabāt tukšu PDF?")
            if not result:
                return

        # Ģenerē unikālu dokumenta ID
        doc_id = str(uuid.uuid4())[:8].upper()  # Īss, unikāls ID
        current_date = datetime.now().strftime("%Y-%m-%d")

        # Klasificē dokumentu un iegūst ieteikto mapi
        first_ocr_text = self.ocr_results[0]["full_text"] if self.ocr_results and self.ocr_results[0] else ""
        document_category, is_sensitive = self.classify_document(first_ocr_text)

        # Pārbauda, vai ir attēli bez OCR rezultātiem
        has_empty_ocr = any(
            self.ocr_results[i] and self.ocr_results[i].get("is_empty_ocr", False)
            for i in range(len(self.ocr_results))
            if self.ocr_results[i] is not None
        )

        # Ja ir tukši OCR rezultāti, iesaka "Bez OCR" mapi
        if has_empty_ocr:
            document_category = "Bez OCR"
            if not auto_save:
                messagebox.showinfo("Informācija",
                                    "Dažiem attēliem nav atrasts teksts. Tie tiks saglabāti mapē 'Bez OCR'.")

        # --- JAUNS KODS SĀKAS ŠEIT ---
        # Pirms saglabāšanas dialoga, piedāvā izvēlēties mapi
        selected_folder_node = None
        if not auto_save:
            # Atver mapes izvēles dialogu
            selected_folder_node = self._show_folder_selection_dialog(document_category)
            if selected_folder_node is None:  # Ja lietotājs atceļ mapes izvēli
                return

        # Ja lietotājs izvēlējās mapi, izmanto to, citādi izmanto klasificēto
        if selected_folder_node:
            target_folder_node = selected_folder_node
            # Atjaunina document_category, ja lietotājs izvēlējās citu mapi
            document_category = target_folder_node["name"]
        else:
            # Ja auto_save vai lietotājs neizvēlējās citu mapi, izmanto klasificēto
            target_folder_node = self.get_or_create_folder(document_category)

        # Izveido fizisko mapi, ja tā neeksistē
        # Šeit ir svarīgi izmantot `self.default_save_path` kā bāzes ceļu
        # un pievienot visu ceļu līdz izvēlētajai mapei iekšējā failu sistēmā.
        # Lai to izdarītu, mums ir jāizveido funkcija, kas atgriež pilnu fizisko ceļu no mapes mezgla.
        physical_category_path = self._get_physical_path_from_node(target_folder_node)
        os.makedirs(physical_category_path, exist_ok=True)

        default_filename = f"{document_category}_{current_date}_{doc_id}.pdf"
        full_save_path = os.path.join(physical_category_path, default_filename)

        out_path = full_save_path
        if not auto_save:
            out_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF faili", "*.pdf")],
                initialdir=physical_category_path,  # Sākotnējais direktorijs ir izvēlētā mape
                initialfile=default_filename,
                title="Saglabāt PDF kā"
            )

        if not out_path:
            return
        # --- JAUNS KODS BEIDZAS ŠEIT ---

        c = canvas.Canvas(out_path)
        fontsize = self.fontsize_var.get()

        jpeg_quality = 85
        if self.pdf_quality == "Zema (60)":
            jpeg_quality = 60
        elif self.pdf_quality == "Augsta (95)":
            jpeg_quality = 95

        for i, img_data in enumerate(self.images):
            if self.stop_processing or i >= len(self.ocr_results) or not self.ocr_results[i]:
                continue

            img_for_pdf = img_data["processed_img"]  # <--- Nomainīts uz processed_img
            img_width, img_height = img_for_pdf.size

            page_size = self.get_page_size(img_width, img_height)
            c.setPageSize(page_size)
            page_width, page_height = page_size

            img_ratio = img_width / img_height
            page_ratio = page_width / page_height

            if img_ratio > page_ratio:
                draw_width = page_width
                draw_height = page_width / img_ratio
            else:
                draw_height = page_height
                draw_width = page_height * img_ratio

            x_offset = (page_width - draw_width) / 2
            y_offset = (page_height - draw_height) / 2

            pil_img = img_for_pdf.convert("RGB")
            temp_img_path = "temp_image_for_pdf.jpg"
            pil_img.save(temp_img_path, "JPEG", quality=jpeg_quality)
            image_reader = ImageReader(temp_img_path)
            c.drawImage(image_reader, x_offset, y_offset, width=draw_width, height=draw_height)
            os.remove(temp_img_path)

            if self.include_text_var.get() and self.ocr_results[i]["word_data"]:
                word_data = self.ocr_results[i]["word_data"]
                processed_img_width, processed_img_height = img_data["processed_img"].size

                scale_x_ocr_to_pdf = draw_width / processed_img_width
                scale_y_ocr_to_pdf = draw_height / processed_img_height

                c.setFillColorRGB(0, 0, 0, 0)  # Caurspīdīgs teksts
                c.setFont("Helvetica", fontsize)

                for word in word_data:
                    x = x_offset + word["x"] * scale_x_ocr_to_pdf
                    y = y_offset + (processed_img_height - word["y"] - word["h"]) * scale_y_ocr_to_pdf

                    word_height_pdf = word["h"] * scale_y_ocr_to_pdf
                    actual_fontsize = max(fontsize, word_height_pdf * 0.8)  # Pielāgo fonta izmēru vārda augstumam
                    c.setFont("Helvetica", actual_fontsize)

                    text_obj = c.beginText(x, y)
                    text_obj.textLine(word["text"])
                    c.drawText(text_obj)

            # Pievieno QR kodu vai svītrkodu ar dokumenta ID, ja iestatīts
            # JAUNS: Pievieno QR kodu vai svītrkodu ar dokumenta ID, ja iestatīts un rāmis ir definēts
            if self.settings.get("add_id_code_to_pdf", False) and self.qr_code_frame_coords:
                id_code_type = self.settings.get("id_code_type", "QR")
                # id_code_position vairs netiek izmantots tieši, jo pozīciju nosaka qr_code_frame_coords

                # Pārbauda, vai ir pieejams doc_id
                current_doc_id = self.ocr_results[i]["doc_id"] if self.ocr_results[i] else None
                if current_doc_id:
                    code_reader = None  # Inicializē code_reader
                    temp_code_path = None  # Inicializē temp_code_path

                    try:
                        if id_code_type == "QR":
                            qr_img = qrcode.make(current_doc_id)
                            qr_pil_img = qr_img.get_image()
                            temp_code_path = os.path.join(tempfile.gettempdir(),
                                                          "temp_qr_code.png")  # Izmanto temp direktoriju
                            qr_pil_img.save(temp_code_path)
                            code_reader = ImageReader(temp_code_path)
                        elif id_code_type == "Barcode":
                            try:
                                from barcode import Code128
                                from barcode.writer import ImageWriter

                                temp_dir = tempfile.gettempdir()
                                temp_code_path = os.path.join(temp_dir, f"temp_barcode_{i}.png")

                                with open(temp_code_path, 'wb') as f:
                                    Code128(current_doc_id, writer=ImageWriter()).write(f)

                                if os.path.exists(temp_code_path):
                                    code_reader = ImageReader(temp_code_path)
                                else:
                                    print(f"Nevarēja atrast ģenerēto svītrkoda failu: {temp_code_path}")
                                    continue
                            except Exception as e:
                                print(f"Svītrkoda ģenerēšanas kļūda: {e}")
                                continue
                        elif id_code_type == "Code39":
                            try:
                                from barcode import Code39
                                from barcode.writer import ImageWriter
                                temp_dir = tempfile.gettempdir()
                                temp_code_path = os.path.join(temp_dir, f"temp_code39_{i}.png")
                                with open(temp_code_path, 'wb') as f:
                                    Code39(current_doc_id, writer=ImageWriter()).write(f)
                                if os.path.exists(temp_code_path):
                                    code_reader = ImageReader(temp_code_path)
                                else:
                                    print(f"Nevarēja atrast ģenerēto Code39 failu: {temp_code_path}")
                                    continue
                            except Exception as e:
                                print(f"Code39 ģenerēšanas kļūda: {e}")
                                continue
                        elif id_code_type == "EAN13":
                            try:
                                from barcode import EAN13
                                from barcode.writer import ImageWriter
                                temp_dir = tempfile.tempdir
                                temp_code_path = os.path.join(temp_dir, f"temp_ean13_{i}.png")
                                if len(current_doc_id) >= 12 and current_doc_id.isdigit():
                                    with open(temp_code_path, 'wb') as f:
                                        EAN13(current_doc_id[:12], writer=ImageWriter()).write(f)
                                    if os.path.exists(temp_code_path):
                                        code_reader = ImageReader(temp_code_path)
                                    else:
                                        print(f"Nevarēja atrast ģenerēto EAN-13 failu: {temp_code_path}")
                                        continue
                                else:
                                    print(f"Nederīgs ID EAN-13 ģenerēšanai: {current_doc_id}. Nepieciešami 12 cipari.")
                                    continue
                            except Exception as e:
                                print(f"EAN-13 ģenerēšanas kļūda: {e}")
                                continue

                        if code_reader:
                            # Izmanto interaktīvi iestatītās koordinātas
                            x1_img, y1_img, x2_img, y2_img = self.qr_code_frame_coords

                            # Pārrēķina koordinātas no attēla pikseļiem uz PDF punktiem
                            # Jāņem vērā, ka PDF lapas izmērs var atšķirties no attēla izmēra
                            # un attēls uz PDF lapas var būt mērogots un centrēts.
                            # Tāpēc ir jāizmanto tie paši mērogošanas faktori un nobīdes, kas tika izmantoti attēla zīmēšanai.

                            # Mērogošanas faktori no oriģinālā attēla uz PDF lapas attēlojumu
                            scale_x_img_to_pdf = draw_width / img_width
                            scale_y_img_to_pdf = draw_height / img_height

                            # QR koda rāmja koordinātas PDF lapas koordinātās
                            code_x_pos = x_offset + x1_img * scale_x_img_to_pdf
                            code_y_pos = y_offset + (img_height - y2_img) * scale_y_img_to_pdf  # Y ass ir apgriezta
                            code_width = (x2_img - x1_img) * scale_x_img_to_pdf
                            code_height = (y2_img - y1_img) * scale_y_img_to_pdf

                            c.drawImage(code_reader, code_x_pos, code_y_pos, width=code_width, height=code_height)
                            os.remove(temp_code_path)  # Dzēš pagaidu failu
                    except ImportError:
                        messagebox.showwarning("Trūkst bibliotēkas",
                                               "Lai ģenerētu svītrkodus, lūdzu, instalējiet 'python-barcode' un 'Pillow' (pip install python-barcode Pillow).")
                    except Exception as code_e:
                        print(f"Koda ģenerēšanas kļūda: {code_e}")

            c.showPage()

        c.save()

        # Pievieno paroles aizsardzību, ja dokuments ir sensitīvs
        if is_sensitive:
            messagebox.showwarning("Datu aizsardzība",
                                   "Dokuments, iespējams, satur sensitīvus datus. Tiks pievienota paroles aizsardzība.")
            password = simpledialog.askstring("Paroles aizsardzība", "Ievadiet paroli PDF failam:", show='*',
                                              parent=self)
            if password:
                try:
                    reader = pypdf.PdfReader(out_path)
                    writer = pypdf.PdfWriter()

                    for page in reader.pages:
                        writer.add_page(page)

                    writer.encrypt(password)

                    # Pagaidu fails, lai saglabātu šifrēto PDF
                    encrypted_pdf_path = out_path.replace(".pdf", "_encrypted.pdf")
                    with open(encrypted_pdf_path, "wb") as f:
                        writer.write(f)

                    os.remove(out_path)  # Dzēš nešifrēto failu
                    os.rename(encrypted_pdf_path, out_path)  # Pārdēvē šifrēto failu uz oriģinālo nosaukumu
                    messagebox.showinfo("Gatavs", f"PDF veiksmīgi saglabāts un šifrēts:\n{out_path}")
                except Exception as e:
                    messagebox.showerror("Kļūda šifrēšanā", f"Neizdevās šifrēt PDF failu: {e}\n"
                                                            "PDF saglabāts bez paroles.")
                    messagebox.showinfo("Gatavs", f"PDF veiksmīgi saglabāts:\n{out_path}")
            else:
                messagebox.showinfo("Gatavs", f"PDF veiksmīgi saglabāts (bez paroles):\n{out_path}")
        else:
            messagebox.showinfo("Gatavs", f"PDF veiksmīgi saglabāts:\n{out_path}")

        # Pievieno failu iekšējai failu sistēmai atbilstošajā mapē
        target_folder_node["contents"].append({
            "type": "file",
            "name": os.path.basename(out_path),
            "filepath": out_path,
            "doc_id": doc_id,
            "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "parent": target_folder_node  # Svarīgi atjaunināt parent atsauci
        })
        self.refresh_pdf_list()  # Atjauno failu sarakstu failu pārvaldības cilnē
        # JAUNS: Automātiskā augšupielāde
        if self.auto_upload_enabled.get():
            upload_target = self.auto_upload_target.get()
            # Izmanto dokumenta kategoriju kā attālo apakšmapi
            self.upload_file_to_remote(out_path, upload_target, document_category)

    def get_page_size(self, img_width, img_height):
        """Aprēķina PDF lapas izmēru, pamatojoties uz izvēlēto orientāciju un attēla izmēriem."""
        orientation = self.orientation_var.get()
        points_per_pixel = 72 / self.dpi_var.get()

        img_width_pt = img_width * points_per_pixel
        img_height_pt = img_height * points_per_pixel

        if orientation == "Tāds pats kā attēls":
            return (img_width_pt, img_height_pt)
        elif orientation == "Auto":
            return (img_width_pt, img_height_pt) if img_width_pt < img_height_pt else landscape(
                (img_width_pt, img_height_pt))
        elif orientation == "Portrets":
            return (min(img_width_pt, img_height_pt), max(img_width_pt, img_height_pt))
        elif orientation == "Ainava":
            return landscape((min(img_width_pt, img_height_pt), max(img_width_pt, img_height_pt)))
        elif orientation == "A4 Portrets":
            return A4
        elif orientation == "A4 Ainava":
            return landscape(A4)
        elif orientation == "Letter Portrets":
            return letter
        elif orientation == "Letter Ainava":
            return landscape(letter)
        else:
            return A4

    def crop_image(self):
        """Atver logu attēla apgriešanai."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, ko apgriezt.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        crop_window = Toplevel(self)
        crop_window.title("Apgriezt attēlu")
        crop_window.geometry("800x600")
        crop_window.transient(self)
        crop_window.grab_set()

        crop_canvas = tk.Canvas(crop_window, bg="gray", cursor="cross")
        crop_canvas.pack(fill="both", expand=True)

        # Pārmēro attēlu, lai tas ietilptu kanvasā
        canvas_width = crop_canvas.winfo_width()
        canvas_height = crop_canvas.winfo_height()
        img_width, img_height = img.size

        # Pagaidām izmantojam fiksētu izmēru, lai iegūtu kanvasa izmērus
        # Pēc tam, kad logs ir parādīts, varam iegūt reālos izmērus
        crop_window.update_idletasks()
        canvas_width = crop_canvas.winfo_width()
        canvas_height = crop_canvas.winfo_height()

        if canvas_width == 0 or canvas_height == 0:  # Ja logs vēl nav pilnībā inicializēts
            canvas_width = 800  # Noklusējuma vērtības
            canvas_height = 600

        ratio = min(canvas_width / img_width, canvas_height / img_height)
        display_w = int(img_width * ratio)
        display_h = int(img_height * ratio)

        display_img = img.resize((display_w, display_h), Image.LANCZOS)
        original_img_tk = ImageTk.PhotoImage(display_img)
        crop_canvas.create_image(0, 0, anchor="nw", image=original_img_tk)
        crop_canvas.image = original_img_tk

        rect_id = None
        start_x = start_y = end_x = end_y = 0

        def on_button_press(event):
            nonlocal start_x, start_y, rect_id
            start_x = crop_canvas.canvasx(event.x)
            start_y = crop_canvas.canvasy(event.y)
            if rect_id:
                crop_canvas.delete(rect_id)
            rect_id = crop_canvas.create_rectangle(start_x, start_y, start_x, start_y, outline="red", width=2)

        def on_mouse_drag(event):
            nonlocal end_x, end_y
            end_x = crop_canvas.canvasx(event.x)
            end_y = crop_canvas.canvasy(event.y)
            crop_canvas.coords(rect_id, start_x, start_y, end_x, end_y)

        def on_button_release(event):
            nonlocal end_x, end_y
            end_x = crop_canvas.canvasx(event.x)
            end_y = crop_canvas.canvasy(event.y)

            x1, y1 = min(start_x, end_x), min(start_y, end_y)
            x2, y2 = max(start_x, end_x), max(start_y, end_y)

            # Pārrēķina koordinātas uz oriģinālā attēla izmēriem
            # Ņem vērā mērogošanas koeficientu, ko izmantoja attēla parādīšanai kanvasā
            original_x1 = int(x1 / ratio)
            original_y1 = int(y1 / ratio)
            original_x2 = int(x2 / ratio)
            original_y2 = int(y2 / ratio)

            try:
                cropped_img = img.crop((original_x1, original_y1, original_x2, original_y2))
                img_data["processed_img"] = cropped_img
                self.show_image_preview(cropped_img)
                crop_window.destroy()
            except Exception as e:
                messagebox.showerror("Apgriešanas kļūda", f"Neizdevās apgriezt attēlu: {e}")

        crop_canvas.bind("<ButtonPress-1>", on_button_press)
        crop_canvas.bind("<B1-Motion>", on_mouse_drag)
        crop_canvas.bind("<ButtonRelease-1>", on_button_release)

        # Pielāgo kanvas izmēru, kad logs tiek parādīts
        crop_canvas.update_idletasks()
        crop_canvas.config(scrollregion=crop_canvas.bbox("all"))

    def rotate_90_degrees(self):
        """Pagriež pašreizējo attēlu par 90 grādiem pulksteņrādītāja virzienā."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return
        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"]
        img = img.rotate(-90, expand=True, fillcolor=(255, 255, 255) if img.mode == 'RGB' else 255)
        img_data["processed_img"] = img
        self.show_image_preview(img)

    def flip_image(self, method):
        """Spoguļo pašreizējo attēlu (horizontāli vai vertikāli)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return
        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"]
        img = img.transpose(method)
        img_data["processed_img"] = img
        self.show_image_preview(img)

    def resize_image_dialog(self):
        """Atver dialoga logu attēla izmēru maiņai."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"]

        resize_window = Toplevel(self)
        resize_window.title("Mainīt attēla izmērus")
        resize_window.transient(self)
        resize_window.grab_set()

        current_width, current_height = img.size

        ttk.Label(resize_window, text=f"Pašreizējie izmēri: {current_width}x{current_height} pikseļi").pack(pady=5)

        ttk.Label(resize_window, text="Jauns platums:").pack(pady=2)
        new_width_var = tk.IntVar(value=current_width)
        ttk.Entry(resize_window, textvariable=new_width_var).pack(pady=2)

        ttk.Label(resize_window, text="Jauns augstums:").pack(pady=2)
        new_height_var = tk.IntVar(value=current_height)
        ttk.Entry(resize_window, textvariable=new_height_var).pack(pady=2)

        keep_aspect_ratio_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(resize_window, text="Saglabāt proporcijas", variable=keep_aspect_ratio_var).pack(pady=5)

        def update_height(*args):
            if keep_aspect_ratio_var.get():
                try:
                    new_w = new_width_var.get()
                    if new_w > 0:
                        new_h = int(new_w * (current_height / current_width))
                        new_height_var.set(new_h)
                except:
                    pass

        def update_width(*args):
            if keep_aspect_ratio_var.get():
                try:
                    new_h = new_height_var.get()
                    if new_h > 0:
                        new_w = int(new_h * (current_width / current_height))
                        new_width_var.set(new_w)
                except:
                    pass

        new_width_var.trace_add("write", update_height)
        new_height_var.trace_add("write", update_width)

        def apply_resize():
            try:
                new_w = new_width_var.get()
                new_h = new_height_var.get()

                if new_w <= 0 or new_h <= 0:
                    messagebox.showwarning("Kļūda", "Platumam un augstumam jābūt pozitīviem skaitļiem.")
                    return

                resized_img = img.resize((new_w, new_h), Image.LANCZOS)
                img_data["processed_img"] = resized_img
                self.show_image_preview(resized_img)
                resize_window.destroy()

            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās mainīt izmērus: {e}")

        ttk.Button(resize_window, text="Mainīt izmērus", command=apply_resize).pack(pady=10)

    def auto_enhance_image(self):
        """Automātiski uzlabo pašreizējā attēla kontrastu un asumu."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        img = ImageOps.autocontrast(img)
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.5)

        img_data["processed_img"] = img
        self.show_image_preview(img)
        messagebox.showinfo("Automātiska uzlabošana", "Attēls ir automātiski uzlabots (kontrasts un asums).")

    def show_image_histogram(self):
        """Parāda pašreizējā attēla krāsu histogrammu."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, lai rādītu histogrammu.")
            return

        img = self.images[self.current_image_index]["processed_img"]
        if img.mode != 'L' and img.mode != 'RGB':
            img = img.convert('RGB')  # Pārvērš uz RGB, ja nav atbalstīts režīms

        hist_window = Toplevel(self)
        hist_window.title("Attēla histogramma")
        hist_window.transient(self)
        hist_window.grab_set()

        canvas_hist = tk.Canvas(hist_window, width=300, height=200, bg="white")
        canvas_hist.pack(padx=10, pady=10)

        if img.mode == 'L':  # Pelēktoņi
            histogram = img.histogram()
            max_val = max(histogram)
            for i, val in enumerate(histogram):
                height = (val / max_val) * 180  # Mērogo augstumu
                canvas_hist.create_rectangle(i, 200 - height, i + 1, 200, fill="gray")
        elif img.mode == 'RGB':  # Krāsains attēls
            r, g, b = img.split()
            hist_r = r.histogram()
            hist_g = g.histogram()
            hist_b = b.histogram()
            max_val = max(max(hist_r), max(hist_g), max(hist_b))

            for i in range(256):
                height_r = (hist_r[i] / max_val) * 180
                height_g = (hist_g[i] / max_val) * 180
                height_b = (hist_b[i] / max_val) * 180
                canvas_hist.create_rectangle(i, 200 - height_r, i + 1, 200, fill="red", outline="red")
                canvas_hist.create_rectangle(i, 200 - height_g, i + 1, 200, fill="green", outline="green")
                canvas_hist.create_rectangle(i, 200 - height_b, i + 1, 200, fill="blue", outline="blue")
        else:
            messagebox.showwarning("Histogramma", "Histogramma nav pieejama šim attēla režīmam.")
            hist_window.destroy()

    def show_image_metadata(self):
        """Parāda pašreizējā attēla metadatus (EXIF utt.)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, lai rādītu metadatus.")
            return

        img = self.images[self.current_image_index]["original_img"]
        metadata = img.info

        if not metadata:
            messagebox.showinfo("Metadati", "Šim attēlam nav pieejami metadati (EXIF utt.).")
            return

        meta_str = "Attēla metadati:\n"
        for key, value in metadata.items():
            meta_str += f"{key}: {value}\n"

        meta_window = Toplevel(self)
        meta_window.title("Attēla metadati")
        meta_window.transient(self)
        meta_window.grab_set()

        text_widget = tk.Text(meta_window, wrap="word", width=60, height=20)
        text_widget.pack(padx=10, pady=10)
        text_widget.insert(tk.END, meta_str)
        text_widget.config(state=DISABLED)  # Padara tekstu nelabojamu

    def show_color_palette(self):
        """Parāda pašreizējā attēla dominējošās krāsas."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, lai rādītu krāsu paleti.")
            return

        img = self.images[self.current_image_index]["processed_img"]
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Samazina attēlu, lai paātrinātu apstrādi
        img_small = img.resize((100, 100))
        colors = img_small.getcolors(img_small.size[0] * img_small.size[1])  # Iegūst visas krāsas un to skaitu

        if not colors:
            messagebox.showinfo("Krāsu palete", "Neizdevās iegūt krāsu informāciju.")
            return

        # Sakārto krāsas pēc biežuma (dominējošās krāsas)
        colors.sort(key=lambda x: x[0], reverse=True)
        top_colors = colors[:10]  # Top 10 dominējošās krāsas

        palette_window = Toplevel(self)
        palette_window.title("Attēla krāsu palete")
        palette_window.transient(self)
        palette_window.grab_set()

        canvas_palette = tk.Canvas(palette_window, width=300, height=150, bg="white")
        canvas_palette.pack(padx=10, pady=10)

        x_offset = 10
        y_offset = 10
        color_block_size = 30

        for count, (r, g, b) in top_colors:
            hex_color = f"#{r:02x}{g:02x}{b:02x}"
            canvas_palette.create_rectangle(x_offset, y_offset,
                                            x_offset + color_block_size, y_offset + color_block_size,
                                            fill=hex_color, outline="black")
            canvas_palette.create_text(x_offset + color_block_size + 5, y_offset + color_block_size / 2,
                                       anchor="w", text=f"{hex_color} (Count: {count})")
            y_offset += color_block_size + 5

    def compare_images(self):
        """Salīdzina divus attēlus un parāda atšķirības."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, ko salīdzināt.")
            return

        img1 = self.images[self.current_image_index]["processed_img"]

        filepaths = filedialog.askopenfilenames(
            title="Izvēlieties otru attēlu salīdzināšanai",
            filetypes=[("Attēlu faili", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("Visi faili", "*.*")]
        )
        if not filepaths:
            return

        try:
            img2 = Image.open(filepaths[0])
            if img1.size != img2.size:
                messagebox.showwarning("Izmēru neatbilstība", "Attēliem jābūt vienāda izmēra salīdzināšanai.")
                return
            if img1.mode != img2.mode:
                img2 = img2.convert(img1.mode)

            diff_img = ImageChops.difference(img1, img2)

            # Parāda atšķirību attēlu
            diff_window = Toplevel(self)
            diff_window.title("Attēlu atšķirības")
            diff_window.transient(self)
            diff_window.grab_set()

            diff_canvas = tk.Canvas(diff_window, bg="black")
            diff_canvas.pack(fill="both", expand=True)

            # Pielāgo attēlu kanvasa izmēram
            diff_window.update_idletasks()  # Atjaunina loga izmērus, lai iegūtu pareizus canvas_width/height
            canvas_width = diff_canvas.winfo_width()
            canvas_height = diff_canvas.winfo_height()

            if canvas_width == 0 or canvas_height == 0:  # Fallback if window not yet rendered
                canvas_width = 600
                canvas_height = 400

            display_diff_img = diff_img.resize((canvas_width, canvas_height), Image.LANCZOS)
            self.diff_photo = ImageTk.PhotoImage(display_diff_img)
            diff_canvas.create_image(0, 0, anchor="nw", image=self.diff_photo)
            diff_canvas.image = self.diff_photo

            messagebox.showinfo("Attēlu salīdzināšana",
                                "Atšķirību attēls ir parādīts jaunā logā. Jo gaišāks pikselis, jo lielāka atšķirība.")

        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās salīdzināt attēlus: {e}")

    def evaluate_image_quality(self):
        """Novērtē pašreizējā attēla kvalitāti (piem., trokšņa līmeni, asumu)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img = self.images[self.current_image_index]["processed_img"]

        # Vienkāršota asuma un trokšņa novērtēšana
        # Asums: Augsta frekvence attēlā (Laplacian variants)
        if img.mode != 'L':
            img_gray = img.convert('L')
        else:
            img_gray = img

        img_np = np.array(img_gray)

        if OPENCV_AVAILABLE:
            # Asums, izmantojot Laplasiāna operatoru
            laplacian_var = cv2.Laplacian(img_np, cv2.CV_64F).var()

            # Trokšņa līmenis (vienkāršots: standarta novirze pelēktoņu attēlā)
            noise_level = np.std(img_np)

            quality_report = f"Attēla kvalitātes novērtējums:\n" \
                             f"Asums (Laplacian Variance): {laplacian_var:.2f}\n" \
                             f"Trokšņa līmenis (Standard Deviation): {noise_level:.2f}\n\n" \
                             f"Augstāka Laplasiāna variācija norāda uz lielāku asumu.\n" \
                             f"Augstāka standarta novirze var norādīt uz lielāku trokšņa līmeni."
        else:
            quality_report = "OpenCV nav pieejams, tāpēc kvalitātes novērtēšana ir ierobežota.\n" \
                             "Lūdzu, instalējiet 'opencv-python' un 'numpy', lai iegūtu pilnu funkcionalitāti."

        messagebox.showinfo("Attēla kvalitāte", quality_report)

    def extract_text_from_region(self):
        """Ļauj lietotājam atlasīt apgabalu attēlā un veikt OCR tikai šajā apgabalā."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img = self.images[self.current_image_index]["processed_img"].copy()

        extract_window = Toplevel(self)
        extract_window.title("Izvilkt tekstu no apgabala")
        extract_window.geometry("800x600")
        extract_window.transient(self)
        extract_window.grab_set()

        extract_canvas = tk.Canvas(extract_window, bg="gray", cursor="cross")
        extract_canvas.pack(fill="both", expand=True)

        # Pārmēro attēlu, lai tas ietilptu kanvasā
        extract_window.update_idletasks()
        canvas_width = extract_canvas.winfo_width()
        canvas_height = extract_canvas.winfo_height()
        img_width, img_height = img.size

        ratio = min(canvas_width / img_width, canvas_height / img_height)
        display_w = int(img_width * ratio)
        display_h = int(img_height * ratio)

        display_img = img.resize((display_w, display_h), Image.LANCZOS)
        original_img_tk = ImageTk.PhotoImage(display_img)
        extract_canvas.create_image(0, 0, anchor="nw", image=original_img_tk)
        extract_canvas.image = original_img_tk

        rect_id = None
        start_x = start_y = end_x = end_y = 0

        def on_button_press(event):
            nonlocal start_x, start_y, rect_id
            start_x = extract_canvas.canvasx(event.x)
            start_y = extract_canvas.canvasy(event.y)
            if rect_id:
                extract_canvas.delete(rect_id)
            rect_id = extract_canvas.create_rectangle(start_x, start_y, start_x, start_y, outline="blue", width=2)

        def on_mouse_drag(event):
            nonlocal end_x, end_y
            end_x = extract_canvas.canvasx(event.x)
            end_y = extract_canvas.canvasy(event.y)
            extract_canvas.coords(rect_id, start_x, start_y, end_x, end_y)

        def on_button_release(event):
            nonlocal end_x, end_y
            end_x = extract_canvas.canvasx(event.x)
            end_y = extract_canvas.canvasy(event.y)

            x1, y1 = min(start_x, end_x), min(start_y, end_y)
            x2, y2 = max(start_x, end_x), max(start_y, end_y)

            original_x1 = int(x1 / ratio)
            original_y1 = int(y1 / ratio)
            original_x2 = int(x2 / ratio)
            original_y2 = int(y2 / ratio)

            try:
                region_img = img.crop((original_x1, original_y1, original_x2, original_y2))
                extracted_text = pytesseract.image_to_string(region_img, lang="lav+eng")  # Pieņemam latviešu un angļu

                messagebox.showinfo("Izvilktais teksts", f"Teksts no atlasītā apgabala:\n\n{extracted_text}")
                extract_window.destroy()
            except Exception as e:
                messagebox.showerror("OCR kļūda", f"Neizdevās izvilkt tekstu no apgabala: {e}")

        extract_canvas.bind("<ButtonPress-1>", on_button_press)
        extract_canvas.bind("<B1-Motion>", on_mouse_drag)  # Changed from draw_mask to on_mouse_drag
        extract_canvas.bind("<ButtonRelease-1>", on_button_release)

    def convert_color_space(self):
        """Konvertē pašreizējā attēla krāsu telpu (piem., uz HSV, CMYK)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        color_space_window = Toplevel(self)
        color_space_window.title("Krāsu telpas konvertēšana")
        color_space_window.transient(self)
        color_space_window.grab_set()

        ttk.Label(color_space_window, text="Izvēlieties krāsu telpu:").pack(pady=5)
        color_space_var = tk.StringVar(value="RGB")
        color_space_options = ["RGB", "L (Grayscale)", "HSV (Requires OpenCV)", "CMYK"]
        color_space_combo = ttk.Combobox(color_space_window, textvariable=color_space_var, values=color_space_options,
                                         state="readonly")
        color_space_combo.pack(pady=5)

        def apply_conversion():
            selected_space = color_space_var.get()
            try:
                converted_img = None
                if selected_space == "RGB":
                    converted_img = img.convert("RGB")
                elif selected_space == "L (Grayscale)":
                    converted_img = img.convert("L")
                elif selected_space == "HSV (Requires OpenCV)":
                    if OPENCV_AVAILABLE:
                        img_np = np.array(img.convert('RGB'))
                        img_hsv = cv2.cvtColor(img_np, cv2.COLOR_RGB2HSV)
                        converted_img = Image.fromarray(img_hsv)
                    else:
                        messagebox.showwarning("Trūkst bibliotēkas", "HSV konvertēšanai nepieciešams 'opencv-python'.")
                        return
                elif selected_space == "CMYK":
                    converted_img = img.convert("CMYK")
                else:
                    return

                if converted_img:
                    img_data["processed_img"] = converted_img
                    self.show_image_preview(converted_img)
                    color_space_window.destroy()
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās konvertēt krāsu telpu: {e}")

        ttk.Button(color_space_window, text="Konvertēt", command=apply_conversion).pack(pady=10)

    def add_watermark(self):
        """Pievieno ūdenszīmi pašreizējam attēlam."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        watermark_text = simpledialog.askstring("Ūdenszīme", "Ievadiet ūdenszīmes tekstu:", parent=self)
        if watermark_text:
            try:
                draw = ImageDraw.Draw(img)
                font_size = 40
                try:
                    # Mēģina ielādēt sistēmas fontu
                    font = ImageFont.truetype("arial.ttf", font_size)
                except IOError:
                    # Ja sistēmas fonts nav pieejams, izmanto noklusējuma fontu
                    font = ImageFont.load_default()

                # Aprēķina teksta izmērus
                bbox = draw.textbbox((0, 0), watermark_text, font=font)
                text_width = bbox[2] - bbox[0]
                text_height = bbox[3] - bbox[1]

                # Centra pozīcija
                x = (img.width - text_width) / 2
                y = (img.height - text_height) / 2

                # Pievieno tekstu ar caurspīdīgumu
                # Lai pievienotu caurspīdīgu tekstu, attēlam jābūt ar alfa kanālu (RGBA)
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')

                temp_draw = ImageDraw.Draw(img)
                temp_draw.text((x, y), watermark_text, font=font,
                               fill=(128, 128, 128, 128))  # Pelēka, daļēji caurspīdīga

                img_data["processed_img"] = img
                self.show_image_preview(img)
                messagebox.showinfo("Ūdenszīme", "Ūdenszīme veiksmīgi pievienota.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pievienot ūdenszīmi: {e}")
        else:
            messagebox.showinfo("Ūdenszīme", "Ūdenszīmes teksts netika ievadīts.")

    def create_image_mosaic(self):
        """Izveido attēla mozaīku no vairākiem attēliem."""
        filepaths = filedialog.askopenfilenames(
            title="Izvēlieties attēlus mozaīkai",
            filetypes=[("Attēlu faili", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("Visi faili", "*.*")]
        )
        if not filepaths:
            return

        images_to_mosaic = []
        for fp in filepaths:
            try:
                images_to_mosaic.append(Image.open(fp))
            except Exception as e:
                messagebox.showwarning("Kļūda", f"Neizdevās ielādēt attēlu {fp}: {e}")
                continue

        if not images_to_mosaic:
            messagebox.showwarning("Nav attēlu", "Netika ielādēts neviens attēls mozaīkai.")
            return

        # Vienkāršota mozaīka: saliek attēlus vienā rindā
        widths, heights = zip(*(i.size for i in images_to_mosaic))
        max_height = max(heights)
        total_width = sum(widths)

        mosaic_img = Image.new('RGB', (total_width, max_height))

        x_offset = 0
        for img in images_to_mosaic:
            mosaic_img.paste(img, (x_offset, 0))
            x_offset += img.width

        # Parāda mozaīku
        mosaic_window = Toplevel(self)
        mosaic_window.title("Attēla mozaīka")
        mosaic_window.transient(self)
        mosaic_window.grab_set()

        mosaic_canvas = tk.Canvas(mosaic_window, bg="black")
        mosaic_canvas.pack(fill="both", expand=True)

        # Pielāgo attēlu kanvasa izmēram
        mosaic_window.update_idletasks()
        canvas_width = mosaic_canvas.winfo_width()
        canvas_height = mosaic_canvas.winfo_height()

        if canvas_width == 0 or canvas_height == 0:
            canvas_width = 600
            canvas_height = 400

        display_mosaic_img = mosaic_img.resize((canvas_width, canvas_height),
                                               Image.LANCZOS)
        self.mosaic_photo = ImageTk.PhotoImage(display_mosaic_img)
        mosaic_canvas.create_image(0, 0, anchor="nw", image=self.mosaic_photo)
        mosaic_canvas.image = self.mosaic_photo

        messagebox.showinfo("Attēla mozaīka", "Attēlu mozaīka izveidota un parādīta jaunā logā.")

    def stitch_images(self):
        """Saliek vairākus attēlus kopā, lai izveidotu panorāmu vai lielāku attēlu (vienkāršota versija)."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas", "Attēlu salikšanai nepieciešams 'opencv-python'.")
            return

        filepaths = filedialog.askopenfilenames(
            title="Izvēlieties attēlus salikšanai (vismaz 2)",
            filetypes=[("Attēlu faili", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("Visi faili", "*.*")]
        )
        if not filepaths or len(filepaths) < 2:
            messagebox.showwarning("Nav pietiekami daudz attēlu", "Lūdzu, atlasiet vismaz divus attēlus salikšanai.")
            return

        images_to_stitch = []
        for fp in filepaths:
            try:
                img_pil = Image.open(fp).convert('RGB')
                images_to_stitch.append(cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR))
            except Exception as e:
                messagebox.showwarning("Kļūda", f"Neizdevās ielādēt attēlu {fp}: {e}")
                continue

        if len(images_to_stitch) < 2:
            messagebox.showwarning("Nav pietiekami daudz attēlu", "Netika ielādēts pietiekami daudz attēlu salikšanai.")
            return

        try:
            # Izmanto OpenCV Sticher klasi
            stitcher = cv2.Stitcher_create()
            status, stitched_image = stitcher.stitch(images_to_stitch)

            if status == cv2.Stitcher.OK:
                stitched_image_pil = Image.fromarray(cv2.cvtColor(stitched_image, cv2.COLOR_BGR2RGB))

                # Parāda salikto attēlu
                stitch_window = Toplevel(self)
                stitch_window.title("Saliktais attēls")
                stitch_window.transient(self)
                stitch_window.grab_set()

                stitch_canvas = tk.Canvas(stitch_window, bg="black")
                stitch_canvas.pack(fill="both", expand=True)

                # Pielāgo attēlu kanvasa izmēram
                stitch_window.update_idletasks()
                canvas_width = stitch_canvas.winfo_width()
                canvas_height = stitch_canvas.winfo_height()

                if canvas_width == 0 or canvas_height == 0:
                    canvas_width = 600
                    canvas_height = 400

                display_stitched_img = stitched_image_pil.resize(
                    (canvas_width, canvas_height), Image.LANCZOS)
                self.stitch_photo = ImageTk.PhotoImage(display_stitched_img)
                stitch_canvas.create_image(0, 0, anchor="nw", image=self.stitch_photo)
                stitch_canvas.image = self.stitch_photo

                messagebox.showinfo("Attēlu salikšana", "Attēli veiksmīgi salikti un parādīti jaunā logā.")
            else:
                messagebox.showerror("Kļūda", f"Neizdevās salikt attēlus. Statuss: {status}")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Kļūda attēlu salikšanas procesā: {e}")

    def image_inpainting(self):
        """Aizpilda trūkstošās vai bojātās attēla daļas (vienkāršota versija)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas",
                                   "Attēla atjaunošanai (inpainting) nepieciešams 'opencv-python'.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        inpainting_window = Toplevel(self)
        inpainting_window.title("Attēla atjaunošana (Inpainting)")
        inpainting_window.geometry("800x600")
        inpainting_window.transient(self)
        inpainting_window.grab_set()

        inpainting_canvas = tk.Canvas(inpainting_window, bg="gray")
        inpainting_canvas.pack(fill="both", expand=True)

        # Pārmēro attēlu, lai tas ietilptu kanvasā
        inpainting_window.update_idletasks()
        canvas_width = inpainting_canvas.winfo_width()
        canvas_height = inpainting_canvas.winfo_height()
        img_width, img_height = img.size

        ratio = min(canvas_width / img_width, canvas_height / img_height)
        display_w = int(img_width * ratio)
        display_h = int(img_height * ratio)

        display_img = img.resize((display_w, display_h), Image.LANCZOS)
        self.inpainting_photo = ImageTk.PhotoImage(display_img)
        inpainting_canvas.create_image(0, 0, anchor="nw", image=self.inpainting_photo)
        inpainting_canvas.image = self.inpainting_photo

        mask = np.zeros((img_height, img_width), dtype=np.uint8)  # Maska, kurā iezīmēs bojātās vietas
        drawing = False
        last_x, last_y = 0, 0

        def start_draw(event):
            nonlocal drawing, last_x, last_y
            drawing = True
            last_x, last_y = event.x, event.y

        def draw_mask(event):
            nonlocal last_x, last_y
            if drawing:
                inpainting_canvas.create_line(last_x, last_y, event.x, event.y,
                                              width=10, fill="red", capstyle=tk.ROUND, smooth=tk.TRUE)
                # Pārrēķina uz oriģinālā attēla koordinātām
                orig_x1, orig_y1 = int(last_x / ratio), int(last_y / ratio)
                orig_x2, orig_y2 = int(event.x / ratio), int(event.y / ratio)
                cv2.line(mask, (orig_x1, orig_y1), (orig_x2, orig_y2), 255, 10)  # Zīmē maskā
                last_x, last_y = event.x, event.y

        def stop_draw(event):
            nonlocal drawing
            drawing = False

        inpainting_canvas.bind("<Button-1>", start_draw)
        inpainting_canvas.bind("<B1-Motion>", draw_mask)
        inpainting_canvas.bind("<ButtonRelease-1>", stop_draw)

        def apply_inpainting():
            try:
                img_np = np.array(img.convert('RGB'))
                # Izmanto OpenCV inpainting funkciju
                # cv2.INPAINT_TELEA vai cv2.INPAINT_NS
                restored_img_np = cv2.inpaint(img_np, mask, 3, cv2.INPAINT_TELEA)
                restored_img_pil = Image.fromarray(cv2.cvtColor(restored_img_np, cv2.COLOR_BGR2RGB))

                img_data["processed_img"] = restored_img_pil
                self.show_image_preview(restored_img_pil)
                inpainting_window.destroy()
                messagebox.showinfo("Attēla atjaunošana", "Attēls veiksmīgi atjaunots.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās veikt attēla atjaunošanu: {e}")

        ttk.Button(inpainting_window, text="Pielietot atjaunošanu", command=apply_inpainting).pack(pady=10)

    def stylize_image(self):
        """Pielieto mākslinieciskus stilus attēlam (vienkāršota versija ar filtriem)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        stylize_window = Toplevel(self)
        stylize_window.title("Attēla stilizācija")
        stylize_window.transient(self)
        stylize_window.grab_set()

        ttk.Label(stylize_window, text="Izvēlieties stilu:").pack(pady=5)
        style_var = tk.StringVar(value="Oriģināls")
        styles = ["Oriģināls", "Zīmējums (FIND_EDGES)", "Reljefs (EMBOSS)", "Kontūras (CONTOUR)", "Gausiāns (BLUR)"]
        style_combo = ttk.Combobox(stylize_window, textvariable=style_var, values=styles, state="readonly")
        style_combo.pack(pady=5)

        def apply_style():
            selected_style = style_var.get()
            styled_img = img.copy()
            try:
                if selected_style == "Zīmējums (FIND_EDGES)":
                    styled_img = styled_img.filter(ImageFilter.FIND_EDGES)
                elif selected_style == "Reljefs (EMBOSS)":
                    styled_img = styled_img.filter(ImageFilter.EMBOSS)
                elif selected_style == "Kontūras (CONTOUR)":
                    styled_img = styled_img.filter(ImageFilter.CONTOUR)
                elif selected_style == "Gausiāns (BLUR)":
                    styled_img = styled_img.filter(ImageFilter.GaussianBlur(radius=2))
                # "Oriģināls" neko nedara

                img_data["processed_img"] = styled_img
                self.show_image_preview(styled_img)
                stylize_window.destroy()
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot stilu: {e}")

        ttk.Button(stylize_window, text="Pielietot stilu", command=apply_style).pack(pady=10)

    def geometric_transformations(self):
        """Veic ģeometriskās transformācijas attēlam (vienkāršota versija)."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return

        img_data = self.images[self.current_image_index]
        img = img_data["processed_img"].copy()

        transform_window = Toplevel(self)
        transform_window.title("Ģeometriskās transformācijas")
        transform_window.transient(self)
        transform_window.grab_set()

        ttk.Label(transform_window, text="Izvēlieties transformāciju:").pack(pady=5)
        transform_type_var = tk.StringVar(value="Nobīde")
        transform_types = ["Nobīde", "Mērogošana", "Bīde (Shear)"]
        transform_type_combo = ttk.Combobox(transform_window, textvariable=transform_type_var, values=transform_types,
                                            state="readonly")
        transform_type_combo.pack(pady=5)

        # Nobīdes parametri
        ttk.Label(transform_window, text="X nobīde (pikseļi):").pack(pady=2)
        offset_x_var = tk.IntVar(value=0)
        ttk.Entry(transform_window, textvariable=offset_x_var).pack(pady=2)

        ttk.Label(transform_window, text="Y nobīde (pikseļi):").pack(pady=2)
        offset_y_var = tk.IntVar(value=0)
        ttk.Entry(transform_window, textvariable=offset_y_var).pack(pady=2)

        # Mērogošanas parametri
        ttk.Label(transform_window, text="Mērogošanas faktors:").pack(pady=2)
        scale_factor_var = tk.DoubleVar(value=1.0)
        ttk.Entry(transform_window, textvariable=scale_factor_var).pack(pady=2)

        # Bīdes parametri
        ttk.Label(transform_window, text="X bīdes faktors:").pack(pady=2)
        shear_x_var = tk.DoubleVar(value=0.0)
        ttk.Entry(transform_window, textvariable=shear_x_var).pack(pady=2)

        ttk.Label(transform_window, text="Y bīdes faktors:").pack(pady=2)
        shear_y_var = tk.DoubleVar(value=0.0)
        ttk.Entry(transform_window, textvariable=shear_y_var).pack(pady=2)

        def apply_transform():
            selected_transform = transform_type_var.get()
            transformed_img = img.copy()
            try:
                if selected_transform == "Nobīde":
                    offset_x = offset_x_var.get()
                    offset_y = offset_y_var.get()
                    transformed_img = ImageChops.offset(transformed_img, offset_x, offset_y)
                elif selected_transform == "Mērogošana":
                    scale_factor = scale_factor_var.get()
                    new_width = int(img.width * scale_factor)
                    new_height = int(img.height * scale_factor)
                    transformed_img = transformed_img.resize((new_width, new_height), Image.LANCZOS)
                elif selected_transform == "Bīde (Shear)":
                    shear_x = shear_x_var.get()
                    shear_y = shear_y_var.get()
                    # Bīdes transformācija ar affine transformāciju
                    # Matrica: [[1, shear_x, 0], [shear_y, 1, 0]]
                    transformed_img = transformed_img.transform(
                        transformed_img.size, Image.AFFINE, (1, shear_x, 0, shear_y, 1, 0)
                    )

                img_data["processed_img"] = transformed_img
                self.show_image_preview(transformed_img)
                transform_window.destroy()
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās veikt ģeometrisko transformāciju: {e}")

        ttk.Button(transform_window, text="Pielietot transformāciju", command=apply_transform).pack(pady=10)

    def check_ocr_languages(self):
        """Pārbauda, vai atlasītās OCR valodas ir pieejamas Tesseract instalācijā."""
        try:
            available_langs = pytesseract.get_languages(config='')
            selected_langs = []
            for lang_name, var in self.lang_vars.items():
                if var.get():
                    selected_langs.append(self.lang_options[lang_name])

            missing_langs = [lang for lang in selected_langs if lang not in available_langs]

            if not selected_langs:
                messagebox.showinfo("OCR valodu pārbaude", "Nav atlasīta neviena OCR valoda.")
            elif not missing_langs:
                messagebox.showinfo("OCR valodu pārbaude",
                                    "Visas atlasītās OCR valodas ir pieejamas Tesseract instalācijā.")
            else:
                messagebox.showwarning("OCR valodu pārbaude",
                                       f"Trūkst šādu atlasīto OCR valodu Tesseract instalācijā:\n{', '.join(missing_langs)}\n"
                                       "Lūdzu, instalējiet tās, lai nodrošinātu pareizu OCR darbību.")
        except pytesseract.TesseractNotFoundError:
            messagebox.showerror("Kļūda", "Tesseract nav atrasts. Lūdzu, pārbaudiet Tesseract ceļu iestatījumos.")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās pārbaudīt OCR valodas: {e}")

    def browse_scan_folder(self):
        """Atver dialogu, lai izvēlētos mapi automātiskai skenēšanai."""
        path = filedialog.askdirectory(title="Izvēlieties mapi automātiskai skenēšanai")
        if path:
            self.scan_folder_path.set(path)
            self.settings["scan_folder_path"] = path  # Uzreiz saglabā iestatījumos
            self.save_app_settings()
            if self.auto_scan_enabled.get():
                self.stop_auto_scan()
                self.start_auto_scan()  # Restartē uzraudzību ar jauno mapi

    def toggle_auto_scan(self):
        """Ieslēdz vai izslēdz automātisko skenēšanu."""
        if self.auto_scan_enabled.get():
            self.start_auto_scan()
        else:
            self.stop_auto_scan()
        self.update_auto_scan_status()
        self.settings["auto_scan_enabled"] = self.auto_scan_enabled.get()  # Uzreiz saglabā iestatījumos
        self.save_app_settings()

    def update_remote_storage_fields(self, event=None):
        """Atjaunina attālinātās glabāšanas iestatījumu lauku redzamību."""
        storage_type = self.remote_storage_type.get()
        if storage_type == "FTP" or storage_type == "SFTP":
            self.ftp_settings_frame.grid()
            self.google_drive_settings_frame.grid_remove()
        elif storage_type == "Google Drive":
            self.ftp_settings_frame.grid_remove()
            self.google_drive_settings_frame.grid()
        else:  # Local
            self.ftp_settings_frame.grid_remove()
            self.google_drive_settings_frame.grid_remove()

    def test_ftp_connection(self):
        """Pārbauda FTP/SFTP savienojumu."""
        host = self.ftp_host.get()
        port = self.ftp_port.get()
        user = self.ftp_user.get()
        password = self.ftp_pass.get()
        use_sftp = self.ftp_use_sftp.get()

        if not host or not user or not password:
            messagebox.showwarning("Trūkst datu", "Lūdzu, ievadiet FTP/SFTP hostu, lietotājvārdu un paroli.")
            return

        try:
            if use_sftp:
                import paramiko
                transport = paramiko.Transport((host, port))
                transport.connect(username=user, password=password)
                sftp = paramiko.SFTPClient.from_transport(transport)
                sftp.close()
                transport.close()
                messagebox.showinfo("Savienojums", "SFTP savienojums veiksmīgs!")
            else:
                import ftplib
                ftp = ftplib.FTP()
                ftp.connect(host, port)
                ftp.login(user, password)
                ftp.quit()
                messagebox.showinfo("Savienojums", "FTP savienojums veiksmīgs!")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās izveidot savienojumu: {e}")

    def browse_google_credentials(self):
        """Atver dialogu Google Drive akreditācijas faila izvēlei."""
        path = filedialog.askopenfilename(title="Izvēlieties Google Drive credentials.json",
                                          filetypes=[("JSON faili", "*.json")])
        if path:
            self.google_drive_credentials_path.set(path)
            self.save_app_settings()

    def browse_google_token(self):
        """Atver dialogu Google Drive token faila izvēlei."""
        path = filedialog.askopenfilename(title="Izvēlieties Google Drive token.json",
                                          filetypes=[("JSON faili", "*.json")])
        if path:
            self.google_drive_token_path.set(path)
            self.save_app_settings()

    def authorize_google_drive(self):
        """Autorizējas Google Drive API."""
        try:
            from google_auth_oauthlib.flow import InstalledAppFlow
            from google.auth.transport.requests import Request
            from google.oauth2.credentials import Credentials

            creds = None
            token_path = self.google_drive_token_path.get()
            credentials_path = self.google_drive_credentials_path.get()

            if os.path.exists(token_path):
                creds = Credentials.from_authorized_user_file(token_path,
                                                              ['https://www.googleapis.com/auth/drive.file'])

            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    if not os.path.exists(credentials_path):
                        messagebox.showerror("Kļūda", f"Akreditācijas fails '{credentials_path}' nav atrasts.\n"
                                                      "Lūdzu, lejupielādējiet 'credentials.json' no Google Cloud Console.")
                        return

                    flow = InstalledAppFlow.from_client_secrets_file(
                        credentials_path, ['https://www.googleapis.com/auth/drive.file'])
                    creds = flow.run_local_server(port=0)

                with open(token_path, 'w') as token:
                    token.write(creds.to_json())

            messagebox.showinfo("Google Drive", "Google Drive autorizācija veiksmīga!")
            # Šeit varētu inicializēt Google Drive servisu, ja nepieciešams tūlītējai lietošanai
            # from googleapiclient.discovery import build
            # self.google_drive_service = build('drive', 'v3', credentials=creds)

        except Exception as e:
            messagebox.showerror("Kļūda", f"Google Drive autorizācijas kļūda: {e}")

    def toggle_auto_upload(self):
        """Ieslēdz vai izslēdz automātisko augšupielādi."""
        self.settings["auto_upload_enabled"] = self.auto_upload_enabled.get()
        self.settings["auto_upload_target"] = self.auto_upload_target.get()
        self.save_app_settings()
        messagebox.showinfo("Automātiskā augšupielāde",
                            f"Automātiskā augšupielāde ir {'ieslēgta' if self.auto_upload_enabled.get() else 'izslēgta'}.")

    def refresh_scanned_docs_list(self):
        """Atjaunina skenēto dokumentu sarakstu automatizācijas cilnē."""
        self.scanned_docs_listbox.delete(0, tk.END)
        scan_path = self.scan_folder_path.get()
        if not os.path.exists(scan_path):
            os.makedirs(scan_path, exist_ok=True)

        for filename in os.listdir(scan_path):
            filepath = os.path.join(scan_path, filename)
            if os.path.isfile(filepath) and filename.lower().endswith(
                    ('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.pdf')):
                # Pārbauda, vai dokuments ir apstrādāts (vienkāršots variants)
                # Reālā sistēmā varētu būt nepieciešams saglabāt apstrādāto failu sarakstu vai pievienot metadatus
                is_processed = False
                for item in self.internal_file_system["contents"]:
                    if item["type"] == "folder":
                        for doc in item["contents"]:
                            if doc["type"] == "file" and os.path.basename(doc["filepath"]) == filename:
                                is_processed = True
                                break
                        if is_processed: break

                status = "Skenēts un apstrādāts" if is_processed else "Gaida apstrādi"
                self.scanned_docs_listbox.insert(tk.END, f"{filename} - [{status}]")

    def on_scanned_doc_select(self, event=None):
        """Apstrādā skenēta dokumenta atlasi sarakstā."""
        pass  # Pašlaik nedara neko, var pievienot priekšskatījumu vai apstrādes opcijas

    def open_scanned_doc_location(self, event=None):
        """Atver skenētā dokumenta atrašanās vietu sistēmā."""
        selection = self.scanned_docs_listbox.curselection()
        if not selection:
            return
        index = selection[0]
        filename_with_status = self.scanned_docs_listbox.get(index)
        filename = filename_with_status.split(" - ")[0]  # Iegūst tikai faila nosaukumu
        filepath = os.path.join(self.scan_folder_path.get(), filename)

        if os.path.exists(filepath):
            try:
                if os.name == 'nt':  # Windows
                    os.startfile(os.path.dirname(filepath))
                elif os.name == 'posix':  # macOS, Linux
                    import subprocess
                    import sys
                    if sys.platform == 'darwin':  # macOS
                        subprocess.Popen(['open', '-R', filepath])
                    else:  # Linux
                        subprocess.Popen(['xdg-open', os.path.dirname(filepath)])
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās atvērt faila atrašanās vietu:\n{e}")
        else:
            messagebox.showwarning("Fails nav atrasts", "Skenētais fails nav atrasts norādītajā vietā.")

    def upload_file_to_remote(self, local_filepath, remote_target, remote_path=None):
        """
        Augšupielādē failu uz attālinātu glabāšanas vietu.
        remote_target var būt "FTP", "SFTP", "Google Drive".
        remote_path ir attālā mape, kurā saglabāt.
        """
        if not self.auto_upload_enabled.get():
            return  # Ja automātiskā augšupielāde nav ieslēgta, neko nedara

        try:
            if remote_target == "FTP":
                import ftplib
                ftp = ftplib.FTP()
                ftp.connect(self.ftp_host.get(), self.ftp_port.get())
                ftp.login(self.ftp_user.get(), self.ftp_pass.get())

                # Pārliecinās, ka attālā mape eksistē vai izveido to
                current_dir = ftp.pwd()
                target_dir = os.path.join(self.ftp_remote_path.get(), remote_path).replace("\\",
                                                                                           "/") if remote_path else self.ftp_remote_path.get()

                # Mēģina mainīt direktoriju, ja neizdodas, mēģina izveidot
                try:
                    ftp.cwd(target_dir)
                except ftplib.error_perm:
                    # Mape neeksistē, mēģina izveidot
                    parts = target_dir.split('/')
                    for i in range(1, len(parts) + 1):
                        sub_dir = '/'.join(parts[:i])
                        try:
                            ftp.mkd(sub_dir)
                        except ftplib.error_perm:
                            pass  # Mape jau eksistē
                    ftp.cwd(target_dir)  # Pēc izveides ieiet mapē

                with open(local_filepath, 'rb') as f:
                    ftp.storbinary(f"STOR {os.path.basename(local_filepath)}", f)
                ftp.quit()
                messagebox.showinfo("Augšupielāde",
                                    f"Fails '{os.path.basename(local_filepath)}' veiksmīgi augšupielādēts uz FTP.")

            elif remote_target == "SFTP":
                import paramiko
                transport = paramiko.Transport((self.ftp_host.get(), self.ftp_port.get()))
                transport.connect(username=self.ftp_user.get(), password=self.ftp_pass.get())
                sftp = paramiko.SFTPClient.from_transport(transport)

                # Pārliecinās, ka attālā mape eksistē vai izveido to
                target_dir = os.path.join(self.ftp_remote_path.get(), remote_path).replace("\\",
                                                                                           "/") if remote_path else self.ftp_remote_path.get()

                # Rekursīvi izveido mapes
                parts = target_dir.split('/')
                current_path = ''
                for part in parts:
                    if part:  # Izvairās no tukšām daļām, ja ceļš sākas ar '/'
                        current_path = os.path.join(current_path, part).replace("\\", "/")
                        try:
                            sftp.stat(current_path)
                        except FileNotFoundError:
                            sftp.mkdir(current_path)

                sftp.put(local_filepath, os.path.join(target_dir, os.path.basename(local_filepath)).replace("\\", "/"))
                sftp.close()
                transport.close()
                messagebox.showinfo("Augšupielāde",
                                    f"Fails '{os.path.basename(local_filepath)}' veiksmīgi augšupielādēts uz SFTP.")

            elif remote_target == "Google Drive":
                from google_auth_oauthlib.flow import InstalledAppFlow
                from google.auth.transport.requests import Request
                from google.oauth2.credentials import Credentials
                from googleapiclient.discovery import build
                from googleapiclient.http import MediaFileUpload

                creds = None
                token_path = self.google_drive_token_path.get()
                credentials_path = self.google_drive_credentials_path.get()

                if os.path.exists(token_path):
                    creds = Credentials.from_authorized_user_file(token_path,
                                                                  ['https://www.googleapis.com/auth/drive.file'])

                if not creds or not creds.valid:
                    if creds and creds.expired and creds.refresh_token:
                        creds.refresh(Request())
                    else:
                        if not os.path.exists(credentials_path):
                            messagebox.showerror("Kļūda", f"Akreditācijas fails '{credentials_path}' nav atrasts.\n"
                                                          "Lūdzu, lejupielādējiet 'credentials.json' no Google Cloud Console.")
                            return
                        flow = InstalledAppFlow.from_client_secrets_file(
                            credentials_path, ['https://www.googleapis.com/auth/drive.file'])
                        creds = flow.run_local_server(port=0)
                    with open(token_path, 'w') as token:
                        token.write(creds.to_json())

                service = build('drive', 'v3', credentials=creds)

                # Pārbauda, vai mērķa mape eksistē, ja nē, izveido to
                target_folder_id = self.google_drive_folder_id.get()
                if not target_folder_id:
                    # Ja nav norādīts mapes ID, mēģina atrast vai izveidot saknes mapi
                    results = service.files().list(
                        q="name='OCR_Documents' and mimeType='application/vnd.google-apps.folder'",
                        fields="files(id, name)").execute()
                    items = results.get('files', [])
                    if not items:
                        file_metadata = {
                            'name': 'OCR_Documents',
                            'mimeType': 'application/vnd.google-apps.folder'
                        }
                        file = service.files().create(body=file_metadata, fields='id').execute()
                        target_folder_id = file.get('id')
                        self.google_drive_folder_id.set(target_folder_id)  # Saglabā jauno ID
                        self.save_app_settings()
                    else:
                        target_folder_id = items[0]['id']
                        self.google_drive_folder_id.set(target_folder_id)  # Saglabā atrasto ID
                        self.save_app_settings()

                # Izveido apakšmapi, ja remote_path ir norādīts (piem., dokumentu kategorija)
                if remote_path:
                    # Pārbauda, vai apakšmape jau eksistē
                    results = service.files().list(
                        q=f"'{target_folder_id}' in parents and name='{remote_path}' and mimeType='application/vnd.google-apps.folder'",
                        fields="files(id, name)").execute()
                    items = results.get('files', [])
                    if not items:
                        file_metadata = {
                            'name': remote_path,
                            'parents': [target_folder_id],
                            'mimeType': 'application/vnd.google-apps.folder'
                        }
                        file = service.files().create(body=file_metadata, fields='id').execute()
                        target_folder_id = file.get('id')  # Tagad mērķa ID ir apakšmapes ID
                    else:
                        target_folder_id = items[0]['id']

                file_metadata = {
                    'name': os.path.basename(local_filepath),
                    'parents': [target_folder_id]
                }
                media = MediaFileUpload(local_filepath, mimetype='application/pdf')  # Pieņemam, ka PDF
                file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                messagebox.showinfo("Augšupielāde",
                                    f"Fails '{os.path.basename(local_filepath)}' veiksmīgi augšupielādēts uz Google Drive (ID: {file.get('id')}).")

            else:
                messagebox.showwarning("Augšupielāde", "Nav izvēlēts derīgs attālinātās glabāšanas veids.")

        except Exception as e:
            messagebox.showerror("Augšupielādes kļūda", f"Neizdevās augšupielādēt failu: {e}")

    def update_auto_scan_status(self):
        """Atjaunina automātiskās skenēšanas statusa etiķeti."""
        if self.auto_scan_enabled.get():
            self.auto_scan_status_label.config(text=f"Statuss: Ieslēgts (Uzrauga: {self.scan_folder_path.get()})",
                                               bootstyle="success")
        else:
            self.auto_scan_status_label.config(text="Statuss: Izslēgts", bootstyle="info")

    def start_auto_scan(self):
        """Sāk failu sistēmas uzraudzību automātiskai skenēšanai."""
        scan_path = self.scan_folder_path.get()
        if not os.path.exists(scan_path):
            os.makedirs(scan_path, exist_ok=True)
            messagebox.showinfo("Mape izveidota", f"Izveidota skenēšanas mape: {scan_path}")

        if self.observer:
            self.observer.stop()
            self.observer.join()

        event_handler = ScanEventHandler(self)
        self.observer = Observer()
        self.observer.schedule(event_handler, scan_path, recursive=False)  # Uzrauga tikai tiešos failus mapē
        self.observer.start()
        print(f"Automātiskā skenēšana sākta mapē: {scan_path}")
        self.update_auto_scan_status()

    def stop_auto_scan(self):
        """Aptur failu sistēmas uzraudzību."""
        if self.observer:
            self.observer.stop()
            self.observer.join()
            self.observer = None
            print("Automātiskā skenēšana apturēta.")
        self.update_auto_scan_status()

    def process_new_scanned_file(self, filepath):
        """Apstrādā jaunu skenētu failu (attēlu vai PDF)."""
        print(f"Jauns fails atrasts: {filepath}")
        if not os.path.exists(filepath):
            print(f"Fails {filepath} vairs neeksistē, ignorē.")
            return

        # Pagaida, kamēr fails ir pilnībā uzrakstīts (īpaši svarīgi lieliem failiem)
        size_before = -1
        for _ in range(10):  # Mēģina 10 reizes ar 0.5s intervālu
            current_size = os.path.getsize(filepath)
            if current_size == size_before:
                break
            size_before = current_size
            time.sleep(0.5)
        else:
            print(f"Brīdinājums: Fails {filepath} varbūt nav pilnībā uzrakstīts.")

        # Pārbauda faila tipu
        file_extension = os.path.splitext(filepath)[1].lower()
        if file_extension in ['.pdf']:
            self.after(100, lambda: self._process_pdf_for_auto_scan(filepath))
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp']:
            self.after(100, lambda: self._process_image_for_auto_scan(filepath))
        else:
            print(f"Neatbalstīts faila tips automātiskai apstrādei: {filepath}")
            self.after(100, lambda: messagebox.showwarning("Automātiskā skenēšana",
                                                           f"Neatbalstīts faila tips: {os.path.basename(filepath)}"))
            self.after(0, self.refresh_scanned_docs_list)  # Atjaunina sarakstu pēc apstrādes

    def _process_image_for_auto_scan(self, filepath):
        """Ielādē un apstrādā attēlu automātiskai skenēšanai."""
        try:
            img = Image.open(filepath)
            self.clear_files()  # Notīra iepriekšējos attēlus
            self.images.append({"filepath": filepath, "original_img": img.copy(), "processed_img": img.copy()})
            self.file_listbox.insert(tk.END, os.path.basename(filepath))
            self.file_listbox.select_set(0)
            self.on_file_select()
            self._camera_scan_in_progress = True  # Izmanto to pašu karogu, lai automātiski saglabātu PDF
            self.start_processing()
        except Exception as e:
            messagebox.showerror("Automātiskā skenēšana",
                                 f"Neizdevās apstrādāt attēlu {os.path.basename(filepath)}: {e}")

    def _process_pdf_for_auto_scan(self, filepath):
        """Ielādē un apstrādā PDF automātiskai skenēšanai."""
        try:
            doc = PDFEditor.open(filepath)
            self.clear_files()  # Notīra iepriekšējos attēlus
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=self.dpi_var.get())  # Izmanto iestatīto DPI
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                self.images.append({"filepath": filepath, "original_img": img.copy(), "processed_img": img.copy()})
                self.file_listbox.insert(tk.END, f"{os.path.basename(filepath)} (Lapa {page_num + 1})")
            doc.close()

            if self.images:
                self.file_listbox.select_set(0)
                self.on_file_select()
                self._camera_scan_in_progress = True  # Izmanto to pašu karogu, lai automātiski saglabātu PDF
                self.start_processing()
            else:
                messagebox.showwarning("Automātiskā skenēšana",
                                       f"PDF dokuments {os.path.basename(filepath)} nesatur attēlus vai lapas.")
        except Exception as e:
            messagebox.showerror("Automātiskā skenēšana", f"Neizdevās apstrādāt PDF {os.path.basename(filepath)}: {e}")

    def show_scan_settings(self, parent_window):
        """Parāda skenēšanas iestatījumu logu."""
        settings_window = Toplevel(parent_window)
        settings_window.title("Detekcijas iestatījumi")
        settings_window.geometry("520x650")
        settings_window.transient(parent_window)
        settings_window.grab_set()

        # Galvenais konteiners
        main_container = ttk.Frame(settings_window)
        main_container.pack(fill="both", expand=True, padx=10, pady=10)

        # Scrollable canvas un scrollbar
        canvas = tk.Canvas(main_container, highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_container, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        # Konfigurē scroll reģionu
        def configure_scroll_region(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        scrollable_frame.bind("<Configure>", configure_scroll_region)

        # Pievieno scrollable_frame uz canvas
        canvas_frame = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        # Konfigurē canvas izmēru
        def configure_canvas(event):
            canvas.itemconfig(canvas_frame, width=event.width)

        canvas.bind('<Configure>', configure_canvas)
        canvas.configure(yscrollcommand=scrollbar.set)

        # Peles rullīša atbalsts ar drošības pārbaudēm
        def on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            except tk.TclError:
                # Canvas ir iznīcināts, atceļam notikumu
                pass

        def bind_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.bind_all("<MouseWheel>", on_mousewheel)
            except tk.TclError:
                pass

        def unbind_mousewheel(event):
            try:
                canvas.unbind_all("<MouseWheel>")
            except tk.TclError:
                pass

        def on_window_destroy():
            """Notīra notikumus, kad logs tiek aizvērts."""
            try:
                canvas.unbind_all("<MouseWheel>")
            except:
                pass
            settings_window.destroy()

        # Piesaista notikumus
        canvas.bind('<Enter>', bind_mousewheel)
        canvas.bind('<Leave>', unbind_mousewheel)

        # Nodrošina, ka peles rullīša notikumi tiek atcelti, kad logs aizveras
        settings_window.protocol("WM_DELETE_WINDOW", on_window_destroy)

        # Iestatījumu saturs
        content_frame = ttk.Frame(scrollable_frame, padding="10")
        content_frame.pack(fill="both", expand=True)

        # Virsraksts
        title_label = ttk.Label(content_frame, text="Dokumenta detekcijas iestatījumi",
                                font=("Arial", 14, "bold"))
        title_label.pack(pady=(0, 20))

        # Gausa izplūšana
        blur_frame = ttk.LabelFrame(content_frame, text="Gausa izplūšana", padding="15")
        blur_frame.pack(fill="x", pady=8)

        ttk.Label(blur_frame, text="Kodola izmērs (nepāra skaitlis):").pack(anchor="w")
        blur_scale = ttk.Scale(blur_frame, from_=1, to=15, variable=self.scan_gaussian_blur_kernel,
                               orient="horizontal")
        blur_scale.pack(fill="x", pady=5)
        blur_value_label = ttk.Label(blur_frame, textvariable=self.scan_gaussian_blur_kernel)
        blur_value_label.pack(anchor="w")

        # Adaptīvā sliekšņošana
        thresh_frame = ttk.LabelFrame(content_frame, text="Adaptīvā sliekšņošana", padding="15")
        thresh_frame.pack(fill="x", pady=8)

        ttk.Label(thresh_frame, text="Bloka izmērs (nepāra skaitlis):").pack(anchor="w")
        block_scale = ttk.Scale(thresh_frame, from_=3, to=31, variable=self.scan_adaptive_thresh_block_size,
                                orient="horizontal")
        block_scale.pack(fill="x", pady=5)
        block_value_label = ttk.Label(thresh_frame, textvariable=self.scan_adaptive_thresh_block_size)
        block_value_label.pack(anchor="w")

        ttk.Label(thresh_frame, text="C konstante:").pack(anchor="w", pady=(10, 0))
        c_scale = ttk.Scale(thresh_frame, from_=0, to=20, variable=self.scan_adaptive_thresh_c,
                            orient="horizontal")
        c_scale.pack(fill="x", pady=5)
        c_value_label = ttk.Label(thresh_frame, textvariable=self.scan_adaptive_thresh_c)
        c_value_label.pack(anchor="w")

        # Canny malu detekcija
        canny_frame = ttk.LabelFrame(content_frame, text="Canny malu detekcija", padding="15")
        canny_frame.pack(fill="x", pady=8)

        ttk.Label(canny_frame, text="Zemākais slieksnis:").pack(anchor="w")
        canny1_scale = ttk.Scale(canny_frame, from_=10, to=200, variable=self.scan_canny_thresh1,
                                 orient="horizontal")
        canny1_scale.pack(fill="x", pady=5)
        canny1_value_label = ttk.Label(canny_frame, textvariable=self.scan_canny_thresh1)
        canny1_value_label.pack(anchor="w")

        ttk.Label(canny_frame, text="Augstākais slieksnis:").pack(anchor="w", pady=(10, 0))
        canny2_scale = ttk.Scale(canny_frame, from_=50, to=300, variable=self.scan_canny_thresh2,
                                 orient="horizontal")
        canny2_scale.pack(fill="x", pady=5)
        canny2_value_label = ttk.Label(canny_frame, textvariable=self.scan_canny_thresh2)
        canny2_value_label.pack(anchor="w")

        # Kontūru filtrēšana
        contour_frame = ttk.LabelFrame(content_frame, text="Kontūru filtrēšana", padding="15")
        contour_frame.pack(fill="x", pady=8)

        ttk.Label(contour_frame, text="Minimālais kontūras laukums:").pack(anchor="w")
        area_scale = ttk.Scale(contour_frame, from_=1000, to=50000, variable=self.scan_min_contour_area,
                               orient="horizontal")
        area_scale.pack(fill="x", pady=5)
        area_value_label = ttk.Label(contour_frame, textvariable=self.scan_min_contour_area)
        area_value_label.pack(anchor="w")

        ttk.Label(contour_frame, text="Min aspekta attiecība:").pack(anchor="w", pady=(10, 0))
        ratio_min_scale = ttk.Scale(contour_frame, from_=0.1, to=2.0, variable=self.scan_aspect_ratio_min,
                                    orient="horizontal")
        ratio_min_scale.pack(fill="x", pady=5)
        ratio_min_value_label = ttk.Label(contour_frame, textvariable=self.scan_aspect_ratio_min)
        ratio_min_value_label.pack(anchor="w")

        ttk.Label(contour_frame, text="Max aspekta attiecība:").pack(anchor="w", pady=(10, 0))
        ratio_max_scale = ttk.Scale(contour_frame, from_=1.0, to=5.0, variable=self.scan_aspect_ratio_max,
                                    orient="horizontal")
        ratio_max_scale.pack(fill="x", pady=5)
        ratio_max_value_label = ttk.Label(contour_frame, textvariable=self.scan_aspect_ratio_max)
        ratio_max_value_label.pack(anchor="w")

        # Pogas
        button_frame = ttk.Frame(content_frame)
        button_frame.pack(fill="x", pady=20)

        ttk.Button(button_frame, text="Atiestatīt uz noklusējumu",
                   command=self.reset_scan_settings, bootstyle="warning").pack(side="left", padx=5)
        ttk.Button(button_frame, text="Aizvērt",
                   command=on_window_destroy, bootstyle="secondary").pack(side="right", padx=5)

        # Ievieto canvas un scrollbar galvenajā konteinera
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Fokusē uz logu
        settings_window.focus_set()

    def reset_scan_settings(self):
        """Atiestatīt skenēšanas iestatījumus uz noklusējuma vērtībām."""
        # Esošie iestatījumi
        self.scan_gaussian_blur_kernel.set(5)
        self.scan_adaptive_thresh_block_size.set(11)
        self.scan_adaptive_thresh_c.set(2)
        self.scan_canny_thresh1.set(50)
        self.scan_canny_thresh2.set(150)
        self.scan_min_contour_area.set(10000)
        self.scan_aspect_ratio_min.set(0.3)
        self.scan_aspect_ratio_max.set(3.0)

        # PIEVIENOJIET ŠĪSRINDAS:
        # Jaunie iestatījumi
        self.scan_brightness.set(0)
        self.scan_contrast.set(0)
        self.scan_saturation.set(0)
        self.scan_gamma.set(1.0)
        self.scan_use_color_detection.set(False)
        self.scan_target_color.set("#FFFFFF")
        self.scan_color_tolerance.set(30)
        self.scan_morphology_enabled.set(False)
        self.scan_morphology_kernel_size.set(3)
        self.scan_edge_dilation.set(2)

    def show_document_detection_menu(self):
        """Parāda dokumentu detekcijas izvēlni."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, ko apstrādāt.")
            return

        # Izveidojam izvēlnes logu
        menu_window = Toplevel(self)
        menu_window.title("Dokumenta detekcijas izvēlne")
        menu_window.geometry("500x400")
        menu_window.transient(self)
        menu_window.grab_set()

        # Galvenais frame
        main_frame = ttk.Frame(menu_window, padding="20")
        main_frame.pack(fill="both", expand=True)

        # Virsraksts
        title_label = ttk.Label(main_frame, text="Izvēlieties dokumenta detekcijas veidu:",
                                font=("Arial", 12, "bold"))
        title_label.pack(pady=(0, 20))

        # Automātiskās detekcijas poga
        auto_btn = ttk.Button(main_frame,
                              text="🤖 Automātiskā detekcija",
                              command=lambda: self.start_document_detection(menu_window, auto=True),
                              bootstyle="success",
                              width=30)
        auto_btn.pack(pady=5, fill="x")

        auto_desc = ttk.Label(main_frame,
                              text="Programma automātiski mēģinās atrast dokumenta robežas",
                              font=("Arial", 9),
                              foreground="gray")
        auto_desc.pack(pady=(0, 15))

        # Manuālās detekcijas poga
        manual_btn = ttk.Button(main_frame,
                                text="✋ Manuālā atlase",
                                command=lambda: self.start_document_detection(menu_window, auto=False),
                                bootstyle="warning",
                                width=30)
        manual_btn.pack(pady=5, fill="x")

        manual_desc = ttk.Label(main_frame,
                                text="Jūs paši varēsiet izvēlēties dokumenta stūrus",
                                font=("Arial", 9),
                                foreground="gray")
        manual_desc.pack(pady=(0, 15))

        # Iestatījumu poga
        settings_btn = ttk.Button(main_frame,
                                  text="⚙️ Detekcijas iestatījumi",
                                  command=lambda: self.show_scan_settings(menu_window),
                                  bootstyle="info",
                                  width=30)
        settings_btn.pack(pady=5, fill="x")

        settings_desc = ttk.Label(main_frame,
                                  text="Pielāgojiet automātiskās detekcijas parametrus",
                                  font=("Arial", 9),
                                  foreground="gray")
        settings_desc.pack(pady=(0, 20))

        # Atcelt poga
        cancel_btn = ttk.Button(main_frame,
                                text="Atcelt",
                                command=menu_window.destroy,
                                bootstyle="secondary",
                                width=30)
        cancel_btn.pack(pady=10, fill="x")

    def start_document_detection(self, menu_window, auto=True):
        """Sāk dokumenta detekciju ar izvēlēto metodi."""
        menu_window.destroy()

        current_image_pil = self.images[self.current_image_index]["processed_img"]
        self.document_scanner.set_image(current_image_pil)

        if auto:
            # Automātiskā detekcija
            self.document_scanner.show_document_detection_preview()
        else:
            # Manuālā atlase - sāk ar tukšiem stūriem
            self.document_scanner.corners = []
            self.document_scanner.show_document_detection_preview()

    def enhance_document_detection(self, img_cv):
        """Uzlabo attēlu dokumenta atpazīšanai."""
        gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

        # Adaptīvs kontrasta uzlabojums
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)

        # Gausa izplūšana
        blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)

        return blurred

    def auto_detect_document(self):
        """Automātiski atpazīst dokumentu attēlā ar uzlabotu algoritmu dažādiem apstākļiem."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas",
                                   "Dokumentu atpazīšanai nepieciešams 'opencv-python'.")
            return

        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu, lai noteiktu dokumentu.")
            return

        try:
            img_data = self.images[self.current_image_index]
            img_pil = img_data["processed_img"]
            img_cv = np.array(img_pil.convert('RGB'))
            original_height, original_width = img_cv.shape[:2]

            # Vairāki mēģinājumi ar dažādiem parametriem
            document_contour = None

            # 1. mēģinājums: Standarta pieeja
            document_contour = self._try_detect_document_method1(img_cv)

            # 2. mēģinājums: Canny edge detection
            if document_contour is None:
                document_contour = self._try_detect_document_method2(img_cv)

            # 3. mēģinājums: Morfoloģiskās operācijas
            if document_contour is None:
                document_contour = self._try_detect_document_method3(img_cv)

            # 4. mēģinājums: Krāsu segmentācija
            if document_contour is None:
                document_contour = self._try_detect_document_method4(img_cv)

            if document_contour is None:
                messagebox.showwarning("Dokumenta noteikšana",
                                       "Neizdevās automātiski noteikt dokumenta kontūru ar nevenu metodi.\n"
                                       "Ieteikumi:\n"
                                       "• Pārliecinieties, ka dokuments ir skaidri redzams\n"
                                       "• Mēģiniet uzlabot attēla kontrastu\n"
                                       "• Izmantojiet manuālo atlasi")
                return

            # Pielieto perspektīvas transformāciju
            processed_img_pil = self._apply_perspective_transform(img_cv, document_contour)

            if processed_img_pil:
                img_data["processed_img"] = processed_img_pil
                self.show_image_preview(processed_img_pil)
                messagebox.showinfo("Dokumenta noteikšana", "Dokuments veiksmīgi noteikts un koriģēts.")
            else:
                messagebox.showwarning("Dokumenta noteikšana", "Neizdevās veikt perspektīvas korekciju.")

        except Exception as e:
            messagebox.showerror("Kļūda dokumenta noteikšanā", f"Neizdevās automātiski noteikt dokumentu: {e}")

    def _try_detect_document_method1(self, img_cv):
        """1. metode: Uzlabota adaptīvā sliekšņošana"""
        try:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

            # Uzlabo kontrastu
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)

            # Gausa izplūšana
            blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)

            # Adaptīvā sliekšņošana ar dažādiem parametriem
            for block_size in [11, 15, 19, 23]:
                for c_value in [2, 5, 10]:
                    thresh = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                                   cv2.THRESH_BINARY, block_size, c_value)

                    contour = self._find_best_contour(thresh, img_cv.shape)
                    if contour is not None:
                        return contour

            return None
        except:
            return None

    def _try_detect_document_method2(self, img_cv):
        """2. metode: Canny edge detection"""
        try:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

            # Uzlabo kontrastu
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)

            # Gausa izplūšana
            blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)

            # Canny edge detection ar dažādiem sliekšņiem
            for low_thresh in [50, 75, 100]:
                for high_thresh in [150, 200, 250]:
                    edges = cv2.Canny(blurred, low_thresh, high_thresh)

                    # Morfoloģiskās operācijas, lai aizvērtu pārtraukumus
                    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
                    edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel)

                    contour = self._find_best_contour(edges, img_cv.shape)
                    if contour is not None:
                        return contour

            return None
        except:
            return None

    def _try_detect_document_method3(self, img_cv):
        """3. metode: Morfoloģiskās operācijas"""
        try:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

            # Binārizācija ar Otsu metodi
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            # Morfoloģiskās operācijas
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))

            # Closing - aizvērt mazos caurums
            closed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel, iterations=2)

            # Opening - noņemt troksni
            opened = cv2.morphologyEx(closed, cv2.MORPH_OPEN, kernel, iterations=1)

            contour = self._find_best_contour(opened, img_cv.shape)
            if contour is not None:
                return contour

            # Mēģinam ar invertētu attēlu
            inverted = cv2.bitwise_not(opened)
            contour = self._find_best_contour(inverted, img_cv.shape)
            return contour

        except:
            return None

    def _try_detect_document_method4(self, img_cv):
        """4. metode: Krāsu segmentācija"""
        try:
            # Konvertē uz HSV krāsu telpu
            hsv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2HSV)

            # Definē baltās krāsas diapazonu (dokumenti bieži ir balti)
            lower_white = np.array([0, 0, 180])
            upper_white = np.array([180, 30, 255])

            # Izveido masku baltajām krāsām
            white_mask = cv2.inRange(hsv, lower_white, upper_white)

            # Morfoloģiskās operācijas
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
            white_mask = cv2.morphologyEx(white_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
            white_mask = cv2.morphologyEx(white_mask, cv2.MORPH_OPEN, kernel, iterations=1)

            contour = self._find_best_contour(white_mask, img_cv.shape)
            if contour is not None:
                return contour

            # Mēģinam ar plašāku krāsu diapazonu
            lower_light = np.array([0, 0, 120])
            upper_light = np.array([180, 50, 255])
            light_mask = cv2.inRange(hsv, lower_light, upper_light)

            light_mask = cv2.morphologyEx(light_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
            contour = self._find_best_contour(light_mask, img_cv.shape)
            return contour

        except:
            return None

    def _find_best_contour(self, binary_img, img_shape):
        """Atrod labāko kontūru, kas varētu būt dokuments"""
        try:
            contours, _ = cv2.findContours(binary_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            if not contours:
                return None

            img_area = img_shape[0] * img_shape[1]

            # Sakārto kontūras pēc laukuma
            contours = sorted(contours, key=cv2.contourArea, reverse=True)

            for contour in contours:
                area = cv2.contourArea(contour)

                # Pārbauda minimālo laukumu (vismaz 5% no attēla)
                if area < img_area * 0.05:
                    continue

                # Pārbauda maksimālo laukumu (ne vairāk kā 95% no attēla)
                if area > img_area * 0.95:
                    continue

                # Aproksimē kontūru
                peri = cv2.arcLength(contour, True)

                # Mēģina ar dažādiem epsilon parametriem
                for epsilon_factor in [0.01, 0.02, 0.03, 0.04, 0.05]:
                    approx = cv2.approxPolyDP(contour, epsilon_factor * peri, True)

                    # Ja ir 4 stūri
                    if len(approx) == 4:
                        # Pārbauda, vai stūri veido saprātīgu taisnstūri
                        if self._is_valid_rectangle(approx, img_shape):
                            return approx

                    # Ja ir vairāk nekā 4 stūri, mēģina atrast 4 galvenos
                    elif len(approx) > 4:
                        # Atrod 4 galvenos stūrus
                        rect_corners = self._find_four_corners(approx)
                        if rect_corners is not None and self._is_valid_rectangle(rect_corners, img_shape):
                            return rect_corners

            return None
        except:
            return None

    def _is_valid_rectangle(self, corners, img_shape):
        """Pārbauda, vai 4 punkti veido derīgu taisnstūri"""
        try:
            if len(corners) != 4:
                return False

            # Pārbauda, vai visi punkti ir attēla robežās
            h, w = img_shape[:2]
            for corner in corners:
                x, y = corner[0]
                if x < 0 or x >= w or y < 0 or y >= h:
                    return False

            # Aprēķina laukumu
            area = cv2.contourArea(corners)
            img_area = h * w

            # Pārbauda laukuma attiecību
            if area < img_area * 0.05 or area > img_area * 0.95:
                return False

            # Pārbauda, vai forma ir pietiekami taisnstūrveida
            # Aprēķina convex hull un salīdzina laukumus
            hull = cv2.convexHull(corners)
            hull_area = cv2.contourArea(hull)

            if hull_area > 0:
                solidity = area / hull_area
                if solidity < 0.8:  # Ja forma nav pietiekami "cieta"
                    return False

            return True
        except:
            return False

    def _find_four_corners(self, contour):
        """Atrod 4 galvenos stūrus no kontūras ar vairāk punktiem"""
        try:
            # Atrod kontūras bounding rectangle
            rect = cv2.minAreaRect(contour)
            box = cv2.boxPoints(rect)
            box = np.int0(box)

            return box.reshape(4, 1, 2)
        except:
            return None

    def _apply_perspective_transform(self, img_cv, document_contour):
        """Pielieto perspektīvas transformāciju"""
        try:
            def order_points(pts):
                """Sakārto punktus: augšējais kreisais, augšējais labais, apakšējais labais, apakšējais kreisais"""
                rect = np.zeros((4, 2), dtype="float32")

                # Summa: augšējais kreisais būs mazākā, apakšējais labais - lielākā
                s = pts.sum(axis=1)
                rect[0] = pts[np.argmin(s)]  # Augšējais kreisais
                rect[2] = pts[np.argmax(s)]  # Apakšējais labais

                # Starpība: augšējais labais būs mazākā, apakšējais kreisais - lielākā
                diff = np.diff(pts, axis=1)
                rect[1] = pts[np.argmin(diff)]  # Augšējais labais
                rect[3] = pts[np.argmax(diff)]  # Apakšējais kreisais

                return rect

            # Sakārto stūrus
            corners = document_contour.reshape(4, 2)
            ordered_corners = order_points(corners)

            # Aprēķina jaunā attēla izmērus
            (tl, tr, br, bl) = ordered_corners

            # Platums
            widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
            widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
            maxWidth = max(int(widthA), int(widthB))

            # Augstums
            heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
            heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
            maxHeight = max(int(heightA), int(heightB))

            # Pārbauda, vai izmēri ir saprātīgi
            if maxWidth < 50 or maxHeight < 50:
                return None

            if maxWidth > 5000 or maxHeight > 5000:
                # Ierobežo maksimālos izmērus
                ratio = min(5000 / maxWidth, 5000 / maxHeight)
                maxWidth = int(maxWidth * ratio)
                maxHeight = int(maxHeight * ratio)

            # Definē mērķa punktus (taisnstūrveida dokuments)
            dst = np.array([
                [0, 0],
                [maxWidth - 1, 0],
                [maxWidth - 1, maxHeight - 1],
                [0, maxHeight - 1]], dtype="float32")

            # Aprēķina perspektīvas transformācijas matricu
            M = cv2.getPerspectiveTransform(ordered_corners, dst)

            # Pielieto perspektīvas transformāciju
            warped = cv2.warpPerspective(img_cv, M, (maxWidth, maxHeight))

            # Konvertē atpakaļ uz PIL attēlu
            processed_img_pil = Image.fromarray(warped)

            return processed_img_pil

        except Exception as e:
            print(f"Kļūda perspektīvas transformācijā: {e}")
            return None

    def manual_document_selection(self):
        """Ļauj lietotājam manuāli atlasīt dokumenta apgabalu."""
        if self.current_image_index == -1:
            return

        messagebox.showinfo("Manuāla atlase",
                            "Izmantojiet peles kreiso pogu, lai iezīmētu dokumenta apgabalu attēla priekšskatījumā.")

    def correct_document_borders(self):
        """Koriģē dokumenta robežas."""
        if self.current_image_index == -1:
            return

        messagebox.showinfo("Robežu korekcija",
                            "Šī funkcija ļaus precizēt dokumenta robežas.")

    def on_closing(self):
        """Apstrādā loga aizvēršanas notikumu, saglabājot iestatījumus un arhīvu."""
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.save_app_settings()
        self.save_scan_settings()  # JAUNS: Saglabā skenēšanas iestatījumus
        self.stop_auto_scan()  # Aptur watchdog observer
        self.save_pdf_archive()

        if self.current_pdf_document:
            self.current_pdf_document.close()
            self.current_pdf_document = None
        self.destroy()

    # --- JAUNAS FUNKCIJAS ---

    def _get_current_image(self):
        """Palīgfunkcija, lai iegūtu pašreizējo apstrādājamo attēlu."""
        if self.current_image_index == -1:
            messagebox.showwarning("Nav attēla", "Lūdzu, vispirms atlasiet attēlu.")
            return None
        return self.images[self.current_image_index]["processed_img"]

    def _update_current_image(self, new_img):
        """Palīgfunkcija, lai atjauninātu pašreizējo attēlu un priekšskatījumu."""
        if self.current_image_index != -1:
            self.images[self.current_image_index]["processed_img"] = new_img
            self.show_image_preview(new_img)

    def convert_to_grayscale(self):
        """Konvertē pašreizējo attēlu pelēktoņos."""
        img = self._get_current_image()
        if img:
            try:
                if img.mode != 'L':
                    img = img.convert('L')
                    self._update_current_image(img)
                    messagebox.showinfo("Konvertēšana", "Attēls konvertēts uz pelēktoņiem.")
                else:
                    messagebox.showinfo("Konvertēšana", "Attēls jau ir pelēktoņos.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās konvertēt uz pelēktoņiem: {e}")

    def apply_thresholding(self):
        """Pielieto attēlam bināro sliekšņošanu."""
        img = self._get_current_image()
        if img:
            try:
                if img.mode != 'L':
                    img = img.convert('L')  # Sliekšņošanai nepieciešami pelēktoņi

                threshold_value = simpledialog.askinteger("Sliekšņošana", "Ievadiet sliekšņa vērtību (0-255):",
                                                          parent=self, minvalue=0, maxvalue=255, initialvalue=128)
                if threshold_value is not None:
                    img = img.point(lambda p: 255 if p > threshold_value else 0)
                    self._update_current_image(img)
                    messagebox.showinfo("Sliekšņošana", "Sliekšņošana veiksmīgi pielietota.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot sliekšņošanu: {e}")

    def apply_gaussian_blur(self):
        """Pielieto Gausa izplūšanu attēlam."""
        img = self._get_current_image()
        if img:
            try:
                radius = simpledialog.askfloat("Gausa izplūšana", "Ievadiet izplūšanas rādiusu (piem., 2.0):",
                                               parent=self, minvalue=0.1, initialvalue=2.0)
                if radius is not None:
                    img = img.filter(ImageFilter.GaussianBlur(radius))
                    self._update_current_image(img)
                    messagebox.showinfo("Gausa izplūšana", "Gausa izplūšana veiksmīgi pielietota.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot Gausa izplūšanu: {e}")

    def apply_median_filter(self):
        """Pielieto mediānas filtru trokšņu samazināšanai."""
        img = self._get_current_image()
        if img:
            try:
                size = simpledialog.askinteger("Mediānas filtrs", "Ievadiet filtra izmēru (nepāra skaitlis, piem., 3):",
                                               parent=self, minvalue=1, initialvalue=3)
                if size is not None and size % 2 == 1:
                    img = img.filter(ImageFilter.MedianFilter(size))
                    self._update_current_image(img)
                    messagebox.showinfo("Mediānas filtrs", "Mediānas filtrs veiksmīgi pielietots.")
                elif size is not None:
                    messagebox.showwarning("Mediānas filtrs", "Filtra izmēram jābūt nepāra skaitlim.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot mediānas filtru: {e}")

    def sharpen_image(self):
        """Uzlabo attēla asumu."""
        img = self._get_current_image()
        if img:
            try:
                factor = simpledialog.askfloat("Asums", "Ievadiet asuma faktoru (1.0 - oriģināls, >1.0 - asāks):",
                                               parent=self, minvalue=0.1, initialvalue=1.5)
                if factor is not None:
                    enhancer = ImageEnhance.Sharpness(img)
                    img = enhancer.enhance(factor)
                    self._update_current_image(img)
                    messagebox.showinfo("Asums", "Attēla asums veiksmīgi uzlabots.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās uzlabot asumu: {e}")

    def rotate_image_by_angle(self):
        """Pagriež attēlu par norādītu leņķi."""
        img = self._get_current_image()
        if img:
            try:
                angle = simpledialog.askfloat("Pagriezt attēlu", "Ievadiet pagriešanas leņķi (grādos):",
                                              parent=self, initialvalue=45.0)
                if angle is not None:
                    img = img.rotate(angle, expand=True, fillcolor=(255, 255, 255) if img.mode == 'RGB' else 255)
                    self._update_current_image(img)
                    messagebox.showinfo("Pagriešana", f"Attēls pagriezts par {angle}°.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pagriezt attēlu: {e}")

    def add_text_overlay(self):
        """Pievieno attēlam teksta pārklājumu."""
        img = self._get_current_image()
        if img:
            text = simpledialog.askstring("Teksta pārklājums", "Ievadiet tekstu:", parent=self)
            if text:
                try:
                    if img.mode != 'RGBA':
                        img = img.convert('RGBA')  # Nepieciešams caurspīdīgumam

                    draw = ImageDraw.Draw(img)
                    font_size = simpledialog.askinteger("Teksta pārklājums", "Ievadiet fonta izmēru:",
                                                        parent=self, minvalue=10, initialvalue=50)
                    if font_size is None: return

                    font_color = simpledialog.askstring("Teksta pārklājums",
                                                        "Ievadiet fonta krāsu (piem., 'red', '#FF0000'):",
                                                        parent=self, initialvalue="black")
                    if font_color is None: return

                    x_pos = simpledialog.askinteger("Teksta pārklājums", "Ievadiet X pozīciju:",
                                                    parent=self, initialvalue=50)
                    if x_pos is None: return

                    y_pos = simpledialog.askinteger("Teksta pārklājums", "Ievadiet Y pozīciju:",
                                                    parent=self, initialvalue=50)
                    if y_pos is None: return

                    try:
                        font = ImageFont.truetype("arial.ttf", font_size)
                    except IOError:
                        font = ImageFont.load_default()

                    draw.text((x_pos, y_pos), text, font=font, fill=font_color)
                    self._update_current_image(img)
                    messagebox.showinfo("Teksta pārklājums", "Teksts veiksmīgi pievienots.")
                except Exception as e:
                    messagebox.showerror("Kļūda", f"Neizdevās pievienot teksta pārklājumu: {e}")

    def draw_rectangle_on_image(self):
        """Zīmē taisnstūri uz attēla."""
        img = self._get_current_image()
        if img:
            try:
                x1 = simpledialog.askinteger("Zīmēt taisnstūri", "Ievadiet X1 koordinātu:", parent=self,
                                             initialvalue=50)
                y1 = simpledialog.askinteger("Zīmēt taisnstūri", "Ievadiet Y1 koordinātu:", parent=self,
                                             initialvalue=50)
                x2 = simpledialog.askinteger("Zīmēt taisnstūri", "Ievadiet X2 koordinātu:", parent=self,
                                             initialvalue=150)
                y2 = simpledialog.askinteger("Zīmēt taisnstūri", "Ievadiet Y2 koordinātu:", parent=self,
                                             initialvalue=150)
                if any(coord is None for coord in [x1, y1, x2, y2]): return

                color = simpledialog.askstring("Zīmēt taisnstūri", "Ievadiet krāsu (piem., 'red', '#FF0000'):",
                                               parent=self, initialvalue="red")
                if color is None: return

                width = simpledialog.askinteger("Zīmēt taisnstūri", "Ievadiet līnijas biezumu:",
                                                parent=self, minvalue=1, initialvalue=3)
                if width is None: return

                draw = ImageDraw.Draw(img)
                draw.rectangle([x1, y1, x2, y2], outline=color, width=width)
                self._update_current_image(img)
                messagebox.showinfo("Zīmēšana", "Taisnstūris veiksmīgi uzzīmēts.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās uzzīmēt taisnstūri: {e}")

    def draw_circle_on_image(self):
        """Zīmē apli uz attēla."""
        img = self._get_current_image()
        if img:
            try:
                x = simpledialog.askinteger("Zīmēt apli", "Ievadiet centra X koordinātu:", parent=self,
                                            initialvalue=100)
                y = simpledialog.askinteger("Zīmēt apli", "Ievadiet centra Y koordinātu:", parent=self,
                                            initialvalue=100)
                radius = simpledialog.askinteger("Zīmēt apli", "Ievadiet rādiusu:", parent=self, minvalue=1,
                                                 initialvalue=50)
                if any(val is None for val in [x, y, radius]): return

                color = simpledialog.askstring("Zīmēt apli", "Ievadiet krāsu (piem., 'green', '#00FF00'):",
                                               parent=self, initialvalue="blue")
                if color is None: return

                width = simpledialog.askinteger("Zīmēt apli", "Ievadiet līnijas biezumu:",
                                                parent=self, minvalue=1, initialvalue=3)
                if width is None: return

                draw = ImageDraw.Draw(img)
                draw.ellipse([x - radius, y - radius, x + radius, y + radius], outline=color, width=width)
                self._update_current_image(img)
                messagebox.showinfo("Zīmēšana", "Aplis veiksmīgi uzzīmēts.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās uzzīmēt apli: {e}")

    def extract_color_channels(self):
        """Izvelk atsevišķus krāsu kanālus (R, G, B) un parāda tos."""
        img = self._get_current_image()
        if img:
            try:
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                r, g, b = img.split()

                # Parāda katru kanālu atsevišķā logā
                for channel_img, channel_name in zip([r, g, b], ["Red", "Green", "Blue"]):
                    channel_window = Toplevel(self)
                    channel_window.title(f"Krāsu kanāls: {channel_name}")
                    channel_window.transient(self)
                    channel_window.grab_set()

                    canvas_channel = tk.Canvas(channel_window, bg="black")
                    canvas_channel.pack(fill="both", expand=True)

                    # Pielāgo attēlu kanvasa izmēram
                    channel_window.update_idletasks()
                    canvas_width = canvas_channel.winfo_width()
                    canvas_height = canvas_channel.winfo_height()

                    if canvas_width == 0 or canvas_height == 0:
                        canvas_width = 400
                        canvas_height = 300

                    display_channel_img = channel_img.resize(
                        (canvas_width, canvas_height), Image.LANCZOS)
                    photo_channel = ImageTk.PhotoImage(display_channel_img)
                    canvas_channel.create_image(0, 0, anchor="nw", image=photo_channel)
                    canvas_channel.image = photo_channel  # Saglabā atsauci

                messagebox.showinfo("Krāsu kanāli", "Krāsu kanāli veiksmīgi izvilkti un parādīti atsevišķos logos.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās izvilkt krāsu kanālus: {e}")

    def merge_color_channels(self):
        """Apvieno krāsu kanālus (R, G, B) no atsevišķiem attēliem."""
        messagebox.showinfo("Apvienot krāsu kanālus",
                            "Lūdzu, atlasiet trīs pelēktoņu attēlus (sarkanajam, zaļajam un zilajam kanālam).")
        filepaths = filedialog.askopenfilenames(
            title="Izvēlieties 3 pelēktoņu attēlus (R, G, B)",
            filetypes=[("Attēlu faili", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("Visi faili", "*.*")]
        )

        if len(filepaths) != 3:
            messagebox.showwarning("Kļūda", "Lūdzu, atlasiet tieši 3 attēlus (pa vienam katram kanālam).")
            return

        try:
            # Ielādē attēlus un pārliecinās, ka tie ir pelēktoņos
            r_img = Image.open(filepaths[0]).convert('L')
            g_img = Image.open(filepaths[1]).convert('L')
            b_img = Image.open(filepaths[2]).convert('L')

            # Pārliecinās, ka attēli ir vienāda izmēra
            if not (r_img.size == g_img.size == b_img.size):
                messagebox.showwarning("Kļūda", "Visiem kanālu attēliem jābūt vienāda izmēra.")
                return

            merged_img = Image.merge('RGB', (r_img, g_img, b_img))
            self._update_current_image(merged_img)
            messagebox.showinfo("Apvienošana", "Krāsu kanāli veiksmīgi apvienoti.")
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās apvienot krāsu kanālus: {e}")

    def apply_sepia_filter(self):
        """Pielieto sēpijas filtru attēlam."""
        img = self._get_current_image()
        if img:
            try:
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                pixels = img.load()
                for y in range(img.size[1]):
                    for x in range(img.size[0]):
                        r, g, b = pixels[x, y]
                        tr = int(0.393 * r + 0.769 * g + 0.189 * b)
                        tg = int(0.349 * r + 0.686 * g + 0.168 * b)
                        tb = int(0.272 * r + 0.534 * g + 0.131 * b)
                        pixels[x, y] = (min(255, tr), min(255, tg), min(255, tb))
                self._update_current_image(img)
                messagebox.showinfo("Sēpijas filtrs", "Sēpijas filtrs veiksmīgi pielietots.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot sēpijas filtru: {e}")

    def apply_vignette_effect(self):
        """Pielieto vinjetes efektu attēlam."""
        img = self._get_current_image()
        if img:
            try:
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                width, height = img.size
                center_x, center_y = width // 2, height // 2
                max_dist = (center_x ** 2 + center_y ** 2) ** 0.5

                pixels = img.load()
                for y in range(height):
                    for x in range(width):
                        dist = ((x - center_x) ** 2 + (y - center_y) ** 2) ** 0.5
                        intensity = 1 - (dist / max_dist) * 0.7  # 0.7 ir vinjetes stiprums
                        r, g, b = pixels[x, y]
                        pixels[x, y] = (int(r * intensity), int(g * intensity), int(b * intensity))
                self._update_current_image(img)
                messagebox.showinfo("Vinjetes efekts", "Vinjetes efekts veiksmīgi pielietots.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pielietot vinjetes efektu: {e}")

    def pixelate_image(self):
        """Pikselizē attēlu."""
        img = self._get_current_image()
        if img:
            try:
                pixel_size = simpledialog.askinteger("Pikselizācija", "Ievadiet pikseļa bloka izmēru (piem., 10):",
                                                     parent=self, minvalue=1, initialvalue=10)
                if pixel_size is not None:
                    small_img = img.resize((img.width // pixel_size, img.height // pixel_size), Image.NEAREST)
                    pixelated_img = small_img.resize(img.size, Image.NEAREST)
                    self._update_current_image(pixelated_img)
                    messagebox.showinfo("Pikselizācija", "Attēls veiksmīgi pikselizēts.")
            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās pikselizēt attēlu: {e}")

    def detect_faces(self):
        """Noteikt sejas attēlā, izmantojot OpenCV."""
        if not OPENCV_AVAILABLE:
            messagebox.showwarning("Trūkst bibliotēkas", "Sejas noteikšanai nepieciešams 'opencv-python'.")
            return

        img = self._get_current_image()
        if img:
            try:
                # Pārvērš PIL attēlu uz OpenCV formātu
                img_cv = np.array(img.convert('RGB'))
                img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)
                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)

                # Ielādē sejas kaskādes klasifikatoru, izmantojot relatīvo ceļu
                # Fails atradīsies mapē "data" blakus .exe failam
                cascade_path = resource_path(os.path.join("data", "haarcascade_frontalface_default.xml"))

                # Pārbauda, vai fails eksistē
                if not os.path.exists(cascade_path):
                    messagebox.showerror("Kļūda", f"Haar kaskādes klasifikators '{cascade_path}' nav atrasts.\n"
                                                  "Lūdzu, pārliecinieties, ka fails 'haarcascade_frontalface_default.xml' atrodas mapē 'data' blakus programmai.")
                    return

                face_cascade = cv2.CascadeClassifier(cascade_path)

                if face_cascade.empty():
                    messagebox.showerror("Kļūda", f"Neizdevās ielādēt Haar kaskādes klasifikatoru no '{cascade_path}'.")
                    return

                faces = face_cascade.detectMultiScale(gray, 1.1, 4)

                if len(faces) == 0:
                    messagebox.showinfo("Sejas noteikšana", "Attēlā netika atrastas sejas.")
                    return

                # Zīmē taisnstūrus ap atrastajām sejām
                for (x, y, w, h) in faces:
                    cv2.rectangle(img_cv, (x, y), (x + w, y + h), (255, 0, 0), 2)  # Zils taisnstūris

                # Pārvērš atpakaļ uz PIL attēlu un atjaunina
                result_img_pil = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))
                self._update_current_image(result_img_pil)
                messagebox.showinfo("Sejas noteikšana", f"Attēlā atrastas {len(faces)} sejas.")

            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās veikt sejas noteikšanu: {e}")

    def _manual_document_selection(self, frame):
        """
        Ļauj lietotājam manuāli atlasīt dokumenta stūrus.
        Atgriež apstrādātu PIL attēlu vai None, ja atlase atcelta.
        Uzlabota pēcapstrāde manuāli atlasītajam apgabalam.
        """
        points = []
        clone = frame.copy()
        window_name = "Manuāla atlase: Noklikšķiniet uz 4 stūriem (Esc, lai atceltu, Enter, lai apstiprinātu)"

        def mouse_callback(event, x, y, flags, param):
            nonlocal points, clone

            if event == cv2.EVENT_LBUTTONDOWN:
                if len(points) < 4:
                    points.append((x, y))
                    cv2.circle(clone, (x, y), 5, (0, 255, 0), -1)
                    cv2.imshow(window_name, clone)
                else:
                    messagebox.showwarning("Atlase",
                                           "Jau ir atlasīti 4 punkti. Nospiediet Enter, lai apstiprinātu, vai Esc, lai atceltu.")

        cv2.namedWindow(window_name, cv2.WINDOW_NORMAL)
        cv2.resizeWindow(window_name, 800, 600)
        cv2.setMouseCallback(window_name, mouse_callback)

        while True:
            cv2.imshow(window_name, clone)
            key = cv2.waitKey(1) & 0xFF

            if key == 27:  # Esc taustiņš
                messagebox.showinfo("Atlase atcelta", "Manuālā atlase atcelta.")
                cv2.destroyWindow(window_name)
                return None
            elif key == 13:  # Enter taustiņš
                if len(points) == 4:
                    break
                else:
                    messagebox.showwarning("Nepietiekami punkti", "Lūdzu, atlasiet visus 4 dokumenta stūrus.")

        cv2.destroyWindow(window_name)

        if len(points) == 4:
            try:
                # Veic perspektīvas transformāciju ar manuāli atlasītajiem punktiem
                warped = self._four_point_transform(frame, np.array(points, dtype="float32"))

                # Uzlabota pēcapstrāde skenētajam attēlam:
                # Pārvērš uz pelēktoņiem
                warped_gray = cv2.cvtColor(warped, cv2.COLOR_BGR2GRAY)

                # Pielieto adaptīvo sliekšņošanu ar lielāku bloka izmēru tīrākam rezultātam
                final_processed_img_cv = cv2.adaptiveThreshold(warped_gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                                               cv2.THRESH_BINARY, 21, 5)

                # Pārliecinās, ka fons ir balts un teksts melns (standarta OCR formāts)
                if np.mean(final_processed_img_cv) < 128:
                    final_processed_img_cv = cv2.bitwise_not(final_processed_img_cv)

                # Saglabā oriģinālo krāsu attēlu
                original_scanned_pil_img = Image.fromarray(cv2.cvtColor(warped, cv2.COLOR_BGR2RGB))

                # Apstrādā attēlu OCR vajadzībām (pelēktoņi, adaptīvā sliekšņošana)
                warped_gray = cv2.cvtColor(warped, cv2.COLOR_BGR2GRAY)
                final_processed_img_cv = cv2.adaptiveThreshold(warped_gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                                               cv2.THRESH_BINARY, 21, 5)
                if np.mean(final_processed_img_cv) < 128:
                    final_processed_img_cv = cv2.bitwise_not(final_processed_img_cv)
                ocr_processed_pil_img = Image.fromarray(final_processed_img_cv)

                messagebox.showinfo("Atlase veiksmīga", "Dokuments veiksmīgi skenēts ar manuālo atlasi.")
                # Atgriežam apstrādāto attēlu, kas tiks izmantots kā "processed_img"
                # "original_img" tiks iestatīts no `original_scanned_pil_img`
                return original_scanned_pil_img, ocr_processed_pil_img

            except Exception as e:
                messagebox.showerror("Kļūda", f"Neizdevās apstrādāt manuāli atlasīto apgabalu: {e}")
                return None
        return None

    def _order_points(self, pts):
        """Sakārto punktus: augšējais kreisais, augšējais labais, apakšējais labais, apakšējais kreisais."""
        rect = np.zeros((4, 2), dtype="float32")

        s = pts.sum(axis=1)
        rect[0] = pts[np.argmin(s)]  # Augšējais kreisais (mazākā summa)
        rect[2] = pts[np.argmax(s)]  # Apakšējais labais (lielākā summa)

        diff = np.diff(pts, axis=1)
        rect[1] = pts[np.argmin(diff)]  # Augšējais labais (mazākā starpība)
        rect[3] = pts[np.argmax(diff)]  # Apakšējais kreisais (lielākā starpība)

        return rect

    def _four_point_transform(self, image, pts):
        """Veic četru punktu perspektīvas transformāciju."""
        rect = self._order_points(pts)
        (tl, tr, br, bl) = rect

        # Aprēķina jaunā attēla platumu
        widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
        widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
        maxWidth = max(int(widthA), int(widthB))

        # Aprēķina jaunā attēla augstumu
        heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
        heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
        maxHeight = max(int(heightA), int(heightB))

        # Galamērķa punkti
        dst = np.array([
            [0, 0],
            [maxWidth - 1, 0],
            [maxWidth - 1, maxHeight - 1],
            [0, maxHeight - 1]], dtype="float32")

        # Aprēķina perspektīvas transformācijas matricu un pielieto to
        M = cv2.getPerspectiveTransform(rect, dst)
        warped = cv2.warpPerspective(image, M, (maxWidth, maxHeight))

        return warped

    def get_camera_frame_hq(self):
        """Iegūst augstas kvalitātes kadru saglabāšanai."""
        if self.camera is None or not self.camera_active:
            return None

        try:
            # Izmet 2-3 kadrus, lai iegūtu jaunāko
            for _ in range(3):
                ret, frame = self.camera.read()
                if not ret:
                    return None

            # Atgriež pilnu izšķirtspēju bez samazināšanas
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            return Image.fromarray(frame_rgb)
        except Exception as e:
            print(f"HQ kadra kļūda: {e}")
            return None


class ScanEventHandler(FileSystemEventHandler):
    """
    Apstrādā failu sistēmas notikumus, lai automātiski apstrādātu jaunus failus.
    """

    def __init__(self, app_instance):
        super().__init__()
        try:
            tesseract_path, tessdata_path = configure_tesseract()
            print(f"Tesseract ceļš: {tesseract_path}")
            print(f"Tessdata ceļš: {tessdata_path}")
        except Exception as e:
            print(f"Tesseract konfigurācijas kļūda: {e}")
        self.app = app_instance
        self.processed_files = set()  # Lai izvairītos no dubultas apstrādes

    def on_created(self, event):
        """Apstrādā faila izveides notikumu."""
        if not event.is_directory:
            filepath = event.src_path
            # Pārbauda, vai fails ir attēls vai PDF un nav jau apstrādāts
            if filepath.lower().endswith(('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp',
                                          '.pdf')) and filepath not in self.processed_files:
                self.processed_files.add(filepath)
                # Izsauc galvenās lietotnes metodi, lai apstrādātu failu
                self.app.after(100, lambda: self.app.process_new_scanned_file(filepath))
                # Pēc apstrādes noņem failu no saraksta, lai to varētu apstrādāt vēlreiz, ja tas tiek modificēts/pārsūtīts
                # Pagaida ilgāku laiku, lai nodrošinātu, ka fails ir pilnībā apstrādāts un augšupielādēts
                self.app.after(10000, lambda: self.processed_files.discard(filepath))


def show_document_detection_menu(self):
    """Parāda dokumentu atlases izvēlni"""
    from tkinter import Toplevel
    menu_window = Toplevel(self)
    menu_window.title("Atlasīt dokumentu no attēla")
    menu_window.geometry("300x200")
    menu_window.transient(self)
    menu_window.grab_set()

    # Centrē logu
    menu_window.update_idletasks()
    x = (menu_window.winfo_screenwidth() // 2) - (menu_window.winfo_width() // 2)
    y = (menu_window.winfo_screenheight() // 2) - (menu_window.winfo_height() // 2)
    menu_window.geometry(f"+{x}+{y}")

    ttk.Label(menu_window, text="Izvēlieties dokumenta avotu:",
              font=("Helvetica", 12, "bold")).pack(pady=20)

    ttk.Button(menu_window, text="1. Atvērt foto no sistēmas",
               command=lambda: self.open_photo_for_detection(menu_window),
               bootstyle="primary").pack(pady=10, padx=20, fill="x")

    ttk.Button(menu_window, text="2. Bildēt foto ar kameru",
               command=lambda: self.capture_photo_for_detection(menu_window),
               bootstyle="success").pack(pady=10, padx=20, fill="x")


def open_photo_for_detection(self, parent_window):
    """Atver foto no sistēmas dokumenta atlasei"""
    parent_window.destroy()

    filepath = filedialog.askopenfilename(
        title="Izvēlieties foto dokumenta atlasei",
        filetypes=[("Attēli", "*.png *.jpg *.jpeg *.tif *.tiff *.bmp"), ("Visi faili", "*.*")]
    )

    if filepath:
        try:
            img = Image.open(filepath)
            self.process_image_for_document_detection(img, filepath)
        except Exception as e:
            messagebox.showerror("Kļūda", f"Neizdevās ielādēt attēlu: {e}")


def capture_photo_for_detection(self, parent_window):
    """Bildē foto ar kameru dokumenta atlasei"""
    parent_window.destroy()
    messagebox.showinfo("Info", "Kameras funkcija tiks pievienota nākamajā versijā")


def process_image_for_document_detection(self, image, source_path):
    """Apstrādā attēlu dokumenta atlasei"""
    if not OPENCV_AVAILABLE:
        messagebox.showwarning("Trūkst bibliotēkas",
                               "Dokumenta atlasei nepieciešams 'opencv-python'.")
        return

    try:
        # Konvertē uz OpenCV formātu
        img_cv = np.array(image)
        if len(img_cv.shape) == 3:
            img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

        # Dokumenta kontūras atrašana
        detected_corners = self.detect_document_corners(img_cv)

        if detected_corners is not None:
            # Parāda rezultātu lietotājam apstiprināšanai
            self.show_detection_result(image, detected_corners, source_path)
        else:
            # Ja automātiskā atlase neizdevās, ļauj lietotājam manuāli iezīmēt
            self.manual_corner_selection(image, source_path)

    except Exception as e:
        messagebox.showerror("Kļūda", f"Dokumenta atlases kļūda: {e}")


def detect_document_corners(self, img_cv):
    """Atrod dokumenta stūrus attēlā"""
    try:
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        edged = cv2.Canny(blurred, 75, 200)

        contours, _ = cv2.findContours(edged, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        contours = sorted(contours, key=cv2.contourArea, reverse=True)

        for contour in contours:
            epsilon = 0.02 * cv2.arcLength(contour, True)
            approx = cv2.approxPolyDP(contour, epsilon, True)

            if len(approx) == 4 and cv2.contourArea(contour) > 10000:
                return approx.reshape(4, 2)

        return None
    except Exception as e:
        print(f"Dokumenta stūru atrašanas kļūda: {e}")
        return None


def show_detection_result(self, image, corners, source_path):
    """Parāda dokumenta atlases rezultātu apstiprināšanai"""
    from tkinter import Toplevel
    result_window = Toplevel(self)
    result_window.title("Dokumenta atlases rezultāts")
    result_window.geometry("800x700")
    result_window.transient(self)
    result_window.grab_set()

    # Attēla kanvass
    canvas = tk.Canvas(result_window, bg="gray")
    canvas.pack(fill="both", expand=True, padx=10, pady=10)

    # Pogu rāmis
    button_frame = ttk.Frame(result_window)
    button_frame.pack(fill="x", pady=10)

    # Zīmē attēlu ar iezīmētiem stūriem
    img_with_corners = np.array(image)
    if len(img_with_corners.shape) == 3:
        img_cv = cv2.cvtColor(img_with_corners, cv2.COLOR_RGB2BGR)
    else:
        img_cv = img_with_corners

    # Zīmē kontūru
    cv2.drawContours(img_cv, [corners], -1, (0, 255, 0), 3)
    for corner in corners:
        cv2.circle(img_cv, tuple(corner), 10, (255, 0, 0), -1)

    # Konvertē atpakaļ uz PIL
    img_result = cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB)
    pil_result = Image.fromarray(img_result)

    # Parāda kanvasā
    def show_image():
        result_window.update_idletasks()
        canvas_width = canvas.winfo_width()
        canvas_height = canvas.winfo_height()

        if canvas_width > 1 and canvas_height > 1:
            display_img = pil_result.resize((canvas_width - 20, canvas_height - 20), Image.LANCZOS)
            photo = ImageTk.PhotoImage(display_img)
            canvas.create_image(10, 10, anchor="nw", image=photo)
            canvas.image = photo

    result_window.after(100, show_image)

    def accept_detection():
        result_window.destroy()
        self.process_detected_document(image, corners, source_path)

    def manual_selection():
        result_window.destroy()
        self.manual_corner_selection(image, source_path)

    ttk.Label(result_window, text="Vai dokumenta atlase izskatās pareiza?",
              font=("Helvetica", 12)).pack(pady=5)

    ttk.Button(button_frame, text="Jā, izskatās labi", command=accept_detection,
               bootstyle="success").pack(side="left", padx=10)
    ttk.Button(button_frame, text="Nē, izvēlēšos pats", command=manual_selection,
               bootstyle="warning").pack(side="left", padx=10)
    ttk.Button(button_frame, text="Atcelt", command=result_window.destroy,
               bootstyle="danger").pack(side="left", padx=10)


def manual_corner_selection(self, image, source_path):
    """Ļauj lietotājam manuāli izvēlēties dokumenta stūrus"""
    from tkinter import Toplevel
    manual_window = Toplevel(self)
    manual_window.title("Manuāla dokumenta atlase")
    manual_window.geometry("900x700")
    manual_window.transient(self)
    manual_window.grab_set()

    canvas = tk.Canvas(manual_window, bg="gray", cursor="cross")
    canvas.pack(fill="both", expand=True, padx=10, pady=10)

    instruction_frame = ttk.Frame(manual_window)
    instruction_frame.pack(fill="x", pady=5)

    ttk.Label(instruction_frame,
              text="Noklikšķiniet uz 4 dokumenta stūriem secībā: augšā pa kreisi, augšā pa labi, apakšā pa labi, apakšā pa kreisi",
              font=("Helvetica", 10)).pack()

    button_frame = ttk.Frame(manual_window)
    button_frame.pack(fill="x", pady=5)

    selected_corners = []
    corner_circles = []
    scale_x = scale_y = 1.0

    def show_image():
        nonlocal scale_x, scale_y
        manual_window.update_idletasks()
        canvas_width = canvas.winfo_width()
        canvas_height = canvas.winfo_height()

        if canvas_width > 1 and canvas_height > 1:
            display_img = image.resize((canvas_width - 20, canvas_height - 20), Image.LANCZOS)
            photo = ImageTk.PhotoImage(display_img)
            canvas.create_image(10, 10, anchor="nw", image=photo)
            canvas.image = photo

            # Aprēķina mērogošanas faktoru
            scale_x = (canvas_width - 20) / image.width
            scale_y = (canvas_height - 20) / image.height

    manual_window.after(100, show_image)

    def on_canvas_click(event):
        if len(selected_corners) < 4:
            x = event.x - 10
            y = event.y - 10

            # Pārveido uz oriģinālā attēla koordinātām
            orig_x = int(x / scale_x) if scale_x > 0 else x
            orig_y = int(y / scale_y) if scale_y > 0 else y

            selected_corners.append([orig_x, orig_y])

            # Zīmē apli uz kanvasa
            circle = canvas.create_oval(x - 5, y - 5, x + 5, y + 5, fill="red", outline="white", width=2)
            corner_circles.append(circle)

            # Zīmē numuru
            canvas.create_text(x, y - 15, text=str(len(selected_corners)), fill="white", font=("Arial", 12, "bold"))

            if len(selected_corners) == 4:
                # Aktivizē apstrādes pogu
                process_btn.config(state="normal")

    def reset_selection():
        selected_corners.clear()
        for circle in corner_circles:
            canvas.delete(circle)
        corner_circles.clear()
        canvas.delete("text")
        process_btn.config(state="disabled")

    def process_manual_selection():
        if len(selected_corners) == 4:
            manual_window.destroy()
            corners_array = np.array(selected_corners, dtype=np.float32)
            self.process_detected_document(image, corners_array, source_path)

    # Pievieno klikšķa notikumu
    canvas.bind("<Button-1>", on_canvas_click)

    # Pogas
    ttk.Button(button_frame, text="Atiestatīt", command=reset_selection,
               bootstyle="warning").pack(side="left", padx=5)

    process_btn = ttk.Button(button_frame, text="Apstrādāt dokumentu",
                             command=process_manual_selection,
                             bootstyle="success", state="disabled")
    process_btn.pack(side="left", padx=5)

    ttk.Button(button_frame, text="Atcelt", command=manual_window.destroy,
               bootstyle="danger").pack(side="left", padx=5)


def process_detected_document(self, image, corners, source_path):
    """Apstrādā dokumentu pēc stūru noteikšanas"""
    try:
        # Vienkāršots risinājums - vienkārši ielādē attēlu programmā
        self.current_image = image
        self.current_image_path = source_path
        self.original_image = image.copy()

        # Atjaunina attēla parādīšanu
        self.display_image()

        # Pārslēdzas uz attēla apstrādes cilni
        self.notebook.select(self.image_processing_tab)

        messagebox.showinfo("Sekmīgi", "Dokuments veiksmīgi ielādēts!")

    except Exception as e:
        messagebox.showerror("Kļūda", f"Dokumenta apstrādes kļūda: {e}")


def order_corners(self, corners):
    """Sakārto stūrus pareizā secībā"""
    # Aprēķina centru
    center_x = np.mean(corners[:, 0])
    center_y = np.mean(corners[:, 1])

    # Klasificē stūrus pēc pozīcijas attiecībā pret centru
    top_left = None
    top_right = None
    bottom_right = None
    bottom_left = None

    for corner in corners:
        x, y = corner
        if x < center_x and y < center_y:
            top_left = corner
        elif x > center_x and y < center_y:
            top_right = corner
        elif x > center_x and y > center_y:
            bottom_right = corner
        elif x < center_x and y > center_y:
            bottom_left = corner

    # Ja kāds stūris nav atrasts, izmanto tuvāko
    ordered_corners = []
    for target_corner in [top_left, top_right, bottom_right, bottom_left]:
        if target_corner is not None:
            ordered_corners.append(target_corner)
        else:
            # Atrod tuvāko neizmantoto stūri
            remaining_corners = [c for c in corners if not any(np.array_equal(c, oc) for oc in ordered_corners)]
            if remaining_corners:
                ordered_corners.append(remaining_corners[0])

    return np.array(ordered_corners, dtype=np.float32)


def correct_perspective(self, img, corners):
    """Veic perspektīvas korekciju"""
    try:
        # Aprēķina jauno attēla izmēru
        width_top = np.linalg.norm(corners[1] - corners[0])
        width_bottom = np.linalg.norm(corners[2] - corners[3])
        width = int(max(width_top, width_bottom))

        height_left = np.linalg.norm(corners[3] - corners[0])
        height_right = np.linalg.norm(corners[2] - corners[1])
        height = int(max(height_left, height_right))

        # Mērķa punkti (taisnstūris)
        dst_corners = np.array([
            [0, 0],
            [width, 0],
            [width, height],
            [0, height]
        ], dtype=np.float32)

        # Perspektīvas transformācijas matrica
        matrix = cv2.getPerspectiveTransform(corners, dst_corners)

        # Veic transformāciju
        corrected = cv2.warpPerspective(img, matrix, (width, height))

        return corrected

    except Exception as e:
        print(f"Perspektīvas korekcijas kļūda: {e}")
        return None


def load_processed_image(self, image, source_path):
    """Ielādē apstrādāto attēlu programmā"""
    try:
        # Saglabā attēlu pagaidu failā
        import tempfile
        import os

        temp_dir = tempfile.gettempdir()
        temp_filename = f"processed_document_{int(time.time())}.png"
        temp_path = os.path.join(temp_dir, temp_filename)

        image.save(temp_path, "PNG")

        # Ielādē attēlu programmā (izmanto esošo funkcionalitāti)
        self.current_image = image
        self.current_image_path = temp_path
        self.original_image = image.copy()

        # Atjaunina attēla parādīšanu
        self.display_image()

        # Pārslēdzas uz attēla apstrādes cilni
        self.notebook.select(self.image_processing_tab)

    except Exception as e:
        messagebox.showerror("Kļūda", f"Neizdevās ielādēt apstrādāto attēlu: {e}")


def try_detect_document_method1(img_cv):
    """1. metode: Uzlabota adaptīvā sliekšņošana"""
    try:
        gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

        # Uzlabo kontrastu
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)

        # Gausa izplūšana
        blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)

        # Adaptīvā sliekšņošana ar dažādiem parametriem
        for block_size in [11, 15, 19, 23]:
            for c_value in [2, 5, 10]:
                thresh = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                               cv2.THRESH_BINARY, block_size, c_value)

                contour = find_best_contour(thresh, img_cv.shape)
                if contour is not None:
                    return contour

        return None
    except:
        return None


def try_detect_document_method2(img_cv):
    """2. metode: Canny edge detection"""
    try:
        gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

        # Uzlabo kontrastu
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)

        # Gausa izplūšana
        blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)

        # Canny edge detection ar dažādiem sliekšņiem
        for low_thresh in [50, 75, 100]:
            for high_thresh in [150, 200, 250]:
                edges = cv2.Canny(blurred, low_thresh, high_thresh)

                # Morfoloģiskās operācijas, lai aizvērtu pārtraukumus
                kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
                edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel)

                contour = find_best_contour(edges, img_cv.shape)
                if contour is not None:
                    return contour

        return None
    except:
        return None


def try_detect_document_method3(img_cv):
    """3. metode: Morfoloģiskās operācijas"""
    try:
        gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)

        # Binārizācija ar Otsu metodi
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Morfoloģiskās operācijas
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))

        # Closing - aizvērt mazos caurums
        closed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel, iterations=2)

        # Opening - noņemt troksni
        opened = cv2.morphologyEx(closed, cv2.MORPH_OPEN, kernel, iterations=1)

        contour = find_best_contour(opened, img_cv.shape)
        if contour is not None:
            return contour

        # Mēģinam ar invertētu attēlu
        inverted = cv2.bitwise_not(opened)
        contour = find_best_contour(inverted, img_cv.shape)
        return contour

    except:
        return None


def try_detect_document_method4(img_cv):
    """4. metode: Krāsu segmentācija"""
    try:
        # Konvertē uz HSV krāsu telpu
        hsv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2HSV)

        # Definē baltās krāsas diapazonu (dokumenti bieži ir balti)
        lower_white = np.array([0, 0, 180])
        upper_white = np.array([180, 30, 255])

        # Izveido masku baltajām krāsām
        white_mask = cv2.inRange(hsv, lower_white, upper_white)

        # Morfoloģiskās operācijas
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
        white_mask = cv2.morphologyEx(white_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
        white_mask = cv2.morphologyEx(white_mask, cv2.MORPH_OPEN, kernel, iterations=1)

        contour = find_best_contour(white_mask, img_cv.shape)
        if contour is not None:
            return contour

        # Mēģinam ar plašāku krāsu diapazonu
        lower_light = np.array([0, 0, 120])
        upper_light = np.array([180, 50, 255])
        light_mask = cv2.inRange(hsv, lower_light, upper_light)

        light_mask = cv2.morphologyEx(light_mask, cv2.MORPH_CLOSE, kernel, iterations=2)
        contour = find_best_contour(light_mask, img_cv.shape)
        return contour

    except:
        return None


def find_best_contour(binary_img, img_shape):
    """Atrod labāko kontūru, kas varētu būt dokuments"""
    try:
        contours, _ = cv2.findContours(binary_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        if not contours:
            return None

        img_area = img_shape[0] * img_shape[1]

        # Sakārto kontūras pēc laukuma
        contours = sorted(contours, key=cv2.contourArea, reverse=True)

        for contour in contours:
            area = cv2.contourArea(contour)

            # Pārbauda minimālo laukumu (vismaz 5% no attēla)
            if area < img_area * 0.05:
                continue

            # Pārbauda maksimālo laukumu (ne vairāk kā 95% no attēla)
            if area > img_area * 0.95:
                continue

            # Aproksimē kontūru
            peri = cv2.arcLength(contour, True)

            # Mēģina ar dažādiem epsilon parametriem
            for epsilon_factor in [0.01, 0.02, 0.03, 0.04, 0.05]:
                approx = cv2.approxPolyDP(contour, epsilon_factor * peri, True)

                # Ja ir 4 stūri
                if len(approx) == 4:
                    # Pārbauda, vai stūri veido saprātīgu taisnstūri
                    if is_valid_rectangle(approx, img_shape):
                        return approx

                # Ja ir vairāk nekā 4 stūri, mēģina atrast 4 galvenos
                elif len(approx) > 4:
                    # Atrod 4 galvenos stūrus
                    rect_corners = find_four_corners(approx)
                    if rect_corners is not None and is_valid_rectangle(rect_corners, img_shape):
                        return rect_corners

        return None
    except:
        return None


def is_valid_rectangle(corners, img_shape):
    """Pārbauda, vai 4 punkti veido derīgu taisnstūri"""
    try:
        if len(corners) != 4:
            return False

        # Pārbauda, vai visi punkti ir attēla robežās
        h, w = img_shape[:2]
        for corner in corners:
            x, y = corner[0]
            if x < 0 or x >= w or y < 0 or y >= h:
                return False

        # Aprēķina laukumu
        area = cv2.contourArea(corners)
        img_area = h * w

        # Pārbauda laukuma attiecību
        if area < img_area * 0.05 or area > img_area * 0.95:
            return False

        # Pārbauda, vai forma ir pietiekami taisnstūrveida
        # Aprēķina convex hull un salīdzina laukumus
        hull = cv2.convexHull(corners)
        hull_area = cv2.contourArea(hull)

        if hull_area > 0:
            solidity = area / hull_area
            if solidity < 0.8:  # Ja forma nav pietiekami "cieta"
                return False

        return True
    except:
        return False


def find_four_corners(contour):
    """Atrod 4 galvenos stūrus no kontūras ar vairāk punktiem"""
    try:
        # Atrod kontūras bounding rectangle
        rect = cv2.minAreaRect(contour)
        box = cv2.boxPoints(rect)
        box = np.int0(box)

        return box.reshape(4, 1, 2)
    except:
        return None


def apply_perspective_transform(img_cv, document_contour):
    """Pielieto perspektīvas transformāciju"""
    try:
        def order_points(pts):
            """Sakārto punktus: augšējais kreisais, augšējais labais, apakšējais labais, apakšējais kreisais"""
            rect = np.zeros((4, 2), dtype="float32")

            # Summa: augšējais kreisais būs mazākā, apakšējais labais - lielākā
            s = pts.sum(axis=1)
            rect[0] = pts[np.argmin(s)]  # Augšējais kreisais
            rect[2] = pts[np.argmax(s)]  # Apakšējais labais

            # Starpība: augšējais labais būs mazākā, apakšējais kreisais - lielākā
            diff = np.diff(pts, axis=1)
            rect[1] = pts[np.argmin(diff)]  # Augšējais labais
            rect[3] = pts[np.argmax(diff)]  # Apakšējais kreisais

            return rect

        # Sakārto stūrus
        corners = document_contour.reshape(4, 2)
        ordered_corners = order_points(corners)

        # Aprēķina jaunā attēla izmērus
        (tl, tr, br, bl) = ordered_corners

        # Platums
        widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
        widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
        maxWidth = max(int(widthA), int(widthB))

        # Augstums
        heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
        heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
        maxHeight = max(int(heightA), int(heightB))

        # Pārbauda, vai izmēri ir saprātīgi
        if maxWidth < 50 or maxHeight < 50:
            print("Pārāk mazi izmēri")
            return None

        if maxWidth > 5000 or maxHeight > 5000:
            # Ierobežo maksimālos izmērus
            ratio = min(5000 / maxWidth, 5000 / maxHeight)
            maxWidth = int(maxWidth * ratio)
            maxHeight = int(maxHeight * ratio)

        # Definē mērķa punktus (taisnstūrveida dokuments)
        dst = np.array([
            [0, 0],
            [maxWidth - 1, 0],
            [maxWidth - 1, maxHeight - 1],
            [0, maxHeight - 1]], dtype="float32")

        # Aprēķina perspektīvas transformācijas matricu
        M = cv2.getPerspectiveTransform(ordered_corners, dst)

        # Pielieto perspektīvas transformāciju
        warped = cv2.warpPerspective(img_cv, M, (maxWidth, maxHeight))

        # Konvertē atpakaļ uz PIL attēlu
        processed_img_pil = Image.fromarray(warped)

        return processed_img_pil

    except Exception as e:
        print(f"Kļūda perspektīvas transformācijā: {e}")
        return None


if __name__ == "__main__":
    # Pārbauda, vai ir instalēti nepieciešamie moduļi
    try:
        from PIL import ImageFilter, ImageChops  # ImageChops priekš attēlu salīdzināšanas
        import urllib.parse  # Nepieciešams e-pasta pielikumiem
        from docx import Document  # Priekš Word dokumentiem
    except ImportError as e:
        messagebox.showwarning("Trūkst bibliotēku", f"Dažas nepieciešamās bibliotēkas nav instalētas: {str(e)}\n"
                                                    "Lūdzu, instalējiet trūkstošās bibliotēkas (pip install pillow python-docx).")
    # Palaiž galveno lietotni
    app = OCRPDFApp()
    app.protocol("WM_DELETE_WINDOW", app.on_closing)
    app.mainloop()
